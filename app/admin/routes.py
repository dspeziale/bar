import secrets as _secrets
from datetime import date, timedelta, datetime as _dt
from functools import wraps
from flask import (render_template, redirect, url_for, flash, request, abort,
                   jsonify, current_app)
from flask_login import login_required, current_user
from app import db, tables_enabled, cesto_enabled, numero_italiano
from app.admin import bp
from app.models import (User, Product, Category, Order, OrderItem,
                        TimeSlot, DailyStock,
                        IngredientCategory, Ingredient,
                        Table, TableReservation, TableTimeBand,
                        Permission, Role, AppSetting, Poll, PollChoice, PollVote,
                        Tenant, Supplier, ConsumableItem, ConsumableMovement,
                        CorporateAccount, CorporateMembership,
                        DailyFixedMeal, CorporateMealBooking, MealConfiguration,
                        Transaction, CustomOrderItem, CustomOrderItemIngredient,
                        BancoItem, BancoSession, PrepLabel,
                        Prenotazione, PrenotazioneItem,
                        CaricoMensile,
                        ALLERGENS)
from app.notifications import (send_telegram, send_telegram_to_user, send_web_push_to_user,
                                send_email_to_all_users, send_supplier_low_stock_alert,
                                send_account_activated_email,
                                telegram_poll_message, email_poll_html, get_setting)


# ── Tenant scope helpers ──────────────────────────────────────────────────────

def _tenant_filter():
    """Per query filter_by: vuoto per il super admin globale, tenant_id per gli altri."""
    if current_user.is_admin:
        return {}
    return {'tenant_id': current_user.tenant_id}


def _active_tenant_id():
    """Ritorna il tenant_id da usare per categorie/prodotti.
    Super admin (tenant_id=None) → tenant 'default'.
    """
    if current_user.tenant_id:
        return current_user.tenant_id
    default_t = Tenant.query.filter_by(slug='default').first()
    return default_t.id if default_t else None


# ── User cascade delete ───────────────────────────────────────────────────────

def _delete_user_cascade(uid):
    """Elimina un utente e tutti i record collegati rispettando i vincoli FK NOT NULL."""
    # ordini: prima i sotto-record, poi gli ordini stessi
    # scalar_subquery() è richiesto da SQLAlchemy 2.x quando usato dentro .in_()
    order_ids = db.session.query(Order.id).filter_by(user_id=uid).scalar_subquery()
    coi_ids   = db.session.query(CustomOrderItem.id).filter(
                    CustomOrderItem.order_id.in_(order_ids)).scalar_subquery()

    CustomOrderItemIngredient.query.filter(
        CustomOrderItemIngredient.custom_item_id.in_(coi_ids)
    ).delete(synchronize_session=False)

    CustomOrderItem.query.filter(
        CustomOrderItem.order_id.in_(order_ids)
    ).delete(synchronize_session=False)

    OrderItem.query.filter(
        OrderItem.order_id.in_(order_ids)
    ).delete(synchronize_session=False)

    # transazioni che referenziano ordini di questo utente: annulla il FK nullable
    Transaction.query.filter(
        Transaction.order_id.in_(order_ids)
    ).update({'order_id': None}, synchronize_session=False)

    Order.query.filter_by(user_id=uid).delete(synchronize_session=False)

    # sessioni banco: customer_id è nullable → azzera; staff_id NOT NULL → elimina
    BancoSession.query.filter_by(customer_id=uid).update(
        {'customer_id': None}, synchronize_session=False)
    BancoSession.query.filter_by(staff_id=uid).delete(synchronize_session=False)

    # tutti gli altri record collegati direttamente all'utente
    Transaction.query.filter_by(user_id=uid).delete(synchronize_session=False)
    TableReservation.query.filter_by(user_id=uid).delete(synchronize_session=False)
    PollVote.query.filter_by(user_id=uid).delete(synchronize_session=False)
    CorporateMealBooking.query.filter_by(user_id=uid).delete(synchronize_session=False)
    CorporateMembership.query.filter_by(user_id=uid).delete(synchronize_session=False)

    user = db.session.get(User, uid)
    if user:
        db.session.delete(user)


# ── Decorators ────────────────────────────────────────────────────────────────

def staff_required(f):
    """Admin o qualsiasi utente con almeno un permesso backoffice."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('auth.login'))
        if not (current_user.is_admin or current_user.is_staff):
            flash('Accesso riservato al personale.', 'danger')
            return redirect(url_for('main.index'))
        return f(*args, **kwargs)
    return login_required(decorated)


def tables_required(f):
    """Blocca la rotta se la gestione tavoli e' disattivata da Impostazioni."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not tables_enabled():
            flash('La gestione tavoli e\' disattivata nelle Impostazioni.', 'warning')
            return redirect(url_for('admin.dashboard'))
        return f(*args, **kwargs)
    return decorated


def cesto_required(f):
    """Blocca la rotta se la gestione del cesto e' disattivata da Impostazioni."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not cesto_enabled():
            flash('La gestione del cesto e\' disattivata nelle Impostazioni.',
                  'warning')
            return redirect(url_for('admin.dashboard'))
        return f(*args, **kwargs)
    return decorated


def require_permission(perm_name):
    """Richiede un permesso specifico; admin bypassa sempre."""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('auth.login'))
            if not current_user.has_permission(perm_name):
                abort(403)
            return f(*args, **kwargs)
        return login_required(decorated_function)
    return decorator


# ── Dashboard ─────────────────────────────────────────────────────────────────

@bp.route('/')
@staff_required
def dashboard():
    today = date.today()
    orders_today = Order.query.filter_by(order_date=today)\
        .filter(Order.status != 'cancelled').all()
    revenue_today = sum(o.total_price for o in orders_today)
    clients_count = User.query.filter_by(is_client=True, is_active=True, **_tenant_filter()).count()
    staff_count   = User.query.filter_by(is_client=False, is_admin=False, **_tenant_filter()).count()
    products_count = Product.query.filter_by(is_active=True).count()
    res_today = TableReservation.query.filter_by(reservation_date=today)\
        .filter(TableReservation.status != 'cancelled').count()
    stock_alerts = [(p, p.available_today())
                    for p in Product.query.filter_by(is_active=True).all()
                    if p.available_today() <= 3]
    wallet_users = User.query.filter_by(is_client=True, is_active=True, **_tenant_filter()).all()
    total_wallet = round(sum(u.wallet_balance for u in wallet_users), 2)
    consumable_alerts = ConsumableItem.query.filter_by(**_tenant_filter())\
        .filter(ConsumableItem.alert_active == True).count()
    # Panoramica pasti aziendali — tutte le opzioni di oggi
    corp_meals_today = DailyFixedMeal.query.filter_by(meal_date=today)\
        .order_by(DailyFixedMeal.corporate_id, DailyFixedMeal.name).all()

    # Clienti in attesa di attivazione
    pending_clients = User.query.filter_by(
        is_client=True, is_active=False, **_tenant_filter()
    ).count()

    # Sondaggio attivo più recente + andamento voti
    active_poll = Poll.query.filter_by(is_active=True)\
        .order_by(Poll.poll_date.desc()).first()
    poll_stats = None
    if active_poll:
        total_votes = active_poll.total_votes()
        poll_stats = {
            'poll': active_poll,
            'total': total_votes,
            'choices': [
                {'text': c.text, 'emoji': c.emoji,
                 'count': c.vote_count,
                 'pct': round(c.vote_count / total_votes * 100, 1) if total_votes else 0}
                for c in sorted(active_poll.choices, key=lambda x: x.vote_count, reverse=True)
            ],
        }

    today_meal_booking    = None
    meal_booking_reminder = False
    membership = getattr(current_user, 'corporate_membership', None)
    if membership and membership.is_active:
        today_meals_corp = DailyFixedMeal.query.filter_by(
            corporate_id=membership.corporate_id,
            meal_date=today,
            is_active=True,
        ).all()
        if today_meals_corp:
            meal_ids = [m.id for m in today_meals_corp]
            today_meal_booking = CorporateMealBooking.query.filter(
                CorporateMealBooking.user_id == current_user.id,
                CorporateMealBooking.meal_id.in_(meal_ids),
                CorporateMealBooking.status != 'cancelled',
            ).first()
            if not today_meal_booking:
                meal_booking_reminder = True

    return render_template('admin/dashboard.html',
                           orders_today=orders_today,
                           pending=[o for o in orders_today if o.status in ('pending', 'confirmed')],
                           revenue_today=revenue_today,
                           clients_count=clients_count,
                           staff_count=staff_count,
                           products_count=products_count,
                           res_today=res_today,
                           stock_alerts=stock_alerts,
                           total_wallet=total_wallet,
                           consumable_alerts=consumable_alerts,
                           today_meal_booking=today_meal_booking,
                           meal_booking_reminder=meal_booking_reminder,
                           corp_meals_today=corp_meals_today,
                           poll_stats=poll_stats,
                           pending_clients=pending_clients)


# ── Prodotti ──────────────────────────────────────────────────────────────────

@bp.route('/products')
@require_permission('manage_products')
def products():
    tid = _active_tenant_id()
    return render_template('admin/products.html',
                           products=Product.query.filter_by(tenant_id=tid).order_by(Product.category_id, Product.name).all(),
                           categories=Category.query.filter_by(tenant_id=tid).order_by(Category.name).all(),
                           allergens=ALLERGENS)


@bp.route('/products/dt')
@require_permission('manage_products')
def products_dt():
    tid    = _active_tenant_id()
    draw   = request.args.get('draw', 1, type=int)
    start  = request.args.get('start', 0, type=int)
    length = request.args.get('length', 25, type=int)
    search = (request.args.get('search[value]') or '').strip()
    col    = request.args.get('order[0][column]', 0, type=int)
    dirn   = request.args.get('order[0][dir]', 'asc')

    q = Product.query.join(Category).filter(Product.tenant_id == tid)
    total = q.count()

    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(
            Product.name.ilike(like),
            Product.description.ilike(like),
            Category.name.ilike(like),
        ))
    filtered = q.count()

    col_map = {0: Product.name, 1: Category.name, 2: Product.price,
               3: Product.daily_quantity, 5: Product.is_active}
    order_expr = col_map.get(col, Product.name)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())

    data = []
    for p in q.offset(start).limit(length).all():
        data.append({
            'id':             p.id,
            'name':           p.name,
            'description':    p.description or '',
            'category_id':    p.category_id,
            'category_name':  p.category.name,
            'category_color': p.category.color,
            'category_icon':  p.category.icon,
            'price':          p.price,
            'daily_quantity': p.daily_quantity,
            'allergens':      p.allergens or '',
            'allergen_list':  [[k, l, i] for k, l, i in p.allergen_list],
            'is_active':      p.is_active,
            'barcode':        p.barcode or '',
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── Helper DT params ──────────────────────────────────────────────────────────
def _dt_params():
    draw   = request.args.get('draw', 1, type=int)
    start  = request.args.get('start', 0, type=int)
    length = request.args.get('length', 25, type=int)
    search = (request.args.get('search[value]') or '').strip()
    col    = request.args.get('order[0][column]', 0, type=int)
    dirn   = request.args.get('order[0][dir]', 'asc')
    return draw, start, length, search, col, dirn


# ── CLIENTS DT ────────────────────────────────────────────────────────────────
@bp.route('/clients/dt')
@require_permission('manage_clients')
def clients_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    q = User.query.filter_by(is_client=True, tenant_id=tid)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(
            User.first_name.ilike(like), User.last_name.ilike(like),
            User.email.ilike(like), User.phone.ilike(like), User.address.ilike(like),
        ))
    filtered = q.count()
    col_map = {0: db.func.concat(User.last_name, ' ', User.first_name),
               1: User.phone, 2: User.birth_date, 5: User.wallet_balance, 6: User.is_active}
    order_expr = col_map.get(col, db.func.concat(User.last_name, ' ', User.first_name))
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    data = []
    for c in q.offset(start).limit(length).all():
        m = c.corporate_membership
        data.append({
            'id':             c.id,
            'full_name':      c.full_name,
            'first_name':     c.first_name or '',
            'last_name':      c.last_name or '',
            'email':          c.email,
            'phone':          c.phone or '',
            'birth_date':     c.birth_date.strftime('%Y-%m-%d') if c.birth_date else '',
            'birth_date_str': c.birth_date.strftime('%d/%m/%Y') if c.birth_date else '',
            'address':        c.address or '',
            'telegram':       c.telegram_chat_id or '',
            'wallet_balance': round(c.wallet_balance, 2),
            'wallet_overdraft': round(c.wallet_overdraft or 0, 2),
            'loyalty_points': c.loyalty_points,
            'is_active':      c.is_active,
            'has_google':     bool(c.google_id),
            'corporate_name': (m.corporate.name if m and m.is_active else ''),
            'toggle_url':     url_for('admin.client_toggle',  uid=c.id),
            'edit_url':       url_for('admin.client_edit',    uid=c.id),
            'topup_url':      url_for('admin.client_topup',   uid=c.id),
            'delete_url':     url_for('admin.client_delete',  uid=c.id),
            'push_test_url':  url_for('admin.push_test',      uid=c.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── ORDERS DT ─────────────────────────────────────────────────────────────────
@bp.route('/orders/dt')
@require_permission('view_orders')
def orders_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    filter_date_str = request.args.get('date', str(date.today()))
    status_filter   = request.args.get('status', '')
    try:
        filter_date_obj = _dt.strptime(filter_date_str, '%Y-%m-%d').date()
    except ValueError:
        filter_date_obj = date.today()
    q = Order.query.join(User, Order.user_id == User.id).filter(
        Order.tenant_id == tid, Order.order_date == filter_date_obj)
    if status_filter:
        q = q.filter(Order.status == status_filter)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(
            Order.order_code.ilike(like), User.username.ilike(like),
            User.first_name.ilike(like), User.last_name.ilike(like),
        ))
    filtered = q.count()
    col_map = {0: Order.created_at, 1: User.username, 5: Order.total_price, 6: Order.status}
    order_expr = col_map.get(col, Order.created_at)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    _STATUS = {'pending':('Ricevuto','warning'), 'confirmed':('Confermato','info'),
               'preparing':('In prep.','primary'), 'ready':('Pronto','success'),
               'completed':('Consegnato','success'), 'cancelled':('Annullato','secondary')}
    data = []
    for o in q.offset(start).limit(length).all():
        lbl = _STATUS.get(o.status, (o.status, 'light'))
        items_parts = []
        for it in o.items:
            items_parts.append(f'{it.quantity}× {it.product.name}')
        for ci in o.custom_items:
            t = '🥪' if ci.builder_type == 'panino' else '🥗'
            items_parts.append(f'{t} Builder')
        data.append({
            'id':          o.id,
            'order_code':  o.order_code or f'#{o.id}',
            'created_at':  o.created_at.strftime('%H:%M') if o.created_at else '',
            'username':    o.user.username,
            'full_name':   o.user.full_name,
            'slot':        o.slot.time_str if o.slot else '—',
            'items':       ' · '.join(items_parts),
            'notes':       o.notes or '',
            'total_price': round(o.total_price, 2),
            'status':      o.status,
            'status_label': lbl[0],
            'status_color': lbl[1],
            'status_url':  url_for('admin.order_status', oid=o.id),
            'slip_url':    url_for('admin.order_slip',   oid=o.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── USERS DT ──────────────────────────────────────────────────────────────────
@bp.route('/users/dt')
@require_permission('manage_users')
def users_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    q = User.query.filter_by(is_admin=False, is_client=False, tenant_id=tid)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(User.username.ilike(like), User.email.ilike(like)))
    filtered = q.count()
    col_map = {0: User.email, 2: User.wallet_balance, 4: User.is_active, 5: User.created_at}
    order_expr = col_map.get(col, User.username)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    data = []
    for u in q.offset(start).limit(length).all():
        data.append({
            'id':            u.id,
            'email':         u.email,
            'username':      u.username,
            'roles':         [{'label': r.label, 'color': r.color} for r in u.roles],
            'wallet_balance': round(u.wallet_balance, 2),
            'loyalty_points': u.loyalty_points,
            'is_active':     u.is_active,
            'created_at':    u.created_at.strftime('%d/%m/%Y') if u.created_at else '',
            'toggle_url':     url_for('admin.user_toggle',    uid=u.id),
            'topup_url':      url_for('admin.user_topup',     uid=u.id),
            'delete_url':     url_for('admin.user_delete',    uid=u.id),
            'roles_url':      url_for('admin.user_roles_assign', uid=u.id),
            'push_test_url':  url_for('admin.push_test',     uid=u.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── FORNITORI DT ──────────────────────────────────────────────────────────────
@bp.route('/fornitori/dt')
@require_permission('manage_stock')
def fornitori_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    q = Supplier.query.filter_by(tenant_id=tid)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(Supplier.name.ilike(like), Supplier.email.ilike(like),
                             Supplier.phone.ilike(like), Supplier.notes.ilike(like)))
    filtered = q.count()
    col_map = {0: Supplier.name, 1: Supplier.email, 2: Supplier.phone}
    order_expr = col_map.get(col, Supplier.name)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    data = []
    for s in q.offset(start).limit(length).all():
        data.append({
            'id':         s.id,
            'name':       s.name,
            'email':      s.email or '',
            'phone':      s.phone or '',
            'notes':      s.notes or '',
            'items':      [i.name for i in s.items],
            'edit_url':   url_for('admin.fornitore_edit',   sid=s.id),
            'delete_url': url_for('admin.fornitore_delete', sid=s.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── MAGAZZINO DT ──────────────────────────────────────────────────────────────
@bp.route('/magazzino/dt')
@require_permission('manage_stock')
def magazzino_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    q = ConsumableItem.query.outerjoin(Supplier).filter(ConsumableItem.tenant_id == tid)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(ConsumableItem.name.ilike(like),
                             Supplier.name.ilike(like)))
    filtered = q.count()
    col_map = {0: ConsumableItem.name, 1: ConsumableItem.unit,
               2: ConsumableItem.quantity, 3: ConsumableItem.min_threshold}
    order_expr = col_map.get(col, ConsumableItem.name)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    data = []
    for item in q.offset(start).limit(length).all():
        st = item.stock_status
        data.append({
            'id':            item.id,
            'name':          item.name,
            'unit':          item.unit,
            'quantity':      round(item.quantity, 1),
            'min_threshold': round(item.min_threshold, 1),
            'status':        st,
            'supplier_name': item.supplier.name if item.supplier else '',
            'alert_active':  item.alert_active,
            'movement_url':  url_for('admin.magazzino_movimento', iid=item.id),
            'edit_url':      url_for('admin.magazzino_edit',      iid=item.id),
            'delete_url':    url_for('admin.magazzino_delete',    iid=item.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── BANCO ITEMS DT ────────────────────────────────────────────────────────────
@bp.route('/banco/items/dt')
@require_permission('manage_products')
def banco_items_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    q = BancoItem.query.filter_by(tenant_id=tid)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(BancoItem.name.ilike(like))
    filtered = q.count()
    col_map = {0: BancoItem.name, 1: BancoItem.price, 2: BancoItem.sort_order, 3: BancoItem.is_active}
    order_expr = col_map.get(col, BancoItem.sort_order)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    data = []
    for item in q.offset(start).limit(length).all():
        data.append({
            'id':         item.id,
            'name':       item.name,
            'price':      round(item.price, 2),
            'icon':       item.icon,
            'color':      item.color,
            'sort_order': item.sort_order,
            'is_active':  item.is_active,
            'edit_url':   url_for('admin.banco_item_edit',   iid=item.id),
            'delete_url': url_for('admin.banco_item_delete', iid=item.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


# ── CATEGORIES DT ─────────────────────────────────────────────────────────────
@bp.route('/categories/dt')
@require_permission('manage_categories')
def categories_dt():
    tid                    = _active_tenant_id()
    draw, start, length, search, col, dirn = _dt_params()
    q = Category.query.filter_by(tenant_id=tid)
    total = q.count()
    if search:
        like = f'%{search}%'
        q = q.filter(db.or_(Category.name.ilike(like), Category.color.ilike(like)))
    filtered = q.count()
    col_map = {0: Category.name, 3: Category.color}
    order_expr = col_map.get(col, Category.name)
    q = q.order_by(order_expr.desc() if dirn == 'desc' else order_expr.asc())
    data = []
    for cat in q.offset(start).limit(length).all():
        data.append({
            'id':        cat.id,
            'name':      cat.name,
            'icon':      cat.icon,
            'color':     cat.color,
            'n_products': cat.products.count(),
            'edit_url':   url_for('admin.category_edit',   cid=cat.id),
            'delete_url': url_for('admin.category_delete', cid=cat.id),
        })
    return jsonify(draw=draw, recordsTotal=total, recordsFiltered=filtered, data=data)


@bp.route('/products/new', methods=['POST'])
@require_permission('manage_products')
def product_new():
    name = request.form.get('name', '').strip()
    price = request.form.get('price', type=float)
    category_id = request.form.get('category_id', type=int)
    daily_quantity = request.form.get('daily_quantity', type=int, default=20)
    description = request.form.get('description', '').strip()
    allergens = ','.join(request.form.getlist('allergens'))
    barcode   = request.form.get('barcode', '').strip() or None
    if not name or not price or price <= 0 or not category_id:
        flash('Compila tutti i campi obbligatori.', 'danger')
        return redirect(url_for('admin.products'))
    db.session.add(Product(name=name, description=description, price=price,
                           category_id=category_id, daily_quantity=daily_quantity,
                           allergens=allergens, barcode=barcode,
                           tenant_id=_active_tenant_id()))
    db.session.commit()
    flash(f'Prodotto "{name}" aggiunto.', 'success')
    return redirect(url_for('admin.products'))


@bp.route('/products/<int:pid>/edit', methods=['POST'])
@require_permission('manage_products')
def product_edit(pid):
    p = db.get_or_404(Product, pid)
    p.name         = request.form.get('name', p.name).strip()
    p.description  = request.form.get('description', p.description).strip()
    p.price        = request.form.get('price', type=float) or p.price
    p.category_id  = request.form.get('category_id', type=int) or p.category_id
    p.daily_quantity = request.form.get('daily_quantity', type=int) or p.daily_quantity
    p.is_active    = 'is_active' in request.form
    p.allergens    = ','.join(request.form.getlist('allergens'))
    p.barcode      = request.form.get('barcode', '').strip() or None
    db.session.commit()
    flash(f'Prodotto "{p.name}" aggiornato.', 'success')
    return redirect(url_for('admin.products'))


@bp.route('/products/<int:pid>/delete', methods=['POST'])
@require_permission('manage_products')
def product_delete(pid):
    p = db.get_or_404(Product, pid)
    DailyStock.query.filter_by(product_id=pid).delete(synchronize_session=False)
    OrderItem.query.filter_by(product_id=pid).delete(synchronize_session=False)
    name = p.name
    db.session.delete(p)
    db.session.commit()
    flash(f'Prodotto "{name}" eliminato.', 'success')
    return redirect(url_for('admin.products'))


@bp.route('/products/<int:pid>/toggle', methods=['POST'])
@require_permission('manage_products')
def product_toggle(pid):
    p = db.get_or_404(Product, pid)
    p.is_active = not p.is_active
    db.session.commit()
    flash(f'"{p.name}" {"attivato" if p.is_active else "disattivato"}.', 'info')
    return redirect(url_for('admin.products'))


# ── Stock giornaliero ─────────────────────────────────────────────────────────

@bp.route('/stock')
@require_permission('manage_stock')
def stock():
    today = date.today()
    products = Product.query.filter_by(is_active=True).all()
    stock_data = []
    for p in products:
        s = DailyStock.query.filter_by(product_id=p.id, stock_date=today).first()
        if not s:
            s = DailyStock(product_id=p.id, stock_date=today,
                           quantity_available=p.daily_quantity, quantity_reserved=0)
            db.session.add(s)
        stock_data.append((p, s))
    db.session.commit()
    return render_template('admin/stock.html', stock_data=stock_data, today=today)


@bp.route('/stock/<int:pid>/update', methods=['POST'])
@require_permission('manage_stock')
def stock_update(pid):
    today = date.today()
    new_qty = request.form.get('quantity_available', type=int)
    s = DailyStock.query.filter_by(product_id=pid, stock_date=today).first()
    if not s:
        p = db.get_or_404(Product, pid)
        s = DailyStock(product_id=pid, stock_date=today,
                       quantity_available=p.daily_quantity, quantity_reserved=0)
        db.session.add(s)
    if new_qty is not None and new_qty >= 0:
        s.quantity_available = new_qty
    db.session.commit()
    flash('Stock aggiornato.', 'success')
    return redirect(url_for('admin.stock'))


# ── Ordini ────────────────────────────────────────────────────────────────────

@bp.route('/orders')
@require_permission('view_orders')
def orders():
    filter_date = request.args.get('date', str(date.today()))
    try:
        from datetime import datetime as dt
        filter_date_obj = dt.strptime(filter_date, '%Y-%m-%d').date()
    except ValueError:
        filter_date_obj = date.today()
        filter_date = str(filter_date_obj)
    status_filter = request.args.get('status', '')
    q = Order.query.filter_by(order_date=filter_date_obj)
    if status_filter:
        q = q.filter_by(status=status_filter)
    orders = q.order_by(Order.slot_id, Order.created_at).all()
    slots = TimeSlot.query.filter_by(is_active=True).order_by(TimeSlot.time_str).all()
    return render_template('admin/orders.html', orders=orders, slots=slots,
                           filter_date=filter_date, status_filter=status_filter)


@bp.route('/orders/<int:oid>/slip')
@require_permission('view_orders')
def order_slip(oid):
    order = db.get_or_404(Order, oid)
    return render_template('admin/order_slip.html', order=order)


@bp.route('/orders/<int:oid>/status', methods=['POST'])
@require_permission('manage_orders')
def order_status(oid):
    order = db.get_or_404(Order, oid)
    new_status = request.form.get('status')
    if new_status not in ('pending', 'confirmed', 'preparing', 'ready', 'completed', 'cancelled'):
        flash('Stato non valido.', 'danger')
        return redirect(url_for('admin.orders'))
    if new_status == 'cancelled' and order.status != 'cancelled':
        order.user.credit_wallet(order.total_price,
                                 f'Rimborso ordine #{order.id} (admin)', order_id=order.id)
        for item in order.items:
            s = DailyStock.query.filter_by(product_id=item.product_id,
                                           stock_date=date.today()).first()
            if s:
                s.quantity_reserved = max(0, s.quantity_reserved - item.quantity)
    order.status = new_status
    db.session.commit()
    flash(f'Ordine #{order.id} → {new_status}', 'success')
    order_ref = order.order_code or f'#{order.id}'
    if new_status == 'preparing':
        has_grill = any(ci.grill_requested for ci in order.custom_items)
        if has_grill:
            send_telegram(
                f'🔥 <b>PANINO SULLA PIASTRA</b>\n'
                f'Ordine <b>{order_ref}</b> — {order.user.full_name}\n'
                + '\n'.join(
                    f'  • {ci.label}'
                    for ci in order.custom_items if ci.grill_requested
                )
            )
        else:
            send_telegram(
                f'👨‍🍳 Ordine <b>{order_ref}</b> in preparazione\n'
                f'👤 {order.user.full_name}'
            )
        send_telegram_to_user(
            order.user,
            f'👨‍🍳 Il tuo ordine <b>{order_ref}</b> è <b>in preparazione</b>!\n'
            f'Ti avvisiamo quando è pronto.'
        )
    elif new_status == 'ready':
        send_telegram(
            f'🔔 Ordine <b>{order_ref}</b> PRONTO — {order.user.full_name}'
        )
        _items_str = ', '.join(
            f'{i.quantity}× {i.product.name}' for i in order.items
        ) or 'panino custom'
        send_telegram_to_user(
            order.user,
            f'🔔 Il tuo ordine <b>{order_ref}</b> è <b>PRONTO</b> per il ritiro!\n'
            f'Vieni a ritirarlo entro qualche minuto.'
        )
        send_web_push_to_user(
            order.user,
            title='🔔 Il tuo ordine è pronto!',
            body=f'{order_ref} • {_items_str} — Vieni al banco!',
            url='/my-orders'
        )
    elif new_status == 'cancelled':
        send_telegram_to_user(
            order.user,
            f'❌ Ordine <b>{order_ref}</b> annullato dall\'amministratore.\n'
            f'Rimborso di <b>{numero_italiano(order.total_price)}€</b> sul tuo wallet.'
        )
    referer = request.referrer or ''
    if 'cucina' in referer:
        return redirect(url_for('admin.cucina'))
    return redirect(url_for('admin.orders'))


@bp.route('/ordini/<int:oid>/notifica-ritiro', methods=['POST'])
@require_permission('manage_orders')
def order_notifica_ritiro(oid):
    """Invia un promemoria (Telegram + Web Push) al cliente per ritirare l'ordine pronto."""
    order = db.get_or_404(Order, oid)
    if order.status != 'ready':
        return jsonify(ok=False, message='Ordine non in stato pronto.'), 400
    order_ref  = order.order_code or f'#{order.id}'
    items_str  = ', '.join(f'{i.quantity}× {i.product.name}' for i in order.items)
    if order.custom_items:
        items_str = (items_str + ', ' if items_str else '') + 'panino custom'
    send_telegram_to_user(
        order.user,
        f'⏰ <b>Promemoria</b>: il tuo ordine <b>{order_ref}</b> è ancora al banco!\n'
        f'🥗 {items_str}\nVieni a ritirarlo.'
    )
    send_web_push_to_user(
        order.user,
        title='⏰ Il tuo ordine ti aspetta!',
        body=f'{order_ref} • {items_str}',
        url='/my-orders'
    )
    return jsonify(ok=True, message=f'Promemoria inviato a {order.user.full_name}.')


@bp.route('/push-test/<int:uid>', methods=['POST'])
@login_required
def push_test(uid):
    """Invia una notifica Web Push di test a un utente specifico (solo admin)."""
    if not current_user.is_admin:
        return jsonify(ok=False, message='Non autorizzato.'), 403
    from app.models import User, PushSubscription
    user = db.get_or_404(User, uid)
    subs_count = PushSubscription.query.filter_by(user_id=uid).count()
    if subs_count == 0:
        return jsonify(ok=False, message=f'{user.full_name} non ha subscription push registrate.')
    send_web_push_to_user(
        user,
        title='🔔 Test QuickLunch',
        body='Le notifiche push funzionano correttamente!',
        url='/prenotazioni'
    )
    return jsonify(ok=True, message=f'Test push inviato a {user.full_name} ({subs_count} subscription).')


# ── Cucina / KDS ──────────────────────────────────────────────────────────────

@bp.route('/cucina')
@require_permission('view_orders')
def cucina():
    today = date.today()
    da_fare  = Order.query.filter_by(order_date=today, status='confirmed')\
                          .order_by(Order.slot_id, Order.created_at).all()
    in_prep  = Order.query.filter_by(order_date=today, status='preparing')\
                          .order_by(Order.slot_id, Order.created_at).all()
    pronti   = Order.query.filter_by(order_date=today, status='ready')\
                          .order_by(Order.slot_id, Order.created_at).all()
    return render_template('admin/cucina.html',
                           orders_da_fare=da_fare,
                           orders_in_prep=in_prep,
                           orders_pronti=pronti,
                           today=today)


# ── Utenti ────────────────────────────────────────────────────────────────────

@bp.route('/users')
@require_permission('manage_users')
def users():
    all_roles = Role.query.order_by(Role.name).all()
    all_users = User.query.filter_by(is_admin=False, is_client=False,
                                     **_tenant_filter()).order_by(User.username).all()
    return render_template('admin/users.html', users=all_users, all_roles=all_roles)


# ── Clienti ───────────────────────────────────────────────────────────────────

@bp.route('/clients/registration-qr')
@require_permission('manage_clients')
def clients_registration_qr():
    """Pagina stampabile con QR code per la registrazione clienti."""
    from flask import request as _req
    join_url = _req.host_url.rstrip('/') + url_for('auth.join')
    tenant = Tenant.query.filter_by(slug='default').first()
    return render_template('admin/registration_qr.html',
                           join_url=join_url, tenant=tenant)


@bp.route('/clients')
@require_permission('manage_clients')
def clients():
    all_clients = User.query.filter_by(is_client=True,
                                       **_tenant_filter()).order_by(User.last_name, User.first_name).all()
    corporates  = CorporateAccount.query.filter_by(**_tenant_filter()).order_by(CorporateAccount.name).all()
    return render_template('admin/clients.html', clients=all_clients, corporates=corporates)


@bp.route('/clients/new', methods=['POST'])
@require_permission('manage_clients')
def client_new():
    import re
    from datetime import datetime as dt
    email      = request.form.get('email', '').strip().lower()
    first_name = request.form.get('first_name', '').strip()
    last_name  = request.form.get('last_name', '').strip()
    if not email or '@' not in email or not first_name or not last_name:
        flash('Email, nome e cognome sono obbligatori.', 'danger')
        return redirect(url_for('admin.clients'))
    if User.query.filter_by(email=email).first():
        flash(f'Email "{email}" già registrata.', 'warning')
        return redirect(url_for('admin.clients'))
    base = re.sub(r'[^a-z0-9]', '.', email.split('@')[0]).strip('.') or 'cliente'
    username, n = base[:30], 1
    while User.query.filter_by(username=username).first():
        username = f'{base[:28]}{n}'; n += 1
    birth_raw = request.form.get('birth_date', '').strip()
    birth_date = None
    if birth_raw:
        try:
            birth_date = dt.strptime(birth_raw, '%Y-%m-%d').date()
        except ValueError:
            pass
    u = User(
        username=username, email=email,
        is_client=True,
        first_name=first_name,
        last_name=last_name,
        phone=request.form.get('phone', '').strip(),
        birth_date=birth_date,
        address=request.form.get('address', '').strip(),
        telegram_chat_id=request.form.get('telegram_chat_id', '').strip(),
    )
    pwd = request.form.get('password', '').strip()
    if pwd:
        u.set_password(pwd)
    db.session.add(u)
    db.session.commit()
    u.apply_registration_bonus()
    flash(f'Cliente "{u.full_name}" creato.', 'success')
    return redirect(url_for('admin.clients'))


@bp.route('/clients/<int:uid>/edit', methods=['POST'])
@require_permission('manage_clients')
def client_edit(uid):
    from datetime import datetime as dt
    u = db.get_or_404(User, uid)
    if not u.is_client:
        abort(404)
    u.first_name       = request.form.get('first_name', u.first_name).strip()
    u.last_name        = request.form.get('last_name',  u.last_name).strip()
    u.phone            = request.form.get('phone',      u.phone or '').strip()
    u.address          = request.form.get('address',    u.address or '').strip()
    u.telegram_chat_id = request.form.get('telegram_chat_id', u.telegram_chat_id or '').strip()
    birth_raw = request.form.get('birth_date', '').strip()
    if birth_raw:
        try:
            u.birth_date = dt.strptime(birth_raw, '%Y-%m-%d').date()
        except ValueError:
            pass
    overdraft_raw = request.form.get('wallet_overdraft', '').strip()
    if overdraft_raw != '':
        try:
            u.wallet_overdraft = max(0.0, round(float(overdraft_raw), 2))
        except ValueError:
            pass
    pwd = request.form.get('password', '').strip()
    if pwd:
        u.set_password(pwd)
    db.session.commit()
    flash(f'Cliente "{u.full_name}" aggiornato.', 'success')
    return redirect(url_for('admin.clients'))


@bp.route('/clients/<int:uid>/toggle', methods=['POST'])
@require_permission('manage_clients')
def client_toggle(uid):
    u = db.get_or_404(User, uid)
    if not u.is_client:
        abort(404)
    activating = not u.is_active
    u.is_active = activating
    if activating:
        corporate_id = request.form.get('corporate_id', '').strip()
        if corporate_id:
            try:
                cid = int(corporate_id)
                mem = CorporateMembership.query.filter_by(user_id=u.id).first()
                if mem:
                    mem.corporate_id = cid
                    mem.is_active    = True
                else:
                    db.session.add(CorporateMembership(user_id=u.id, corporate_id=cid, is_active=True))
            except (ValueError, TypeError):
                pass
    db.session.commit()
    if activating:
        send_telegram_to_user(u,
            '✅ Il tuo account è stato attivato!\n'
            'Puoi ora accedere al servizio.')
        # L'email e' l'unico canale che raggiunge un cliente appena registrato,
        # che non ha ancora collegato Telegram.
        send_account_activated_email(
            u, login_url=url_for('auth.login', _external=True))
    else:
        send_telegram_to_user(u,
            '🔒 Il tuo account è stato sospeso.\n'
            'Contatta l\'amministratore per informazioni.')
    flash(f'Cliente {u.full_name} {"attivato" if activating else "sospeso"}.', 'info')
    return redirect(url_for('admin.clients'))


@bp.route('/clients/<int:uid>/delete', methods=['POST'])
@require_permission('manage_clients')
def client_delete(uid):
    u = db.get_or_404(User, uid)
    if not u.is_client:
        abort(404)
    name = u.full_name
    _delete_user_cascade(uid)
    db.session.commit()
    flash(f'Cliente "{name}" eliminato.', 'info')
    return redirect(url_for('admin.clients'))


@bp.route('/clients/<int:uid>/topup', methods=['POST'])
@require_permission('manage_clients')
def client_topup(uid):
    u = db.get_or_404(User, uid)
    if not u.is_client:
        abort(404)
    amount = request.form.get('amount', type=float)
    note = request.form.get('note', 'Ricarica manuale').strip() or 'Ricarica manuale'
    if not amount or amount <= 0:
        flash('Importo non valido.', 'danger')
        return redirect(url_for('admin.clients'))
    u.credit_wallet(amount, note)
    db.session.commit()
    flash(f'+{numero_italiano(amount)}€ aggiunti al wallet di {u.full_name}.', 'success')
    send_telegram_to_user(u,
        f'💳 Ricarica wallet: <b>+{numero_italiano(amount)}€</b>\n'
        f'💰 Saldo attuale: <b>{numero_italiano(u.wallet_balance)}€</b>\n'
        f'📝 {note}'
    )
    return redirect(url_for('admin.clients'))


@bp.route('/users/new', methods=['POST'])
@require_permission('manage_users')
def user_new():
    import re
    email    = request.form.get('email', '').strip().lower()
    password = request.form.get('password', '').strip()
    if not email or '@' not in email or not password:
        flash('Email e password obbligatori.', 'danger')
        return redirect(url_for('admin.users'))
    if User.query.filter_by(email=email).first():
        flash(f'Email "{email}" già registrata.', 'warning')
        return redirect(url_for('admin.users'))
    base = re.sub(r'[^a-z0-9]', '.', email.split('@')[0]).strip('.') or 'utente'
    username, n = base[:30], 1
    while User.query.filter_by(username=username).first():
        username = f'{base[:28]}{n}'; n += 1
    u = User(username=username, email=email)
    u.set_password(password)
    db.session.add(u)
    db.session.commit()
    flash(f'Utente "{email}" creato.', 'success')
    return redirect(url_for('admin.users'))


@bp.route('/users/<int:uid>/topup', methods=['POST'])
@require_permission('manage_users')
def user_topup(uid):
    user = db.get_or_404(User, uid)
    amount = request.form.get('amount', type=float)
    note = request.form.get('note', 'Ricarica manuale').strip() or 'Ricarica manuale'
    if not amount or amount <= 0:
        flash('Importo non valido.', 'danger')
        return redirect(url_for('admin.users'))
    user.credit_wallet(amount, note)
    db.session.commit()
    flash(f'+{numero_italiano(amount)}€ aggiunti al wallet di {user.username}.', 'success')
    return redirect(url_for('admin.users'))


@bp.route('/users/<int:uid>/toggle', methods=['POST'])
@require_permission('manage_users')
def user_toggle(uid):
    user = db.get_or_404(User, uid)
    user.is_active = not user.is_active
    db.session.commit()
    if user.is_active:
        send_telegram_to_user(user,
            f'\U00002705 Il tuo account è stato attivato!\n'
            f'Puoi ora accedere al servizio.')
    flash(f'Utente {user.username} {"attivato" if user.is_active else "sospeso"}.', 'info')
    return redirect(url_for('admin.users'))


@bp.route('/users/<int:uid>/delete', methods=['POST'])
@require_permission('manage_users')
def user_delete(uid):
    user = db.get_or_404(User, uid)
    if user.is_admin:
        flash('Non puoi eliminare un amministratore.', 'danger')
        return redirect(url_for('admin.users'))
    name = user.username
    _delete_user_cascade(uid)
    db.session.commit()
    flash(f'Utente "{name}" eliminato.', 'info')
    return redirect(url_for('admin.users'))


@bp.route('/users/<int:uid>/roles', methods=['POST'])
@require_permission('manage_users')
def user_roles_assign(uid):
    user = db.get_or_404(User, uid)
    role_ids = request.form.getlist('role_ids', type=int)
    user.roles = Role.query.filter(Role.id.in_(role_ids)).all() if role_ids else []
    db.session.commit()
    flash(f'Ruoli di {user.username} aggiornati.', 'success')
    return redirect(url_for('admin.users'))


# ── Categorie prodotto ────────────────────────────────────────────────────────

@bp.route('/categories')
@require_permission('manage_categories')
def categories():
    tid = _active_tenant_id()
    cats = Category.query.filter_by(tenant_id=tid).order_by(Category.name).all()
    return render_template('admin/categories.html', categories=cats)


@bp.route('/categories/new', methods=['POST'])
@require_permission('manage_categories')
def category_new():
    name = request.form.get('name', '').strip()
    if not name:
        flash('Nome obbligatorio.', 'danger')
        return redirect(url_for('admin.categories'))
    tid = _active_tenant_id()
    if Category.query.filter_by(name=name, tenant_id=tid).first():
        flash('Categoria già esistente.', 'warning')
        return redirect(url_for('admin.categories'))
    db.session.add(Category(name=name,
                            icon=request.form.get('icon', 'fa-utensils'),
                            color=request.form.get('color', 'secondary'),
                            tenant_id=tid))
    db.session.commit()
    flash(f'Categoria "{name}" creata.', 'success')
    return redirect(url_for('admin.categories'))


@bp.route('/categories/<int:cid>/edit', methods=['POST'])
@require_permission('manage_categories')
def category_edit(cid):
    cat = Category.query.get_or_404(cid)
    name = request.form.get('name', '').strip()
    if not name:
        flash('Nome obbligatorio.', 'danger')
        return redirect(url_for('admin.categories'))
    existing = Category.query.filter_by(name=name, tenant_id=cat.tenant_id).first()
    if existing and existing.id != cid:
        flash(f'Esiste già una categoria con il nome "{name}".', 'warning')
        return redirect(url_for('admin.categories'))
    cat.name  = name
    cat.icon  = request.form.get('icon', cat.icon)
    cat.color = request.form.get('color', cat.color)
    db.session.commit()
    flash(f'Categoria "{name}" aggiornata.', 'success')
    return redirect(url_for('admin.categories'))


@bp.route('/categories/<int:cid>/delete', methods=['POST'])
@require_permission('manage_categories')
def category_delete(cid):
    cat = Category.query.get_or_404(cid)
    if cat.products:
        flash(f'Impossibile eliminare "{cat.name}": contiene {len(cat.products)} prodott{"o" if len(cat.products)==1 else "i"}. Spostali prima.', 'danger')
        return redirect(url_for('admin.categories'))
    name = cat.name
    db.session.delete(cat)
    db.session.commit()
    flash(f'Categoria "{name}" eliminata.', 'success')
    return redirect(url_for('admin.categories'))


# ── Slot orari ────────────────────────────────────────────────────────────────

@bp.route('/slots')
@require_permission('manage_slots')
def slots():
    return redirect(url_for('admin.tavoli', tab='slot'))


@bp.route('/slots/<int:sid>/toggle', methods=['POST'])
@require_permission('manage_slots')
def slot_toggle(sid):
    slot = db.get_or_404(TimeSlot, sid)
    slot.is_active = not slot.is_active
    db.session.commit()
    flash(f'Slot {slot.time_str} {"attivato" if slot.is_active else "disattivato"}.', 'info')
    return redirect(url_for('admin.tavoli', tab='slot'))


@bp.route('/slots/<int:sid>/capacity', methods=['POST'])
@require_permission('manage_slots')
def slot_capacity(sid):
    slot = db.get_or_404(TimeSlot, sid)
    cap = request.form.get('max_orders', type=int)
    if cap and cap > 0:
        slot.max_orders = cap
        db.session.commit()
        flash(f'Capacità slot {slot.time_str} → {cap}.', 'success')
    return redirect(url_for('admin.tavoli', tab='slot'))


@bp.route('/slots/new', methods=['POST'])
@require_permission('manage_slots')
def slot_new():
    time_str = (request.form.get('time_str') or '').strip()
    max_orders = request.form.get('max_orders', 20, type=int)
    if not time_str:
        flash('Orario obbligatorio.', 'danger')
        return redirect(url_for('admin.tavoli', tab='slot'))
    import re
    if not re.match(r'^\d{2}:\d{2}$', time_str):
        flash('Formato orario non valido (HH:MM).', 'danger')
        return redirect(url_for('admin.tavoli', tab='slot'))
    tid = _active_tenant_id()
    if TimeSlot.query.filter_by(time_str=time_str, tenant_id=tid).first():
        flash(f'Lo slot {time_str} esiste già.', 'warning')
        return redirect(url_for('admin.tavoli', tab='slot'))
    db.session.add(TimeSlot(time_str=time_str, max_orders=max(1, max_orders), tenant_id=tid))
    db.session.commit()
    flash(f'Slot {time_str} creato.', 'success')
    return redirect(url_for('admin.tavoli', tab='slot'))


@bp.route('/slots/<int:sid>/delete', methods=['POST'])
@require_permission('manage_slots')
def slot_delete(sid):
    slot = db.get_or_404(TimeSlot, sid)
    if slot.orders:
        flash(f'Impossibile eliminare lo slot {slot.time_str}: ha ordini associati.', 'danger')
        return redirect(url_for('admin.tavoli', tab='slot'))
    db.session.delete(slot)
    db.session.commit()
    flash(f'Slot {slot.time_str} eliminato.', 'info')
    return redirect(url_for('admin.tavoli', tab='slot'))


# ── Banco POS ─────────────────────────────────────────────────────────────────

@bp.route('/banco')
@staff_required
def banco():
    tid   = _active_tenant_id()
    items = BancoItem.query.filter_by(is_active=True, tenant_id=tid)\
                    .order_by(BancoItem.sort_order, BancoItem.name).all()
    users = User.query.filter_by(is_admin=False, is_staff=False)\
                      .order_by(User.username).all()
    import json
    items_json = json.dumps([
        {'id': i.id, 'name': i.name, 'price': i.price,
         'icon': i.icon, 'color': i.color}
        for i in items
    ])
    return render_template('admin/banco_pos.html',
                           items=items, users=users, items_json=items_json)


@bp.route('/banco/session', methods=['POST'])
@staff_required
def banco_session_new():
    import json
    cart_json = request.form.get('cart_json', '[]')
    try:
        cart = json.loads(cart_json)
    except Exception:
        cart = []
    if not cart:
        return jsonify({'error': 'Nessun articolo'}), 400
    total = round(sum(i['price'] * i['qty'] for i in cart), 2)
    token = _secrets.token_hex(16)
    now   = _dt.utcnow()
    sess  = BancoSession(
        token=token, staff_id=current_user.id,
        items_json=cart_json, total=total,
        status='pending',
        created_at=now, expires_at=now + timedelta(minutes=10),
        tenant_id=_active_tenant_id(),
    )
    db.session.add(sess)
    db.session.commit()
    pay_url = request.host_url.rstrip('/') + '/banco/pay/' + token
    return jsonify({'token': token, 'pay_url': pay_url, 'total': total, 'expires_in': 600})


@bp.route('/banco/session/<token>/status')
@staff_required
def banco_session_status(token):
    sess = BancoSession.query.filter_by(token=token).first_or_404()
    if sess.status == 'pending' and _dt.utcnow() > sess.expires_at:
        sess.status = 'expired'
        db.session.commit()
    customer_name = None
    if sess.customer:
        customer_name = sess.customer.full_name or sess.customer.username
    return jsonify({'status': sess.status, 'customer': customer_name})


@bp.route('/banco/session/<token>/cancel', methods=['POST'])
@staff_required
def banco_session_cancel(token):
    sess = BancoSession.query.filter_by(token=token).first_or_404()
    if sess.status == 'pending':
        sess.status = 'cancelled'
        db.session.commit()
    return jsonify({'ok': True})


@bp.route('/banco/items')
@staff_required
def banco_items():
    tid   = _active_tenant_id()
    items = BancoItem.query.filter_by(tenant_id=tid)\
                    .order_by(BancoItem.sort_order, BancoItem.name).all()
    return render_template('admin/banco_items.html', items=items)


@bp.route('/banco/items/new', methods=['POST'])
@staff_required
def banco_item_new():
    name  = request.form.get('name', '').strip()
    price = request.form.get('price', type=float)
    if not name or not price or price <= 0:
        flash('Nome e prezzo obbligatori.', 'danger')
        return redirect(url_for('admin.banco_items'))
    tid = _active_tenant_id()
    db.session.add(BancoItem(
        name=name, price=price,
        icon=request.form.get('icon', 'fa-mug-hot').strip(),
        color=request.form.get('color', 'info').strip(),
        sort_order=request.form.get('sort_order', 0, type=int),
        tenant_id=tid,
    ))
    db.session.commit()
    flash(f'Articolo "{name}" aggiunto.', 'success')
    return redirect(url_for('admin.banco_items'))


@bp.route('/banco/items/<int:iid>/edit', methods=['POST'])
@staff_required
def banco_item_edit(iid):
    item = db.get_or_404(BancoItem, iid)
    item.name       = request.form.get('name', item.name).strip()
    item.price      = request.form.get('price', type=float) or item.price
    item.icon       = request.form.get('icon', item.icon).strip()
    item.color      = request.form.get('color', item.color).strip()
    item.sort_order = request.form.get('sort_order', item.sort_order, type=int)
    item.is_active  = request.form.get('is_active') == '1'
    db.session.commit()
    flash(f'Articolo "{item.name}" aggiornato.', 'success')
    return redirect(url_for('admin.banco_items'))


@bp.route('/banco/items/<int:iid>/delete', methods=['POST'])
@staff_required
def banco_item_delete(iid):
    item = db.get_or_404(BancoItem, iid)
    name = item.name
    db.session.delete(item)
    db.session.commit()
    flash(f'Articolo "{name}" eliminato.', 'info')
    return redirect(url_for('admin.banco_items'))


# ── BANCO: consegna pasto aziendale ──────────────────────────────────────────

@bp.route('/banco/pasto-lookup')
@staff_required
def banco_pasto_lookup():
    token   = request.args.get('token', '').strip().upper()
    if not token:
        return jsonify({'error': 'Token mancante'}), 400
    booking = CorporateMealBooking.query.filter_by(pickup_token=token).first()
    if not booking:
        return jsonify({'error': 'Codice non trovato'}), 404
    label, badge = booking.label()
    return jsonify({
        'id':        booking.id,
        'status':    booking.status,
        'status_label': label,
        'status_badge': badge,
        'user_name': booking.user.full_name or booking.user.username,
        'meal_name': booking.meal.name,
        'meal_date': booking.meal.meal_date.strftime('%d/%m/%Y'),
        'slot':      booking.slot.time_str if booking.slot else None,
        'quantity':  booking.quantity,
    })


@bp.route('/banco/pasto-consegna', methods=['POST'])
@staff_required
def banco_pasto_consegna():
    bid = request.form.get('booking_id', type=int)
    booking = CorporateMealBooking.query.get_or_404(bid)
    if booking.status == 'booked':
        booking.status = 'consumed'
        db.session.commit()
        send_telegram_to_user(
            booking.user,
            f'✅ Pasto ritirato: <b>{booking.meal.name}</b>. Buon appetito!')
        return jsonify({'ok': True, 'message': f'Consegnato a {booking.user.full_name}.'})
    return jsonify({'ok': False, 'message': f'Stato attuale: {booking.label()[0]}'}), 409


# ── CESTO CUCINA ──────────────────────────────────────────────────────────────

@bp.route('/cesto')
@require_permission('manage_cesto')
@cesto_required
def cesto():
    from collections import defaultdict
    from datetime import datetime as _dtm, timedelta
    tid    = _active_tenant_id()
    today  = date.today()
    start  = _dtm(today.year, today.month, today.day)
    end    = start + timedelta(days=1)
    labels = (PrepLabel.query
              .filter_by(tenant_id=tid)
              .filter(PrepLabel.prepared_at >= start, PrepLabel.prepared_at < end)
              .order_by(PrepLabel.prepared_at.desc())
              .all())
    summary = defaultdict(lambda: {'name': '', 'price': 0.0, 'ready': 0, 'sold': 0, 'expired': 0})
    for lb in labels:
        d = summary[lb.product_id]
        d['name']   = lb.product.name
        d['price']  = lb.product.price
        d[lb.status] = d[lb.status] + 1
    products = (Product.query
                .filter_by(tenant_id=tid, is_active=True)
                .order_by(Product.name).all())
    return render_template('admin/cesto.html',
                           labels=labels, summary=dict(summary),
                           products=products, today=today)


@bp.route('/cesto/genera', methods=['POST'])
@require_permission('manage_cesto')
@cesto_required
def cesto_genera():
    import secrets as _sec
    CHARS      = 'ABCDEFGHJKLMNPQRSTUVWXYZ23456789'
    tid        = _active_tenant_id()
    product_id = request.form.get('product_id', type=int)
    qty        = min(max(request.form.get('qty', type=int, default=1), 1), 50)
    if not product_id:
        flash('Seleziona un prodotto.', 'danger')
        return redirect(url_for('admin.cesto'))
    product  = db.get_or_404(Product, product_id)
    batch_id = _dt.now().strftime('%y%m%d%H%M') + str(current_user.id)
    for _ in range(qty):
        for _attempt in range(30):
            code = 'CESTO-' + ''.join(_sec.choice(CHARS) for _ in range(6))
            if not PrepLabel.query.filter_by(code=code).first():
                break
        db.session.add(PrepLabel(
            code=code, product_id=product_id, batch_id=batch_id,
            tenant_id=tid, prepared_by=current_user.id,
        ))
    db.session.commit()
    flash(f'{qty} etichet{"ta" if qty == 1 else "te"} generate per "{product.name}".', 'success')
    return redirect(url_for('admin.cesto_stampa', bid=batch_id))


@bp.route('/cesto/stampa/<bid>')
@require_permission('manage_cesto')
@cesto_required
def cesto_stampa(bid):
    tid    = _active_tenant_id()
    labels = (PrepLabel.query
              .filter_by(batch_id=bid, tenant_id=tid)
              .order_by(PrepLabel.id).all())
    if not labels:
        flash('Lotto non trovato.', 'warning')
        return redirect(url_for('admin.cesto'))
    co_name  = get_setting('company_name') or 'QuickLunch'
    base_url = request.host_url.rstrip('/')
    return render_template('admin/cesto_stampa.html',
                           labels=labels, co_name=co_name, base_url=base_url)


@bp.route('/cesto/<code>/annulla', methods=['POST'])
@require_permission('manage_cesto')
@cesto_required
def cesto_annulla(code):
    tid = _active_tenant_id()
    lb  = PrepLabel.query.filter_by(code=code, tenant_id=tid).first_or_404()
    if lb.status == 'sold':
        flash('Etichetta già venduta, non annullabile.', 'danger')
    else:
        lb.status = 'expired'
        db.session.commit()
        flash(f'Etichetta {code} annullata.', 'info')
    return redirect(url_for('admin.cesto'))


@bp.route('/cesto/annulla-tutto', methods=['POST'])
@require_permission('manage_cesto')
@cesto_required
def cesto_annulla_tutto():
    from datetime import datetime as _dtm, timedelta
    tid   = _active_tenant_id()
    today = date.today()
    start = _dtm(today.year, today.month, today.day)
    end   = start + timedelta(days=1)
    n = (PrepLabel.query
         .filter_by(tenant_id=tid)
         .filter(PrepLabel.prepared_at >= start, PrepLabel.prepared_at < end)
         .delete(synchronize_session=False))
    db.session.commit()
    flash(f'{n} etichet{"ta eliminata" if n == 1 else "te eliminate"}.', 'info')
    return redirect(url_for('admin.cesto'))


# ── Prenotazioni cucina ───────────────────────────────────────────────────────

@bp.route('/cucina/prenotazioni')
@require_permission('manage_cesto')
def cucina_prenotazioni():
    from datetime import datetime as _dtm, timedelta as _td
    tid   = _active_tenant_id()
    today = date.today()

    # Prenotazioni di oggi e future non cancellate
    prens = (Prenotazione.query
             .filter_by(tenant_id=tid)
             .filter(Prenotazione.pickup_date >= today)
             .filter(Prenotazione.status != 'cancelled')
             .order_by(Prenotazione.pickup_date.asc(), Prenotazione.slot_id.asc())
             .all())

    # Conta PrepLabel ready per prodotto (solo oggi)
    start = _dtm(today.year, today.month, today.day)
    end   = start + _td(days=1)
    ready_labels = (PrepLabel.query
                    .filter_by(tenant_id=tid, status='ready')
                    .filter(PrepLabel.prepared_at >= start, PrepLabel.prepared_at < end)
                    .all())
    # { product_id: count }
    cesto_qty = {}
    for lb in ready_labels:
        cesto_qty[lb.product_id] = cesto_qty.get(lb.product_id, 0) + 1

    # Raggruppa prenotazioni per data
    by_date = {}
    for p in prens:
        by_date.setdefault(p.pickup_date, []).append(p)

    # Per ogni data, calcola fabbisogno per prodotto
    # { date: { product_id: {product, needed, in_cesto} } }
    fabbisogno = {}
    for d, plist in by_date.items():
        prod_map = {}
        for p in plist:
            for item in p.items:
                pid = item.product_id
                if pid not in prod_map:
                    prod_map[pid] = {'product': item.product, 'needed': 0}
                prod_map[pid]['needed'] += item.quantity
        # cesto qty è rilevante solo per oggi
        for pid, info in prod_map.items():
            info['in_cesto'] = cesto_qty.get(pid, 0) if d == today else 0
        fabbisogno[d] = prod_map

    return render_template('admin/cucina_prenotazioni.html',
                           by_date=by_date, fabbisogno=fabbisogno,
                           today=today, cesto_qty=cesto_qty)


@bp.route('/cucina/prenotazioni/<int:pid>/pronto', methods=['POST'])
@require_permission('manage_cesto')
def cucina_prenotazione_pronto(pid):
    """Marca la prenotazione come pronta e notifica il cliente via Telegram."""
    tid  = _active_tenant_id()
    pren = Prenotazione.query.filter_by(id=pid, tenant_id=tid).first_or_404()
    if pren.status not in ('pending', 'confirmed'):
        return jsonify(ok=False, message='Stato non valido per questa operazione.'), 400

    pren.status = 'ready'
    db.session.commit()

    prodotti_str = ', '.join(
        f"{it.quantity}× {it.product.name}" for it in pren.items
    )
    slot_str = pren.slot.time_str if pren.slot else ''
    msg_utente = (
        f'🔔 <b>Il tuo ordine è pronto!</b>\n'
        f'👤 {pren.user.full_name}\n'
        f'🥗 {prodotti_str}\n'
        + (f'⏰ Ritiro alle {slot_str}\n' if slot_str else '')
        + f'📌 Codice: <code>{pren.code}</code>\n\n'
        f'Presentati al banco con il codice per ritirare il tuo pranzo.'
    )
    send_telegram_to_user(pren.user, msg_utente)
    send_web_push_to_user(
        pren.user,
        title='🍽️ Il tuo ordine è pronto!',
        body=f'{prodotti_str} • Presentati al banco con il codice {pren.code}',
        url='/prenotazioni'
    )

    send_telegram(
        f'✅ Ordine pronto — notifica inviata a {pren.user.full_name}\n'
        f'📌 {pren.code}  |  {prodotti_str}'
    )

    return jsonify(ok=True, message='Notifica inviata al cliente.')


# ── Tavoli ────────────────────────────────────────────────────────────────────

@bp.route('/tables')
@require_permission('manage_tables_admin')
@tables_required
def tables():
    return redirect(url_for('admin.tavoli', tab='tavoli'))


@bp.route('/tables/new', methods=['POST'])
@require_permission('manage_tables_admin')
@tables_required
def table_new():
    number = request.form.get('number', type=int)
    seats = request.form.get('seats', type=int, default=4)
    location = request.form.get('location', '').strip()
    if not number:
        flash('Numero tavolo obbligatorio.', 'danger')
        return redirect(url_for('admin.tables'))
    if Table.query.filter_by(number=number).first():
        flash(f'Tavolo {number} già esistente.', 'warning')
        return redirect(url_for('admin.tables'))
    db.session.add(Table(number=number, seats=seats, location=location))
    db.session.commit()
    flash(f'Tavolo {number} aggiunto.', 'success')
    return redirect(url_for('admin.tavoli', tab='tavoli'))


@bp.route('/tables/<int:tid>/edit', methods=['POST'])
@require_permission('manage_tables_admin')
@tables_required
def table_edit(tid):
    t = db.get_or_404(Table, tid)
    t.seats = request.form.get('seats', type=int) or t.seats
    t.location = request.form.get('location', t.location).strip()
    t.is_active = 'is_active' in request.form
    db.session.commit()
    flash(f'Tavolo {t.number} aggiornato.', 'success')
    return redirect(url_for('admin.tavoli', tab='tavoli'))


# ── Prenotazioni tavoli ───────────────────────────────────────────────────────

@bp.route('/reservations')
@require_permission('manage_reservations_admin')
@tables_required
def reservations():
    d = request.args.get('date', '')
    return redirect(url_for('admin.tavoli', date=d or None))


@bp.route('/reservations/<int:rid>/cancel', methods=['POST'])
@require_permission('manage_reservations_admin')
@tables_required
def reservation_cancel(rid):
    res = db.get_or_404(TableReservation, rid)
    res.status = 'cancelled'
    db.session.commit()
    flash(f'Prenotazione tavolo {res.table.number} annullata.', 'info')
    send_telegram_to_user(
        res.user,
        f'❌ La tua prenotazione del tavolo <b>{res.table.number}</b> '
        f'per le <b>{res.session_start}</b> del '
        f'{res.reservation_date.strftime("%d/%m/%Y")} è stata annullata dall\'amministratore.'
    )
    return redirect(url_for('admin.tavoli'))


# ── Gestione Tavoli unificata ─────────────────────────────────────────────────

@bp.route('/tavoli')
@require_permission('manage_tables_admin')
@tables_required
def tavoli():
    from datetime import datetime as dt
    tab = request.args.get('tab', 'panoramica')

    raw_date = request.args.get('date', str(date.today()))
    try:
        sel_date = dt.strptime(raw_date, '%Y-%m-%d').date()
    except ValueError:
        sel_date = date.today()
        raw_date = str(sel_date)

    prev_day = (sel_date - timedelta(days=1)).isoformat()
    next_day = (sel_date + timedelta(days=1)).isoformat()

    all_tables = Table.query.filter_by(is_active=True).order_by(Table.number).all()
    all_tables_with_inactive = Table.query.order_by(Table.number).all()
    bands = TableTimeBand.query.order_by(TableTimeBand.sort_order, TableTimeBand.start_time).all()
    order_slots = TimeSlot.query.order_by(TimeSlot.time_str).all()

    # Per la panoramica: prenotazioni del giorno indicizzate per (table_id, session_start)
    day_reservations = (TableReservation.query
                        .filter_by(reservation_date=sel_date)
                        .filter(TableReservation.status != 'cancelled')
                        .all())
    # index: (table_id, session_start) → reservation
    res_index = {(r.table_id, r.session_start): r for r in day_reservations}

    # Tutte le prenotazioni del giorno (incluse annullate) per la lista completa
    all_day_res = (TableReservation.query
                   .filter_by(reservation_date=sel_date)
                   .order_by(TableReservation.session_start, TableReservation.table_id)
                   .all())

    return render_template('admin/tavoli.html',
                           tab=tab, sel_date=raw_date, today=str(date.today()),
                           prev_day=prev_day, next_day=next_day,
                           all_tables=all_tables,
                           all_tables_with_inactive=all_tables_with_inactive,
                           bands=bands, order_slots=order_slots,
                           res_index=res_index,
                           all_day_res=all_day_res)


@bp.route('/tavoli/bande/new', methods=['POST'])
@require_permission('manage_tables_admin')
@tables_required
def band_new():
    start = request.form.get('start_time', '').strip()
    end   = request.form.get('end_time',   '').strip()
    dur   = request.form.get('duration_minutes', type=int)
    if not start or not end or not dur or dur < 1:
        flash('Compila tutti i campi della fascia oraria.', 'danger')
        return redirect(url_for('admin.tavoli', tab='fasce'))
    tid = current_user.tenant_id if not current_user.is_admin else None
    # calcola sort_order come posizione temporale
    existing = TableTimeBand.query.count()
    band = TableTimeBand(start_time=start, end_time=end,
                         duration_minutes=dur, sort_order=existing,
                         tenant_id=tid)
    db.session.add(band)
    db.session.commit()
    flash(f'Fascia {start}–{end} ({dur} min) aggiunta.', 'success')
    return redirect(url_for('admin.tavoli', tab='fasce'))


@bp.route('/tavoli/bande/<int:bid>/delete', methods=['POST'])
@require_permission('manage_tables_admin')
@tables_required
def band_delete(bid):
    band = TableTimeBand.query.get_or_404(bid)
    label = band.label
    if band.reservations:
        flash(f'Impossibile eliminare la fascia "{label}": ha prenotazioni collegate.', 'danger')
        return redirect(url_for('admin.tavoli', tab='fasce'))
    db.session.delete(band)
    db.session.commit()
    flash(f'Fascia "{label}" eliminata.', 'success')
    return redirect(url_for('admin.tavoli', tab='fasce'))


@bp.route('/tables/<int:tid>/delete', methods=['POST'])
@require_permission('manage_tables_admin')
@tables_required
def table_delete(tid):
    t = Table.query.get_or_404(tid)
    active_res = TableReservation.query.filter_by(
        table_id=tid, status='confirmed').count()
    if active_res:
        flash(f'Tavolo {t.number} ha {active_res} prenotazioni attive: annullale prima.', 'danger')
        return redirect(url_for('admin.tavoli', tab='tavoli'))
    db.session.delete(t)
    db.session.commit()
    flash(f'Tavolo {t.number} eliminato.', 'success')
    return redirect(url_for('admin.tavoli', tab='tavoli'))


# ── Ingredienti builder ───────────────────────────────────────────────────────

@bp.route('/ingredients')
@require_permission('manage_ingredients')
def ingredients():
    cats = IngredientCategory.query.order_by(
        IngredientCategory.builder_type, IngredientCategory.sort_order).all()
    return render_template('admin/ingredients.html', categories=cats)


@bp.route('/ingredients/<int:iid>/stock', methods=['POST'])
@require_permission('manage_ingredients')
def ingredient_stock(iid):
    ing = db.get_or_404(Ingredient, iid)
    op  = request.form.get('op', 'set')
    try:
        qty = float(request.form.get('qty', 0))
    except (ValueError, TypeError):
        qty = 0.0
    if op == 'add':
        ing.stock_qty = (ing.stock_qty or 0.0) + qty
    elif op == 'sub':
        ing.stock_qty = max(0.0, (ing.stock_qty or 0.0) - qty)
    else:
        ing.stock_qty = qty if qty >= 0 else None
    db.session.commit()
    flash(f'Giacenza "{ing.name}" aggiornata: {numero_italiano(ing.stock_qty, 0)}g.', 'success')
    return redirect(url_for('admin.ingredients'))


@bp.route('/ingredients/new', methods=['POST'])
@require_permission('manage_ingredients')
def ingredient_new():
    name = request.form.get('name', '').strip()
    category_id = request.form.get('category_id', type=int)
    price_extra = request.form.get('price_extra', type=float, default=0.0)
    is_vegetarian = 'is_vegetarian' in request.form
    allergens = request.form.get('allergens', '').strip()
    gps_raw = request.form.get('grams_per_serving', '').strip()
    grams_per_serving = float(gps_raw) if gps_raw else None
    if not name or not category_id:
        flash('Nome e categoria obbligatori.', 'danger')
        return redirect(url_for('admin.ingredients'))
    db.session.add(Ingredient(name=name, category_id=category_id,
                              price_extra=price_extra or 0.0,
                              is_vegetarian=is_vegetarian, allergens=allergens,
                              grams_per_serving=grams_per_serving))
    db.session.commit()
    flash(f'Ingrediente "{name}" aggiunto.', 'success')
    return redirect(url_for('admin.ingredients'))


@bp.route('/ingredients/<int:iid>/toggle', methods=['POST'])
@require_permission('manage_ingredients')
def ingredient_toggle(iid):
    ing = db.get_or_404(Ingredient, iid)
    ing.is_active = not ing.is_active
    db.session.commit()
    flash(f'"{ing.name}" {"attivato" if ing.is_active else "disattivato"}.', 'info')
    return redirect(url_for('admin.ingredients'))


@bp.route('/ingredients/<int:iid>/edit', methods=['POST'])
@require_permission('manage_ingredients')
def ingredient_edit(iid):
    ing = db.get_or_404(Ingredient, iid)
    ing.name        = request.form.get('name', ing.name).strip()
    ing.price_extra = request.form.get('price_extra', type=float, default=ing.price_extra) or 0.0
    gps_raw = request.form.get('grams_per_serving', '').strip()
    ing.grams_per_serving = float(gps_raw) if gps_raw else None
    ing.is_vegetarian = 'is_vegetarian' in request.form
    ing.allergens = request.form.get('allergens', '').strip()
    db.session.commit()
    flash(f'Ingrediente "{ing.name}" aggiornato.', 'success')
    return redirect(url_for('admin.ingredients'))


# ── Report ────────────────────────────────────────────────────────────────────

@bp.route('/report')
@require_permission('view_reports')
def report():
    from datetime import timedelta
    from sqlalchemy import func
    today = date.today()
    days = [(today - timedelta(days=i)) for i in range(29, -1, -1)]
    revenue_by_day, orders_by_day = {}, {}
    for d in days:
        res = db.session.query(func.sum(Order.total_price), func.count(Order.id))\
            .filter(Order.order_date == d, Order.status != 'cancelled').first()
        revenue_by_day[str(d)] = round(res[0] or 0, 2)
        orders_by_day[str(d)] = res[1] or 0

    top_products = db.session.query(
        Product.name, func.sum(OrderItem.quantity).label('sold')
    ).join(OrderItem, Product.id == OrderItem.product_id)\
     .join(Order, Order.id == OrderItem.order_id)\
     .filter(Order.status != 'cancelled')\
     .group_by(Product.id)\
     .order_by(func.sum(OrderItem.quantity).desc()).limit(10).all()

    return render_template('admin/report.html',
                           days=[str(d) for d in days],
                           revenue_by_day=revenue_by_day,
                           orders_by_day=orders_by_day,
                           top_products=top_products)


# ── Ruoli & Permessi ──────────────────────────────────────────────────────────

@bp.route('/roles')
@require_permission('manage_roles')
def roles():
    all_roles = Role.query.order_by(Role.name).all()
    all_perms = Permission.query.order_by(Permission.category, Permission.label).all()
    # group permissions by category
    from itertools import groupby
    cats = {}
    for p in all_perms:
        cats.setdefault(p.category, []).append(p)
    return render_template('admin/roles.html',
                           roles=all_roles,
                           all_perms=all_perms,
                           perm_cats=cats)


@bp.route('/roles/new', methods=['POST'])
@require_permission('manage_roles')
def role_new():
    name  = request.form.get('name', '').strip().lower().replace(' ', '_')
    label = request.form.get('label', '').strip()
    color = request.form.get('color', 'secondary')
    if not name or not label:
        flash('Nome e etichetta obbligatori.', 'danger')
        return redirect(url_for('admin.roles'))
    if Role.query.filter_by(name=name).first():
        flash(f'Ruolo "{name}" già esistente.', 'warning')
        return redirect(url_for('admin.roles'))
    role = Role(name=name, label=label, color=color, is_system=False)
    perm_ids = request.form.getlist('perm_ids', type=int)
    role.permissions = Permission.query.filter(Permission.id.in_(perm_ids)).all() if perm_ids else []
    db.session.add(role)
    db.session.commit()
    flash(f'Ruolo "{label}" creato.', 'success')
    return redirect(url_for('admin.roles'))


@bp.route('/roles/<int:rid>/edit', methods=['POST'])
@require_permission('manage_roles')
def role_edit(rid):
    role = db.get_or_404(Role, rid)
    role.label = request.form.get('label', role.label).strip()
    role.color = request.form.get('color', role.color)
    perm_ids = request.form.getlist('perm_ids', type=int)
    role.permissions = Permission.query.filter(Permission.id.in_(perm_ids)).all() if perm_ids else []
    db.session.commit()
    flash(f'Ruolo "{role.label}" aggiornato.', 'success')
    return redirect(url_for('admin.roles'))


@bp.route('/roles/<int:rid>/delete', methods=['POST'])
@require_permission('manage_roles')
def role_delete(rid):
    role = db.get_or_404(Role, rid)
    if role.is_system:
        flash('I ruoli di sistema non possono essere eliminati.', 'warning')
        return redirect(url_for('admin.roles'))
    db.session.delete(role)
    db.session.commit()
    flash(f'Ruolo "{role.label}" eliminato.', 'info')
    return redirect(url_for('admin.roles'))


# ── Impostazioni (Telegram / Gmail) ───────────────────────────────────────────

@bp.route('/settings')
@require_permission('manage_settings')
def settings():
    all_keys = [
        'company_name', 'company_address', 'company_city',
        'company_vat', 'company_phone', 'company_email',
        'telegram_bot_token', 'telegram_chat_id',
        'gmail_user', 'gmail_app_password',
        'google_client_id', 'google_client_secret',
        'registration_bonus',
        'loyalty_points_per_euro', 'loyalty_reward_points', 'loyalty_reward_amount',
        'builder_price_panino', 'builder_price_insalata', 'builder_price_poke',
        'table_reminder_minutes', 'order_reminder_minutes', 'meal_reminder_minutes',
        'tables_enabled', 'cesto_enabled',
        'sim_pasti_min', 'sim_pasti_max', 'sim_snack_min', 'sim_snack_max',
        'sim_caffe_min', 'sim_caffe_max', 'sim_builder_min', 'sim_builder_max',
    ]
    cfg = {k: get_setting(k) for k in all_keys}

    # Sezione "Dati": elenco dei carichi generati, solo per il super admin
    carichi = []
    if current_user.is_admin:
        carichi = (CaricoMensile.query
                   .order_by(CaricoMensile.creato_il.desc()).all())
    oggi = date.today()
    return render_template('admin/settings.html', cfg=cfg, carichi=carichi,
                           mese_corrente='%04d-%02d' % (oggi.year, oggi.month))


@bp.route('/settings/save', methods=['POST'])
@require_permission('manage_settings')
def settings_save():
    # Ogni sezione del form include un campo "keys" con i propri tasti (comma-separated)
    keys_raw = request.form.get('keys', '')
    keys = [k.strip() for k in keys_raw.split(',') if k.strip()]
    if not keys:
        flash('Nessun campo da salvare.', 'warning')
        return redirect(url_for('admin.settings'))
    for k in keys:
        # getlist + ultimo valore: le checkbox inviano un hidden '0' seguito da '1'
        # quando sono selezionate, e deve vincere lo stato effettivo della casella.
        _vals = request.form.getlist(k)
        val = (_vals[-1] if _vals else '').strip()
        s   = AppSetting.query.filter_by(key=k).first()
        if s:
            s.value = val
        else:
            db.session.add(AppSetting(key=k, value=val))
    db.session.commit()
    flash('Impostazioni salvate.', 'success')
    return redirect(url_for('admin.settings'))


@bp.route('/settings/test-telegram', methods=['POST'])
@require_permission('manage_settings')
def settings_test_telegram():
    ok, msg = send_telegram('✅ <b>Test QuickLunch</b>\nConnessione Telegram funzionante!')
    flash(f'Telegram: {msg}', 'success' if ok else 'danger')
    return redirect(url_for('admin.settings'))


@bp.route('/settings/test-email', methods=['POST'])
@require_permission('manage_settings')
def settings_test_email():
    from app.notifications import send_email
    dest = get_setting('gmail_user')
    ok, msg = send_email(dest, 'Test QuickLunch', '<b>Test email QuickLunch</b> — configurazione Gmail funzionante!')
    flash(f'Email: {msg}', 'success' if ok else 'danger')
    return redirect(url_for('admin.settings'))


@bp.route('/settings/broadcast-telegram', methods=['POST'])
@require_permission('send_notifications')
def broadcast_telegram():
    text = request.form.get('message', '').strip()
    if not text:
        flash('Messaggio vuoto.', 'warning')
        return redirect(url_for('admin.settings'))
    ok, msg = send_telegram(text)
    flash(f'Telegram: {msg}', 'success' if ok else 'danger')
    return redirect(url_for('admin.settings'))


@bp.route('/settings/broadcast-email', methods=['POST'])
@require_permission('send_notifications')
def broadcast_email():
    subject = request.form.get('subject', 'Comunicazione QuickLunch').strip()
    text    = request.form.get('message', '').strip()
    if not text:
        flash('Messaggio vuoto.', 'warning')
        return redirect(url_for('admin.settings'))
    html = f'<div style="font-family:sans-serif;max-width:520px;">{text}</div>'
    sent, failed, _ = send_email_to_all_users(subject, html)
    flash(f'Email inviate: {sent} ok, {failed} fallite.', 'success' if not failed else 'warning')
    return redirect(url_for('admin.settings'))


# ── Sondaggi (Polls) ──────────────────────────────────────────────────────────

@bp.route('/polls')
@require_permission('manage_polls')
def polls():
    from datetime import timedelta
    all_polls = Poll.query.order_by(Poll.poll_date.desc()).all()
    products  = Product.query.filter_by(is_active=True).order_by(Product.name).all()
    tomorrow  = str(date.today() + timedelta(days=1))
    return render_template('admin/polls.html', polls=all_polls, products=products,
                           tomorrow=tomorrow)


@bp.route('/polls/new', methods=['POST'])
@require_permission('manage_polls')
def poll_new():
    from datetime import datetime as dt
    title     = request.form.get('title', '').strip()
    poll_date = request.form.get('poll_date', '').strip()
    if not title or not poll_date:
        flash('Titolo e data obbligatori.', 'danger')
        return redirect(url_for('admin.polls'))
    try:
        pd = dt.strptime(poll_date, '%Y-%m-%d').date()
    except ValueError:
        flash('Data non valida.', 'danger')
        return redirect(url_for('admin.polls'))

    poll = Poll(title=title, poll_date=pd, is_active=True)
    db.session.add(poll)
    db.session.flush()

    # Scelte da prodotti selezionati
    product_ids = request.form.getlist('product_ids', type=int)
    for pid in product_ids:
        p = Product.query.get(pid)
        if p:
            db.session.add(PollChoice(poll_id=poll.id, text=p.name, emoji='🍽️'))

    # Scelte custom
    for i in range(1, 8):
        ct = request.form.get(f'custom_text_{i}', '').strip()
        ce = request.form.get(f'custom_emoji_{i}', '🍽️').strip() or '🍽️'
        if ct:
            db.session.add(PollChoice(poll_id=poll.id, text=ct, emoji=ce))

    db.session.commit()
    flash(f'Sondaggio "{title}" creato con {len(product_ids)} scelte.', 'success')
    return redirect(url_for('admin.polls'))


@bp.route('/polls/<int:pid>/toggle', methods=['POST'])
@require_permission('manage_polls')
def poll_toggle(pid):
    poll = db.get_or_404(Poll, pid)
    poll.is_active = not poll.is_active
    db.session.commit()
    flash(f'Sondaggio {"attivato" if poll.is_active else "chiuso"}.', 'info')
    return redirect(url_for('admin.polls'))


@bp.route('/polls/<int:pid>/delete', methods=['POST'])
@require_permission('manage_polls')
def poll_delete(pid):
    poll = db.get_or_404(Poll, pid)
    db.session.delete(poll)
    db.session.commit()
    flash('Sondaggio eliminato.', 'info')
    return redirect(url_for('admin.polls'))


@bp.route('/polls/<int:pid>/notify-telegram', methods=['POST'])
@require_permission('send_notifications')
def poll_notify_telegram(pid):
    poll     = db.get_or_404(Poll, pid)
    base_url = request.host_url.rstrip('/')
    msg      = telegram_poll_message(poll, base_url)
    ok, info = send_telegram(msg)
    flash(f'Telegram: {info}', 'success' if ok else 'danger')
    return redirect(url_for('admin.polls'))


@bp.route('/polls/<int:pid>/notify-email', methods=['POST'])
@require_permission('send_notifications')
def poll_notify_email(pid):
    poll     = db.get_or_404(Poll, pid)
    base_url = request.host_url.rstrip('/')
    html     = email_poll_html(poll, base_url)
    subject  = f'📊 Vota: {poll.title}'
    sent, failed, _ = send_email_to_all_users(subject, html)
    flash(f'Email inviate: {sent} ok, {failed} fallite.', 'success' if not failed else 'warning')
    return redirect(url_for('admin.polls'))


# ── Tenant management (solo superadmin) ───────────────────────────────────────

def _superadmin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            abort(403)
        return f(*args, **kwargs)
    return login_required(decorated)


@bp.route('/seed-demo', methods=['POST'])
@_superadmin_required
def seed_demo():
    from app.demo_seed import reset_demo_data, seed_demo_data
    reset_demo_data()
    ok, msg = seed_demo_data()
    flash(msg, 'success' if ok else 'warning')
    return redirect(url_for('admin.dashboard'))


@bp.route('/superadmin/guadagni')
@_superadmin_required
def ds_guadagni():
    from calendar import monthrange
    from app.notifications import get_numeric_setting

    try:
        year  = int(request.args.get('year',  _dt.now().year))
        month = int(request.args.get('month', _dt.now().month))
        if not (1 <= month <= 12):
            raise ValueError
    except (ValueError, TypeError):
        year, month = _dt.now().year, _dt.now().month

    start_date = date(year, month, 1)
    end_date   = date(year, month, monthrange(year, month)[1])

    fee_pct      = get_numeric_setting('platform_fee_percentage', 0.0) / 100.0
    monthly_fee  = get_numeric_setting('tenant_monthly_fee', 0.0)

    tenants = Tenant.query.order_by(Tenant.name).all()

    rows = []
    for t in tenants:
        ord_sum = db.session.query(
            db.func.coalesce(db.func.sum(Order.total_price), 0.0)
        ).filter(
            Order.tenant_id == t.id,
            Order.status == 'completed',
            Order.order_date >= start_date,
            Order.order_date <= end_date,
        ).scalar() or 0.0

        banco_sum = db.session.query(
            db.func.coalesce(db.func.sum(BancoSession.total), 0.0)
        ).filter(
            BancoSession.tenant_id == t.id,
            BancoSession.status == 'paid',
            db.func.date(BancoSession.created_at) >= start_date,
            db.func.date(BancoSession.created_at) <= end_date,
        ).scalar() or 0.0

        # Cesto: la vendita non crea un ordine ne' una sessione banco, e gli
        # eventuali extra (lattine, snack) sono legati solo al movimento di
        # wallet. L'importo realmente incassato si legge quindi dalle
        # transazioni, che portano il prezzo applicato al momento della vendita.
        # Attenzione: dipende dal testo delle descrizioni scritte in
        # main.cesto_acquista ('Cesto: ...' e 'Cesto extra: ...').
        cesto_sum = db.session.query(
            db.func.coalesce(db.func.sum(-Transaction.amount), 0.0)
        ).select_from(Transaction).join(
            User, Transaction.user_id == User.id
        ).filter(
            User.tenant_id == t.id,
            Transaction.ttype == 'payment',
            db.or_(Transaction.description.like('Cesto: %'),
                   Transaction.description.like('Cesto extra: %')),
            db.func.date(Transaction.created_at) >= start_date,
            db.func.date(Transaction.created_at) <= end_date,
        ).scalar() or 0.0

        meals_sum = 0.0
        for ca in CorporateAccount.query.filter_by(tenant_id=t.id).all():
            for meal in ca.daily_meals:
                if start_date <= meal.meal_date <= end_date:
                    for b in meal.bookings:
                        if b.status != 'cancelled':
                            meals_sum += meal.price * (b.quantity or 1)

        t_total       = ord_sum + banco_sum + cesto_sum + meals_sum
        t_excl_vat    = round(t_total / 1.10, 2)   # imponibile (scorporo IVA 10%)
        t_fee         = round(t_excl_vat * fee_pct + monthly_fee, 2)

        rows.append({
            'tenant':    t,
            'orders':    round(ord_sum,   2),
            'banco':     round(banco_sum, 2),
            'cesto':     round(cesto_sum, 2),
            'meals':     round(meals_sum, 2),
            'total':     round(t_total,   2),
            'excl_vat':  t_excl_vat,
            'fee':       t_fee,
        })

    grand_total = round(sum(r['total'] for r in rows), 2)
    grand_fee   = round(sum(r['fee']   for r in rows), 2)

    _it_m = ['Gennaio','Febbraio','Marzo','Aprile','Maggio','Giugno',
             'Luglio','Agosto','Settembre','Ottobre','Novembre','Dicembre']

    return render_template('admin/superadmin_guadagni.html',
        rows=rows, year=year, month=month,
        month_label=_it_m[month - 1],
        fee_pct=round(fee_pct * 100, 2),
        monthly_fee=monthly_fee,
        grand_total=grand_total,
        grand_fee=grand_fee,
    )


# ── Magazzino materiali di consumo ────────────────────────────────────────────

@bp.route('/magazzino')
@require_permission('manage_stock')
def magazzino():
    tf = _tenant_filter()
    items = ConsumableItem.query.filter_by(**tf).order_by(ConsumableItem.name).all()
    alerts = [i for i in items if i.is_below_threshold]
    return render_template('admin/magazzino.html', items=items, alerts=alerts)


@bp.route('/magazzino/nuovo', methods=['GET', 'POST'])
@bp.route('/magazzino/<int:iid>/modifica', methods=['GET', 'POST'])
@require_permission('manage_stock')
def magazzino_edit(iid=None):
    tf = _tenant_filter()
    item = ConsumableItem.query.get_or_404(iid) if iid else None
    suppliers = Supplier.query.filter_by(**tf).order_by(Supplier.name).all()

    if request.method == 'POST':
        name          = request.form.get('name', '').strip()
        unit          = request.form.get('unit', 'pz').strip()
        quantity      = float(request.form.get('quantity', 0) or 0)
        min_threshold = float(request.form.get('min_threshold', 0) or 0)
        supplier_id   = request.form.get('supplier_id') or None
        if supplier_id:
            supplier_id = int(supplier_id)

        if not name:
            flash('Il nome è obbligatorio.', 'danger')
            return render_template('admin/magazzino_edit.html', item=item, suppliers=suppliers)

        if item is None:
            item = ConsumableItem(tenant_id=current_user.tenant_id if not current_user.is_admin else None)
            db.session.add(item)

        item.name          = name
        item.unit          = unit
        item.quantity      = quantity
        item.min_threshold = min_threshold
        item.supplier_id   = supplier_id

        # reset alert se la giacenza torna sopra soglia
        if not item.is_below_threshold:
            item.alert_active = False

        db.session.commit()
        flash(f'Materiale "{name}" salvato.', 'success')
        return redirect(url_for('admin.magazzino'))

    return render_template('admin/magazzino_edit.html', item=item, suppliers=suppliers)


@bp.route('/magazzino/<int:iid>/movimento', methods=['POST'])
@require_permission('manage_stock')
def magazzino_movimento(iid):
    item  = ConsumableItem.query.get_or_404(iid)
    delta = float(request.form.get('delta', 0) or 0)
    notes = request.form.get('notes', '').strip()

    if delta == 0:
        flash('Inserisci un valore diverso da zero.', 'warning')
        return redirect(url_for('admin.magazzino'))

    item.quantity = round(item.quantity + delta, 3)
    mv = ConsumableMovement(item_id=item.id, delta=delta,
                             notes=notes, user_id=current_user.id)
    db.session.add(mv)

    # gestione alert soglia
    if item.is_below_threshold and not item.alert_active:
        ok, msg = send_supplier_low_stock_alert(item)
        if ok:
            flash(f'⚠️ Soglia raggiunta per "{item.name}" — email inviata al fornitore.', 'warning')
        elif item.supplier and item.supplier.email:
            flash(f'⚠️ Soglia raggiunta per "{item.name}" — invio email fallito: {msg}', 'danger')
        else:
            flash(f'⚠️ Soglia raggiunta per "{item.name}" — nessun fornitore con email configurata.', 'warning')
        item.alert_active = True
    elif not item.is_below_threshold:
        item.alert_active = False

    db.session.commit()
    return redirect(url_for('admin.magazzino'))


@bp.route('/magazzino/<int:iid>/elimina', methods=['POST'])
@require_permission('manage_stock')
def magazzino_delete(iid):
    item = ConsumableItem.query.get_or_404(iid)
    name = item.name
    db.session.delete(item)
    db.session.commit()
    flash(f'Materiale "{name}" eliminato.', 'success')
    return redirect(url_for('admin.magazzino'))


# ── Fornitori ─────────────────────────────────────────────────────────────────

@bp.route('/fornitori')
@require_permission('manage_stock')
def fornitori():
    tf = _tenant_filter()
    sups = Supplier.query.filter_by(**tf).order_by(Supplier.name).all()
    return render_template('admin/fornitori.html', suppliers=sups)


@bp.route('/fornitori/nuovo', methods=['GET', 'POST'])
@bp.route('/fornitori/<int:sid>/modifica', methods=['GET', 'POST'])
@require_permission('manage_stock')
def fornitore_edit(sid=None):
    sup = Supplier.query.get_or_404(sid) if sid else None

    if request.method == 'POST':
        name  = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        phone = request.form.get('phone', '').strip()
        notes = request.form.get('notes', '').strip()

        if not name:
            flash('Il nome è obbligatorio.', 'danger')
            return render_template('admin/fornitore_edit.html', supplier=sup)

        if sup is None:
            sup = Supplier(tenant_id=current_user.tenant_id if not current_user.is_admin else None)
            db.session.add(sup)

        sup.name  = name
        sup.email = email
        sup.phone = phone
        sup.notes = notes
        db.session.commit()
        flash(f'Fornitore "{name}" salvato.', 'success')
        return redirect(url_for('admin.fornitori'))

    return render_template('admin/fornitore_edit.html', supplier=sup)


@bp.route('/fornitori/<int:sid>/elimina', methods=['POST'])
@require_permission('manage_stock')
def fornitore_delete(sid):
    sup = Supplier.query.get_or_404(sid)
    name = sup.name
    # scollega gli articoli prima di eliminare
    for item in sup.items:
        item.supplier_id = None
    db.session.delete(sup)
    db.session.commit()
    flash(f'Fornitore "{name}" eliminato.', 'success')
    return redirect(url_for('admin.fornitori'))

# â”€â”€ Convenzioni aziendali â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

@bp.route('/convenzioni')
@require_permission('manage_products')
def convenzioni():
    tf = _tenant_filter()
    corps = CorporateAccount.query.filter_by(**tf).order_by(CorporateAccount.name).all()
    from collections import defaultdict as _dd
    today_meals: dict = _dd(list)
    for m in DailyFixedMeal.query.filter_by(meal_date=date.today(), **tf).all():
        today_meals[m.corporate_id].append(m)
    oggi = date.today()
    return render_template('admin/convenzioni.html', corps=corps,
                           today_meals=today_meals,
                           mese_corrente='%04d-%02d' % (oggi.year, oggi.month))


@bp.route('/convenzioni/nuovo', methods=['GET', 'POST'])
@bp.route('/convenzioni/<int:cid>/modifica', methods=['GET', 'POST'])
@require_permission('manage_products')
def convenzione_edit(cid=None):
    tf = _tenant_filter()
    corp = CorporateAccount.query.get_or_404(cid) if cid else None
    all_users = User.query.filter_by(is_admin=False, **tf).order_by(User.username).all()

    if request.method == 'POST':
        name             = request.form.get('name', '').strip()
        contact_email    = request.form.get('contact_email', '').strip()
        daily_price      = float(request.form.get('daily_price', 7.0) or 7.0)
        max_daily_covers = int(request.form.get('max_daily_covers', 60) or 60)
        notes            = request.form.get('notes', '').strip()
        member_ids       = set(int(x) for x in request.form.getlist('member_ids') if x)

        if not name:
            flash('Il nome Ã¨ obbligatorio.', 'danger')
            return render_template('admin/convenzione_edit.html', corp=corp, all_users=all_users)

        if corp is None:
            corp = CorporateAccount(
                tenant_id=current_user.tenant_id if not current_user.is_admin else None)
            db.session.add(corp)

        corp.name             = name
        corp.contact_email    = contact_email
        corp.daily_price      = daily_price
        corp.max_daily_covers = max_daily_covers
        corp.notes            = notes
        corp.is_active        = 'is_active' in request.form
        db.session.flush()

        existing = {m.user_id: m for m in corp.memberships}
        for uid in member_ids:
            if uid not in existing:
                db.session.add(CorporateMembership(user_id=uid, corporate_id=corp.id))
            else:
                existing[uid].is_active = True
        for uid, m in existing.items():
            if uid not in member_ids:
                m.is_active = False

        db.session.commit()
        flash(f'Azienda "{name}" salvata.', 'success')
        return redirect(url_for('admin.convenzioni'))

    return render_template('admin/convenzione_edit.html', corp=corp, all_users=all_users)


@bp.route('/convenzioni/<int:cid>/elimina', methods=['POST'])
@require_permission('manage_products')
def convenzione_delete(cid):
    corp = CorporateAccount.query.get_or_404(cid)
    name = corp.name
    db.session.delete(corp)
    db.session.commit()
    flash(f'Azienda "{name}" eliminata.', 'success')
    return redirect(url_for('admin.convenzioni'))


@bp.route('/convenzioni/<int:cid>/pasto', methods=['GET', 'POST'])
@require_permission('manage_products')
def convenzione_pasto(cid):
    corp = CorporateAccount.query.get_or_404(cid)
    today = date.today()

    try:
        _d_raw = request.args.get('d') or request.form.get('d') or ''
        sel_date = date.fromisoformat(_d_raw) if _d_raw else today
    except ValueError:
        sel_date = today
    prev_date = sel_date - timedelta(days=1)
    next_date = sel_date + timedelta(days=1)

    def _d_param():
        return sel_date.isoformat() if sel_date != today else None

    if request.method == 'POST':
        action  = request.form.get('action', 'add')
        meal_id = request.form.get('meal_id', type=int)

        # ── Elimina opzione ────────────────────────────────────────────────────
        if action == 'delete' and meal_id:
            meal = DailyFixedMeal.query.get_or_404(meal_id)
            db.session.delete(meal)
            db.session.commit()
            flash('Opzione eliminata.', 'info')
            return redirect(url_for('admin.convenzione_pasto', cid=cid, d=_d_param()))

        # ── Attiva/disattiva opzione ───────────────────────────────────────────
        if action == 'toggle' and meal_id:
            meal = DailyFixedMeal.query.get_or_404(meal_id)
            meal.is_active = not meal.is_active
            db.session.commit()
            stato = 'attivata' if meal.is_active else 'disattivata'
            flash(f'Opzione "{meal.name}" {stato}.', 'success')
            return redirect(url_for('admin.convenzione_pasto', cid=cid, d=_d_param()))

        # ── Aggiunge / modifica opzione ───────────────────────────────────────
        name = request.form.get('name', '').strip()
        if not name:
            flash('Il nome è obbligatorio.', 'danger')
            return redirect(url_for('admin.convenzione_pasto', cid=cid, d=_d_param()))

        price    = request.form.get('price', '').strip()
        max_book = request.form.get('max_bookings', '').strip()

        if action == 'edit' and meal_id:
            meal = DailyFixedMeal.query.get_or_404(meal_id)
        else:
            meal = DailyFixedMeal(
                corporate_id=cid, meal_date=sel_date,
                tenant_id=current_user.tenant_id if not current_user.is_admin else None)
            db.session.add(meal)

        meal.name         = name
        meal.description  = request.form.get('description', '').strip()
        meal.allergens    = ','.join(request.form.getlist('allergens'))
        meal.primo        = request.form.get('primo',    '').strip()
        meal.secondo      = request.form.get('secondo',   '').strip()
        meal.contorno     = request.form.get('contorno',  '').strip()
        meal.bevanda      = request.form.get('bevanda',   '').strip()
        meal.caffe        = request.form.get('caffe',     '').strip()
        meal.price        = float(price)    if price    else corp.daily_price
        meal.max_bookings = int(max_book)   if max_book else corp.max_daily_covers
        meal.is_active    = 'is_active' in request.form
        db.session.commit()
        flash(f'Opzione "{name}" salvata.', 'success')
        return redirect(url_for('admin.convenzione_pasto', cid=cid, d=_d_param()))

    meals = DailyFixedMeal.query.filter_by(corporate_id=cid, meal_date=sel_date)\
        .order_by(DailyFixedMeal.name).all()
    bookings_by_meal = {
        m.id: CorporateMealBooking.query.filter_by(meal_id=m.id).all()
        for m in meals
    }

    import json as _json
    configs_json = _json.dumps({
        cfg.id: {
            'name':         cfg.name,
            'primo':        cfg.primo        or '',
            'secondo':      cfg.secondo      or '',
            'contorno':     cfg.contorno     or '',
            'bevanda':      cfg.bevanda      or '',
            'caffe':        cfg.caffe        or '',
            'description':  cfg.description  or '',
            'allergens':    [a.strip() for a in (cfg.allergens or '').split(',') if a.strip()],
            'price':        cfg.price,
            'max_bookings': cfg.max_bookings,
        }
        for cfg in corp.configurations
    })
    return render_template('admin/convenzione_pasto.html',
                           corp=corp, meals=meals, sel_date=sel_date, today=today,
                           prev_date=prev_date, next_date=next_date,
                           bookings_by_meal=bookings_by_meal, allergens_list=ALLERGENS,
                           configs_json=configs_json)


@bp.route('/convenzioni/<int:cid>/pasto/<int:bid>/consuma', methods=['POST'])
@require_permission('manage_products')
def convenzione_consuma(cid, bid):
    booking = CorporateMealBooking.query.get_or_404(bid)
    booking.status = 'consumed'
    db.session.commit()
    d = request.args.get('d', '')
    flash(f'Pasto di {booking.user.username} segnato come consumato.', 'success')
    return redirect(url_for('admin.convenzione_pasto', cid=cid, d=d or None))


# ── Configurazioni pasto ──────────────────────────────────────────────────────

@bp.route('/convenzioni/<int:cid>/configurazioni', methods=['GET', 'POST'])
@require_permission('manage_products')
def convenzione_configurazioni(cid):
    corp = CorporateAccount.query.get_or_404(cid)
    if request.method == 'POST':
        name    = request.form.get('name', '').strip()
        if not name:
            flash('Il nome è obbligatorio.', 'danger')
            return redirect(url_for('admin.convenzione_configurazioni', cid=cid))
        cfg = MealConfiguration(
            corporate_id = cid,
            name         = name,
            primo        = request.form.get('primo',   '').strip(),
            secondo      = request.form.get('secondo',  '').strip(),
            contorno     = request.form.get('contorno', '').strip(),
            bevanda      = request.form.get('bevanda',  '').strip(),
            caffe        = request.form.get('caffe',    '').strip(),
            description  = request.form.get('description', '').strip(),
            allergens    = ','.join(request.form.getlist('allergens')),
            price        = float(p) if (p := request.form.get('price', '').strip()) else None,
            max_bookings = int(m) if (m := request.form.get('max_bookings', '').strip()) else None,
            sort_order   = int(request.form.get('sort_order', 0) or 0),
            tenant_id    = None if current_user.is_admin else current_user.tenant_id,
        )
        db.session.add(cfg)
        db.session.commit()
        flash(f'Configurazione "{name}" creata.', 'success')
        return redirect(url_for('admin.convenzione_configurazioni', cid=cid))

    configs = MealConfiguration.query.filter_by(corporate_id=cid)\
        .order_by(MealConfiguration.sort_order, MealConfiguration.name).all()
    return render_template('admin/convenzione_configurazioni.html',
                           corp=corp, configs=configs, allergens_list=ALLERGENS)


@bp.route('/convenzioni/<int:cid>/configurazioni/<int:cfg_id>/edit', methods=['POST'])
@require_permission('manage_products')
def configurazione_edit(cid, cfg_id):
    cfg = MealConfiguration.query.get_or_404(cfg_id)
    name = request.form.get('name', '').strip()
    if not name:
        flash('Il nome è obbligatorio.', 'danger')
        return redirect(url_for('admin.convenzione_configurazioni', cid=cid))
    cfg.name         = name
    cfg.primo        = request.form.get('primo',   '').strip()
    cfg.secondo      = request.form.get('secondo',  '').strip()
    cfg.contorno     = request.form.get('contorno', '').strip()
    cfg.bevanda      = request.form.get('bevanda',  '').strip()
    cfg.caffe        = request.form.get('caffe',    '').strip()
    cfg.description  = request.form.get('description', '').strip()
    cfg.allergens    = ','.join(request.form.getlist('allergens'))
    p = request.form.get('price', '').strip()
    cfg.price        = float(p) if p else None
    m = request.form.get('max_bookings', '').strip()
    cfg.max_bookings = int(m) if m else None
    cfg.sort_order   = int(request.form.get('sort_order', 0) or 0)
    db.session.commit()
    flash(f'Configurazione "{cfg.name}" aggiornata.', 'success')
    return redirect(url_for('admin.convenzione_configurazioni', cid=cid))


@bp.route('/convenzioni/<int:cid>/configurazioni/<int:cfg_id>/delete', methods=['POST'])
@require_permission('manage_products')
def configurazione_delete(cid, cfg_id):
    cfg = MealConfiguration.query.get_or_404(cfg_id)
    name = cfg.name
    db.session.delete(cfg)
    db.session.commit()
    flash(f'Configurazione "{name}" eliminata.', 'info')
    return redirect(url_for('admin.convenzione_configurazioni', cid=cid))


@bp.route('/convenzioni/<int:cid>/configurazioni/<int:cfg_id>/json')
@require_permission('manage_products')
def configurazione_json(cid, cfg_id):
    cfg = MealConfiguration.query.get_or_404(cfg_id)
    return jsonify({
        'name':         cfg.name,
        'primo':        cfg.primo        or '',
        'secondo':      cfg.secondo      or '',
        'contorno':     cfg.contorno     or '',
        'bevanda':      cfg.bevanda      or '',
        'caffe':        cfg.caffe        or '',
        'description':  cfg.description  or '',
        'allergens':    [a.strip() for a in (cfg.allergens or '').split(',') if a.strip()],
        'price':        cfg.price,
        'max_bookings': cfg.max_bookings,
    })


# ── Manutenzione ──────────────────────────────────────────────────────────────

@bp.route('/maintenance', methods=['GET', 'POST'])
@login_required
def maintenance():
    if not current_user.is_admin:
        abort(403)

    if request.method == 'POST':
        op = request.form.get('operation', '')

        if op == 'clear_orders':
            db.session.execute(db.text('UPDATE transactions SET order_id = NULL WHERE order_id IS NOT NULL'))
            CustomOrderItemIngredient.query.delete(synchronize_session=False)
            CustomOrderItem.query.delete(synchronize_session=False)
            OrderItem.query.delete(synchronize_session=False)
            CorporateMealBooking.query.delete(synchronize_session=False)
            Order.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Tutti gli ordini eliminati.', 'success')

        elif op == 'clear_reservations':
            TableReservation.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Tutte le prenotazioni tavoli eliminate.', 'success')

        elif op == 'clear_transactions':
            Transaction.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Tutte le transazioni eliminate.', 'success')

        elif op == 'clear_polls':
            PollVote.query.delete(synchronize_session=False)
            PollChoice.query.delete(synchronize_session=False)
            Poll.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Tutti i sondaggi eliminati.', 'success')

        elif op == 'clear_stock':
            DailyStock.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Stock giornaliero eliminato.', 'success')

        elif op == 'clear_movements':
            ConsumableMovement.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Movimenti magazzino eliminati.', 'success')

        elif op == 'clear_clients':
            client_ids = [u.id for u in User.query.filter_by(is_client=True).all()]
            if client_ids:
                # null nullable FK: consumable_movements.user_id
                db.session.execute(
                    db.text('UPDATE consumable_movements SET user_id = NULL WHERE user_id = ANY(:ids)' if db.engine.url.drivername.startswith('postgresql')
                            else 'UPDATE consumable_movements SET user_id = NULL WHERE user_id IN ({})'.format(','.join(str(i) for i in client_ids))),
                    {'ids': client_ids} if db.engine.url.drivername.startswith('postgresql') else {}
                )
                # child tables of orders
                order_ids = [o.id for o in Order.query.filter(Order.user_id.in_(client_ids)).all()]
                if order_ids:
                    db.session.execute(
                        db.text('UPDATE transactions SET order_id = NULL WHERE order_id = ANY(:ids)' if db.engine.url.drivername.startswith('postgresql')
                                else 'UPDATE transactions SET order_id = NULL WHERE order_id IN ({})'.format(','.join(str(i) for i in order_ids))),
                        {'ids': order_ids} if db.engine.url.drivername.startswith('postgresql') else {}
                    )
                    CustomOrderItemIngredient.query.filter(
                        CustomOrderItemIngredient.custom_item_id.in_(
                            db.session.query(CustomOrderItem.id).filter(CustomOrderItem.order_id.in_(order_ids))
                        )
                    ).delete(synchronize_session=False)
                    CustomOrderItem.query.filter(CustomOrderItem.order_id.in_(order_ids)).delete(synchronize_session=False)
                    OrderItem.query.filter(OrderItem.order_id.in_(order_ids)).delete(synchronize_session=False)
                    Order.query.filter(Order.id.in_(order_ids)).delete(synchronize_session=False)
                # other tables directly linked to client users
                PollVote.query.filter(PollVote.user_id.in_(client_ids)).delete(synchronize_session=False)
                CorporateMealBooking.query.filter(CorporateMealBooking.user_id.in_(client_ids)).delete(synchronize_session=False)
                CorporateMembership.query.filter(CorporateMembership.user_id.in_(client_ids)).delete(synchronize_session=False)
                TableReservation.query.filter(TableReservation.user_id.in_(client_ids)).delete(synchronize_session=False)
                Transaction.query.filter(Transaction.user_id.in_(client_ids)).delete(synchronize_session=False)
                User.query.filter_by(is_client=True).delete(synchronize_session=False)
            db.session.commit()
            flash('Tutti i clienti e i loro dati eliminati.', 'success')

        elif op == 'clear_catalog':
            # cancella prodotti e categorie (assume ordini già svuotati o li svuota)
            db.session.execute(db.text('UPDATE transactions SET order_id = NULL WHERE order_id IS NOT NULL'))
            CustomOrderItemIngredient.query.delete(synchronize_session=False)
            CustomOrderItem.query.delete(synchronize_session=False)
            OrderItem.query.delete(synchronize_session=False)
            CorporateMealBooking.query.delete(synchronize_session=False)
            Order.query.delete(synchronize_session=False)
            DailyStock.query.delete(synchronize_session=False)
            PollChoice.query.delete(synchronize_session=False)
            Poll.query.delete(synchronize_session=False)
            Product.query.delete(synchronize_session=False)
            Category.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Prodotti e categorie eliminati.', 'success')

        elif op == 'clear_ingredients':
            CustomOrderItemIngredient.query.delete(synchronize_session=False)
            CustomOrderItem.query.delete(synchronize_session=False)
            Ingredient.query.delete(synchronize_session=False)
            IngredientCategory.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Ingredienti eliminati.', 'success')

        elif op == 'clear_consumables':
            ConsumableMovement.query.delete(synchronize_session=False)
            ConsumableItem.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Consumabili e movimenti eliminati.', 'success')

        elif op == 'clear_suppliers':
            ConsumableMovement.query.delete(synchronize_session=False)
            ConsumableItem.query.delete(synchronize_session=False)
            Supplier.query.delete(synchronize_session=False)
            db.session.commit()
            flash('Fornitori, consumabili e movimenti eliminati.', 'success')

        elif op == 'reset_all':
            # step 1: null nullable FKs
            db.session.execute(db.text('UPDATE transactions SET order_id = NULL WHERE order_id IS NOT NULL'))
            db.session.execute(db.text('UPDATE consumable_movements SET user_id = NULL WHERE user_id IS NOT NULL'))
            # step 2: leaf / child tables
            PollVote.query.delete(synchronize_session=False)
            CorporateMealBooking.query.delete(synchronize_session=False)
            CorporateMembership.query.delete(synchronize_session=False)
            CustomOrderItemIngredient.query.delete(synchronize_session=False)
            CustomOrderItem.query.delete(synchronize_session=False)
            OrderItem.query.delete(synchronize_session=False)
            # step 3: transaction / reservation tables
            Order.query.delete(synchronize_session=False)
            TableReservation.query.delete(synchronize_session=False)
            Transaction.query.delete(synchronize_session=False)
            # step 4: catalog / content tables
            DailyStock.query.delete(synchronize_session=False)
            PollChoice.query.delete(synchronize_session=False)
            Poll.query.delete(synchronize_session=False)
            Product.query.delete(synchronize_session=False)
            Category.query.delete(synchronize_session=False)
            Ingredient.query.delete(synchronize_session=False)
            IngredientCategory.query.delete(synchronize_session=False)
            # step 5: magazzino
            ConsumableMovement.query.delete(synchronize_session=False)
            ConsumableItem.query.delete(synchronize_session=False)
            Supplier.query.delete(synchronize_session=False)
            # step 6: banco sessions (customer_id → users FK)
            BancoSession.query.delete(synchronize_session=False)
            # step 7: users
            is_pg = db.engine.url.drivername.startswith('postgresql')
            client_ids = [r[0] for r in db.session.query(User.id).filter_by(is_client=True).all()]
            if client_ids:
                db.session.flush()
                if is_pg:
                    db.session.execute(db.text('DELETE FROM user_roles WHERE user_id = ANY(:ids)'), {'ids': client_ids})
                    db.session.execute(db.text('DELETE FROM users WHERE id = ANY(:ids)'), {'ids': client_ids})
                else:
                    ids_str = ','.join(str(i) for i in client_ids)
                    db.session.execute(db.text(f'DELETE FROM user_roles WHERE user_id IN ({ids_str})'))
                    db.session.execute(db.text(f'DELETE FROM users WHERE id IN ({ids_str})'))
            db.session.commit()
            flash('Reset completo eseguito.', 'success')

        elif op == 'run_demo_seed':
            from app.demo_seed import reset_demo_data, seed_demo_data
            reset_demo_data()
            ok, msg = seed_demo_data()
            flash(f'Seed demo: {msg}', 'success' if ok else 'warning')
            return redirect(url_for('admin.maintenance'))

        elif op == 'reset_admin':
            import re as _re
            email    = request.form.get('admin_email', '').strip().lower()
            password = request.form.get('admin_password', '').strip()
            username = request.form.get('admin_username', 'admin').strip()
            if not email or '@' not in email or len(password) < 6:
                flash('Email valida e password (min 6 caratteri) obbligatorie.', 'danger')
            else:
                user = User.query.filter_by(email=email).first()
                if user:
                    user.is_admin  = True
                    user.is_active = True
                    user.set_password(password)
                    db.session.commit()
                    flash(f'Admin "{email}" aggiornato.', 'success')
                else:
                    base  = _re.sub(r'[^a-z0-9]', '.', username.lower()).strip('.') or 'admin'
                    uname, n = base[:30], 1
                    while User.query.filter_by(username=uname).first():
                        uname = f'{base[:28]}{n}'; n += 1
                    new_admin = User(username=uname, email=email,
                                     is_admin=True, is_active=True)
                    new_admin.set_password(password)
                    db.session.add(new_admin)
                    db.session.commit()
                    flash(f'Admin "{email}" creato con username "{uname}".', 'success')

        return redirect(url_for('admin.maintenance'))

    stats = {
        'orders':       Order.query.count(),
        'reservations': TableReservation.query.count(),
        'transactions': Transaction.query.count(),
        'polls':        Poll.query.count(),
        'stock':        DailyStock.query.count(),
        'movements':    ConsumableMovement.query.count(),
        'clients':      User.query.filter_by(is_client=True).count(),
        'products':     Product.query.count(),
        'categories':   Category.query.count(),
        'ingredients':  Ingredient.query.count(),
        'consumables':  ConsumableItem.query.count(),
        'suppliers':    Supplier.query.count(),
        'admins':       User.query.filter_by(is_admin=True).all(),
    }
    return render_template('admin/maintenance.html', stats=stats)


@bp.route('/convenzioni/<int:cid>/presenze')
@require_permission('manage_products')
def convenzione_presenze(cid):
    from sqlalchemy import extract
    corp = CorporateAccount.query.get_or_404(cid)

    try:
        mese_str = request.args.get('mese', '')
        if mese_str:
            anno, mese = int(mese_str[:4]), int(mese_str[5:7])
        else:
            anno, mese = date.today().year, date.today().month
    except (ValueError, IndexError):
        anno, mese = date.today().year, date.today().month

    meals = (DailyFixedMeal.query
             .filter_by(corporate_id=cid)
             .filter(extract('year', DailyFixedMeal.meal_date) == anno)
             .filter(extract('month', DailyFixedMeal.meal_date) == mese)
             .order_by(DailyFixedMeal.meal_date.desc())
             .all())

    days = []
    for m in meals:
        consumed  = [b for b in m.bookings if b.status == 'consumed']
        booked    = [b for b in m.bookings if b.status == 'booked']
        cancelled = [b for b in m.bookings if b.status == 'cancelled']
        days.append({'meal': m, 'consumed': consumed,
                     'booked': booked, 'cancelled': cancelled})

    if mese == 1:
        prev_anno, prev_mese = anno - 1, 12
    else:
        prev_anno, prev_mese = anno, mese - 1
    if mese == 12:
        next_anno, next_mese = anno + 1, 1
    else:
        next_anno, next_mese = anno, mese + 1

    MESI = ['', 'Gennaio', 'Febbraio', 'Marzo', 'Aprile', 'Maggio', 'Giugno',
            'Luglio', 'Agosto', 'Settembre', 'Ottobre', 'Novembre', 'Dicembre']

    return render_template('admin/convenzione_presenze.html',
                           corp=corp, days=days, anno=anno, mese=mese,
                           nome_mese=MESI[mese],
                           prev=f'{prev_anno}-{prev_mese:02d}',
                           nxt=f'{next_anno}-{next_mese:02d}')


# ── Ritiro pasti aziendali ────────────────────────────────────────────────────

@bp.route('/pasti/ritiro', methods=['GET', 'POST'])
@require_permission('manage_products')
def meal_ritiro():
    booking = None
    token   = ''
    if request.method == 'POST':
        token = request.form.get('token', '').strip().upper()
        if token:
            booking = CorporateMealBooking.query.filter_by(
                pickup_token=token, status='booked').first()
            if not booking:
                flash('Codice non trovato o pasto già consegnato.', 'warning')
    return render_template('admin/meal_ritiro.html', booking=booking, token=token)


@bp.route('/pasti/consegna', methods=['POST'])
@require_permission('manage_products')
def meal_consegna():
    bid = request.form.get('booking_id', type=int)
    booking = CorporateMealBooking.query.get_or_404(bid)
    if booking.status == 'booked':
        booking.status = 'consumed'
        db.session.commit()
        send_telegram_to_user(
            booking.user,
            f'✅ Pasto ritirato: <b>{booking.meal.name}</b>. Buon appetito!')
        flash(f'Consegnato a {booking.user.full_name}.', 'success')
    else:
        flash('Prenotazione non in stato "prenotato".', 'warning')
    next_url = request.form.get('next') or url_for('admin.meal_ritiro')
    return redirect(next_url)


# ── Report giornaliero pasti aziendali ────────────────────────────────────────

@bp.route('/convenzioni/report')
@require_permission('manage_products')
def convenzioni_report():
    tf = _tenant_filter()
    try:
        sel_date = date.fromisoformat(request.args.get('d', ''))
    except (ValueError, TypeError):
        sel_date = date.today()

    corps = CorporateAccount.query.filter_by(is_active=True, **tf)\
        .order_by(CorporateAccount.name).all()

    report = []
    for corp in corps:
        meals   = DailyFixedMeal.query.filter_by(corporate_id=corp.id, meal_date=sel_date).all()
        entries = []
        for meal in meals:
            for b in meal.bookings:
                if b.status != 'cancelled':
                    entries.append({'booking': b, 'meal': meal, 'user': b.user})
        entries.sort(key=lambda x: (x['user'].last_name or '', x['user'].first_name or '',
                                    x['user'].username))
        report.append({'corp': corp, 'entries': entries})

    prev_date = sel_date - timedelta(days=1)
    next_date = sel_date + timedelta(days=1)
    return render_template('admin/convenzioni_report.html',
                           report=report, sel_date=sel_date, today=date.today(),
                           prev_date=prev_date, next_date=next_date)


@bp.route('/convenzioni/<int:cid>/report-pdf')
@require_permission('manage_products')
def convenzione_report_pdf(cid):
    import os
    from io import BytesIO
    from pathlib import Path
    from flask import send_file
    from fpdf import FPDF

    BRAND  = (233,  69,  96)
    DARK   = ( 26,  26,  46)
    DGRAY  = ( 80,  80,  95)
    LGRAY  = (200, 200, 210)
    VLIGHT = (248, 248, 252)
    WHITE  = (255, 255, 255)
    GREEN  = ( 39, 174,  96)
    FONT   = 'PTSansNarrow'

    _IT_MONTHS = ['gennaio','febbraio','marzo','aprile','maggio','giugno',
                  'luglio','agosto','settembre','ottobre','novembre','dicembre']
    _IT_DAYS   = ['Lunedì','Martedì','Mercoledì','Giovedì','Venerdì','Sabato','Domenica']

    corp = CorporateAccount.query.get_or_404(cid)
    try:
        sel_date = date.fromisoformat(request.args.get('d', ''))
    except (ValueError, TypeError):
        sel_date = date.today()

    meals = DailyFixedMeal.query.filter_by(corporate_id=corp.id, meal_date=sel_date).all()

    entries    = []
    n_consumed = 0
    n_booked   = 0
    for meal in meals:
        for b in meal.bookings:
            if b.status == 'cancelled' or not b.user:
                continue
            qty = b.quantity or 1
            entries.append({
                'full_name': (b.user.full_name or b.user.username or '')[:38],
                'meal_name': meal.name[:38],
                'qty':       qty,
                'status':    b.status,
                'slot':      b.slot.time_str if b.slot else '',
                'price':     float(meal.price),
            })
            if b.status == 'consumed':
                n_consumed += qty
            else:
                n_booked   += qty
    entries.sort(key=lambda x: x['full_name'])
    n_total = n_consumed + n_booked

    _co_name    = get_setting('company_name')    or 'QuickLunch Bar'
    _co_address = get_setting('company_address') or ''
    _co_city    = get_setting('company_city')    or ''
    _co_vat     = get_setting('company_vat')     or ''
    _co_phone   = get_setting('company_phone')   or ''
    _co_email   = get_setting('company_email')   or ''

    date_str = (f'{_IT_DAYS[sel_date.weekday()]} {sel_date.day}'
                f' {_IT_MONTHS[sel_date.month - 1]} {sel_date.year}')

    _font_dir = Path(os.path.abspath(__file__)).parent.parent / 'static' / 'fonts'

    class ReportPDF(FPDF):
        def header(self):
            self.set_fill_color(*BRAND)
            self.rect(0, 0, 210, 11, 'F')
            self.set_fill_color(*DARK)
            self.rect(0, 0, 8, 11, 'F')
            self.set_font(FONT, 'B', 9)
            self.set_text_color(*WHITE)
            self.set_xy(12, 1.5)
            self.cell(130, 8,
                      f'REPORT PASTI AZIENDALI  \xb7  {corp.name.upper()}', ln=0)
            self.set_font(FONT, '', 9)
            self.set_x(-58)
            self.cell(46, 8, date_str.upper(), align='R')
            self.set_text_color(*DARK)
            self.ln(14)

        def footer(self):
            self.set_y(-14)
            self.set_font(FONT, '', 8)
            self.set_text_color(*LGRAY)
            self.set_draw_color(*LGRAY)
            self.set_line_width(0.25)
            self.line(12, self.get_y(), 198, self.get_y())
            self.ln(1.5)
            now_str = _dt.now().strftime('%d/%m/%Y  %H:%M')
            self.set_x(12)
            self.cell(155, 5,
                      f'Generato il {now_str}  \xb7  Documento riservato  \xb7  {_co_name}',
                      ln=0)
            self.set_x(-25)
            self.cell(13, 5, f'Pag. {self.page_no()}', align='R')

    pdf = ReportPDF(orientation='P', unit='mm', format='A4')
    pdf.set_margins(12, 16, 12)
    pdf.set_auto_page_break(True, margin=20)
    pdf.add_font(FONT, '',  str(_font_dir / 'PTSansNarrow-Regular.ttf'))
    pdf.add_font(FONT, 'B', str(_font_dir / 'PTSansNarrow-Bold.ttf'))
    pdf.add_page()

    # ── Mittente (intestazione bar) ───────────────────────────────────────────
    pdf.set_font(FONT, 'B', 13)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 7, _co_name, ln=True)
    co_parts = [p for p in [
        _co_address, _co_city, _co_phone, _co_email,
        (f'P.IVA {_co_vat}' if _co_vat else ''),
    ] if p]
    if co_parts:
        pdf.set_font(FONT, '', 9)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 5, '  \xb7  '.join(co_parts), ln=True)
    pdf.ln(4)

    # Separatore BRAND
    pdf.set_draw_color(*BRAND)
    pdf.set_line_width(0.8)
    pdf.line(12, pdf.get_y(), 198, pdf.get_y())
    pdf.ln(6)

    # ── Destinatario / oggetto ────────────────────────────────────────────────
    pdf.set_font(FONT, '', 9)
    pdf.set_text_color(*DGRAY)
    pdf.cell(0, 5, 'A:', ln=True)
    pdf.set_font(FONT, 'B', 18)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 9, corp.name, ln=True)
    pdf.set_font(FONT, 'B', 10)
    pdf.set_text_color(*BRAND)
    pdf.cell(0, 6, f'LISTA PASTI  \xb7  {date_str.upper()}', ln=True)
    pdf.ln(6)

    # ── KPI boxes ─────────────────────────────────────────────────────────────
    kpi_y = pdf.get_y()

    def _kpi(x, label, value, color):
        pdf.set_fill_color(*VLIGHT)
        pdf.set_draw_color(*LGRAY)
        pdf.set_line_width(0.3)
        pdf.rect(x, kpi_y, 59, 26, 'FD')
        pdf.set_fill_color(*color)
        pdf.rect(x, kpi_y, 59, 4, 'F')
        pdf.set_font(FONT, 'B', 26)
        pdf.set_text_color(*color)
        pdf.set_xy(x, kpi_y + 5)
        pdf.cell(59, 12, str(value), align='C')
        pdf.set_font(FONT, '', 8.5)
        pdf.set_text_color(*DGRAY)
        pdf.set_xy(x, kpi_y + 18.5)
        pdf.cell(59, 5, label, align='C')

    _kpi(12,  'TOTALE PASTI',  n_total,    DARK)
    _kpi(76,  'CONSUMATI',     n_consumed, GREEN)
    _kpi(140, 'IN ATTESA',     n_booked,   BRAND)
    pdf.set_y(kpi_y + 33)

    # ── Opzioni del giorno ────────────────────────────────────────────────────
    if meals:
        pdf.set_font(FONT, 'B', 10)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 7, 'OPZIONI DEL GIORNO', ln=True)
        pdf.set_draw_color(*LGRAY)
        pdf.set_line_width(0.25)
        pdf.line(12, pdf.get_y(), 198, pdf.get_y())
        pdf.ln(3)

        for meal in meals:
            n_bk = sum(b.quantity or 1 for b in meal.bookings if b.status != 'cancelled')
            pdf.set_font(FONT, 'B', 10)
            pdf.set_text_color(*DARK)
            pdf.set_x(12)
            pdf.cell(143, 6, meal.name, ln=0)
            pdf.set_font(FONT, '', 9)
            pdf.set_text_color(*DGRAY)
            pdf.cell(43, 6,
                     f'{n_bk} pren.  \xb7  € {numero_italiano(meal.price)}',
                     align='R', ln=True)
            if meal.courses:
                pdf.set_font(FONT, '', 8.5)
                pdf.set_text_color(*DGRAY)
                pdf.set_x(16)
                cstr = '  \xb7  '.join(
                    f'{lbl}: {val}' for lbl, _icon, val in meal.courses if val)
                if cstr:
                    pdf.multi_cell(180, 4.5, cstr)
            pdf.ln(2)
        pdf.ln(3)

    # ── Elenco prenotazioni ───────────────────────────────────────────────────
    pdf.set_font(FONT, 'B', 10)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 7, 'ELENCO PRENOTAZIONI', ln=True)
    pdf.set_draw_color(*LGRAY)
    pdf.set_line_width(0.25)
    pdf.line(12, pdf.get_y(), 198, pdf.get_y())
    pdf.ln(3)

    if not entries:
        pdf.set_font(FONT, '', 10)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 8, 'Nessuna prenotazione per questa data.', ln=True)
    else:
        # Colonne: # | Nominativo | Pasto | Orario | Qtà | Stato = 186mm
        CW  = [7, 67, 63, 17, 12, 20]
        HDR = ['#', 'NOMINATIVO', 'PASTO', 'ORARIO', 'QTÀ', 'STATO']
        ALN = ['C', 'L', 'L', 'C', 'C', 'C']
        TH  = 7

        # Intestazione tabella
        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        pdf.set_font(FONT, 'B', 8.5)
        for w, ht, al in zip(CW, HDR, ALN):
            pdf.cell(w, TH, ht, align=al, fill=True)
        pdf.ln()
        pdf.set_draw_color(*BRAND)
        pdf.set_line_width(0.5)
        pdf.line(12, pdf.get_y(), 198, pdf.get_y())

        STATUS_IT  = {'consumed': 'Consumato', 'booked': 'Prenotato'}
        STATUS_CLR = {'consumed': GREEN, 'booked': BRAND}

        for idx, e in enumerate(entries):
            fill = idx % 2 == 0
            pdf.set_fill_color(*(VLIGHT if fill else WHITE))
            pdf.set_draw_color(*LGRAY)
            pdf.set_line_width(0.15)
            RH = 6.5

            pdf.set_font(FONT, '', 9)
            pdf.set_text_color(*DARK)
            pdf.cell(CW[0], RH, str(idx + 1),    align='C', fill=fill, border='B')
            pdf.cell(CW[1], RH, e['full_name'],               fill=fill, border='B')
            pdf.cell(CW[2], RH, e['meal_name'],               fill=fill, border='B')
            pdf.cell(CW[3], RH, e['slot'],        align='C', fill=fill, border='B')
            pdf.cell(CW[4], RH, str(e['qty']),    align='C', fill=fill, border='B')
            pdf.set_text_color(*STATUS_CLR.get(e['status'], DGRAY))
            pdf.set_font(FONT, 'B', 8.5)
            pdf.cell(CW[5], RH,
                     STATUS_IT.get(e['status'], e['status']),
                     align='C', fill=fill, border='B')
            pdf.ln()

        # Riga totale
        span = CW[0] + CW[1] + CW[2] + CW[3]
        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        pdf.set_font(FONT, 'B', 9)
        pdf.cell(span, TH, f'TOTALE  ({len(entries)} dipendenti)', fill=True)
        pdf.cell(CW[4], TH, str(n_total), align='C', fill=True)
        pdf.cell(CW[5], TH, '', fill=True)
        pdf.ln()

        # Fatturabile
        total_fat = sum(e['qty'] * e['price'] for e in entries)
        pdf.ln(4)
        pdf.set_font(FONT, '', 10)
        pdf.set_text_color(*DGRAY)
        pdf.set_x(12)
        pdf.cell(50, 6, 'Totale fatturabile:', ln=0)
        pdf.set_font(FONT, 'B', 10)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 6, f'€ {numero_italiano(total_fat)}', ln=True)

    buf = BytesIO(bytes(pdf.output()))
    buf.seek(0)
    filename = f'pasti_{corp.name.replace(" ", "_")}_{sel_date.isoformat()}.pdf'
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/pdf')

@bp.route('/convenzioni/abstract-docx')
@require_permission('manage_products')
def convenzioni_abstract_docx():
    """Genera il documento descrittivo del modulo Pasti Aziendali."""
    from io import BytesIO
    from flask import send_file
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BRAND = '#E94560'
    DARK  = '#1a1a2e'
    GRAY  = '#666666'

    def _rgb(h):
        h = h.lstrip('#')
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    FONT   = 'PT Sans Narrow'
    LS     = Pt(13)   # interlinea compatta uniforme
    SA     = Pt(3)    # spazio dopo paragrafo ridotto
    SB     = Pt(0)    # spazio prima paragrafo

    _IT_MONTHS_D = ['gennaio','febbraio','marzo','aprile','maggio','giugno',
                    'luglio','agosto','settembre','ottobre','novembre','dicembre']

    def _it_mese_anno(dt):
        return f'{_IT_MONTHS_D[dt.month - 1]} {dt.year}'

    def _set_spacing(p, ls=LS, sa=SA, sb=SB):
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing  = ls
        pf.space_after   = sa
        pf.space_before  = sb

    def _fnt(run, fname, size_pt=None):
        """Imposta il font via XML diretto, eliminando i riferimenti al tema."""
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rFonts.set(qn(attr), fname)
        for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
            rFonts.attrib.pop(qn(attr), None)
        if size_pt is not None:
            run.font.size = size_pt

    def _fix_style_font(style, fname):
        """Imposta il font sul rPr dello stile tramite XML."""
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rFonts.set(qn(attr), fname)
        for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
            rFonts.attrib.pop(qn(attr), None)

    def _heading(doc, text, level=1, color=BRAND):
        p = doc.add_heading(text, level)
        pf = p.paragraph_format
        pf.space_before = Pt(6) if level == 1 else Pt(4)
        pf.space_after  = Pt(2)
        for run in p.runs:
            _fnt(run, FONT)
            run.font.color.rgb = _rgb(color)
        return p

    def _body(doc, text):
        p = doc.add_paragraph(text)
        _set_spacing(p)
        for run in p.runs:
            _fnt(run, FONT, Pt(11))
        return p

    def _bullet(doc, text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        _set_spacing(p, sa=Pt(2))
        if bold_prefix:
            r = p.add_run(bold_prefix + ' ')
            r.bold = True
            _fnt(r, FONT, Pt(11))
            r.font.color.rgb = _rgb(BRAND)
        r2 = p.add_run(text)
        _fnt(r2, FONT, Pt(11))

    def _step(doc, num, title, body):
        p = doc.add_paragraph()
        _set_spacing(p, sa=Pt(3))
        p.paragraph_format.left_indent = Inches(0.15)
        r1 = p.add_run(f'{num}. {title}  ')
        r1.bold = True
        _fnt(r1, FONT, Pt(11))
        r1.font.color.rgb = _rgb(BRAND)
        r2 = p.add_run(body)
        _fnt(r2, FONT, Pt(11))

    def _gap(doc, h=Pt(4)):
        p = doc.add_paragraph()
        _set_spacing(p, ls=h, sa=Pt(0))

    def _divider(doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'E94560')
        pBdr.append(bot)
        pPr.append(pBdr)

    _co_name    = get_setting('company_name')    or ''
    _co_address = get_setting('company_address') or ''
    _co_city    = get_setting('company_city')    or ''
    _co_vat     = get_setting('company_vat')     or ''
    _co_phone   = get_setting('company_phone')   or ''
    _co_email   = get_setting('company_email')   or ''

    doc = Document()

    # Font + interlinea su tutti gli stili principali (via XML per bypassare il tema)
    for sname in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                  'List Bullet', 'List Bullet 2'):
        try:
            st = doc.styles[sname]
            _fix_style_font(st, FONT)
            st.font.name = FONT
            if sname == 'Normal':
                st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                st.paragraph_format.line_spacing  = LS
                st.paragraph_format.space_after   = SA
                st.paragraph_format.space_before  = SB
        except Exception:
            pass

    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Copertina ─────────────────────────────────────────────────────────────
    for _ in range(2):
        _gap(doc, Pt(8))
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run('QuickLunch — Bar Self-Service')
    _fnt(r, FONT, Pt(26)); r.bold = True; r.font.color.rgb = _rgb(BRAND)
    _set_spacing(tp, sa=Pt(4))
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run('Modulo Pasti Aziendali')
    _fnt(r, FONT, Pt(18)); r.font.color.rgb = _rgb(DARK)
    _set_spacing(sp, sa=Pt(6))
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run('Documento descrittivo per le aziende convenzionate')
    _fnt(r, FONT, Pt(12)); r.italic = True; r.font.color.rgb = _rgb(GRAY)
    _set_spacing(dp, sa=Pt(8))
    datp = doc.add_paragraph()
    datp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = datp.add_run(_it_mese_anno(_dt.now()))
    _fnt(r, FONT, Pt(11)); r.font.color.rgb = _rgb(GRAY)
    _set_spacing(datp)
    if _co_name:
        _gap(doc, Pt(6))
        cop = doc.add_paragraph()
        cop.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cop.add_run(_co_name)
        _fnt(r, FONT, Pt(11)); r.bold = True
        _set_spacing(cop, sa=Pt(2))
        co_details = []
        if _co_address: co_details.append(_co_address)
        if _co_city:    co_details.append(_co_city)
        if _co_vat:     co_details.append(f'P.IVA {_co_vat}')
        if _co_phone:   co_details.append(_co_phone)
        if _co_email:   co_details.append(_co_email)
        if co_details:
            co_dp = doc.add_paragraph()
            co_dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = co_dp.add_run('  |  '.join(co_details))
            _fnt(r, FONT, Pt(9)); r.font.color.rgb = _rgb(GRAY)
            _set_spacing(co_dp)
    doc.add_page_break()

    # ── 1. Introduzione ───────────────────────────────────────────────────────
    _heading(doc, '1. Introduzione')
    _divider(doc)
    _body(doc,
        "QuickLunch è il sistema digitale di gestione del bar e della mensa self-service. "
        "Il modulo Pasti Aziendali consente alle aziende convenzionate di offrire ai propri "
        "dipendenti un servizio di prenotazione pasto strutturato, tracciato e amministrabile "
        "in tempo reale.")
    _body(doc,
        "Ogni azienda dispone di un proprio spazio con configurazioni personalizzate "
        "(prezzo del pasto, numero massimo di coperti, portate). I dipendenti prenotano "
        "autonomamente tramite applicazione web da qualsiasi dispositivo, senza installazione.")
    _gap(doc)

    # ── 2. Attori ─────────────────────────────────────────────────────────────
    _heading(doc, '2. Attori del sistema')
    _divider(doc)
    _heading(doc, 'Amministratore / Staff del bar', 2, color=DARK)
    _bullet(doc, 'Gestisce le convenzioni aziendali (creazione, modifica, disattivazione).')
    _bullet(doc, 'Pubblica il menu giornaliero per ciascuna azienda (portate, prezzo, massimo coperti).')
    _bullet(doc, 'Visualizza e scarica il report giornaliero delle presenze in formato PDF/Word.')
    _bullet(doc, 'Segna i pasti come Consumati in tempo reale durante il servizio, '
                 'anche direttamente dal terminale Banco tramite il pannello dedicato.')
    _bullet(doc, 'Consulta lo storico mensile delle presenze per convenzione.')
    _gap(doc)
    _heading(doc, 'Dipendente (utente finale)', 2, color=DARK)
    _bullet(doc, "Si registra nell'applicazione (email + password oppure account Google OAuth 2.0).")
    _bullet(doc, "Viene associato alla propria azienda convenzionata dall'amministratore.")
    _bullet(doc, 'Visualizza il menu del giorno e prenota con pochi clic da smartphone, tablet o PC.')
    _bullet(doc, "Può annullare la prenotazione in autonomia entro i termini configurati.")
    _bullet(doc, 'Riceve notifiche Telegram di conferma prenotazione, reminder e conferma ritiro.')
    _gap(doc)

    # ── 3. Ciclo completo ─────────────────────────────────────────────────────
    _heading(doc, '3. Ciclo completo: dalla prenotazione alla consegna')
    _divider(doc)
    _body(doc,
        "Il ciclo di vita di ogni pasto aziendale si articola in cinque fasi: pubblicazione "
        "del menu, prenotazione, generazione del codice di ritiro, consegna al banco e "
        "riconciliazione.")
    _gap(doc)

    _heading(doc, 'Fase 1 — Pubblicazione del menu (amministratore)', 2, color=DARK)
    _step(doc, 1, 'Accesso:', "ogni mattina l'amministratore accede a Convenzioni e seleziona l'azienda.")
    _step(doc, 2, 'Creazione del menu:', "inserisce nome del pasto, portate (primo, secondo, contorno, "
          "bevanda, caffè), prezzo e numero massimo di coperti. Supporta più opzioni per la stessa "
          "giornata e template riutilizzabili per menu ricorrenti.")
    _step(doc, 3, 'Pubblicazione:', "il menu diventa immediatamente visibile ai dipendenti dell'azienda.")
    _gap(doc)

    _heading(doc, 'Fase 2 — Prenotazione (dipendente)', 2, color=DARK)
    _step(doc, 1, 'Accesso:', "il dipendente apre il browser e accede con email/password o account Google.")
    _step(doc, 2, 'Visualizzazione menu:', "nella sezione Pasto Aziendale vede il menu del giorno "
          "con portate, prezzo e posti ancora disponibili.")
    _step(doc, 3, 'Selezione e conferma:', "sceglie l'opzione, le porzioni e l'orario di ritiro "
          "(se configurato), quindi preme Prenota. Il sistema verifica la disponibilità e registra "
          "la prenotazione in stato Prenotato.")
    _step(doc, 4, 'Conferma immediata:', "compare la conferma con il riepilogo del pasto scelto. "
          "Se il Telegram è configurato, il dipendente riceve un messaggio di conferma in tempo reale.")
    _step(doc, 5, 'Annullamento (opzionale):', "il dipendente può annullare in autonomia entro "
          "i termini configurati. Oltre tale soglia il sistema blocca l'annullamento.")
    _gap(doc)

    _heading(doc, 'Fase 3 — Codice di ritiro', 2, color=DARK)
    _body(doc,
        "Al momento della prenotazione il sistema genera automaticamente un codice di ritiro "
        "univoco (es. A7K3M2). Il codice identifica la prenotazione senza necessità di cercare "
        "il nominativo manualmente e rimane visibile al dipendente fino alla consegna.")
    _step(doc, 1, 'Visualizzazione:', "il codice appare in grande nella pagina Pasto Aziendale "
          "del dipendente, sempre accessibile da smartphone.")
    _step(doc, 2, 'Reminder automatico:', "N minuti prima del ritiro (configurabile in Impostazioni) "
          "il sistema invia un reminder Telegram con codice e nome del pasto.")
    _step(doc, 3, 'Utilizzo al banco:', "il dipendente mostra il codice sullo schermo del proprio "
          "dispositivo allo staff oppure lo comunica verbalmente.")
    _gap(doc)

    _heading(doc, 'Fase 4 — Consegna al banco (staff)', 2, color=DARK)
    _body(doc, "Lo staff dispone di tre modalità equivalenti per registrare la consegna:")
    _step(doc, 1, 'Pannello Banco POS (novità):',
          "direttamente dal terminale cassa (schermata Banco), il pannello "
          "\"Pasto Aziendale\" consente di digitare o scansionare (QR da fotocamera) "
          "il codice di ritiro. Il sistema mostra immediatamente nome del dipendente, "
          "pasto, data e orario di slot. Un click su \"Conferma Consegna\" chiude la "
          "pratica, aggiorna lo stato in Consumato e invia la notifica Telegram al dipendente.")
    _step(doc, 2, 'Sezione Ritiro Pasti (admin):',
          "da Convenzioni → Ritiro Pasti, lo staff digita il codice a caratteri mostrato "
          "dal dipendente. Utile da PC o tablet senza accesso al Banco POS.")
    _step(doc, 3, 'Report giornaliero:',
          "da Convenzioni → Report Giornaliero ogni riga Prenotato riporta un pulsante Consegna. "
          "Utile per spuntare più prenotazioni in sequenza o per correzioni a fine servizio.")
    _gap(doc)

    _heading(doc, 'Fase 5 — Riconciliazione (amministratore)', 2, color=DARK)
    _step(doc, 1, 'Scarica il report PDF:', "a fine servizio il pulsante Scarica PDF genera "
          "un documento per ciascuna azienda con: intestazione azienda/data, totale prenotazioni "
          "e tabella Nominativo / Pasto / Quantità / Stato.")
    _step(doc, 2, "Invio all'azienda:", "il documento viene inviato al referente aziendale "
          "per la verifica e la riconciliazione contabile del mese.")
    _step(doc, 3, 'Storico mensile:', "la sezione Storico Presenze mantiene l'archivio completo "
          "con contatori di prenotati, consumati e annullati per ogni giornata.")
    _gap(doc)

    # ── 4. Stati prenotazione ─────────────────────────────────────────────────
    _heading(doc, '4. Stati della prenotazione')
    _divider(doc)
    _body(doc, "Ogni prenotazione transita attraverso i seguenti stati:")
    tbl_s = doc.add_table(rows=1, cols=3)
    tbl_s.style = 'Table Grid'
    for i, h in enumerate(['Stato', 'Descrizione', 'Transizione successiva']):
        r = tbl_s.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True; _fnt(r, FONT, Pt(10))
    for stato, descr, trans in [
        ('Prenotato',
         'Prenotazione attiva. Codice di ritiro visibile al dipendente.',
         'Consumato (consegna confermata) oppure Annullato'),
        ('Consumato',
         'Pasto ritirato e confermato dallo staff al banco.',
         'Stato finale — incluso nel report'),
        ('Annullato',
         'Dipendente ha annullato entro i termini.',
         'Stato finale — escluso dal conteggio pasti'),
    ]:
        row = tbl_s.add_row().cells
        for i, val in enumerate([stato, descr, trans]):
            p2 = row[i].paragraphs[0]
            _set_spacing(p2, sa=Pt(2))
            r = p2.add_run(val)
            _fnt(r, FONT, Pt(10))
            if i == 0: r.bold = True
    _gap(doc)

    # ── 5. Configurazione ─────────────────────────────────────────────────────
    _heading(doc, '5. Configurazione della convenzione')
    _divider(doc)
    _body(doc, 'Ogni azienda viene configurata con i seguenti parametri:')
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    h0, h1 = tbl.rows[0].cells
    for cell, txt in [(h0, 'Parametro'), (h1, 'Descrizione')]:
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True; _fnt(r, FONT, Pt(10))
    for pname, pdesc in [
        ('Nome azienda',           "Ragione sociale dell'azienda convenzionata."),
        ('Email di contatto',      'Referente aziendale per le comunicazioni.'),
        ('Prezzo giornaliero (€)', 'Prezzo standard del pasto, sovrascrivibile per singola giornata.'),
        ('Massimo coperti/giorno', 'Numero massimo di prenotazioni accettabili per giornata.'),
        ('Dipendenti associati',   'Utenti registrati associati alla convenzione.'),
        ('Template menu',          'Configurazioni riutilizzabili per menu ricorrenti.'),
    ]:
        row = tbl.add_row().cells
        for cell, txt, bold in [(row[0], pname, True), (row[1], pdesc, False)]:
            p2 = cell.paragraphs[0]
            _set_spacing(p2, sa=Pt(2))
            r = p2.add_run(txt)
            r.bold = bold; _fnt(r, FONT, Pt(10))
    _gap(doc)

    # ── 6. Report ─────────────────────────────────────────────────────────────
    _heading(doc, '6. Report e documenti disponibili')
    _divider(doc)
    _heading(doc, 'Report Giornaliero (pagina web)', 2, color=DARK)
    _body(doc, "Accessibile da Convenzioni → Report Giornaliero: lista dipendenti con prenotazione "
          "del giorno. Navigazione tra date con tasti freccia o selettore calendario.")
    _heading(doc, 'Report Giornaliero (PDF scaricabile)', 2, color=DARK)
    _body(doc, "Il pulsante Scarica PDF genera un documento con intestazione azienda/data, "
          "totale prenotazioni e tabella: Nominativo, Pasto, Quantità, Stato.")
    _heading(doc, 'Storico mensile presenze', 2, color=DARK)
    _body(doc, "Accessibile da Convenzioni → [Azienda] → Storico Presenze. Mostra mese per mese "
          "il dettaglio con contatori di prenotati, consumati e annullati.")
    _gap(doc)

    # ── 7. Notifiche ──────────────────────────────────────────────────────────
    _heading(doc, '7. Notifiche automatiche (Telegram)')
    _divider(doc)
    _body(doc, "Opzionale — richiede Telegram Chat ID configurato nel profilo utente:")
    _bullet(doc, 'Messaggio di conferma al momento della prenotazione.', 'Conferma prenotazione:')
    _bullet(doc, "Reminder N minuti prima del ritiro (configurabile in Impostazioni).", 'Reminder:')
    _bullet(doc, 'Messaggio di conferma alla registrazione della consegna da parte dello staff.',
            'Conferma ritiro:')
    _bullet(doc, 'Messaggio di conferma alla cancellazione della prenotazione.', 'Annullamento:')
    _gap(doc)

    # ── 8. Sicurezza ──────────────────────────────────────────────────────────
    _heading(doc, '8. Accesso e sicurezza')
    _divider(doc)
    _bullet(doc, 'Email + password oppure account Google (OAuth 2.0).', 'Autenticazione:')
    _bullet(doc, 'Autenticazione a due fattori opzionale (TOTP — Google Authenticator).', 'MFA:')
    _bullet(doc, 'Ciascun dipendente vede solo le informazioni della propria azienda.', 'Isolamento dati:')
    _bullet(doc, 'Tutte le comunicazioni avvengono su connessione cifrata HTTPS.', 'Trasporto:')
    _gap(doc)

    # ── 9. GDPR e conservazione dei dati ─────────────────────────────────────
    _heading(doc, '9. Conformità GDPR e conservazione dei dati')
    _divider(doc)
    _body(doc,
        "Il sistema QuickLunch è progettato per essere conforme al Regolamento Europeo sulla "
        "protezione dei dati personali (GDPR — Reg. UE 2016/679). Di seguito i principi "
        "applicati e le misure adottate.")
    _gap(doc)

    _heading(doc, 'Dati trattati', 2, color=DARK)
    _body(doc, "Il sistema tratta le seguenti categorie di dati personali dei dipendenti:")
    _bullet(doc, 'Dati anagrafici: nome, cognome, indirizzo e-mail.')
    _bullet(doc, 'Dati di autenticazione: password cifrata (bcrypt) — mai memorizzata in chiaro; '
                 'oppure token OAuth Google.')
    _bullet(doc, 'Dati operativi: prenotazioni pasto, cronologia acquisti, saldo wallet.')
    _bullet(doc, 'Identificativi facoltativi: Telegram Chat ID (solo se fornito volontariamente).')
    _body(doc,
        "Non vengono raccolti dati sensibili (sanitari, biometrici, politici) né vengono "
        "effettuate profilazioni automatizzate con effetti giuridici.")
    _gap(doc)

    _heading(doc, 'Base giuridica del trattamento', 2, color=DARK)
    _bullet(doc, 'Esecuzione di un contratto (prenotazione pasto, erogazione del servizio).', 'Art. 6.1.b:')
    _bullet(doc, 'Adempimento di obblighi contabili e fiscali.', 'Art. 6.1.c:')
    _bullet(doc, 'Notifiche Telegram e funzionalità opzionali attivate dall\'interessato.', 'Art. 6.1.a (consenso):')
    _gap(doc)

    _heading(doc, 'Diritti degli interessati', 2, color=DARK)
    _body(doc, "Ogni dipendente può esercitare i seguenti diritti contattando il titolare del trattamento:")
    _bullet(doc, 'Accesso ai propri dati personali (art. 15).')
    _bullet(doc, 'Rettifica di dati inesatti (art. 16).')
    _bullet(doc, 'Cancellazione — «diritto all\'oblio» (art. 17).')
    _bullet(doc, 'Limitazione del trattamento (art. 18).')
    _bullet(doc, 'Portabilità dei dati in formato strutturato (art. 20).')
    _bullet(doc, 'Opposizione al trattamento (art. 21).')
    _gap(doc)

    _heading(doc, 'Conservazione dei dati', 2, color=DARK)
    _body(doc,
        "I dati sono conservati solo per il tempo strettamente necessario alle finalità per cui "
        "sono stati raccolti, nel rispetto dell'art. 5.1.e GDPR (principio di limitazione della "
        "conservazione):")
    tbl_g = doc.add_table(rows=1, cols=3)
    tbl_g.style = 'Table Grid'
    for i, h in enumerate(['Categoria dato', 'Periodo di conservazione', 'Motivazione']):
        r = tbl_g.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True; _fnt(r, FONT, Pt(10))
    for cat, period, reason in [
        ('Prenotazioni pasto e storico ordini',
         '5 anni',
         'Obblighi contabili e fiscali (art. 2220 C.C.)'),
        ('Movimenti wallet (accrediti/addebiti)',
         '5 anni',
         'Documentazione contabile obbligatoria'),
        ('Log di accesso e autenticazione',
         '12 mesi',
         'Sicurezza informatica e audit trail'),
        ('Account utente attivo',
         'Fino a richiesta di cancellazione o cessazione convenzione',
         'Necessità operativa del servizio'),
        ('Account utente cancellato',
         'Anonimizzazione immediata',
         'Diritto all\'oblio — record storici de-identificati'),
        ('Telegram Chat ID',
         'Fino a revoca del consenso o cancellazione account',
         'Consenso revocabile in qualsiasi momento dal profilo utente'),
    ]:
        row = tbl_g.add_row().cells
        for i, (val, bold) in enumerate([(cat, True), (period, False), (reason, False)]):
            p2 = row[i].paragraphs[0]
            _set_spacing(p2, sa=Pt(2))
            r = p2.add_run(val)
            r.bold = bold; _fnt(r, FONT, Pt(10))
    _gap(doc)

    _heading(doc, 'Misure di sicurezza tecniche e organizzative', 2, color=DARK)
    _bullet(doc, 'Tutte le comunicazioni su HTTPS (TLS 1.2+) — dati cifrati in transito.', 'Cifratura in transito:')
    _bullet(doc, 'Password cifrate con bcrypt (fattore di costo adattivo) — mai in chiaro.', 'Password:')
    _bullet(doc, 'Ogni azienda opera in uno spazio logicamente separato (tenant_id su ogni record).', 'Isolamento multi-tenant:')
    _bullet(doc, 'Accesso alle funzioni amministrative limitato tramite ruoli RBAC con permessi granulari.', 'Controllo accessi:')
    _bullet(doc, 'MFA opzionale tramite TOTP (es. Google Authenticator) disponibile per tutti gli utenti.', 'Autenticazione a due fattori:')
    _bullet(doc, 'Backup giornalieri automatici sull\'infrastruttura di hosting.', 'Backup:')
    _gap(doc)

    _heading(doc, 'Titolare e responsabile del trattamento', 2, color=DARK)
    _body(doc,
        "Il titolare del trattamento è il gestore del bar/caffetteria che ha adottato il sistema "
        "QuickLunch. L'azienda convenzionata agisce in qualità di responsabile esterno del "
        "trattamento per i dati dei propri dipendenti, ai sensi dell'art. 28 GDPR. "
        "È raccomandata la stipula di un apposito DPA (Data Processing Agreement) tra le parti.")
    _gap(doc)

    # ── Colophon ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _it_m2 = ['gennaio','febbraio','marzo','aprile','maggio','giugno',
              'luglio','agosto','settembre','ottobre','novembre','dicembre']
    _it_d2 = f'{_dt.now().day} {_it_m2[_dt.now().month-1]} {_dt.now().year}'
    r = cp.add_run(f'QuickLunch Bar Self-Service  —  Documento generato il {_it_d2}')
    _fnt(r, FONT, Pt(9)); r.italic = True; r.font.color.rgb = _rgb(GRAY)
    _set_spacing(cp)

    buf2 = BytesIO()
    doc.save(buf2)
    buf2.seek(0)
    fname = f'QuickLunch_Pasti_Aziendali_Abstract_{_dt.now().strftime("%Y%m")}.docx'
    return send_file(buf2, as_attachment=True, download_name=fname,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@bp.route('/convenzioni/presentazione-pdf')
@require_permission('manage_products')
def convenzioni_presentazione_pdf():
    import os
    from io import BytesIO
    from pathlib import Path
    from flask import send_file
    from fpdf import FPDF

    _IT_MONTHS = ['gennaio','febbraio','marzo','aprile','maggio','giugno',
                  'luglio','agosto','settembre','ottobre','novembre','dicembre']

    def _it_mese_anno(dt):
        return f'{_IT_MONTHS[dt.month - 1]} {dt.year}'

    BRAND  = (233, 69,  96)
    DARK   = (26,  26,  46)
    DGRAY  = (80,  80,  95)
    LGRAY  = (200, 200, 210)
    VLIGHT = (248, 248, 252)
    WHITE  = (255, 255, 255)
    FONT   = 'PTSansNarrow'

    class Brochure(FPDF):
        def header(self):
            if self.page_no() == 1:
                return
            self.set_fill_color(*BRAND)
            self.rect(0, 0, 210, 11, 'F')
            self.set_font(FONT, 'B', 9)
            self.set_text_color(*WHITE)
            self.set_xy(12, 1.5)
            self.cell(130, 8, 'QuickLunch  -  Modulo Pasti Aziendali', ln=0)
            self.set_x(-30)
            self.cell(18, 8, 'quicklunch.app', align='R')
            self.set_text_color(*DARK)
            self.ln(14)

        def footer(self):
            if self.page_no() == 1:
                return
            self.set_y(-13)
            self.set_font(FONT, '', 8.5)
            self.set_text_color(*LGRAY)
            self.set_draw_color(*LGRAY)
            self.line(12, self.get_y(), 198, self.get_y())
            self.ln(1)
            self.cell(0, 6,
                f'QuickLunch Bar Self-Service  -  Documento riservato  -  Pag. {self.page_no()-1}',
                align='C')

        def section_title(self, text):
            self.set_font(FONT, 'B', 14)
            self.set_text_color(*BRAND)
            self.cell(0, 9, text, ln=True)
            self.set_draw_color(*BRAND)
            self.set_line_width(0.5)
            self.line(self.get_x(), self.get_y(), self.get_x() + 186, self.get_y())
            self.ln(4)
            self.set_text_color(*DARK)

        def sub_title(self, text):
            self.set_font(FONT, 'B', 12)
            self.set_text_color(*DARK)
            self.cell(0, 7, text, ln=True)
            self.set_text_color(*DGRAY)

        def body(self, text, indent=0):
            self.set_font(FONT, '', 11)
            self.set_text_color(*DGRAY)
            self.set_x(12 + indent)
            self.multi_cell(186 - indent, 6, text)
            self.ln(1)
            self.set_x(12)

        def bullet(self, text, color=BRAND):
            self.set_font(FONT, 'B', 12)
            self.set_text_color(*color)
            self.set_x(14)
            self.cell(6, 6, '-', ln=0)
            self.set_font(FONT, '', 11)
            self.set_text_color(*DGRAY)
            self.multi_cell(176, 6, text)
            self.set_x(12)

        def step_box(self, num, title, desc):
            x0 = self.get_x()
            y0 = self.get_y()
            self.set_fill_color(*BRAND)
            self.ellipse(x0 + 12, y0, 10, 10, 'F')
            self.set_font(FONT, 'B', 10)
            self.set_text_color(*WHITE)
            self.set_xy(x0 + 12, y0 + 0.5)
            self.cell(10, 9, str(num), align='C')
            self.set_font(FONT, 'B', 11)
            self.set_text_color(*DARK)
            self.set_xy(x0 + 26, y0 + 0.5)
            self.cell(160, 5.5, title)
            self.set_font(FONT, '', 11)
            self.set_text_color(*DGRAY)
            self.set_xy(x0 + 26, y0 + 7)
            self.multi_cell(160, 5.5, desc)
            self.ln(3)
            self.set_x(12)

        def kpi_box(self, label, value, x, y, w=55, h=28):
            self.set_fill_color(*VLIGHT)
            self.set_draw_color(*LGRAY)
            self.set_line_width(0.3)
            self.rect(x, y, w, h, 'FD')
            self.set_font(FONT, 'B', 20)
            self.set_text_color(*BRAND)
            self.set_xy(x, y + 4)
            self.cell(w, 10, value, align='C')
            self.set_font(FONT, '', 10)
            self.set_text_color(*DGRAY)
            self.set_xy(x, y + 16)
            self.cell(w, 6, label, align='C')

    _font_dir = Path(os.path.abspath(__file__)).parent.parent / 'static' / 'fonts'
    pdf = Brochure(orientation='P', unit='mm', format='A4')
    pdf.set_margins(12, 14, 12)
    pdf.set_auto_page_break(True, margin=18)
    pdf.add_font(FONT, '',  str(_font_dir / 'PTSansNarrow-Regular.ttf'))
    pdf.add_font(FONT, 'B', str(_font_dir / 'PTSansNarrow-Bold.ttf'))

    # Copertina
    pdf.add_page()
    pdf.set_fill_color(*DARK)
    pdf.rect(0, 0, 210, 297, 'F')
    pdf.set_fill_color(*BRAND)
    pdf.rect(0, 0, 8, 297, 'F')
    pdf.rect(20, 110, 170, 1.5, 'F')

    pdf.set_font(FONT, 'B', 12)
    pdf.set_text_color(180, 180, 195)
    pdf.set_xy(20, 70)
    pdf.cell(0, 8, 'QuickLunch  -  Bar Self-Service')

    pdf.set_font(FONT, 'B', 40)
    pdf.set_text_color(*WHITE)
    pdf.set_xy(20, 82)
    pdf.cell(0, 18, 'Pasti Aziendali')

    pdf.set_font(FONT, '', 17)
    pdf.set_text_color(*BRAND)
    pdf.set_xy(20, 102)
    pdf.cell(0, 8, 'Il servizio mensa digitale per la tua azienda')

    pdf.set_font(FONT, '', 12)
    pdf.set_text_color(170, 170, 190)
    pdf.set_xy(20, 125)
    pdf.multi_cell(160, 7,
        'Una soluzione semplice e moderna per gestire le prenotazioni pasto '
        'dei tuoi dipendenti: niente code, niente liste cartacee, tutto '
        'tracciato in tempo reale.')

    pdf.kpi_box('Prenotazioni online', '100%',   20, 165)
    pdf.kpi_box('Accesso da mobile',   'App Web',  80, 165)
    pdf.kpi_box('Report giornaliero',  'DOCX',    140, 165)

    pdf.set_font(FONT, '', 10)
    pdf.set_text_color(100, 100, 120)
    pdf.set_xy(20, 275)
    pdf.cell(0, 6, 'Documento di presentazione  -  ' + _it_mese_anno(_dt.now()))

    # Il servizio
    pdf.add_page()
    pdf.section_title('Il servizio')
    pdf.body(
        "QuickLunch Pasti Aziendali e' il modulo dedicato alle convenzioni mensa. "
        'Consente ai dipendenti di prenotare il proprio pasto giornaliero direttamente '
        'dal telefono o dal computer, senza bisogno di installare nulla. '
        'Il gestore del bar pubblica ogni mattina il menu e gestisce le presenze in tempo reale.')
    pdf.ln(3)

    pdf.section_title("Perche' sceglierlo")
    pdf.bullet('Eliminazione delle liste cartacee e delle code al banco.')
    pdf.bullet("Tracciabilita' completa: chi ha prenotato, chi ha consumato, chi ha annullato.")
    pdf.bullet('Report giornaliero scaricabile in Word per la riconciliazione contabile.')
    pdf.bullet('Notifiche automatiche Telegram ai dipendenti (conferma, reminder orario).')
    pdf.bullet('Accesso da qualsiasi dispositivo senza app da installare.')
    pdf.bullet('Configurazione personalizzata per ogni azienda: prezzi, coperti, portate.')
    pdf.ln(5)

    pdf.section_title("A chi e' rivolto")
    pdf.body(
        "Il servizio e' pensato per aziende, enti pubblici, studi professionali e "
        'qualsiasi organizzazione che voglia offrire ai propri dipendenti un servizio '
        'mensa o pasto convenzionato in modo semplice, digitale e verificabile.')
    pdf.ln(3)

    pdf.set_fill_color(*VLIGHT)
    pdf.set_draw_color(*LGRAY)
    pdf.set_line_width(0.3)
    pdf.rect(12, pdf.get_y(), 186, 32, 'FD')
    pdf.set_font(FONT, 'B', 12)
    pdf.set_text_color(*BRAND)
    pdf.set_xy(18, pdf.get_y() + 5)
    pdf.cell(0, 6, '"Meno burocrazia, piu\' tempo per il pranzo."')
    pdf.set_font(FONT, '', 11)
    pdf.set_text_color(*DGRAY)
    pdf.set_xy(18, pdf.get_y() + 8)
    pdf.multi_cell(174, 6,
        'Con QuickLunch il dipendente prenota in 10 secondi, il gestore '
        "ha il conteggio esatto in tempo reale e l'azienda riceve il report pronto alla firma.")

    # Come funziona
    pdf.add_page()
    pdf.section_title('Come funziona: il dipendente')
    pdf.body('Tutto quello che il dipendente deve fare ogni giorno:')
    pdf.ln(2)
    pdf.step_box(1, "Accede all'applicazione",
        'Apre il browser sul telefono o PC e accede con email/password oppure con il proprio account Google.')
    pdf.step_box(2, 'Consulta il menu del giorno',
        "Vede il menu pubblicato dal gestore: nome del pasto, portate (primo, secondo, contorno, bevanda, caffe').")
    pdf.step_box(3, 'Prenota in un clic',
        "Seleziona l'opzione desiderata, sceglie l'orario di ritiro e conferma. Tempo medio: 15 secondi.")
    pdf.step_box(4, 'Riceve la conferma',
        'La prenotazione viene registrata immediatamente. Con Telegram configurato, '
        'riceve anche un messaggio diretto di conferma.')
    pdf.step_box(5, 'Annulla se necessario',
        "Puo' annullare la propria prenotazione in autonomia fino al momento del ritiro.")
    pdf.ln(4)

    pdf.section_title('Come funziona: il gestore del bar')
    pdf.ln(1)
    pdf.step_box(1, 'Pubblica il menu',
        'Ogni mattina inserisce il menu del giorno: portate, prezzo e numero massimo di coperti. '
        "Puo' usare template salvati per i menu ricorrenti.")
    pdf.step_box(2, 'Monitora le prenotazioni',
        'Il Report Giornaliero mostra in tempo reale la lista dei prenotati con i relativi dettagli.')
    pdf.step_box(3, 'Segna i pasti consumati',
        'Durante il servizio spunta ogni prenotazione come Consumato. '
        'Il sistema aggiorna immediatamente i contatori.')
    pdf.step_box(4, 'Scarica il report',
        'A fine servizio scarica il report del giorno in formato Word (.docx), '
        "pronto per essere inviato all'azienda.")

    # Funzionalita
    pdf.add_page()
    pdf.section_title("Funzionalita' principali")
    pdf.ln(1)

    features = [
        ('Prenotazione online',
         'I dipendenti prenotano da browser senza installare niente. '
         'Funziona su smartphone, tablet e PC.'),
        ('Menu giornaliero configurabile',
         'Il gestore pubblica ogni giorno il pasto con portate, prezzo e posti disponibili. '
         "Supporto per piu' opzioni menu nella stessa giornata."),
        ('Template menu riutilizzabili',
         'Salva i menu ricorrenti come template e riutilizzali con un clic.'),
        ('Report presenze in tempo reale',
         'Pagina web con la lista aggiornata di chi ha prenotato, suddivisa per azienda.'),
        ('Esportazione DOCX',
         'Report giornaliero scaricabile in Word, pronto per la riconciliazione contabile.'),
        ('Storico mensile',
         'Archivio completo mese per mese: prenotati, consumati e annullati per ogni giornata.'),
        ('Notifiche Telegram',
         "Reminder automatico al dipendente N minuti prima dell'orario di ritiro."),
        ('Accesso Google',
         'Login con account Google tramite OAuth 2.0. Nessuna nuova password da ricordare.'),
        ('Autenticazione a due fattori',
         'Supporto MFA opzionale con Google Authenticator per maggiore sicurezza.'),
        ('Multi-azienda',
         "Gestione di piu' convenzioni in parallelo, ciascuna con proprie configurazioni e report."),
    ]

    col_w  = 87
    col2_x = 111
    _saved_y = pdf.get_y()
    for i, (title, desc) in enumerate(features):
        x0 = 12 if i % 2 == 0 else col2_x
        if i % 2 == 0:
            _saved_y = pdf.get_y()
        card_h = 26
        pdf.set_fill_color(*VLIGHT)
        pdf.set_draw_color(*LGRAY)
        pdf.set_line_width(0.3)
        pdf.rect(x0, _saved_y, col_w, card_h, 'FD')
        pdf.set_fill_color(*BRAND)
        pdf.rect(x0, _saved_y, 2.5, card_h, 'F')
        pdf.set_font(FONT, 'B', 11)
        pdf.set_text_color(*DARK)
        pdf.set_xy(x0 + 5, _saved_y + 3)
        pdf.cell(col_w - 7, 5.5, title)
        pdf.set_font(FONT, '', 10)
        pdf.set_text_color(*DGRAY)
        pdf.set_xy(x0 + 5, _saved_y + 10)
        pdf.multi_cell(col_w - 7, 4.5, desc)
        if i % 2 != 0:
            pdf.set_xy(12, _saved_y + card_h + 3)

    if len(features) % 2 != 0:
        pdf.set_xy(12, _saved_y + card_h + 3)

    # Attivazione
    pdf.add_page()
    pdf.section_title('Avviare la convenzione')
    pdf.body(
        "L'attivazione richiede pochi minuti. Il gestore del bar crea la scheda "
        'della tua azienda nel sistema, inserisce il prezzo del pasto e il numero '
        'massimo di coperti giornalieri. I dipendenti si registrano autonomamente '
        'tramite link di invito o QR code dedicato.')
    pdf.ln(3)

    pdf.sub_title('Cosa serve per iniziare')
    pdf.bullet('Ragione sociale e email del referente aziendale.')
    pdf.bullet('Numero massimo di coperti giornalieri desiderati.')
    pdf.bullet('Prezzo concordato del pasto (modificabile in qualsiasi momento).')
    pdf.bullet('Lista dei dipendenti da abilitare (oppure link/QR di auto-registrazione).')
    pdf.ln(5)

    pdf.section_title('Sicurezza e privacy')
    pdf.body(
        'Tutti i dati sono gestiti su infrastruttura sicura. '
        'Ogni dipendente vede esclusivamente le informazioni relative alla propria azienda. '
        "L'accesso avviene sempre su connessione cifrata HTTPS. "
        'Supporto opzionale alla autenticazione a due fattori (MFA).')
    pdf.ln(5)

    bx_y = pdf.get_y()
    pdf.set_fill_color(*DARK)
    pdf.set_draw_color(*DARK)
    pdf.rect(12, bx_y, 186, 52, 'F')
    pdf.set_font(FONT, 'B', 14)
    pdf.set_text_color(*WHITE)
    pdf.set_xy(20, bx_y + 8)
    pdf.cell(0, 7, 'Contatti e informazioni')
    pdf.set_draw_color(*BRAND)
    pdf.set_line_width(1)
    pdf.line(20, bx_y + 17, 190, bx_y + 17)
    pdf.set_font(FONT, '', 12)
    pdf.set_text_color(200, 200, 215)
    pdf.set_xy(20, bx_y + 21)
    pdf.multi_cell(170, 6.5,
        "Per attivare la tua convenzione o richiedere una dimostrazione del sistema, "
        "contatta direttamente il gestore del bar. Il servizio e' attivo e operativo: "
        "la tua azienda puo' essere online in pochi minuti.")
    pdf.set_font(FONT, 'B', 11)
    pdf.set_text_color(*BRAND)
    pdf.set_xy(20, bx_y + 41)
    pdf.cell(0, 6, 'QuickLunch  -  Bar Self-Service')

    buf = BytesIO(bytes(pdf.output()))
    buf.seek(0)
    fname = f'QuickLunch_PastiAziendali_{_dt.now().strftime("%Y%m")}.pdf'
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype='application/pdf')


@bp.route('/slots/<int:sid>/durata', methods=['POST'])
@require_permission('manage_slots')
@tables_required
def slot_durata(sid):
    slot = TimeSlot.query.get_or_404(sid)
    slot.seat_duration_minutes = int(request.form.get('seat_duration_minutes', 0) or 0)
    db.session.commit()
    flash(f'Durata slot {slot.time_str} aggiornata a {slot.seat_duration_minutes} min.', 'success')
    return redirect(url_for('admin.tavoli', tab='slot'))


# â”€â”€ Check-in tavolo â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

@bp.route('/prenotazioni/<int:rid>/checkin', methods=['POST'])
@require_permission('manage_reservations_admin')
@tables_required
def reservation_checkin(rid):
    from datetime import datetime as _dt
    res = TableReservation.query.get_or_404(rid)
    res.checkin_at       = _dt.utcnow()
    res.table_alert_sent = False
    db.session.commit()
    flash(f'Check-in tavolo {res.table.number} — {res.user.username}.', 'success')
    send_telegram(
        f'🟢 <b>Tavolo {res.table.number} OCCUPATO</b>\n'
        f'👤 {res.user.full_name}\n'
        f'🕐 Dalle <b>{res.session_start}</b> — '
        f'{res.reservation_date.strftime("%d/%m/%Y")}'
        + (f'\n👥 {res.party_size} persone' if getattr(res, 'party_size', None) else '')
    )
    return redirect(url_for('admin.tavoli'))


@bp.route('/tavoli/ping-alerts')
@staff_required
@tables_required
def tavoli_ping_alerts():
    import json
    from app.notifications import send_telegram
    WARN_MINUTES = 10
    alerts = []
    active = TableReservation.query\
        .filter(TableReservation.checkin_at.isnot(None))\
        .filter(TableReservation.status == 'confirmed')\
        .filter(TableReservation.table_alert_sent == False)\
        .all()
    for res in active:
        mins_left = res.minutes_remaining
        if mins_left is None:
            continue
        if mins_left <= WARN_MINUTES:
            msg = (f'â° Tavolo <b>{res.table.number}</b> â€” <b>{res.user.username}</b>\n'
                   f'Tempo rimanente: <b>{int(mins_left)} min</b>. Prego liberare il tavolo.')
            send_telegram(msg)
            res.table_alert_sent = True
            alerts.append({'table': res.table.number, 'user': res.user.username,
                           'mins_left': round(mins_left, 1)})
    if alerts:
        db.session.commit()
    return json.dumps({'alerts': alerts}), 200, {'Content-Type': 'application/json'}


# ── Strumenti sui dati: carico mensile, reset totale, backup ──────────────────

@bp.route('/dati/carico', methods=['POST'])
@_superadmin_required
def dati_carico_genera():
    """Genera dati di prova su un mese intero."""
    from app.data_tools import genera_carico, leggi_intervalli, DatiInsufficienti

    raw = (request.form.get('mese') or '').strip()      # atteso 'YYYY-MM'
    try:
        anno, mese = int(raw[:4]), int(raw[5:7])
        if not (1 <= mese <= 12) or not (2000 <= anno <= 2100):
            raise ValueError
    except (ValueError, IndexError):
        flash('Mese non valido. Formato atteso: anno-mese.', 'danger')
        return redirect(url_for('admin.settings'))

    tid = _active_tenant_id()
    solo_lav = 'solo_lavorativi' in request.form

    try:
        carico = genera_carico(anno, mese, tid, utente_id=current_user.id,
                               intervalli=leggi_intervalli(),
                               solo_lavorativi=solo_lav)
    except DatiInsufficienti as exc:
        flash(f'Impossibile generare: {exc}.', 'danger')
        return redirect(url_for('admin.settings'))
    except Exception as exc:
        db.session.rollback()
        flash(f'Generazione interrotta: {exc}', 'danger')
        return redirect(url_for('admin.settings'))

    flash(f'Carico di {carico.etichetta} creato: {carico.n_pasti} pasti, '
          f'{carico.n_snack} panini, {carico.n_caffe} caffe, '
          f'{carico.n_builder} builder su {carico.giorni} giorni '
          f'({numero_italiano(carico.incasso)} € di incassi).', 'success')
    return redirect(url_for('admin.settings'))


@bp.route('/dati/carico/<int:cid>/elimina', methods=['POST'])
@_superadmin_required
def dati_carico_elimina(cid):
    """Annulla un carico eliminando ogni riga che aveva creato."""
    from app.data_tools import elimina_carico

    carico = db.get_or_404(CaricoMensile, cid)
    etichetta = carico.etichetta
    try:
        n = elimina_carico(carico)
    except Exception as exc:
        db.session.rollback()
        flash(f'Eliminazione interrotta: {exc}', 'danger')
        return redirect(url_for('admin.settings'))
    flash(f'Carico di {etichetta} eliminato: {n} righe rimosse.', 'info')
    return redirect(url_for('admin.settings'))


@bp.route('/dati/reset', methods=['POST'])
@_superadmin_required
def dati_reset_totale():
    """Svuota tutte le tabelle e ricrea i dati di base."""
    from app.data_tools import reset_totale

    if (request.form.get('conferma') or '').strip().upper() != 'AZZERA':
        flash('Reset annullato: per procedere scrivi AZZERA nel campo di '
              'conferma.', 'warning')
        return redirect(url_for('admin.settings'))
    try:
        n = reset_totale()
    except Exception as exc:
        db.session.rollback()
        flash(f'Reset interrotto: {exc}', 'danger')
        return redirect(url_for('admin.settings'))
    flash(f'Database azzerato: {n} righe eliminate e dati di base ricreati. '
          f'Le credenziali del super admin sono tornate a quelle predefinite.',
          'success')
    return redirect(url_for('auth.login'))


@bp.route('/dati/backup')
@_superadmin_required
def dati_backup():
    """Scarica una copia integrale del database in formato JSON."""
    import json as _json
    from flask import Response
    from app.data_tools import esporta_backup

    dati = esporta_backup()
    nome = f'quicklunch-backup-{_dt.now().strftime("%Y%m%d-%H%M")}.json'
    corpo = _json.dumps(dati, ensure_ascii=False, indent=1)
    righe = sum(len(v) for v in dati['tabelle'].values())
    current_app.logger.info('Backup scaricato: %d righe, %d tabelle',
                            righe, len(dati['tabelle']))
    return Response(
        corpo, mimetype='application/json',
        headers={'Content-Disposition': f'attachment; filename="{nome}"'})


@bp.route('/dati/restore', methods=['POST'])
@_superadmin_required
def dati_restore():
    """Sostituisce il contenuto del database con quello di un file di backup."""
    import json as _json
    from app.data_tools import importa_backup

    if (request.form.get('conferma') or '').strip().upper() != 'RIPRISTINA':
        flash('Ripristino annullato: per procedere scrivi RIPRISTINA nel campo '
              'di conferma.', 'warning')
        return redirect(url_for('admin.settings'))

    f = request.files.get('backup')
    if not f or not f.filename:
        flash('Nessun file selezionato.', 'danger')
        return redirect(url_for('admin.settings'))
    try:
        dati = _json.loads(f.read().decode('utf-8'))
    except (ValueError, UnicodeDecodeError) as exc:
        flash(f'File non leggibile: {exc}', 'danger')
        return redirect(url_for('admin.settings'))

    try:
        n, note = importa_backup(dati)
    except ValueError as exc:
        flash(str(exc), 'danger')
        return redirect(url_for('admin.settings'))
    except Exception as exc:
        db.session.rollback()
        flash(f'Ripristino interrotto: {exc}', 'danger')
        return redirect(url_for('admin.settings'))

    msg = f'Ripristino completato: {n} righe caricate.'
    if note:
        msg += ' ' + ' '.join(note) + '.'
    flash(msg, 'success')
    return redirect(url_for('auth.login'))


@bp.route('/convenzioni/<int:cid>/report-mensile-pdf')
@require_permission('manage_products')
def convenzione_report_mensile_pdf(cid):
    """Riepilogo mensile dei pasti di una convenzione, pronto da allegare alla fattura."""
    import os
    from calendar import monthrange
    from io import BytesIO
    from pathlib import Path
    from flask import send_file
    from fpdf import FPDF

    BRAND  = (233,  69,  96)
    DARK   = ( 26,  26,  46)
    DGRAY  = ( 80,  80,  95)
    LGRAY  = (200, 200, 210)
    VLIGHT = (248, 248, 252)
    WHITE  = (255, 255, 255)
    GREEN  = ( 39, 174,  96)
    FONT   = 'PTSansNarrow'

    _IT_MONTHS = ['gennaio', 'febbraio', 'marzo', 'aprile', 'maggio', 'giugno',
                  'luglio', 'agosto', 'settembre', 'ottobre', 'novembre',
                  'dicembre']
    _IT_DAYS = ['Lun', 'Mar', 'Mer', 'Gio', 'Ven', 'Sab', 'Dom']

    corp = CorporateAccount.query.get_or_404(cid)

    # Mese richiesto: 'YYYY-MM', con il mese corrente come ripiego
    raw = (request.args.get('m') or '').strip()
    try:
        anno, mese = int(raw[:4]), int(raw[5:7])
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, IndexError):
        oggi = date.today()
        anno, mese = oggi.year, oggi.month

    primo = date(anno, mese, 1)
    ultimo = date(anno, mese, monthrange(anno, mese)[1])

    meals = (DailyFixedMeal.query
             .filter(DailyFixedMeal.corporate_id == corp.id,
                     DailyFixedMeal.meal_date >= primo,
                     DailyFixedMeal.meal_date <= ultimo)
             .order_by(DailyFixedMeal.meal_date).all())

    # Aggregazione per dipendente e per giorno
    per_persona = {}
    per_giorno = {}
    n_totale = 0
    importo_totale = 0.0
    for meal in meals:
        for b in meal.bookings:
            if b.status == 'cancelled' or not b.user:
                continue
            qty = b.quantity or 1
            prezzo = float(meal.price or 0.0)
            # Cognome prima del nome, come nel report a schermo: e' un
            # allegato alla fattura, va ordinato per cognome.
            cog = (b.user.last_name or '').strip()
            nom = (b.user.first_name or '').strip()
            nome = (f'{cog} {nom}'.strip() or b.user.username or '')[:38]
            riga = per_persona.setdefault(
                nome, {'qty': 0, 'importo': 0.0, 'giorni': set(),
                       'ordine': (cog.lower(), nom.lower())})
            riga['qty'] += qty
            riga['importo'] += qty * prezzo
            riga['giorni'].add(meal.meal_date)
            g = per_giorno.setdefault(meal.meal_date,
                                      {'qty': 0, 'importo': 0.0, 'persone': set()})
            g['qty'] += qty
            g['importo'] += qty * prezzo
            g['persone'].add(nome)
            n_totale += qty
            importo_totale += qty * prezzo

    persone = sorted(per_persona.items(),
                     key=lambda kv: (kv[1]['ordine'], kv[0]))
    giorni = sorted(per_giorno.items())
    n_giorni = len(giorni)
    media = (n_totale / n_giorni) if n_giorni else 0

    _co_name    = get_setting('company_name')    or 'QuickLunch Bar'
    _co_address = get_setting('company_address') or ''
    _co_city    = get_setting('company_city')    or ''
    _co_vat     = get_setting('company_vat')     or ''
    _co_phone   = get_setting('company_phone')   or ''
    _co_email   = get_setting('company_email')   or ''

    periodo = f'{_IT_MONTHS[mese - 1].capitalize()} {anno}'
    _font_dir = Path(os.path.abspath(__file__)).parent.parent / 'static' / 'fonts'

    class ReportPDF(FPDF):
        def header(self):
            self.set_fill_color(*BRAND)
            self.rect(0, 0, 210, 11, 'F')
            self.set_fill_color(*DARK)
            self.rect(0, 0, 8, 11, 'F')
            self.set_font(FONT, 'B', 9)
            self.set_text_color(*WHITE)
            self.set_xy(12, 1.5)
            self.cell(130, 8,
                      f'RIEPILOGO MENSILE  \xb7  {corp.name.upper()}', ln=0)
            self.set_font(FONT, '', 9)
            self.set_x(-58)
            self.cell(46, 8, periodo.upper(), align='R')
            self.set_text_color(*DARK)
            self.ln(14)

        def footer(self):
            self.set_y(-14)
            self.set_font(FONT, '', 8)
            self.set_text_color(*LGRAY)
            self.set_draw_color(*LGRAY)
            self.set_line_width(0.25)
            self.line(12, self.get_y(), 198, self.get_y())
            self.ln(1.5)
            now_str = _dt.now().strftime('%d/%m/%Y  %H:%M')
            self.set_x(12)
            self.cell(155, 5,
                      f'Generato il {now_str}  \xb7  Documento riservato  \xb7  {_co_name}',
                      ln=0)
            self.set_x(-25)
            self.cell(13, 5, f'Pag. {self.page_no()}', align='R')

    pdf = ReportPDF(orientation='P', unit='mm', format='A4')
    pdf.set_margins(12, 16, 12)
    pdf.set_auto_page_break(True, margin=20)
    pdf.add_font(FONT, '',  str(_font_dir / 'PTSansNarrow-Regular.ttf'))
    pdf.add_font(FONT, 'B', str(_font_dir / 'PTSansNarrow-Bold.ttf'))
    pdf.add_page()

    # ── Intestazione del bar ─────────────────────────────────────────────
    pdf.set_font(FONT, 'B', 13)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 7, _co_name, ln=True)
    co_parts = [x for x in [_co_address, _co_city, _co_phone, _co_email,
                            (f'P.IVA {_co_vat}' if _co_vat else '')] if x]
    if co_parts:
        pdf.set_font(FONT, '', 9)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 5, '  \xb7  '.join(co_parts), ln=True)
    pdf.ln(4)
    pdf.set_draw_color(*BRAND)
    pdf.set_line_width(0.8)
    pdf.line(12, pdf.get_y(), 198, pdf.get_y())
    pdf.ln(6)

    # ── Destinatario e oggetto ───────────────────────────────────────────
    pdf.set_font(FONT, '', 9)
    pdf.set_text_color(*DGRAY)
    pdf.cell(0, 5, 'A:', ln=True)
    pdf.set_font(FONT, 'B', 18)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 9, corp.name, ln=True)
    pdf.set_font(FONT, 'B', 10)
    pdf.set_text_color(*BRAND)
    pdf.cell(0, 6, f'RIEPILOGO PASTI  \xb7  {periodo.upper()}', ln=True)
    pdf.set_font(FONT, '', 9)
    pdf.set_text_color(*DGRAY)
    pdf.cell(0, 5, f'Periodo dal {primo.strftime("%d/%m/%Y")} '
                   f'al {ultimo.strftime("%d/%m/%Y")}', ln=True)
    pdf.ln(5)

    # ── Riquadri di sintesi ──────────────────────────────────────────────
    kpi_y = pdf.get_y()

    def _kpi(x, label, value, color):
        pdf.set_fill_color(*VLIGHT)
        pdf.set_draw_color(*LGRAY)
        pdf.set_line_width(0.3)
        pdf.rect(x, kpi_y, 43.5, 26, 'FD')
        pdf.set_fill_color(*color)
        pdf.rect(x, kpi_y, 43.5, 4, 'F')
        pdf.set_font(FONT, 'B', 20)
        pdf.set_text_color(*color)
        pdf.set_xy(x, kpi_y + 5)
        pdf.cell(43.5, 12, str(value), align='C')
        pdf.set_font(FONT, '', 8.5)
        pdf.set_text_color(*DGRAY)
        pdf.set_xy(x, kpi_y + 17)
        pdf.cell(43.5, 6, label, align='C')

    _kpi(12,    'Pasti erogati',    n_totale,            BRAND)
    _kpi(58.2,  'Giorni di servizio', n_giorni,          DARK)
    _kpi(104.4, 'Media al giorno',  numero_italiano(media, 1),      GREEN)
    _kpi(150.6, 'Dipendenti',       len(persone),        DGRAY)
    pdf.set_y(kpi_y + 32)

    if not persone:
        pdf.set_font(FONT, '', 11)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 8, f'Nessun pasto registrato in {periodo}.', ln=True)
    else:
        # ── Tabella per dipendente ───────────────────────────────────────
        pdf.set_font(FONT, 'B', 11)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 7, 'Dettaglio per dipendente', ln=True)
        pdf.ln(1)

        CW = [86, 28, 30, 42]
        TH = 7
        pdf.set_font(FONT, 'B', 9)
        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        for w, t, al in zip(CW, ['DIPENDENTE', 'GIORNI', 'PASTI', 'IMPORTO'],
                            ['L', 'C', 'C', 'R']):
            pdf.cell(w, TH, t, align=al, fill=True)
        pdf.ln()

        pdf.set_font(FONT, '', 9)
        for i, (nome, r) in enumerate(persone):
            if pdf.get_y() > 250:
                pdf.add_page()
                pdf.set_font(FONT, 'B', 9)
                pdf.set_fill_color(*DARK)
                pdf.set_text_color(*WHITE)
                for w, t, al in zip(CW, ['DIPENDENTE', 'GIORNI', 'PASTI',
                                         'IMPORTO'], ['L', 'C', 'C', 'R']):
                    pdf.cell(w, TH, t, align=al, fill=True)
                pdf.ln()
                pdf.set_font(FONT, '', 9)
            pdf.set_fill_color(*(VLIGHT if i % 2 == 0 else WHITE))
            pdf.set_text_color(*DGRAY)
            pdf.cell(CW[0], TH, nome, fill=True)
            pdf.cell(CW[1], TH, str(len(r['giorni'])), align='C', fill=True)
            pdf.cell(CW[2], TH, str(r['qty']), align='C', fill=True)
            pdf.set_text_color(*DARK)
            pdf.cell(CW[3], TH, f'\u20ac {numero_italiano(r["importo"])}', align='R', fill=True)
            pdf.ln()

        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        pdf.set_font(FONT, 'B', 9)
        pdf.cell(CW[0], TH, f'TOTALE  ({len(persone)} dipendenti)', fill=True)
        pdf.cell(CW[1], TH, str(n_giorni), align='C', fill=True)
        pdf.cell(CW[2], TH, str(n_totale), align='C', fill=True)
        pdf.cell(CW[3], TH, f'\u20ac {numero_italiano(importo_totale)}', align='R', fill=True)
        pdf.ln(11)

        # ── Tabella per giorno ───────────────────────────────────────────
        if pdf.get_y() > 210:
            pdf.add_page()
        pdf.set_font(FONT, 'B', 11)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 7, 'Dettaglio per giorno', ln=True)
        pdf.ln(1)

        DW = [46, 40, 40, 60]
        pdf.set_font(FONT, 'B', 9)
        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        for w, t, al in zip(DW, ['GIORNO', 'DIPENDENTI', 'PASTI', 'IMPORTO'],
                            ['L', 'C', 'C', 'R']):
            pdf.cell(w, TH, t, align=al, fill=True)
        pdf.ln()

        pdf.set_font(FONT, '', 9)
        for i, (g, r) in enumerate(giorni):
            if pdf.get_y() > 250:
                pdf.add_page()
                pdf.set_font(FONT, 'B', 9)
                pdf.set_fill_color(*DARK)
                pdf.set_text_color(*WHITE)
                for w, t, al in zip(DW, ['GIORNO', 'DIPENDENTI', 'PASTI',
                                         'IMPORTO'], ['L', 'C', 'C', 'R']):
                    pdf.cell(w, TH, t, align=al, fill=True)
                pdf.ln()
                pdf.set_font(FONT, '', 9)
            pdf.set_fill_color(*(VLIGHT if i % 2 == 0 else WHITE))
            pdf.set_text_color(*DGRAY)
            pdf.cell(DW[0], TH,
                     f'{_IT_DAYS[g.weekday()]} {g.strftime("%d/%m/%Y")}',
                     fill=True)
            pdf.cell(DW[1], TH, str(len(r['persone'])), align='C', fill=True)
            pdf.cell(DW[2], TH, str(r['qty']), align='C', fill=True)
            pdf.set_text_color(*DARK)
            pdf.cell(DW[3], TH, f'\u20ac {numero_italiano(r["importo"])}', align='R', fill=True)
            pdf.ln()

        # ── Totale fatturabile ───────────────────────────────────────────
        pdf.ln(4)
        pdf.set_font(FONT, '', 10)
        pdf.set_text_color(*DGRAY)
        pdf.set_x(12)
        pdf.cell(50, 6, 'Totale fatturabile:', ln=0)
        pdf.set_font(FONT, 'B', 12)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 6, f'\u20ac {numero_italiano(importo_totale)}', ln=True)
        pdf.set_font(FONT, '', 8.5)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 5, 'Importi al netto di IVA di legge. Sono esclusi i pasti '
                       'annullati.', ln=True)

    buf = BytesIO(bytes(pdf.output()))
    buf.seek(0)
    nome_file = (f'pasti_{corp.name.replace(" ", "_")}_'
                 f'{anno}-{mese:02d}.pdf')
    return send_file(buf, as_attachment=True, download_name=nome_file,
                     mimetype='application/pdf')


def _raccogli_scontrini(year, month, tenant_id=None):
    """Raccoglie i singoli incassi del mese con imponibile e provvigione.

    Unica fonte dei numeri: la usano sia la pagina sia il PDF, così i due non
    possono divergere. I totali seguono la formula della pagina Guadagni —
    imponibile e provvigione sull'importo complessivo, non somma di valori
    arrotondati riga per riga.
    """
    from calendar import monthrange
    from app.notifications import get_numeric_setting

    start_date = date(year, month, 1)
    end_date = date(year, month, monthrange(year, month)[1])

    fee_pct = get_numeric_setting('platform_fee_percentage', 0.0) / 100.0
    monthly_fee = get_numeric_setting('tenant_monthly_fee', 0.0)

    tenants = Tenant.query.order_by(Tenant.name).all()
    selezionati = [x for x in tenants if (tenant_id is None or x.id == tenant_id)]

    def _riga(quando, tipo, riferimento, cliente, lordo, tenant):
        imponibile = round(lordo / 1.10, 2)
        return {
            'quando': quando, 'tipo': tipo, 'riferimento': riferimento,
            'cliente': cliente, 'lordo': round(lordo, 2),
            'imponibile': imponibile,
            'provvigione': round(imponibile * fee_pct, 2),
            'tenant': tenant,
        }

    righe = []
    for tn in selezionati:
        # Ordini completati
        for o in (Order.query
                  .filter(Order.tenant_id == tn.id,
                          Order.status == 'completed',
                          Order.order_date >= start_date,
                          Order.order_date <= end_date)
                  .order_by(Order.order_date, Order.id).all()):
            righe.append(_riga(
                o.created_at or _dt.combine(o.order_date, _dt.min.time()),
                'Ordine', o.order_code or ('#%d' % o.id),
                (o.user.full_name if o.user else '—'), o.total_price or 0.0, tn))

        # Sessioni banco pagate
        for s in (BancoSession.query
                  .filter(BancoSession.tenant_id == tn.id,
                          BancoSession.status == 'paid',
                          db.func.date(BancoSession.created_at) >= start_date,
                          db.func.date(BancoSession.created_at) <= end_date)
                  .order_by(BancoSession.created_at).all()):
            righe.append(_riga(
                s.created_at, 'Banco QR', s.token[:10].upper(),
                (s.customer.full_name if s.customer else '—'),
                s.total or 0.0, tn))

        # Vendite dal cesto: stanno nei movimenti di wallet
        for tx in (Transaction.query
                   .join(User, Transaction.user_id == User.id)
                   .filter(User.tenant_id == tn.id,
                           Transaction.ttype == 'payment',
                           db.or_(Transaction.description.like('Cesto: %'),
                                  Transaction.description.like('Cesto extra: %')),
                           db.func.date(Transaction.created_at) >= start_date,
                           db.func.date(Transaction.created_at) <= end_date)
                   .order_by(Transaction.created_at).all()):
            righe.append(_riga(
                tx.created_at, 'Cesto QR', tx.description[:40],
                (tx.user.full_name if tx.user else '—'),
                abs(tx.amount or 0.0), tn))

        # Pasti aziendali prenotati
        for ca in CorporateAccount.query.filter_by(tenant_id=tn.id).all():
            for meal in ca.daily_meals:
                if not (start_date <= meal.meal_date <= end_date):
                    continue
                for b in meal.bookings:
                    if b.status == 'cancelled':
                        continue
                    qta = b.quantity or 1
                    righe.append(_riga(
                        _dt.combine(meal.meal_date, _dt.min.time()),
                        'Pasto aziendale',
                        '%s (%s)' % (meal.name[:24], ca.name[:18]),
                        (b.user.full_name if b.user else '—'),
                        (meal.price or 0.0) * qta, tn))

    righe.sort(key=lambda r: (r['quando'] or _dt.min, r['tipo']))

    tot_lordo = round(sum(r['lordo'] for r in righe), 2)
    tot_imponibile = round(tot_lordo / 1.10, 2)
    tot_provvigioni = round(tot_imponibile * fee_pct, 2)
    somma_righe = round(sum(r['provvigione'] for r in righe), 2)
    canoni = round(monthly_fee * len(selezionati), 2)

    per_tipo = {}
    for r in righe:
        d = per_tipo.setdefault(r['tipo'], {'n': 0, 'lordo': 0.0, 'prov': 0.0})
        d['n'] += 1
        d['lordo'] = round(d['lordo'] + r['lordo'], 2)
        d['prov'] = round(d['prov'] + r['provvigione'], 2)

    _it_m = ['Gennaio', 'Febbraio', 'Marzo', 'Aprile', 'Maggio', 'Giugno',
             'Luglio', 'Agosto', 'Settembre', 'Ottobre', 'Novembre', 'Dicembre']

    return {
        'righe': righe, 'per_tipo': per_tipo, 'tenants': tenants,
        'selezionati': selezionati, 'tenant_sel': tenant_id,
        'year': year, 'month': month, 'month_label': _it_m[month - 1],
        'start_date': start_date, 'end_date': end_date,
        'fee_pct': round(fee_pct * 100, 2), 'monthly_fee': monthly_fee,
        'tot_lordo': tot_lordo, 'tot_imponibile': tot_imponibile,
        'tot_provvigioni': tot_provvigioni, 'canoni': canoni,
        'somma_righe': somma_righe,
        'scarto': round(somma_righe - tot_provvigioni, 2),
        'tot_dovuto': round(tot_provvigioni + canoni, 2),
        'n_tenant': len(selezionati),
    }


def _mese_richiesto():
    """Anno e mese dai parametri, col mese corrente come ripiego."""
    try:
        year = int(request.args.get('year', _dt.now().year))
        month = int(request.args.get('month', _dt.now().month))
        if not (1 <= month <= 12):
            raise ValueError
    except (ValueError, TypeError):
        year, month = _dt.now().year, _dt.now().month
    return year, month


@bp.route('/superadmin/guadagni/scontrini')
@_superadmin_required
def ds_scontrini():
    """Elenco dei singoli incassi del mese con la provvigione di ciascuno."""
    year, month = _mese_richiesto()
    dati = _raccogli_scontrini(year, month, request.args.get('t', type=int))
    return render_template('admin/superadmin_scontrini.html', **dati)


@bp.route('/superadmin/guadagni/scontrini/pdf')
@_superadmin_required
def ds_scontrini_pdf():
    """Le transazioni del mese in PDF, con la provvigione di ciascuna."""
    import os
    from io import BytesIO
    from pathlib import Path
    from flask import send_file
    from fpdf import FPDF

    BRAND  = (233,  69,  96)
    DARK   = ( 26,  26,  46)
    DGRAY  = ( 80,  80,  95)
    LGRAY  = (200, 200, 210)
    VLIGHT = (248, 248, 252)
    WHITE  = (255, 255, 255)
    GREEN  = ( 39, 174,  96)
    FONT   = 'PTSansNarrow'

    year, month = _mese_richiesto()
    d = _raccogli_scontrini(year, month, request.args.get('t', type=int))

    _co_name = get_setting('company_name') or 'QuickLunch'
    periodo = f"{d['month_label']} {d['year']}"
    ambito = (d['selezionati'][0].name if d['tenant_sel'] and d['selezionati']
              else 'Tutti i tenant')
    _font_dir = Path(os.path.abspath(__file__)).parent.parent / 'static' / 'fonts'

    class ReportPDF(FPDF):
        def header(self):
            self.set_fill_color(*BRAND)
            self.rect(0, 0, 297, 11, 'F')
            self.set_fill_color(*DARK)
            self.rect(0, 0, 8, 11, 'F')
            self.set_font(FONT, 'B', 9)
            self.set_text_color(*WHITE)
            self.set_xy(12, 1.5)
            self.cell(180, 8, f'TRANSAZIONI E PROVVIGIONI  \xb7  {ambito.upper()}',
                      ln=0)
            self.set_font(FONT, '', 9)
            self.set_x(-70)
            self.cell(58, 8, periodo.upper(), align='R')
            self.set_text_color(*DARK)
            self.ln(15)

        def footer(self):
            self.set_y(-13)
            self.set_font(FONT, '', 8)
            self.set_text_color(*LGRAY)
            self.set_draw_color(*LGRAY)
            self.set_line_width(0.25)
            self.line(12, self.get_y(), 285, self.get_y())
            self.ln(1.5)
            self.set_x(12)
            self.cell(240, 5,
                      f'Generato il {_dt.now().strftime("%d/%m/%Y  %H:%M")}'
                      f'  \xb7  Documento riservato  \xb7  {_co_name}', ln=0)
            self.set_x(-28)
            self.cell(16, 5, f'Pag. {self.page_no()}', align='R')

    pdf = ReportPDF(orientation='L', unit='mm', format='A4')
    pdf.set_margins(12, 16, 12)
    pdf.set_auto_page_break(True, margin=18)
    pdf.add_font(FONT, '',  str(_font_dir / 'PTSansNarrow-Regular.ttf'))
    pdf.add_font(FONT, 'B', str(_font_dir / 'PTSansNarrow-Bold.ttf'))
    pdf.add_page()

    # ── Intestazione ─────────────────────────────────────────────────────
    pdf.set_font(FONT, 'B', 14)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 7, _co_name, ln=True)
    pdf.set_font(FONT, '', 9.5)
    pdf.set_text_color(*DGRAY)
    pdf.cell(0, 5, f"Periodo dal {d['start_date'].strftime('%d/%m/%Y')} "
                   f"al {d['end_date'].strftime('%d/%m/%Y')}  \xb7  "
                   f"{ambito}  \xb7  aliquota {numero_italiano(d['fee_pct'], 1)}%",
             ln=True)
    pdf.ln(2)
    pdf.set_draw_color(*BRAND)
    pdf.set_line_width(0.8)
    pdf.line(12, pdf.get_y(), 285, pdf.get_y())
    pdf.ln(5)

    # ── Riquadri di sintesi ──────────────────────────────────────────────
    ky = pdf.get_y()

    def _kpi(x, etichetta, valore, colore):
        pdf.set_fill_color(*VLIGHT)
        pdf.set_draw_color(*LGRAY)
        pdf.set_line_width(0.3)
        pdf.rect(x, ky, 65, 22, 'FD')
        pdf.set_fill_color(*colore)
        pdf.rect(x, ky, 65, 3.5, 'F')
        pdf.set_font(FONT, 'B', 17)
        pdf.set_text_color(*colore)
        pdf.set_xy(x, ky + 4)
        pdf.cell(65, 11, valore, align='C')
        pdf.set_font(FONT, '', 8.5)
        pdf.set_text_color(*DGRAY)
        pdf.set_xy(x, ky + 14)
        pdf.cell(65, 6, etichetta, align='C')

    _kpi(12,  'Transazioni', str(len(d['righe'])), DARK)
    _kpi(81,  'Incassato (IVA incl.)',
         f"{numero_italiano(d['tot_lordo'])} \u20ac", DGRAY)
    _kpi(150, 'Imponibile',
         f"{numero_italiano(d['tot_imponibile'])} \u20ac", GREEN)
    _kpi(219, 'Dovuto a DS Consulting',
         f"{numero_italiano(d['tot_dovuto'])} \u20ac", BRAND)
    pdf.set_y(ky + 27)

    # ── Come si compone il dovuto ────────────────────────────────────────
    pdf.set_font(FONT, '', 9.5)
    pdf.set_text_color(*DGRAY)
    pdf.cell(0, 5,
             f"Provvigioni {numero_italiano(d['tot_provvigioni'])} \u20ac "
             f"({numero_italiano(d['fee_pct'], 1)}% dell'imponibile)  +  "
             f"canoni {numero_italiano(d['canoni'])} \u20ac "
             f"({d['n_tenant']} \u00d7 {numero_italiano(d['monthly_fee'])} \u20ac)"
             f"  =  {numero_italiano(d['tot_dovuto'])} \u20ac", ln=True)
    if d['scarto']:
        pdf.set_font(FONT, '', 8.5)
        pdf.cell(0, 4.5,
                 'Il totale e calcolato sull\'imponibile complessivo: la somma '
                 'delle provvigioni di riga '
                 f"({numero_italiano(d['somma_righe'])} \u20ac) differisce di "
                 f"{numero_italiano(d['scarto'])} \u20ac per arrotondamenti.",
                 ln=True)
    pdf.ln(3)

    if not d['righe']:
        pdf.set_font(FONT, '', 11)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 8, f'Nessuna transazione registrata in {periodo}.', ln=True)
    else:
        # ── Riepilogo per tipo ───────────────────────────────────────────
        pdf.set_font(FONT, 'B', 11)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 6, 'Per tipo di incasso', ln=True)
        TW = [60, 30, 45, 45]
        TH = 6.5
        pdf.set_font(FONT, 'B', 8.5)
        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        for w, t, al in zip(TW, ['TIPO', 'N.', 'INCASSATO', 'PROVVIGIONI'],
                            ['L', 'C', 'R', 'R']):
            pdf.cell(w, TH, t, align=al, fill=True)
        pdf.ln()
        pdf.set_font(FONT, '', 9)
        for i, (tipo, v) in enumerate(sorted(d['per_tipo'].items())):
            pdf.set_fill_color(*(VLIGHT if i % 2 == 0 else WHITE))
            pdf.set_text_color(*DGRAY)
            pdf.cell(TW[0], TH, tipo, fill=True)
            pdf.cell(TW[1], TH, str(v['n']), align='C', fill=True)
            pdf.cell(TW[2], TH, f"{numero_italiano(v['lordo'])} \u20ac",
                     align='R', fill=True)
            pdf.set_text_color(*BRAND)
            pdf.cell(TW[3], TH, f"{numero_italiano(v['prov'])} \u20ac",
                     align='R', fill=True)
            pdf.ln()
        pdf.ln(6)

        # ── Dettaglio delle transazioni ──────────────────────────────────
        pdf.set_font(FONT, 'B', 11)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 6, 'Dettaglio delle transazioni', ln=True)

        mostra_tenant = not d['tenant_sel']
        CW = ([32, 30, 62, 52] + ([34] if mostra_tenant else [])
              + [28, 28, 30])
        intestazioni = (['DATA E ORA', 'TIPO', 'RIFERIMENTO', 'CLIENTE']
                        + (['TENANT'] if mostra_tenant else [])
                        + ['INCASSATO', 'IMPONIBILE', 'PROVVIGIONE'])
        allinea = (['L', 'L', 'L', 'L'] + (['L'] if mostra_tenant else [])
                   + ['R', 'R', 'R'])

        def _testata():
            pdf.set_font(FONT, 'B', 8.5)
            pdf.set_fill_color(*DARK)
            pdf.set_text_color(*WHITE)
            for w, t, al in zip(CW, intestazioni, allinea):
                pdf.cell(w, TH, t, align=al, fill=True)
            pdf.ln()
            pdf.set_font(FONT, '', 8.5)

        _testata()
        for i, r in enumerate(d['righe']):
            if pdf.get_y() > 185:
                pdf.add_page()
                _testata()
            pdf.set_fill_color(*(VLIGHT if i % 2 == 0 else WHITE))
            pdf.set_text_color(*DGRAY)
            quando = r['quando'].strftime('%d/%m/%Y %H:%M') if r['quando'] else ''
            valori = ([quando, r['tipo'], str(r['riferimento'])[:34],
                       str(r['cliente'])[:28]]
                      + ([r['tenant'].name[:18]] if mostra_tenant else []))
            for w, v, al in zip(CW, valori, allinea):
                pdf.cell(w, TH, v, align=al, fill=True)
            pdf.cell(CW[-3], TH, f"{numero_italiano(r['lordo'])} \u20ac",
                     align='R', fill=True)
            pdf.cell(CW[-2], TH, f"{numero_italiano(r['imponibile'])} \u20ac",
                     align='R', fill=True)
            pdf.set_text_color(*BRAND)
            pdf.cell(CW[-1], TH, f"{numero_italiano(r['provvigione'])} \u20ac",
                     align='R', fill=True)
            pdf.ln()

        # ── Riga di totale ───────────────────────────────────────────────
        span = sum(CW[:-3])
        pdf.set_fill_color(*DARK)
        pdf.set_text_color(*WHITE)
        pdf.set_font(FONT, 'B', 9)
        pdf.cell(span, TH, f"TOTALE  ({len(d['righe'])} transazioni)", fill=True)
        pdf.cell(CW[-3], TH, f"{numero_italiano(d['tot_lordo'])} \u20ac",
                 align='R', fill=True)
        pdf.cell(CW[-2], TH, f"{numero_italiano(d['tot_imponibile'])} \u20ac",
                 align='R', fill=True)
        pdf.cell(CW[-1], TH, f"{numero_italiano(d['tot_provvigioni'])} \u20ac",
                 align='R', fill=True)
        pdf.ln(10)

        pdf.set_font(FONT, '', 9)
        pdf.set_text_color(*DGRAY)
        pdf.cell(60, 6, 'Totale dovuto a DS Consulting:', ln=0)
        pdf.set_font(FONT, 'B', 12)
        pdf.set_text_color(*DARK)
        pdf.cell(0, 6, f"{numero_italiano(d['tot_dovuto'])} \u20ac", ln=True)
        pdf.set_font(FONT, '', 8.5)
        pdf.set_text_color(*DGRAY)
        pdf.cell(0, 5, 'Importi al netto di IVA di legge. Sono esclusi gli '
                       'ordini annullati e le prenotazioni disdette.', ln=True)

    buf = BytesIO(bytes(pdf.output()))
    buf.seek(0)
    suffisso = ('_' + ambito.replace(' ', '_') if d['tenant_sel'] else '')
    nome = f'transazioni{suffisso}_{year}-{month:02d}.pdf'
    return send_file(buf, as_attachment=True, download_name=nome,
                     mimetype='application/pdf')
