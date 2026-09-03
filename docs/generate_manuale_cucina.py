#!/usr/bin/env python3
"""Genera docs/manuali/manuale_cucina.docx — manuale operativo del reparto cucina.

Tre procedure: cesto self-service, pasti aziendali convenzionati, ordini dal
builder visuale. Le voci di menu e i percorsi citati corrispondono
all'applicazione: se cambiano le rotte o i permessi, aggiornare qui.

    python docs/generate_manuale_cucina.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Palette (coerente con gli altri documenti in docs/) ──────────────────────
RED    = RGBColor(0xe9, 0x45, 0x60)
NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x16, 0x21, 0x3e)
DGRAY  = RGBColor(0x34, 0x3a, 0x40)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREEN  = RGBColor(0x27, 0xae, 0x60)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
BLUE   = RGBColor(0x34, 0x98, 0xdb)
PURPLE = RGBColor(0x8e, 0x44, 0xad)
SLATE  = RGBColor(0x7f, 0x8c, 0x8d)

HEX_RED    = 'E94560'
HEX_NAVY   = '0F3460'
HEX_GREEN  = '27AE60'
HEX_ORANGE = 'E67E22'
HEX_BLUE   = '3498DB'
HEX_PURPLE = '8E44AD'
HEX_SLATE  = '7F8C8D'
HEX_LIGHT  = 'F8F9FA'
HEX_RULE   = 'E8EBEE'
HEX_WARN   = 'FDF6E7'
HEX_STOP   = 'FDEEF0'

FONT = 'PT Sans Narrow'
BODY_FONT = 'Calibri'

OUT = os.path.join(os.path.dirname(__file__), 'manuali', 'manuale_cucina.docx')


# ── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None, sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        if color:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color.lstrip('#'))
        else:
            el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), val)
    tcPr.append(el)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _p_spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def _run_font(run, size=10.5, bold=False, italic=False, color=None, font=BODY_FONT):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    _p_spacing(p, before=0, after=0)


# ── Blocchi di contenuto ─────────────────────────────────────────────────────

def heading(doc, text, level=1, color=NAVY, before=16):
    sizes = {1: 19, 2: 14.5, 3: 12}
    p = doc.add_paragraph()
    _p_spacing(p, before=before, after=5)
    _run_font(p.add_run(text), size=sizes.get(level, 12), bold=True,
              color=color, font=FONT)
    return p


def body(doc, text, size=10.5, color=DGRAY, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    _run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def rich(doc, parts, size=10.5, after=6, indent=None):
    """Paragrafo con porzioni in grassetto: parts = [(testo, bold), ...]."""
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def bullet(doc, parts, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    _p_spacing(p, before=0, after=3)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def rule(doc, color=HEX_RULE, after=8):
    """Filetto orizzontale come tabella a una cella."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_border(cell, bottom=color, sz='8')
    _cell_margins(cell, top=0, bottom=0, left=0, right=0)
    _p_spacing(cell.paragraphs[0], before=0, after=0)
    _run_font(cell.paragraphs[0].add_run(''), size=1)
    spacer(doc, after)


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(''), size=pts / 2)


def box(doc, label, paragraphs, accent=HEX_RED, fill=HEX_LIGHT, label_color=RED):
    """Riquadro con bordo sinistro colorato."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, fill)
    _cell_border(cell, top=HEX_RULE, bottom=HEX_RULE, right=HEX_RULE,
                 left=accent, sz='18')
    _cell_margins(cell, top=90, bottom=90, left=140, right=110)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=3)
    _run_font(p.add_run(label.upper()), size=8.5, bold=True, color=label_color,
              font=FONT)

    for i, parts in enumerate(paragraphs):
        pp = cell.add_paragraph()
        _p_spacing(pp, before=0, after=0 if i == len(paragraphs) - 1 else 4)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 10)
    return tbl


def step(doc, num, title, paragraphs, accent=HEX_RED):
    """Passo numerato: numero in pastiglia colorata + contenuto."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_width(tbl, 0, 1.0)
    _set_col_width(tbl, 1, 15.5)

    nc = tbl.rows[0].cells[0]
    _cell_shd(nc, accent)
    _cell_margins(nc, top=60, bottom=60, left=30, right=30)
    _cell_valign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run('%02d' % num), size=12, bold=True, color=WHITE, font=FONT)

    cc = tbl.rows[0].cells[1]
    _cell_margins(cc, top=50, bottom=50, left=150, right=60)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=2)
    _run_font(pt.add_run(title), size=11.5, bold=True, color=DARK, font=FONT)

    for parts in paragraphs:
        pp = cc.add_paragraph()
        _p_spacing(pp, before=0, after=3)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 6)
    return tbl


def table_grid(doc, headers, rows, widths, head_fill=HEX_NAVY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, w in enumerate(widths):
        _set_col_width(tbl, i, w)

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _cell_shd(c, head_fill)
        _cell_margins(c, top=60, bottom=60, left=110, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(h.upper()), size=8.5, bold=True, color=WHITE, font=FONT)

    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            c = tbl.rows[r].cells[i]
            _cell_shd(c, HEX_LIGHT if r % 2 else 'FFFFFF')
            _cell_border(c, bottom=HEX_RULE, sz='4')
            _cell_margins(c, top=55, bottom=55, left=110, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            text, color, bold = val if isinstance(val, tuple) else (val, DGRAY, False)
            _run_font(p.add_run(text), size=10, bold=bold, color=color)
    spacer(doc, 10)
    return tbl


# ── Documento ────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)

    # ══ Copertina ═════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, HEX_NAVY)
    _cell_margins(cell, top=300, bottom=300, left=320, right=320)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=4)
    _run_font(p.add_run('QUICKLUNCH  ·  PROCEDURE DI REPARTO'), size=9.5,
              bold=True, color=RGBColor(0xb2, 0xc2, 0xd9), font=FONT)

    p2 = cell.add_paragraph()
    _p_spacing(p2, before=0, after=6)
    _run_font(p2.add_run('Manuale operativo cucina'), size=30, bold=True,
              color=WHITE, font=FONT)

    p3 = cell.add_paragraph()
    _p_spacing(p3, before=0, after=0)
    _run_font(p3.add_run(
        'Tre flussi di lavoro distinti, dalla preparazione alla consegna: '
        'il cesto self-service, i pasti aziendali convenzionati e gli ordini '
        'composti dal builder visuale.'), size=11,
        color=RGBColor(0xd6, 0xdf, 0xea))
    spacer(doc, 16)

    # ══ Regola fissa ══════════════════════════════════════════════════════
    box(doc, 'Regola fissa · vale su ogni ordine', [
        [('Quando un ordine è preparato, ', False),
         ('allega al prodotto lo scontrino di cassa insieme all\'etichetta', True),
         ('. Prodotto senza scontrino non esce dalla cucina.', False)],
        [('Lo scontrino è quello emesso dal ', False),
         ('registratore di cassa', True),
         (', non un foglio stampato da QuickLunch: è il documento fiscale che '
          'accompagna la vendita. Il tagliando che l\'applicazione stampa dal '
          'display cucina è un\'altra cosa — serve a voi per preparare, e non '
          'sostituisce lo scontrino.', False)],
    ], accent=HEX_RED, fill='FDEEF0', label_color=RED)

    # ══ I documenti ═══════════════════════════════════════════════════════
    heading(doc, 'I documenti e da dove escono', 1)
    body(doc, 'Quattro documenti diversi, e solo uno e fiscale. Prima di iniziare '
              'il turno, verifica di sapere quale stampante serve a cosa.',
         color=GRAY)
    spacer(doc, 4)

    table_grid(doc,
               ['Stampato', 'Dove si genera', 'Contenuto'],
               [
                   [('Scontrino di cassa', DARK, True),
                    'Registratore di cassa — fuori da QuickLunch',
                    'Documento fiscale della vendita. È QUESTO che va allegato '
                    'al prodotto insieme all\'etichetta.'],
                   [('Tagliando ordine', DARK, True),
                    'Cucina · icona stampante sulla scheda ordine',
                    'Codice ordine, data, cliente, orario di ritiro, note, '
                    'articoli e — per i prodotti del builder — l\'elenco completo '
                    'degli ingredienti. Documento interno di preparazione.'],
                   [('Etichetta QR', DARK, True),
                    'Cucina · Cesto Cucina · Genera',
                    'Una per pezzo: nome dell\'attività, prodotto, prezzo, codice '
                    'CESTO-______ e QR che porta il cliente alla pagina di acquisto.'],
                   [('Lista pasti aziendali', DARK, True),
                    'Convenzioni · Report del giorno',
                    'Elenco nominativo per azienda e per data, ordinato per '
                    'cognome. È la lista di produzione e la base per etichettare '
                    'le porzioni.'],
               ],
               widths=[3.6, 4.6, 8.3])

    box(doc, 'Nota', [
        [('Per i pasti aziendali l\'applicazione ', False),
         ('non produce un\'etichetta per porzione', True),
         (': nome e codice di ritiro si scrivono a mano copiandoli dalla lista '
          'di produzione.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    _page_break(doc)

    # ══ 01 · CESTO ════════════════════════════════════════════════════════
    heading(doc, '01 — Cesto: panini e tramezzini pronti', 1, color=GREEN, before=0)
    body(doc, 'Prodotti preparati in anticipo e messi nel cesto a libero servizio. '
              'Qui non esiste un ordine: il cliente prende il pezzo, inquadra il QR '
              'dell\'etichetta col telefono e paga dal proprio wallet — oppure, se '
              'il portafoglio prepagato è disattivato, la vendita si registra e il '
              'conto si salda in cassa. L\'etichetta è il documento di vendita.',
         color=GRAY)
    rule(doc)

    step(doc, 1, 'Prepara i pezzi', [
        [('Componi panini e tramezzini in lotti omogenei: un lotto per tipo di '
          'prodotto. Le etichette si generano per prodotto, quindi conviene '
          'chiudere un tipo alla volta.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'Apri la pagina del cesto', [
        [('Menu ', False), ('Cucina › Cesto Cucina', True),
         ('  (/admin/cesto). La pagina mostra il riepilogo di giornata: quanti '
          'pezzi pronti, venduti e scaduti per ogni prodotto.', False)],
    ], accent=HEX_GREEN)

    step(doc, 3, 'Genera le etichette', [
        [('Scegli il prodotto dall\'elenco, indica la quantità preparata e premi ',
          False), ('Genera', True), ('. Massimo 50 etichette per volta.', False)],
        [('Ogni etichetta riceve un codice irripetibile: due pezzi non possono '
          'mai essere venduti con lo stesso codice.', False)],
    ], accent=HEX_GREEN)

    step(doc, 4, 'Stampa e taglia', [
        [('Dopo ', False), ('Genera', True),
         (' si apre da sola la pagina di stampa del lotto. Premi ', False),
         ('Stampa', True), (', poi taglia le etichette lungo i bordi.', False)],
    ], accent=HEX_GREEN)

    step(doc, 5, 'Applica un\'etichetta per pezzo', [
        [('Applica l\'etichetta sull\'incarto, con il QR ben visibile e non '
          'deformato dalla piega: se il telefono non lo legge, il pezzo non è '
          'vendibile.', False)],
        [('Le etichette stampate e non applicate vanno distrutte, non riutilizzate '
          'il giorno dopo.', False)],
    ], accent=HEX_GREEN)

    step(doc, 6, 'Metti nel cesto', [
        [('Dal momento in cui il pezzo è nel cesto la vendita è automatica: il '
          'cliente inquadra, conferma e l\'importo viene scalato dal suo wallet '
          '(con il portafoglio disattivato la vendita viene comunque registrata '
          'e il cliente paga in cassa). L\'etichetta passa da ', False),
         ('pronta', True), (' a ', False),
         ('venduta', True), (' e nessun altro può riacquistarla.', False)],
    ], accent=HEX_GREEN)

    spacer(doc, 6)
    box(doc, 'Perché qui non c\'è scontrino al momento della preparazione', [
        [('La regola fissa parla di ordini preparati. Nel cesto, quando prepari, '
          'la vendita non e ancora avvenuta: nasce quando il cliente inquadra il '
          'QR e paga dal telefono. Per questo sul pezzo va ', False),
         ('solo l\'etichetta', True),
         (', che contiene prodotto e prezzo.', False)],
        [('Lo scontrino di cassa per queste vendite si gestisce come per gli '
          'altri incassi del banco, secondo le vostre procedure fiscali: non '
          'passa dalla cucina.', False)],
    ], accent=HEX_SLATE, fill=HEX_LIGHT, label_color=SLATE)

    heading(doc, 'Gestione della giornata', 3, color=DARK)
    bullet(doc, [
        ('Pezzo ritirato dal cesto', True),
        (' (caduto, danneggiato, invenduto a fine turno): trova la sua riga e '
         'premi ', False), ('Annulla', True),
        ('. L\'etichetta diventa scaduta e il QR non è più acquistabile. '
         'Un\'etichetta già venduta non si può annullare.', False),
    ])
    bullet(doc, [
        ('Scadenza automatica: ', True),
        ('un\'etichetta più vecchia di 24 ore scade da sola alla prima scansione. '
         'Non contare su questo per la rotazione: ritira fisicamente '
         'l\'invenduto.', False),
    ])
    bullet(doc, [
        ('Riepilogo di fine turno: ', True),
        ('la pagina del cesto mostra pronti, venduti e scaduti di oggi. È il dato '
         'da confrontare con i pezzi rimasti nel cesto.', False),
    ])
    spacer(doc, 10)

    box(doc, 'Non usare · Annulla tutto', [
        [('Il pulsante ', False), ('Annulla tutto', True),
         (' cancella dal registro ', False),
         ('tutte le etichette di oggi, comprese quelle già vendute', True),
         ('. Le vendite spariscono dal riepilogo e l\'operazione non è '
          'reversibile.', False)],
        [('Serve solo a ripulire un lotto generato per errore prima che qualsiasi '
          'pezzo sia stato venduto. In ogni altro caso annulla le etichette una '
          'per una.', False)],
    ], accent=HEX_RED, fill=HEX_STOP, label_color=RED)

    _page_break(doc)

    # ══ 02 · PASTI AZIENDALI ══════════════════════════════════════════════
    heading(doc, '02 — Pasti aziendali convenzionati', 1, color=PURPLE, before=0)
    body(doc, 'Pasto fisso del giorno per i dipendenti delle aziende convenzionate. '
              'Prenotano dal telefono e ricevono un codice di ritiro di sei '
              'caratteri: quel codice è l\'unica chiave che sblocca la consegna.',
         color=GRAY)
    rule(doc)

    box(doc, 'Serve un accesso da responsabile', [
        [('Le pagine di questo flusso — ', False), ('Convenzioni', True),
         (' e ', False), ('Ritiro pasti', True),
         (' — ', False),
         ('non sono accessibili con un profilo di sola cucina', True),
         (': rispondono "accesso negato". Occorre un profilo con i permessi di '
          'gestione prodotti (responsabile o amministratore).', False)],
        [('Se questo flusso deve stare in mano al reparto cucina, va aggiunto il '
          'permesso al ruolo: vedi "Chi può fare cosa" a fine manuale.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    step(doc, 1, 'Pubblica il menu del giorno', [
        [('Convenzioni', True), (' › scegli l\'azienda › ', False),
         ('Pasto del giorno', True),
         ('. Compila le portate — primo, secondo, contorno, bevanda, caffè — gli '
          'allergeni, il prezzo e il numero massimo di prenotazioni.', False)],
        [('Se il menu si ripete, richiamalo da ', False), ('Configurazioni', True),
         (': sono modelli salvati che compilano il modulo in un colpo. Puoi '
          'pubblicare più opzioni per la stessa giornata.', False)],
    ], accent=HEX_PURPLE)

    step(doc, 2, 'Lascia chiudere le prenotazioni', [
        [('I dipendenti prenotano dal telefono, anche più porzioni a testa, e '
          'possono modificare o disdire ', False),
         ('fino a 30 minuti prima', True),
         (' dell\'orario di ritiro scelto. Prima di quel margine il numero non è '
          'definitivo.', False)],
    ], accent=HEX_PURPLE)

    step(doc, 3, 'Stampa la lista di produzione', [
        [('Convenzioni › Report', True),
         (', sulla data di oggi (/admin/convenzioni/report). Ottieni l\'elenco '
          'nominativo diviso per azienda e ordinato per cognome, con le porzioni '
          'per persona. Per una singola convenzione puoi scaricare il PDF.', False)],
        [('Questa lista è il tuo conteggio di produzione e, insieme, la base per '
          'etichettare le porzioni.', False)],
    ], accent=HEX_PURPLE)

    step(doc, 4, 'Prepara e componi i vassoi', [
        [('Prepara le porzioni seguendo la lista. Raggruppa per orario di ritiro, '
          'non per azienda: al banco arrivano per fascia oraria.', False)],
    ], accent=HEX_PURPLE)

    step(doc, 5, 'Etichetta ogni porzione', [
        [('Su ogni porzione va ', False), ('cognome e nome', True),
         (' del prenotante e il ', False), ('codice di ritiro', True),
         (' dalla lista. Senza il nome sul contenitore, al banco bisogna aprire i '
          'vassoi per capire di chi è cosa.', False)],
    ], accent=HEX_PURPLE)

    step(doc, 6, 'Consegna verificando il codice', [
        [('Pasti › Ritiro', True),
         ('  (/admin/pasti/ritiro): inserisci il codice che il dipendente mostra '
          'sul telefono. La pagina conferma nome, pasto, azienda, porzioni e '
          'orario.', False)],
        [('Confronta con l\'etichetta, poi premi ', False), ('Consegna', True),
         ('. Il dipendente riceve la conferma su Telegram. La stessa verifica si '
          'può fare dal banco POS.', False)],
    ], accent=HEX_PURPLE)

    spacer(doc, 6)
    box(doc, 'Doppio ritiro: già impedito', [
        [('Dopo ', False), ('Consegna', True),
         (' il codice non viene più trovato dalla ricerca. Se un codice "non '
          'risulta", le possibilità sono due: il pasto è già stato ritirato, '
          'oppure il codice è stato digitato male. Chiedi al dipendente di '
          'rimostrarlo prima di cercare altrove.', False)],
    ], accent=HEX_SLATE, fill=HEX_LIGHT, label_color=SLATE)

    box(doc, 'Manca l\'etichetta stampata', [
        [('Per i pasti aziendali l\'applicazione ', False),
         ('non produce un\'etichetta per porzione', True),
         (': la si scrive a mano copiando nome e codice dalla lista di '
          'produzione. È il punto più lento e più esposto a errori di tutto il '
          'flusso, ed è anche il più facile da automatizzare — una pagina di '
          'stampa etichette come quella del cesto risolverebbe il problema.',
          False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    _page_break(doc)

    # ══ 03 · BUILDER ══════════════════════════════════════════════════════
    heading(doc, '03 — Ordini dal builder visuale', 1, color=BLUE, before=0)
    body(doc, 'Panini, insalate e poke composti dal cliente ingrediente per '
              'ingrediente. Con il portafoglio attivo arrivano in cucina già '
              'pagati (l\'importo è scalato alla conferma) e non c\'è nulla da '
              'incassare; con il portafoglio disattivato il cliente paga alla '
              'cassa al ritiro. In entrambi i casi, qui si produce '
              'nell\'orario giusto.', color=GRAY)
    rule(doc)

    heading(doc, 'Il display cucina', 3, color=DARK)
    rich(doc, [('Menu ', False), ('Cucina', True),
               ('  (/admin/cucina). Tre colonne, che sono i tre stati '
                'dell\'ordine: si avanza sempre da sinistra a destra.', False)])
    spacer(doc, 4)

    table_grid(doc,
               ['Colonna', 'Significato', 'Cosa sa il cliente'],
               [
                   [('Da preparare', ORANGE, True),
                    'Ordini pagati e confermati, non ancora presi in carico.',
                    'Nulla: attende.'],
                   [('In preparazione', BLUE, True),
                    'Ordine preso in carico dalla cucina.',
                    'Avvisato che stai lavorando il suo ordine.'],
                   [('Pronti', GREEN, True),
                    'Prodotto finito, in attesa di ritiro al banco.',
                    'Riceve Telegram e notifica push.'],
                   [('Consegnato', SLATE, True),
                    'Stato finale, impostato alla consegna in mano.',
                    'Esce dal display.'],
               ],
               widths=[3.4, 6.4, 6.7])

    rich(doc, [('La pagina si ', False),
               ('ricarica da sola ogni 20 secondi', True),
               (' e segnala i nuovi ordini con un avviso visivo e un suono: non '
                'serve aggiornarla a mano, ma serve tenerla in primo piano.',
                False)])
    spacer(doc, 8)

    heading(doc, 'Procedura', 3, color=DARK)

    step(doc, 1, 'Leggi la scheda per intero', [
        [('Ogni scheda porta il codice ordine, l\'', False),
         ('orario di ritiro', True),
         (' in alto a destra, il nome del cliente, gli articoli di menu e — per '
          'ogni prodotto del builder — l\'elenco completo degli ingredienti '
          'scelti.', False)],
        [('Lavora in ordine di orario di ritiro, non di arrivo.', False)],
    ], accent=HEX_BLUE)

    step(doc, 2, 'Controlla il contrassegno piastra', [
        [('Il contrassegno rosso ', False), ('PIASTRA', True),
         (' significa panino da scaldare. Questi vanno fatti per ultimi nella '
          'sequenza di ritiro, così arrivano caldi al banco.', False)],
    ], accent=HEX_BLUE)

    step(doc, 3, 'Prendi in carico', [
        [('Premi ', False), ('In preparazione', True),
         ('. Il cliente riceve la notifica che l\'ordine è in lavorazione; se c\'è '
          'la piastra, sul canale del personale compare anche l\'avviso '
          'dedicato.', False)],
        [('Premilo quando inizi davvero: è l\'informazione con cui il cliente '
          'decide quando scendere.', False)],
    ], accent=HEX_BLUE)

    step(doc, 4, 'Prepara seguendo l\'elenco ingredienti', [
        [('Segui la scheda ingrediente per ingrediente. Se qualcosa è esaurito '
          'non sostituirlo di iniziativa: l\'ordine è già pagato e composto dal '
          'cliente. Avvisa il banco.', False)],
    ], accent=HEX_BLUE)

    step(doc, 5, 'Allega scontrino di cassa ed etichetta', [
        [('Sul prodotto vanno lo ', False), ('scontrino di cassa', True),
         (' e l\'etichetta con il codice ordine, leggibile senza aprire '
          'l\'incarto. Lo scontrino arriva dal registratore di cassa: se non e '
          'ancora stato emesso, chiedilo al banco prima di chiudere il '
          'pacchetto.', False)],
        [('Se ti serve il dettaglio degli ingredienti sotto mano, premi l\'',
          False), ('icona stampante', True),
         (' sulla scheda: stampa il tagliando dell\'ordine, che resta un '
          'documento interno.', False)],
    ], accent=HEX_RED)

    box(doc, 'Regola fissa · è questo il passo', [
        [('Scontrino di cassa ed etichetta allegati al prodotto', True),
         (', prima di dichiarare l\'ordine pronto.', False)],
    ], accent=HEX_RED, fill='FDEEF0', label_color=RED)

    step(doc, 6, 'Dichiara pronto', [
        [('Premi ', False), ('Pronto', True),
         ('. Il cliente riceve Telegram e notifica push sul telefono. Da qui '
          'l\'ordine è in attesa di ritiro al banco.', False)],
    ], accent=HEX_BLUE)

    step(doc, 7, 'Chiudi alla consegna', [
        [('Alla consegna in mano premi ', False), ('Consegnato', True),
         (': l\'ordine esce dal display.', False)],
        [('Se il cliente non si presenta, ', False), ('Sollecita ritiro', True),
         (' gli manda un promemoria senza cambiare stato dell\'ordine.', False)],
    ], accent=HEX_BLUE)

    spacer(doc, 6)
    box(doc, 'Ordini "subito al banco"', [
        [('Un ordine con codice che inizia per BANCO- è stato fatto per il '
          'consumo immediato e ', False), ('non ha un orario di ritiro', True),
         (': sullo scontrino quella riga resta vuota. Trattalo come priorità '
          'immediata.', False)],
    ], accent=HEX_SLATE, fill=HEX_LIGHT, label_color=SLATE)

    box(doc, 'Annullare un ordine', [
        [('Portare un ordine ad "annullato" ', False),
         ('rimborsa automaticamente il wallet del cliente', True),
         (' (quando il portafoglio è attivo)', False),
         (' e libera le scorte impegnate. È corretto quando il prodotto non è '
          'realizzabile, ma è un movimento di denaro: concordalo con il banco, '
          'non farlo per ripulire il display.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    _page_break(doc)

    # ══ Accessi ═══════════════════════════════════════════════════════════
    heading(doc, 'Chi può fare cosa', 1, before=0)
    body(doc, 'Accessi effettivi dei profili predefiniti. Da verificare quando si '
              'assegna un turno: due dei tre flussi stanno in mano alla cucina, '
              'il terzo no.', color=GRAY)
    spacer(doc, 4)

    SI = ('sì', GREEN, True)
    NO = ('no', RED, True)
    table_grid(doc,
               ['Pagina', 'Cuoco', 'Responsabile', 'Cassiere'],
               [
                   ['Cucina — display ordini', SI, SI, SI],
                   ['Cesto Cucina — etichette QR', SI, SI, NO],
                   ['Convenzioni — menu del giorno', NO, SI, NO],
                   ['Convenzioni — report di produzione', NO, SI, NO],
                   ['Ritiro pasti aziendali', NO, SI, NO],
               ],
               widths=[7.6, 2.9, 3.4, 2.7])

    box(doc, 'Da decidere', [
        [('Con i profili predefiniti, la procedura ', False),
         ('02 · Pasti aziendali non è eseguibile dalla cucina', True),
         (': richiede un profilo da responsabile. Le strade sono due — lasciare '
          'quel flusso al responsabile di sala, oppure concedere al ruolo cuoco '
          'il permesso di gestione prodotti dalla pagina Ruoli e permessi.',
          False)],
        [('La seconda strada apre alla cucina anche la modifica del catalogo '
          'prodotti: se serve un accesso più circoscritto, la via pulita è un '
          'permesso dedicato al ritiro pasti.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    spacer(doc, 12)
    rule(doc, color=HEX_NAVY)
    body(doc, 'Assistenza:  Daniele Speziale — DS Consulting', size=11,
         color=DARK, bold=True, after=1)
    body(doc, 'dspeziale@gmail.com   ·   +39 352 0150489', size=11, color=RED, bold=True, after=8)
    body(doc, 'QuickLunch · Manuale operativo cucina · Le voci di menu e i percorsi '
              'citati corrispondono all\'applicazione in uso.', size=8.5, color=GRAY,
         after=2)
    body(doc, '© 2024–26 DS Consulting', size=8.5, color=GRAY)

    return doc


def main():
    doc = build()
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] Documento salvato in: {OUT}')


if __name__ == '__main__':
    main()
