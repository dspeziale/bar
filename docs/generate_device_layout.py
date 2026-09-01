#!/usr/bin/env python3
"""Genera docs/device_layout.docx — layout dispositivi QuickLunch."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0f, 0x34, 0x60)
RED    = RGBColor(0xe9, 0x45, 0x60)
GREEN  = RGBColor(0x28, 0xa7, 0x45)
TEAL   = RGBColor(0x17, 0xa2, 0xb8)
PURPLE = RGBColor(0x6f, 0x42, 0xc1)
GOLD   = RGBColor(0xf0, 0xa5, 0x00)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)
DARK   = RGBColor(0x2d, 0x34, 0x36)
LGRAY  = RGBColor(0xf5, 0xf6, 0xfa)

HEX_NAVY   = '0F3460'
HEX_RED    = 'E94560'
HEX_GREEN  = '28A745'
HEX_TEAL   = '17A2B8'
HEX_PURPLE = '6F42C1'
HEX_GOLD   = 'F0A500'
HEX_GRAY   = '6C757D'
HEX_WHITE  = 'FFFFFF'
HEX_LGRAY  = 'F5F6FA'
HEX_LIGHT  = 'F8F9FC'

FONT = 'Calibri'
OUT  = os.path.join(os.path.dirname(__file__), 'manuali', 'device_layout.docx')


# ── XML helpers ───────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None,
                 sz='8', space='0'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        if color:
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    sz)
            el.set(qn('w:color'), color.lstrip('#'))
            el.set(qn('w:space'), space)
        else:
            el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_all_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcW')):
            tcPr.remove(old)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    if line:
        sp.set(qn('w:line'),     str(line))
        sp.set(qn('w:lineRule'), 'exact')
    pPr.append(sp)


def _cell_vAlign(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(old)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), align)
    tcPr.append(va)


def _cell_margins(cell, top=80, bottom=80, left=113, right=113):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _left_bar(cell, color_hex):
    """Bordo sinistro colorato come accent bar."""
    _cell_border(cell, left=color_hex, sz='24')


def run_fmt(run, bold=False, italic=False, size=None, color=None, font=None):
    run.bold   = bold
    run.italic = italic
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.font.name = font or FONT


def add_run(para, text, bold=False, italic=False, size=None, color=None, font=None):
    r = para.add_run(text)
    run_fmt(r, bold=bold, italic=italic, size=size, color=color, font=font)
    return r


def heading(doc, text, level=1, color=NAVY):
    sizes = {1: 18, 2: 14, 3: 11}
    p = doc.add_paragraph()
    _para_spacing(p, before=240, after=80)
    r = p.add_run(text)
    run_fmt(r, bold=True, size=sizes.get(level, 12), color=color)
    return p


def body(doc, text, size=10, color=DARK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    _para_spacing(p, before=0, after=60)
    r = p.add_run(text)
    run_fmt(r, bold=bold, size=size, color=color)
    return p


def spacer(doc, before=100, after=0):
    p = doc.add_paragraph()
    _para_spacing(p, before=before, after=after)


# ── Contenuto ────────────────────────────────────────────────────────────────

DEVICES = [
    {
        'icon': '🖥️', 'name': 'Display Cucina (KDS)', 'color': HEX_GOLD, 'rgb': GOLD,
        'pos': 'Appeso o appoggiato in cucina, visibile dalla postazione di lavoro',
        'required': 'Necessario',
        'screen': '10" minimo, 15–24" consigliato',
        'os': 'Qualsiasi browser (/admin/cucina a schermo pieno)',
        'functions': [
            'Mostra gli ordini in tre colonne: da preparare, in preparazione, pronti',
            'Si ricarica da sola ogni 20 secondi e segnala i nuovi ordini con un suono',
            'Da qui si stampa lo scontrino da allegare al prodotto',
            'Genera e stampa le etichette QR del cesto (/admin/cesto)',
        ],
    },
    {
        'icon': '📱', 'name': 'Tablet Banco POS', 'color': HEX_GREEN, 'rgb': GREEN,
        'pos': 'Sul bancone, schermo orientabile verso il cliente',
        'required': 'Necessario',
        'screen': '10"+ consigliato',
        'os': 'Android o iOS — richiede fotocamera posteriore',
        'functions': [
            'Compone la vendita al banco (caffè, brioche, bevande) e genera il QR',
            'Mostra il QR di pagamento: il cliente lo inquadra dal display',
            'Consegna i pasti aziendali leggendo il codice di ritiro, a mano o con la camera',
            'Chiude gli ordini ritirati con "Consegnato"',
        ],
    },
    {
        'icon': '🧾', 'name': 'Registratore di cassa', 'color': HEX_RED, 'rgb': RED,
        'pos': 'Alla cassa — dispositivo esterno a QuickLunch',
        'required': 'Necessario',
        'screen': '—',
        'os': 'Indipendente dall\'applicazione',
        'functions': [
            'Emette lo scontrino fiscale della vendita',
            'È lo scontrino che va allegato al prodotto insieme all\'etichetta',
            'Non si integra con QuickLunch: i due sistemi restano separati',
        ],
    },
    {
        'icon': '🖨️', 'name': 'Stampante termica 80 mm', 'color': HEX_GRAY, 'rgb': GRAY,
        'pos': 'In cucina, accanto al display ordini',
        'required': 'Consigliato',
        'screen': 'Rotolo 80 mm',
        'os': 'Di rete con AirPrint/Mopria se si stampa da tablet',
        'functions': [
            'Stampa il tagliando dell\'ordine: documento interno di preparazione',
            'Il tagliando è impostato su 80 mm di larghezza, contenuto a 74 mm',
            'Porta codice ordine, orario di ritiro e composizione completa',
            'Non sostituisce lo scontrino fiscale del registratore di cassa',
        ],
    },
    {
        'icon': '🖨️', 'name': 'Stampante A4 comune', 'color': HEX_GRAY, 'rgb': GRAY,
        'pos': 'In cucina; l\'ufficio vi stampa attraverso la rete',
        'required': 'Necessario',
        'screen': 'A4',
        'os': 'Di rete con AirPrint/Mopria se si stampa da tablet',
        'functions': [
            'Stampa le etichette QR del cesto: griglia a 3 colonne su foglio, da tagliare',
            'Stampa la lista di produzione dei pasti aziendali',
            'Non sostituibile con la termica: la griglia non entra in 80 mm',
        ],
    },
    {
        'icon': '🖥️', 'name': 'PC Backoffice', 'color': HEX_PURPLE, 'rgb': PURPLE,
        'pos': 'Ufficio del responsabile',
        'required': 'Necessario',
        'screen': '13"+, sconsigliato sotto 11"',
        'os': 'Qualsiasi PC/Mac con browser',
        'functions': [
            'Configura menu, prodotti, prezzi, slot di ritiro e impostazioni',
            'Pubblica il menu del giorno delle convenzioni e i suoi allergeni',
            'Stampa la lista di produzione dei pasti aziendali',
            'Gestisce utenti, ricariche wallet, ruoli e permessi; legge i report',
        ],
    },
    {
        'icon': '📲', 'name': 'Smartphone del cliente', 'color': HEX_RED, 'rgb': RED,
        'pos': 'In possesso del cliente — non è hardware da acquistare',
        'required': 'Prerequisito',
        'screen': '4.5"+ qualsiasi',
        'os': 'Browser moderno + account QuickLunch',
        'functions': [
            'Ordina dal menu e compone panini, insalate e poke dal builder',
            'Acquista dal cesto inquadrando il QR dell\'etichetta',
            'Paga al banco inquadrando il QR mostrato dal tablet cassa',
            'Mostra il codice di ritiro del pasto aziendale',
        ],
    },
]

# Conteggio riassuntivo: è la risposta alla domanda "quanti device servono".
COUNT_MIN = [
    ('1', 'Display o tablet in cucina', 'cucina@bar.local  ·  ruolo cuoco'),
    ('2', 'Tablet al banco / cassa', 'banco@bar.local  ·  ruolo cassiere'),
    ('3', 'PC backoffice', 'admin@bar.local oppure sala@bar.local'),
    ('4', 'Registratore di cassa', 'esterno a QuickLunch, nessun account'),
]

COUNT_PRINTERS = [
    ('Termica 80 mm', 'Tagliandi ordine per la cucina', 'In cucina'),
    ('A4 comune', 'Etichette cesto e liste pasti', 'In cucina, condivisa in rete'),
]

QR_STEPS = [
    ('🛒', HEX_GREEN, 'Selezione articoli',
     'Staff', 'Lo staff tocca gli articoli sul tablet banco (caffè, brioche, ecc.) e compone il totale.'),
    ('⬛', HEX_TEAL, 'Genera QR',
     'Sistema', "Il sistema crea una sessione di pagamento con scadenza 10 min e mostra il QR sul tablet."),
    ('📸', HEX_RED, 'Scansione QR',
     'Cliente', 'Il cliente punta la fotocamera del telefono verso lo schermo del banco e apre il link.'),
    ('✅', HEX_PURPLE, 'Conferma pagamento',
     'Cliente', 'Il telefono mostra riepilogo articoli e totale — il cliente tocca "Paga ora".'),
    ('💰', HEX_GOLD, 'Addebito wallet',
     'Sistema', 'Il wallet viene scalato automaticamente. Il tablet banco mostra "✓ Pagato da [nome]".'),
]

REQS = [
    ('🖥️ Display Cucina',      'Wi-Fi',        'Cuoco',            '10"',  'Sì'),
    ('📱 Tablet Banco POS',    'Wi-Fi',        'Cassiere',         '10"',  'Sì'),
    ('🧾 Registratore cassa',  '—',            '—',                '—',    'Sì'),
    ('🖨️ Termica 80 mm',       'Wi-Fi / USB',  '—',                '—',    'Consigliato'),
    ('🖨️ Stampante A4',        'Wi-Fi / USB',  '—',                '—',    'Sì'),
    ('🖥️ PC Backoffice',       'Wi-Fi / LAN',  'Admin o manager',  '13"',  'Sì'),
    ('📲 Smartphone cliente',  'Wi-Fi / 4G',   'Cliente',          '4.5"', 'Prerequisito'),
    ('📱 Secondo tablet banco', 'Wi-Fi',       'Cassiere',         '8"',   'Consigliato'),
]

FLOOR_AREAS = [
    ('🍳 CUCINA',        'Preparazione, cesto, stampa',
     ['Display Cucina (KDS)', 'Stampante termica 80 mm', 'Stampante A4']),
    ('☕ BANCO',         'Cassa QR, scontrini e consegne',
     ['Tablet Banco POS', 'Registratore di cassa',
      'Eventuale secondo tablet di riserva']),
    ('🍽️ SALA',          'Clienti',
     ['Smartphone del cliente (in mano al cliente)']),
    ('🖥️ UFFICIO',       'Configurazione e report',
     ['PC Backoffice']),
]


# ── Build document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Margini pagina
    for sec in doc.sections:
        sec.page_width   = Cm(21)
        sec.page_height  = Cm(29.7)
        sec.left_margin  = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin   = Cm(2)
        sec.bottom_margin = Cm(2)

    # ── Cover band ────────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    _remove_all_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _cell_shd(cell, HEX_NAVY)
    _cell_margins(cell, top=280, bottom=280, left=340, right=340)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('QuickLunch')
    run_fmt(r, bold=True, size=28, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p2, before=60, after=0)
    r2 = p2.add_run('Layout Dispositivi — Quanti device servono e dove')
    run_fmt(r2, size=13, color=RGBColor(0xb2, 0xc2, 0xd9))

    spacer(doc, 300)

    # ── SEZIONE 1: Quanti dispositivi servono ────────────────────────────────
    heading(doc, '1 — Quanti dispositivi servono', 1, NAVY)
    body(doc,
         'Configurazione minima per coprire tutti i flussi attivi: vendita al banco, '
         'cesto self-service, pasti aziendali e ordini dal builder visuale.',
         size=10, color=GRAY)
    spacer(doc, 80)

    # Tre device + account
    tbl_c = doc.add_table(rows=1 + len(COUNT_MIN), cols=3)
    _remove_all_borders(tbl_c)
    for ci, w in enumerate([1.4, 7.4, 8.2]):
        _set_col_width(tbl_c, ci, w)
    for ci, lbl in enumerate(['#', 'Dispositivo', 'Account dedicato']):
        cell = tbl_c.cell(0, ci)
        _cell_shd(cell, HEX_NAVY)
        _cell_margins(cell, top=100, bottom=100, left=113, right=113)
        _cell_vAlign(cell)
        run_fmt(cell.paragraphs[0].add_run(lbl), bold=True, size=9, color=WHITE)
    for ri, (num, name, account) in enumerate(COUNT_MIN):
        bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        cells = tbl_c.rows[ri + 1].cells
        for ci, val in enumerate([num, name, account]):
            cell = cells[ci]
            _cell_shd(cell, bg)
            _cell_margins(cell, top=90, bottom=90, left=113, right=113)
            _cell_vAlign(cell, 'center')
            pp = cell.paragraphs[0]
            if ci == 0:
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(pp, val, bold=True, size=11, color=RED)
            else:
                add_run(pp, val, bold=(ci == 1), size=9.5, color=DARK)

    spacer(doc, 160)

    # Due stampanti
    body(doc, 'Piu due stampanti, con formati non interscambiabili:',
         size=10, color=DARK, bold=True)
    spacer(doc, 60)

    tbl_p = doc.add_table(rows=1 + len(COUNT_PRINTERS), cols=3)
    _remove_all_borders(tbl_p)
    for ci, w in enumerate([4.6, 6.4, 6.0]):
        _set_col_width(tbl_p, ci, w)
    for ci, lbl in enumerate(['Stampante', 'Serve per', 'Posizione']):
        cell = tbl_p.cell(0, ci)
        _cell_shd(cell, HEX_GRAY)
        _cell_margins(cell, top=100, bottom=100, left=113, right=113)
        _cell_vAlign(cell)
        run_fmt(cell.paragraphs[0].add_run(lbl), bold=True, size=9, color=WHITE)
    for ri, row in enumerate(COUNT_PRINTERS):
        bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        cells = tbl_p.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            _cell_shd(cell, bg)
            _cell_margins(cell, top=90, bottom=90, left=113, right=113)
            _cell_vAlign(cell, 'center')
            add_run(cell.paragraphs[0], val, bold=(ci == 0), size=9.5, color=DARK)

    spacer(doc, 180)

    # Scontrino fiscale vs tagliando dell'applicazione
    tbl_sc = doc.add_table(rows=1, cols=1)
    _remove_all_borders(tbl_sc)
    sc = tbl_sc.cell(0, 0)
    _cell_shd(sc, 'FDEEF0')
    _cell_border(sc, left=HEX_RED, sz='20')
    _cell_margins(sc, top=120, bottom=120, left=170, right=140)
    psc = sc.paragraphs[0]
    add_run(psc, '🧾  Due documenti diversi: ', bold=True, size=10, color=RED)
    add_run(psc, 'lo scontrino da allegare al prodotto è quello fiscale del '
                 'registratore di cassa, che resta un dispositivo esterno e non '
                 'comunica con QuickLunch. Il foglio che l\'applicazione stampa è '
                 'il tagliando dell\'ordine, utile alla cucina per preparare ma '
                 'privo di valore fiscale: per questo la stampante termica è '
                 'consigliata, non necessaria.',
            size=9, color=DARK)

    spacer(doc, 100)

    # Vincolo di stampa da tablet
    tbl_w = doc.add_table(rows=1, cols=1)
    _remove_all_borders(tbl_w)
    wc = tbl_w.cell(0, 0)
    _cell_shd(wc, 'FDF6E7')
    _cell_border(wc, left=HEX_GOLD, sz='20')
    _cell_margins(wc, top=120, bottom=120, left=170, right=140)
    pw = wc.paragraphs[0]
    add_run(pw, '⚠️  Stampa da tablet: ', bold=True, size=10, color=GOLD)
    add_run(pw, 'la stampa parte dal browser, quindi da Android o iOS funziona solo con '
                'stampanti di rete AirPrint / Mopria. Una termica USB ESC/POS non si pilota '
                'dal browser di un tablet: in quel caso il dispositivo in cucina deve essere '
                'un mini-PC, con le stampanti collegate via USB.',
            size=9, color=DARK)

    spacer(doc, 100)

    # Ridondanza
    tbl_rd = doc.add_table(rows=1, cols=1)
    _remove_all_borders(tbl_rd)
    rc = tbl_rd.cell(0, 0)
    _cell_shd(rc, 'EAF0FB')
    _cell_border(rc, left=HEX_TEAL, sz='20')
    _cell_margins(rc, top=120, bottom=120, left=170, right=140)
    pr_ = rc.paragraphs[0]
    add_run(pr_, 'ℹ️  Un quarto dispositivo per sicurezza: ', bold=True, size=10, color=TEAL)
    add_run(pr_, 'il tablet al banco è l\'unico senza alternativa — se si guasta si fermano '
                 'insieme i pagamenti QR e le consegne dei pasti. Un secondo tablet, anche '
                 'piccolo, o uno smartphone di servizio copre il rischio. La cucina in '
                 'emergenza può invece lavorare dal PC backoffice.',
            size=9, color=DARK)

    spacer(doc, 100)

    # Cosa non serve
    tbl_nn = doc.add_table(rows=1, cols=1)
    _remove_all_borders(tbl_nn)
    nnc = tbl_nn.cell(0, 0)
    _cell_shd(nnc, HEX_LIGHT)
    _cell_border(nnc, left=HEX_GRAY, sz='20')
    _cell_margins(nnc, top=120, bottom=120, left=170, right=140)
    pnn = nnc.paragraphs[0]
    add_run(pnn, '—  Cosa non serve: ', bold=True, size=10, color=GRAY)
    add_run(pnn, 'nessun tablet in sala, perché la gestione tavoli è disattivata dalle '
                 'Impostazioni; e nessun lettore barcode fisso, perché la scansione la fa '
                 'sempre il telefono del cliente o la fotocamera del tablet cassa.',
            size=9, color=DARK)

    spacer(doc, 200)

    # ── SEZIONE 2: Posizionamento nel locale ─────────────────────────────────
    heading(doc, '2 — Posizionamento dispositivi nel locale', 1, NAVY)
    body(doc,
         'La tabella seguente indica le aree fisiche del locale e i dispositivi '
         'che devono essere presenti in ciascuna di esse.',
         size=10, color=GRAY)
    spacer(doc, 100)

    tbl2 = doc.add_table(rows=1 + len(FLOOR_AREAS), cols=3)
    _remove_all_borders(tbl2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [4.5, 4.0, 8.5]
    for ci, w in enumerate(widths):
        _set_col_width(tbl2, ci, w)

    # header row
    hdr_labels = ['Area', 'Funzione', 'Dispositivi presenti']
    for ci, lbl in enumerate(hdr_labels):
        cell = tbl2.cell(0, ci)
        _cell_shd(cell, HEX_NAVY)
        _cell_margins(cell, top=100, bottom=100, left=113, right=113)
        _cell_vAlign(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(lbl)
        run_fmt(r, bold=True, size=10, color=WHITE)

    area_colors = [HEX_GOLD, HEX_GREEN, HEX_TEAL, HEX_PURPLE]
    for ri, (area, func, devs) in enumerate(FLOOR_AREAS):
        row_bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        cells = tbl2.rows[ri + 1].cells
        _cell_shd(cells[0], row_bg)
        _cell_shd(cells[1], row_bg)
        _cell_shd(cells[2], row_bg)
        _cell_border(cells[0], left=area_colors[ri], sz='20')
        for ci in range(3):
            _cell_margins(cells[ci], top=90, bottom=90, left=113, right=113)
            _cell_vAlign(cells[ci], 'top')

        p0 = cells[0].paragraphs[0]
        add_run(p0, area, bold=True, size=10, color=NAVY)

        p1 = cells[1].paragraphs[0]
        add_run(p1, func, size=9, color=GRAY, italic=True)

        p2 = cells[2].paragraphs[0]
        for di, dev in enumerate(devs):
            if di > 0:
                p2.add_run('\n')
            add_run(p2, '• ' + dev, size=9, color=DARK)

    spacer(doc, 200)

    # ── SEZIONE 3: Schede dispositivi ────────────────────────────────────────
    heading(doc, '3 — Schede dispositivi', 1, NAVY)
    body(doc, 'Per ogni dispositivo: posizione fisica, sistema operativo richiesto e funzioni nell\'app.',
         size=10, color=GRAY)
    spacer(doc, 80)

    for dev in DEVICES:
        tbl_d = doc.add_table(rows=1, cols=2)
        _remove_all_borders(tbl_d)
        _set_col_width(tbl_d, 0, 2.0)
        _set_col_width(tbl_d, 1, 15.0)

        # icon cell
        ic = tbl_d.cell(0, 0)
        _cell_shd(ic, dev['color'])
        _cell_margins(ic, top=140, bottom=140, left=140, right=100)
        _cell_vAlign(ic, 'top')
        pi = ic.paragraphs[0]
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pi, dev['icon'], size=22, color=WHITE)

        # content cell
        cc = tbl_d.cell(0, 1)
        _cell_shd(cc, 'F8F9FC')
        _cell_border(cc, left=dev['color'], sz='16')
        _cell_margins(cc, top=120, bottom=120, left=170, right=140)
        _cell_vAlign(cc, 'top')

        # name
        pn = cc.paragraphs[0]
        add_run(pn, dev['name'], bold=True, size=12, color=NAVY)
        _para_spacing(pn, before=0, after=30)

        # pos + required
        pp = cc.add_paragraph()
        _para_spacing(pp, before=0, after=60)
        add_run(pp, '📍 ' + dev['pos'] + '   ', size=9, color=GRAY, italic=True)
        badge_color = GREEN if dev['required'] == 'Necessario' else (GOLD if dev['required'] == 'Consigliato' else GRAY)
        add_run(pp, '[' + dev['required'] + ']', bold=True, size=9, color=badge_color)

        # screen + OS
        ptech = cc.add_paragraph()
        _para_spacing(ptech, before=0, after=60)
        add_run(ptech, 'Schermo: ', bold=True, size=9, color=DARK)
        add_run(ptech, dev['screen'] + '    ', size=9, color=DARK)
        add_run(ptech, 'OS / Browser: ', bold=True, size=9, color=DARK)
        add_run(ptech, dev['os'], size=9, color=DARK)

        # functions
        for fn in dev['functions']:
            pf = cc.add_paragraph()
            _para_spacing(pf, before=10, after=10)
            add_run(pf, '▸  ' + fn, size=9, color=DARK)

        spacer(doc, 120)

    # ── SEZIONE 4: Flusso QR ─────────────────────────────────────────────────
    heading(doc, '4 — Flusso pagamento QR banco', 1, NAVY)
    body(doc, 'Sequenza completa dall\'inizio della transazione alla conferma di pagamento.',
         size=10, color=GRAY)
    spacer(doc, 80)

    for step_n, (icon, color, title, actor, desc) in enumerate(QR_STEPS, 1):
        tbl_s = doc.add_table(rows=1, cols=3)
        _remove_all_borders(tbl_s)
        _set_col_width(tbl_s, 0, 1.4)
        _set_col_width(tbl_s, 1, 2.4)
        _set_col_width(tbl_s, 2, 13.2)

        # numero step
        cn = tbl_s.cell(0, 0)
        _cell_shd(cn, color)
        _cell_vAlign(cn, 'center')
        _cell_margins(cn, top=120, bottom=120, left=80, right=80)
        pn = cn.paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pn, str(step_n), bold=True, size=20, color=WHITE)

        # icona + attore
        ca = tbl_s.cell(0, 1)
        _cell_shd(ca, 'F0F4FF')
        _cell_vAlign(ca, 'center')
        _cell_margins(ca, top=100, bottom=100, left=100, right=100)
        pa = ca.paragraphs[0]
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pa, icon + '\n', size=18)
        add_run(pa, actor, bold=True, size=8, color=GRAY)

        # titolo + desc
        cd = tbl_s.cell(0, 2)
        _cell_shd(cd, HEX_WHITE)
        _cell_border(cd, left=color, sz='16')
        _cell_vAlign(cd, 'center')
        _cell_margins(cd, top=100, bottom=100, left=170, right=140)
        pt = cd.paragraphs[0]
        add_run(pt, title, bold=True, size=11, color=NAVY)
        _para_spacing(pt, before=0, after=40)
        pd = cd.add_paragraph()
        add_run(pd, desc, size=9, color=DARK)
        _para_spacing(pd, before=0, after=0)

        spacer(doc, 100)

    # ── SEZIONE 5: Requisiti tecnici ─────────────────────────────────────────
    heading(doc, '5 — Requisiti tecnici', 1, NAVY)
    body(doc, 'Riepilogo connessione, credenziali e dimensione schermo per ogni dispositivo.',
         size=10, color=GRAY)
    spacer(doc, 80)

    req_headers = ['Dispositivo', 'Connessione', 'Account', 'Schermo min.', 'Necessità']
    tbl_r = doc.add_table(rows=1 + len(REQS), cols=5)
    _remove_all_borders(tbl_r)
    req_widths = [5.0, 3.2, 4.0, 2.8, 2.8]
    for ci, w in enumerate(req_widths):
        _set_col_width(tbl_r, ci, w)

    for ci, lbl in enumerate(req_headers):
        cell = tbl_r.cell(0, ci)
        _cell_shd(cell, HEX_NAVY)
        _cell_margins(cell, top=100, bottom=100, left=113, right=113)
        _cell_vAlign(cell)
        p = cell.paragraphs[0]
        r = p.add_run(lbl)
        run_fmt(r, bold=True, size=9, color=WHITE)

    req_badge_colors = {
        'Sì': HEX_GREEN, 'Consigliato': HEX_GOLD,
        'Prerequisito': HEX_TEAL, 'Opzionale': HEX_GRAY
    }
    req_badge_rgb = {
        'Sì': GREEN, 'Consigliato': GOLD,
        'Prerequisito': TEAL, 'Opzionale': GRAY
    }
    for ri, (device, conn, account, screen, req) in enumerate(REQS):
        bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        row_cells = tbl_r.rows[ri + 1].cells
        values = [device, conn, account, screen, req]
        for ci, val in enumerate(values):
            cell = row_cells[ci]
            _cell_shd(cell, bg)
            _cell_margins(cell, top=80, bottom=80, left=113, right=113)
            _cell_vAlign(cell, 'center')
            p = cell.paragraphs[0]
            if ci == 4:
                col = req_badge_rgb.get(val, GRAY)
                add_run(p, val, bold=True, size=9, color=col)
            else:
                add_run(p, val, size=9, color=DARK)

    spacer(doc, 200)

    # ── Nota rete ────────────────────────────────────────────────────────────
    tbl_n = doc.add_table(rows=1, cols=1)
    _remove_all_borders(tbl_n)
    nc = tbl_n.cell(0, 0)
    _cell_shd(nc, 'EAF0FB')
    _cell_border(nc, left=HEX_TEAL, sz='20')
    _cell_margins(nc, top=120, bottom=120, left=170, right=140)
    pnote = nc.paragraphs[0]
    add_run(pnote, 'ℹ️  Requisito di rete: ', bold=True, size=10, color=TEAL)
    add_run(pnote, 'tutti i dispositivi devono raggiungere il server QuickLunch sulla stessa LAN '
            '(o via internet se il server è hosted in cloud). '
            'Connessione minima consigliata: Wi-Fi 802.11n o superiore.',
            size=9, color=DARK)

    spacer(doc, 80)

    # ── Footer ────────────────────────────────────────────────────────────────
    p_foot = doc.add_paragraph()
    _para_spacing(p_foot, before=200, after=0)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_foot, 'QuickLunch — Documento tecnico — © 2024–26 DS Consulting', size=8, color=GRAY)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] salvato in: {OUT}')


if __name__ == '__main__':
    build()
