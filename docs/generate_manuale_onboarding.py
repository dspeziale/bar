#!/usr/bin/env python3
"""Genera docs/manuali/manuale_onboarding_cliente.docx.

Percorso completo del cliente: dai prerequisiti del gestore fino al primo ordine
ritirato. Le voci di menu, i percorsi e i valori predefiniti citati corrispondono
all'applicazione: se cambiano, aggiornare qui.

    python docs/generate_manuale_onboarding.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Palette (coerente con gli altri documenti in docs/) ──────────────────────
RED    = RGBColor(0xe9, 0x45, 0x60)
NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x16, 0x21, 0x3e)
DGRAY  = RGBColor(0x34, 0x3a, 0x40)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREEN  = RGBColor(0x27, 0xae, 0x60)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
BLUE   = RGBColor(0x34, 0x98, 0xdb)
PURPLE = RGBColor(0x8e, 0x44, 0xad)
SLATE  = RGBColor(0x7f, 0x8c, 0x8d)

HEX_RED    = 'E94560'
HEX_NAVY   = '0F3460'
HEX_GREEN  = '27AE60'
HEX_ORANGE = 'E67E22'
HEX_BLUE   = '3498DB'
HEX_PURPLE = '8E44AD'
HEX_SLATE  = '7F8C8D'
HEX_LIGHT  = 'F8F9FA'
HEX_RULE   = 'E8EBEE'
HEX_WARN   = 'FDF6E7'
HEX_STOP   = 'FDEEF0'

FONT = 'PT Sans Narrow'
BODY_FONT = 'Calibri'

OUT = os.path.join(os.path.dirname(__file__), 'manuali',
                   'manuale_onboarding_cliente.docx')


# ── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None, sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        if color:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color.lstrip('#'))
        else:
            el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), val)
    tcPr.append(el)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _p_spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def _run_font(run, size=10.5, bold=False, italic=False, color=None, font=BODY_FONT):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    _p_spacing(p, before=0, after=0)


# ── Blocchi di contenuto ─────────────────────────────────────────────────────

def heading(doc, text, level=1, color=NAVY, before=16):
    sizes = {1: 19, 2: 14.5, 3: 12}
    p = doc.add_paragraph()
    _p_spacing(p, before=before, after=5)
    _run_font(p.add_run(text), size=sizes.get(level, 12), bold=True,
              color=color, font=FONT)
    return p


def body(doc, text, size=10.5, color=DGRAY, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    _run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def rich(doc, parts, size=10.5, after=6, indent=None):
    """Paragrafo con porzioni in grassetto: parts = [(testo, bold), ...]."""
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def bullet(doc, parts, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    _p_spacing(p, before=0, after=3)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def rule(doc, color=HEX_RULE, after=8):
    """Filetto orizzontale come tabella a una cella."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_border(cell, bottom=color, sz='8')
    _cell_margins(cell, top=0, bottom=0, left=0, right=0)
    _p_spacing(cell.paragraphs[0], before=0, after=0)
    _run_font(cell.paragraphs[0].add_run(''), size=1)
    spacer(doc, after)


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(''), size=pts / 2)


def box(doc, label, paragraphs, accent=HEX_RED, fill=HEX_LIGHT, label_color=RED):
    """Riquadro con bordo sinistro colorato."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, fill)
    _cell_border(cell, top=HEX_RULE, bottom=HEX_RULE, right=HEX_RULE,
                 left=accent, sz='18')
    _cell_margins(cell, top=90, bottom=90, left=140, right=110)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=3)
    _run_font(p.add_run(label.upper()), size=8.5, bold=True, color=label_color,
              font=FONT)

    for i, parts in enumerate(paragraphs):
        pp = cell.add_paragraph()
        _p_spacing(pp, before=0, after=0 if i == len(paragraphs) - 1 else 4)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 10)
    return tbl


def step(doc, num, title, paragraphs, accent=HEX_RED):
    """Passo numerato: numero in pastiglia colorata + contenuto."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_width(tbl, 0, 1.0)
    _set_col_width(tbl, 1, 15.5)

    nc = tbl.rows[0].cells[0]
    _cell_shd(nc, accent)
    _cell_margins(nc, top=60, bottom=60, left=30, right=30)
    _cell_valign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run('%02d' % num), size=12, bold=True, color=WHITE, font=FONT)

    cc = tbl.rows[0].cells[1]
    _cell_margins(cc, top=50, bottom=50, left=150, right=60)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=2)
    _run_font(pt.add_run(title), size=11.5, bold=True, color=DARK, font=FONT)

    for parts in paragraphs:
        pp = cc.add_paragraph()
        _p_spacing(pp, before=0, after=3)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 6)
    return tbl


def table_grid(doc, headers, rows, widths, head_fill=HEX_NAVY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, w in enumerate(widths):
        _set_col_width(tbl, i, w)

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _cell_shd(c, head_fill)
        _cell_margins(c, top=60, bottom=60, left=110, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(h.upper()), size=8.5, bold=True, color=WHITE, font=FONT)

    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            c = tbl.rows[r].cells[i]
            _cell_shd(c, HEX_LIGHT if r % 2 else 'FFFFFF')
            _cell_border(c, bottom=HEX_RULE, sz='4')
            _cell_margins(c, top=55, bottom=55, left=110, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            text, color, bold = val if isinstance(val, tuple) else (val, DGRAY, False)
            _run_font(p.add_run(text), size=10, bold=bold, color=color)
    spacer(doc, 10)
    return tbl


# ── Documento ────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)

    # ══ Copertina ═════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, HEX_NAVY)
    _cell_margins(cell, top=300, bottom=300, left=320, right=320)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=4)
    _run_font(p.add_run('QUICKLUNCH  ·  ACCOGLIENZA CLIENTI'), size=9.5,
              bold=True, color=RGBColor(0xb2, 0xc2, 0xd9), font=FONT)

    p2 = cell.add_paragraph()
    _p_spacing(p2, before=0, after=6)
    _run_font(p2.add_run('Onboarding cliente'), size=30, bold=True,
              color=WHITE, font=FONT)

    p3 = cell.add_paragraph()
    _p_spacing(p3, before=0, after=0)
    _run_font(p3.add_run(
        'Dal primo contatto al primo ordine ritirato: cosa fa il cliente, cosa '
        'deve fare lo staff e in quale momento. Sei fasi, piu i prerequisiti da '
        'sistemare una volta sola.'), size=11,
        color=RGBColor(0xd6, 0xdf, 0xea))
    spacer(doc, 16)

    # ══ Il percorso in sintesi ════════════════════════════════════════════
    heading(doc, 'Il percorso in sintesi', 1)
    body(doc, 'Le fasi in cui il cliente e lo staff si alternano. Il primo ordine '
              'non e possibile se una delle prime quattro non e completa.',
         color=GRAY)
    spacer(doc, 4)

    table_grid(doc,
               ['Fase', 'Chi agisce', 'Risultato'],
               [
                   [('0 · Prerequisiti', DARK, True), 'Gestore',
                    'Menu popolato e slot di ritiro attivi'],
                   [('1 · Invito', DARK, True), 'Staff',
                    'QR di registrazione affisso e visibile'],
                   [('2 · Registrazione', DARK, True), 'Cliente',
                    'Account creato, in attesa di approvazione'],
                   [('3 · Attivazione', DARK, True), 'Staff',
                    'Account abilitato ad accedere'],
                   [('4 · Primo credito', DARK, True), 'Staff',
                    'Wallet capiente (si salta se il portafoglio e spento)'],
                   [('5 · Primo ordine', DARK, True), 'Cliente',
                    'Ordine pagato, in coda in cucina'],
                   [('6 · Rifiniture', DARK, True), 'Cliente',
                    'Notifiche attive per gli ordini successivi'],
               ],
               widths=[4.2, 3.6, 8.7])

    box(doc, 'Il punto di attrito', [
        [('Tra la fase 2 e la fase 3 il cliente resta in attesa: si registra, '
          'ottiene un account ', False), ('non attivo', True),
         (' e il login gli viene rifiutato finche non lo abilitate. Voi ricevete '
          'l\'avviso su Telegram, lui riceve ', False),
         ('un\'email automatica appena lo attivate', True),
         ('. Attivate comunque in fretta: finche restate voi il collo di '
          'bottiglia, il cliente e fermo.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    _page_break(doc)

    # ══ FASE 0 · PREREQUISITI ═════════════════════════════════════════════
    heading(doc, 'Fase 0 — Prerequisiti del gestore', 1, color=ORANGE, before=0)
    body(doc, 'Da sistemare una volta sola, prima di invitare il primo cliente. '
              'Senza questi passi il cliente arriva fino al carrello e poi non '
              'trova nulla da ordinare.', color=GRAY)
    rule(doc)

    step(doc, 1, 'Controlla il catalogo di partenza', [
        [('L\'installazione — e ogni azzeramento del database — crea un '
          'catalogo iniziale gia pronto: ', False),
         ('18 categorie di prodotto e un listino di circa 75 voci', True),
         (', gli slot di ritiro, i tavoli, gli articoli del banco e gli '
          'ingredienti del builder. Tutto visibile subito, senza riavviare '
          'nulla.', False)],
        [('Apri ', False), ('Categorie', True),
         (' e togli quelle che non ti servono: e piu rapido che crearle da '
          'zero. Di ognuna puoi cambiare nome, icona e colore.', False)],
    ], accent=HEX_ORANGE)

    step(doc, 2, 'Adatta il listino al tuo bar', [
        [('L\'installazione crea anche un ', False),
         ('listino di partenza', True),
         (' da bar caffetteria con servizio mensa: circa 75 prodotti nelle '
          '18 categorie — caffetteria, colazione, panini e tramezzini, primi '
          'e secondi, contorni, frutta, dolci, bevande — ognuno con prezzo, '
          'quantita giornaliera e allergeni gia compilati.', False)],
        [('Il lavoro e di ', False), ('correzione, non di creazione', True),
         (': in Prodotti disattiva quello che non vendi, allinea i prezzi ai '
          'tuoi e ritocca le quantita sulla tua capacita di produzione. '
          'Aggiungi le tue specialita con Prodotti › Nuovo.', False)],
        [('Controlla gli ', False), ('allergeni', True),
         (' dei prodotti che tieni: quelli predefiniti valgono per una '
          'ricetta tipica, ma la tua puo essere diversa.', False)],
    ], accent=HEX_ORANGE)

    step(doc, 3, 'Verifica gli slot di ritiro', [
        [('Slot orari: ne trovi otto preimpostati da ', False),
         ('11:45 a 13:30', True),
         (' ogni quarto d\'ora, con capienza di 20 ordini ciascuno. Disattiva '
          'quelli che non usi e correggi la capienza sulla tua reale capacita '
          'di produzione.', False)],
        [('Serve almeno uno slot attivo: senza slot il carrello non si puo '
          'confermare, salvo l\'opzione "subito al banco".', False)],
    ], accent=HEX_ORANGE)

    step(doc, 4, 'Compila i dati dell\'attivita', [
        [('Impostazioni › Azienda: la ', False), ('ragione sociale', True),
         (' compare in testa agli scontrini e alle etichette del cesto. Se resta '
          'vuota, i documenti stampano "QuickLunch".', False)],
    ], accent=HEX_ORANGE)

    step(doc, 5, 'Decidi la politica di credito', [
        [('Tre leve, in Impostazioni:', False)],
        [('Bonus benvenuto', True),
         (' — importo accreditato automaticamente al primo accesso al servizio. '
          'Vale per ', False), ('tutti i nuovi clienti', True),
         (': registrazione pubblica, Google, link di un\'azienda convenzionata e '
          'anche creazione manuale dal backoffice. Il valore predefinito e ',
          False), ('0', True), (': finche resta tale, nessun bonus viene dato. '
          'Con il portafoglio prepagato disattivato il bonus non esiste.',
                                False)],
        [('Punti fedelta', True),
         (' — di serie 10 punti per ogni euro speso, 100 punti riscattabili per '
          '1,00 €.', False)],
        [('Fido', True),
         (' — impostabile per singolo cliente dalla sua scheda: consente di '
          'ordinare andando in rosso fino alla soglia scelta.', False)],
    ], accent=HEX_ORANGE)

    spacer(doc, 6)
    box(doc, 'Consigliato, non obbligatorio', [
        [('Configura ', False), ('Telegram', True),
         (' (token del bot e chat del canale) in Impostazioni › Notifiche: senza '
          'di esso perdete gli avvisi di nuova registrazione e di nuovo ordine, '
          'e il cliente non riceve le conferme.', False)],
    ], accent=HEX_SLATE, fill=HEX_LIGHT, label_color=SLATE)

    _page_break(doc)

    # ══ FASE 1 · INVITO ═══════════════════════════════════════════════════
    heading(doc, 'Fase 1 — Portare il cliente alla registrazione', 1,
            color=BLUE, before=0)
    body(doc, 'Il cliente si registra da solo, dal proprio telefono. Il vostro '
              'compito e mettergli davanti il punto di accesso.', color=GRAY)
    rule(doc)

    step(doc, 1, 'Stampa il QR di registrazione', [
        [('Clienti › ', False), ('QR registrazione', True),
         ('  (/admin/clients/registration-qr): pagina già impaginata per la '
          'stampa, con il QR e le istruzioni. Premi Stampa.', False)],
    ], accent=HEX_BLUE)

    step(doc, 2, 'Affiggilo dove si forma la fila', [
        [('Al banco, sulla cassa e sui tavoli. Il QR porta a una pagina che '
          'offre due strade: registrazione ', False), ('con Google', True),
         (' in un tocco, oppure ', False), ('email e password', True), ('.', False)],
    ], accent=HEX_BLUE)

    step(doc, 3, 'Spiega la regola in una frase', [
        [('"Inquadra, registrati, poi passa in cassa a caricare il credito." '
          'Il wallet e prepagato: senza credito l\'ordine non parte, ed e la '
          'sorpresa piu comune per chi arriva la prima volta.', False)],
    ], accent=HEX_BLUE)

    # ══ FASE 2 · REGISTRAZIONE ════════════════════════════════════════════
    heading(doc, 'Fase 2 — Il cliente si registra', 1, color=GREEN)
    body(doc, 'Quello che accade sul telefono del cliente, per sapere cosa '
              'aspettarsi quando vi chiede aiuto.', color=GRAY)
    rule(doc)

    step(doc, 1, 'Compila o entra con Google', [
        [('Con email e password servono nome, cognome, email e una password di ',
          False), ('almeno 6 caratteri', True),
         (', ripetuta due volte. Con Google non serve scegliere password.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'Arriva sulla pagina di attesa', [
        [('L\'account nasce ', False), ('non attivo', True),
         ('. Il cliente vede una schermata che gli chiede di attendere '
          'l\'abilitazione e, se prova a entrare, il login viene rifiutato.',
          False)],
        [('Nello stesso momento sul canale Telegram dello staff arriva '
          '"Nuovo cliente in attesa" con nome ed email.', False)],
    ], accent=HEX_GREEN)

    spacer(doc, 6)
    box(doc, 'Se preferite registrarlo voi', [
        [('Clienti › Nuovo cliente crea l\'account ', False),
         ('già attivo', True),
         (', salvando anche telefono, data di nascita e indirizzo. Due cautele: '
          'assegna sempre una ', False), ('password', True),
         (', altrimenti il cliente non potrà accedere se non tramite Google; e '
          'verifica che compaia poi nella lista clienti.', False)],
        [('Su questo secondo punto c\'e un difetto noto: vedi "Una cosa da '
          'sapere" a fine manuale.', False)],
    ], accent=HEX_SLATE, fill=HEX_LIGHT, label_color=SLATE)

    _page_break(doc)

    # ══ FASE 3 · ATTIVAZIONE ══════════════════════════════════════════════
    heading(doc, 'Fase 3 — Attivare l\'account', 1, color=PURPLE, before=0)
    body(doc, 'E il passaggio che sblocca tutto, e va fatto in fretta: il '
              'cliente ha ricevuto la conferma di registrazione con la guida '
              'allegata, ma finche non lo attivi non puo entrare.', color=GRAY)
    rule(doc)

    box(doc, 'Le due email al cliente', [
        [('Appena si registra', True),
         (' — da qualunque via: email, Google, link dell\'azienda — il '
          'cliente riceve "Registrazione ricevuta", con la ', False),
         ('guida del cliente in PDF', True),
         (' allegata e il pulsante per collegare Telegram.', False)],
        [('Quando lo attivi', True),
         (' riceve la seconda email, "Il tuo account e attivo", con i primi '
          'passi e di nuovo la guida.', False)],
        [('Perche partano serve ', False), ('Gmail configurata', True),
         (' in Impostazioni › Notifiche: senza quella nessuna delle due '
          'viene inviata, e non compare alcun errore.', False)],
    ], accent=HEX_GREEN, fill=HEX_LIGHT, label_color=GREEN)

    step(doc, 1, 'Apri la lista clienti', [
        [('Clienti: i nomi in attesa sono quelli marcati come non attivi. Il '
          'contatore dei clienti da attivare compare anche sulla dashboard.',
          False)],
    ], accent=HEX_PURPLE)

    step(doc, 2, 'Attiva', [
        [('Premi ', False), ('Attiva', True),
         (' sulla riga del cliente. Se il cliente appartiene a un\'azienda '
          'convenzionata, in questo momento puoi anche ', False),
         ('associarlo alla convenzione', True),
         (': da lì in avanti vedrà il pasto del giorno e potrà prenotarlo.',
          False)],
    ], accent=HEX_PURPLE)

    step(doc, 3, 'Il cliente riceve l\'email di conferma', [
        [('All\'attivazione parte automaticamente un\'', False),
         ('email al cliente', True),
         (': gli dice che l\'account e attivo, come accedere e che il credito si '
          'ricarica in cassa. Chi ha collegato Telegram riceve anche l\'avviso '
          'sul telefono.', False)],
        [('L\'email richiede che Gmail sia configurata in Impostazioni › '
          'Notifiche. Se non lo e, nessun avviso parte: in quel caso ditelo al '
          'cliente di persona.', False)],
    ], accent=HEX_PURPLE)

    # ══ FASE 4 · PRIMO CREDITO ════════════════════════════════════════════
    heading(doc, 'Fase 4 — Il primo credito (solo con portafoglio attivo)',
            1, color=GREEN)
    body(doc, 'Il wallet e prepagato e il controllo alla conferma dell\'ordine e '
              'netto: saldo piu fido devono coprire il totale, altrimenti '
              'l\'ordine viene rifiutato. Se invece il portafoglio e '
              'disattivato nelle Impostazioni, questa fase si salta per '
              'intero: il cliente ordina subito e paga alla cassa al ritiro.',
         color=GRAY)
    rule(doc)

    step(doc, 1, 'Ricarica dalla scheda cliente', [
        [('Clienti › ', False), ('Ricarica', True),
         (': inserisci l\'importo consegnato in cassa. Il saldo si aggiorna subito '
          'e il movimento resta nello storico del cliente.', False)],
        [('Indicazione pratica: ', False), ('10–20 €', True),
         (' coprono diversi pranzi ed evitano che il cliente torni in cassa il '
          'giorno dopo.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'Oppure concedi un fido', [
        [('Dalla scheda del cliente puoi impostare un ', False), ('fido', True),
         (': la soglia di rosso consentita. Utile per il personale interno o per '
          'chi salda a fine mese, evita di bloccare l\'ordine per pochi '
          'centesimi.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'Telegram: si collega col proprio codice', [
        [('Entrambe le email contengono il pulsante ', False),
         ('Collega Telegram', True),
         (': porta a una pagina che mostra al cliente un codice personale '
          'e il pulsante per aprire il bot. Il cliente invia il codice in '
          'chat, torna sulla pagina e conferma.', False)],
        [('Non deve cercare nessun numero di identificazione: il codice '
          'basta, e la stessa pagina si raggiunge dal ', False),
         ('Profilo', True), (' del cliente.', False)],
        [('Da collegato riceve gli avvisi sul telefono e puo confermare o '
          'disdire il pasto direttamente dal promemoria.', False)],
    ], accent=HEX_GREEN)

    step(doc, 3, 'Spiegagli come comparira il suo nome', [
        [('Nelle liste, sui tagliandi e nei report il cliente e indicato con '
          'il nome per esteso e il cognome puntato: ', False),
         ('Mario R.', True),
         ('. Nella sua area personale, invece, si legge per intero.', False)],
    ], accent=HEX_GREEN)

    step(doc, 3, 'Verifica che il cliente veda il saldo', [
        [('Il cliente trova il saldo nella sua pagina Wallet, con lo storico dei '
          'movimenti e i punti fedelta accumulati. Se il saldo e a zero, il '
          'primo ordine non partira.', False)],
    ], accent=HEX_GREEN)

    _page_break(doc)

    # ══ FASE 5 · PRIMO ORDINE ═════════════════════════════════════════════
    heading(doc, 'Fase 5 — Il primo ordine', 1, color=RED, before=0)
    body(doc, 'Da qui il cliente e autonomo. Questi sono i passi che compie, utili '
              'per guidarlo la prima volta stando accanto a lui.', color=GRAY)
    rule(doc)

    step(doc, 1, 'Accede', [
        [('Con email e password, oppure con Google. Atterra sulla propria '
          'dashboard: saldo, ordini di oggi e scorciatoie.', False)],
    ], accent=HEX_RED)

    step(doc, 2, 'Scegle cosa mangiare', [
        [('Menu', True),
         (' per i prodotti pronti; oppure il ', False), ('builder', True),
         (' per comporre panino, insalata o poke ingrediente per ingrediente, con '
          'gli extra a prezzo variabile. Per il panino puo chiedere la ', False),
         ('piastra', True), ('.', False)],
        [('C\'e anche il ', False), ('cesto', True),
         (': i pezzi già preparati, che si acquistano inquadrando il QR '
          'dell\'etichetta senza passare dal carrello.', False)],
    ], accent=HEX_RED)

    step(doc, 3, 'Conferma dal carrello', [
        [('Nel carrello sceglie l\'', False), ('orario di ritiro', True),
         (' tra gli slot disponibili, oppure ', False),
         ('subito al banco', True),
         (' per il consumo immediato, e aggiunge eventuali note.', False)],
    ], accent=HEX_RED)

    step(doc, 4, 'L\'ordine viene pagato', [
        [('Alla conferma il totale viene ', False),
         ('scalato dal wallet', True),
         (' e vengono accreditati i punti fedelta. L\'ordine riceve un codice: '
          'QL- con data e orario di ritiro, oppure BANCO- se e per subito.',
          False)],
        [('Se il saldo non basta, l\'ordine non viene creato e il cliente viene '
          'rimandato alla pagina del wallet.', False)],
    ], accent=HEX_RED)

    step(doc, 5, 'La cucina lo prende in carico', [
        [('L\'ordine compare sul display cucina nella colonna "da preparare". Il '
          'cliente riceve un avviso quando passa ', False),
         ('in preparazione', True), (' e un altro quando e ', False),
         ('pronto', True), ('.', False)],
    ], accent=HEX_RED)

    step(doc, 6, 'Ritira al banco', [
        [('Il cliente si presenta col proprio nome o codice ordine. Sul prodotto '
          'trova lo scontrino con codice, orario e composizione. Lo staff chiude '
          'con ', False), ('Consegnato', True), ('.', False)],
        [('Primo ordine completato: da qui in avanti il cliente non ha piu '
          'bisogno di voi.', False)],
    ], accent=HEX_RED)

    # ══ FASE 6 · RIFINITURE ═══════════════════════════════════════════════
    heading(doc, 'Fase 6 — Rifiniture dopo il primo ordine', 1, color=SLATE)
    body(doc, 'Due minuti spesi qui riducono le domande al banco nei giorni '
              'successivi.', color=GRAY)
    rule(doc)

    step(doc, 1, 'Attiva le notifiche sul telefono', [
        [('Il cliente tocca l\'icona della ', False), ('campanella', True),
         (' in alto e concede il permesso: da quel momento riceve la notifica '
          '"ordine pronto" anche con l\'applicazione chiusa.', False)],
    ], accent=HEX_SLATE)

    step(doc, 2, 'Collega Telegram, se vuole', [
        [('Dal proprio ', False), ('Profilo', True),
         (' il cliente apre Collega Telegram, invia al bot il codice che '
          'la pagina gli mostra e riceve lì conferme d\'ordine, promemoria '
          'di ritiro e avvisi di ricarica.', False)],
        [('Non e obbligatorio: chi non collega Telegram riceve i promemoria di '
          'ritiro ', False), ('per email', True),
         ('. Telegram resta piu immediato, perche arriva come messaggio.',
          False)],
    ], accent=HEX_SLATE)

    step(doc, 3, 'Mostragli la guida', [
        [('La voce ', False), ('Guida', True),
         (' nel menu del cliente spiega ordini, wallet, punti, sondaggi e pasto '
          'aziendale. E il posto dove mandarlo invece di rispondere due volte '
          'alla stessa domanda.', False)],
    ], accent=HEX_SLATE)

    _page_break(doc)

    # ══ Una cosa da sapere ════════════════════════════════════════════════
    heading(doc, 'Una cosa da sapere', 1, before=0)
    body(doc, 'Un comportamento verificato dell\'applicazione che incide su questo '
              'percorso. Vale la pena conoscerlo prima di trovarselo davanti al '
              'banco.', color=GRAY)
    spacer(doc, 6)

    box(doc, 'I bottoni Telegram rispondono solo se attivate le risposte', [
        [('Il promemoria del pasto porta i bottoni Sì e No, ma la risposta '
          'del cliente torna all\'applicazione solo dopo aver premuto una '
          'volta ', False),
         ('Impostazioni › Notifiche › Attiva le risposte', True),
         ('. Senza quel passaggio i bottoni compaiono e non fanno nulla: il '
          'cliente crede di aver disdetto e la cucina prepara comunque.',
          False)],
        [('Dalla stessa pagina il pulsante ', False),
         ('Invia una domanda di prova', True),
         (' manda una domanda con i due bottoni e ne legge la risposta: e '
          'il modo di accertarsi che il canale funzioni in entrambe le '
          'direzioni.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    spacer(doc, 10)

    # ══ Riepilogo ruoli ═══════════════════════════════════════════════════
    heading(doc, 'Chi fa cosa, in ordine', 1)
    spacer(doc, 4)
    table_grid(doc,
               ['Momento', 'Cliente', 'Staff'],
               [
                   ['Prima di tutto', '—',
                    'Riavvio, prodotti, slot, dati attivita'],
                   ['Invito', 'Inquadra il QR', 'Stampa e affigge il QR'],
                   ['Registrazione', 'Compila o entra con Google',
                    'Riceve l\'avviso su Telegram'],
                   ['Attivazione', 'Riceve l\'email di conferma',
                    'Attiva dalla lista clienti'],
                   ['Credito', 'Consegna il contante',
                    'Ricarica il wallet o concede il fido'],
                   ['Ordine', 'Ordina e paga dal wallet',
                    'Prepara, stampa lo scontrino, consegna'],
                   ['Dopo', 'Attiva notifiche e Telegram',
                    'Rimanda alla Guida'],
               ],
               widths=[3.8, 6.2, 6.5])

    spacer(doc, 12)
    rule(doc, color=HEX_NAVY)
    body(doc, 'Assistenza:  Daniele Speziale — DS Consulting', size=11,
         color=DARK, bold=True, after=1)
    body(doc, 'dspeziale@gmail.com   ·   +39 352 0150489', size=11, color=RED, bold=True, after=8)
    body(doc, 'QuickLunch · Onboarding cliente · Le voci di menu, i percorsi e i '
              'valori predefiniti citati corrispondono all\'applicazione in uso.',
         size=8.5, color=GRAY, after=2)
    body(doc, '© 2024–26 DS Consulting', size=8.5, color=GRAY)

    return doc


def main():
    doc = build()
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] Documento salvato in: {OUT}')


if __name__ == '__main__':
    main()
