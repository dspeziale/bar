#!/usr/bin/env python3
"""Genera docs/manuali/manuale_gestore_giornata.docx.

La giornata del gestore con orari limite: dal pomeriggio precedente fino al
primo ordine servito. Gli orari sono tarati su un servizio di ritiro
11:45-13:30 (gli slot predefiniti dell'applicazione), sul banco che chiude
al pubblico alle 17:00 e sulla giornata di lavoro che finisce alle 17:30:
dopo quell'ora non e' prevista alcuna attivita'. Se cambiano gli orari,
spostare ORA_PRIMO_SLOT, ORA_FINE_SERVIZIO, ORA_CHIUSURA_BANCO e
ORA_FINE_LAVORO e ritarare la scaletta.

    python docs/generate_manuale_gestore.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Palette (coerente con gli altri documenti in docs/) ──────────────────────
RED    = RGBColor(0xe9, 0x45, 0x60)
NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x16, 0x21, 0x3e)
DGRAY  = RGBColor(0x34, 0x3a, 0x40)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREEN  = RGBColor(0x27, 0xae, 0x60)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
BLUE   = RGBColor(0x34, 0x98, 0xdb)
PURPLE = RGBColor(0x8e, 0x44, 0xad)
SLATE  = RGBColor(0x7f, 0x8c, 0x8d)

HEX_RED    = 'E94560'
HEX_NAVY   = '0F3460'
HEX_GREEN  = '27AE60'
HEX_ORANGE = 'E67E22'
HEX_BLUE   = '3498DB'
HEX_PURPLE = '8E44AD'
HEX_SLATE  = '7F8C8D'
HEX_LIGHT  = 'F8F9FA'
HEX_RULE   = 'E8EBEE'
HEX_WARN   = 'FDF6E7'
HEX_STOP   = 'FDEEF0'

FONT = 'PT Sans Narrow'
BODY_FONT = FONT           # tutto il documento e' in PT Sans Narrow

EMAIL = 'dspeziale@gmail.com'
CELL = '+39 352 0150489'
CONTATTI = 'DS Consulting  ·  Daniele Speziale  ·  %s  ·  %s' % (EMAIL, CELL)

OUT = os.path.join(os.path.dirname(__file__), 'manuali',
                   'manuale_gestore_giornata.docx')


# ── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None, sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        if color:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color.lstrip('#'))
        else:
            el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), val)
    tcPr.append(el)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _p_spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def _run_font(run, size=10.5, bold=False, italic=False, color=None, font=BODY_FONT):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    _p_spacing(p, before=0, after=0)


# ── Fogli di stile ───────────────────────────────────────────────────────────

def imposta_stili(doc):
    """Definisce gli stili del documento, cosi' che in Word si possa ritoccare
    l'aspetto da un punto solo invece che paragrafo per paragrafo."""
    from docx.enum.style import WD_STYLE_TYPE

    # Normale: e' la radice da cui tutto eredita
    normale = doc.styles['Normal']
    normale.font.name = FONT
    normale.font.size = Pt(10.5)
    normale.font.color.rgb = DGRAY
    normale.paragraph_format.space_after = Pt(6)
    normale.paragraph_format.line_spacing = 1.08
    # Anche per i caratteri non latini, altrimenti Word ricade su Calibri
    rpr = normale.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), FONT)

    # Titoli: gerarchia esplicita, tutti in PT Sans Narrow
    for nome, dim, colore, prima, dopo in [
        ('Heading 1', 19, NAVY, 16, 5),
        ('Heading 2', 14.5, NAVY, 14, 4),
        ('Heading 3', 12, DARK, 12, 3),
    ]:
        try:
            st = doc.styles[nome]
        except KeyError:
            st = doc.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st.font.size = Pt(dim)
        st.font.bold = True
        st.font.color.rgb = colore
        st.paragraph_format.space_before = Pt(prima)
        st.paragraph_format.space_after = Pt(dopo)
        st.paragraph_format.keep_with_next = True

    # Stile dedicato alle didascalie e alle note
    try:
        nota = doc.styles['Nota QuickLunch']
    except KeyError:
        nota = doc.styles.add_style('Nota QuickLunch', WD_STYLE_TYPE.PARAGRAPH)
    nota.base_style = doc.styles['Normal']
    nota.font.name = FONT
    nota.font.size = Pt(9.5)
    nota.font.color.rgb = GRAY


# ── Blocchi di contenuto ─────────────────────────────────────────────────────

def heading(doc, text, level=1, color=None, before=None):
    """Titolo che usa lo stile Titolo N: il colore passato lo sovrascrive solo
    quando serve distinguere una sezione."""
    p = doc.add_paragraph(style='Heading %d' % min(max(level, 1), 3))
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    if color is not None:
        r.font.color.rgb = color
    return p


def body(doc, text, size=10.5, color=DGRAY, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    _run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def rich(doc, parts, size=10.5, after=6, indent=None):
    """Paragrafo con porzioni in grassetto: parts = [(testo, bold), ...]."""
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def bullet(doc, parts, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    _p_spacing(p, before=0, after=3)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def rule(doc, color=HEX_RULE, after=8):
    """Filetto orizzontale come tabella a una cella."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_border(cell, bottom=color, sz='8')
    _cell_margins(cell, top=0, bottom=0, left=0, right=0)
    _p_spacing(cell.paragraphs[0], before=0, after=0)
    _run_font(cell.paragraphs[0].add_run(''), size=1)
    spacer(doc, after)


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(''), size=pts / 2)


def box(doc, label, paragraphs, accent=HEX_RED, fill=HEX_LIGHT, label_color=RED):
    """Riquadro con bordo sinistro colorato."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, fill)
    _cell_border(cell, top=HEX_RULE, bottom=HEX_RULE, right=HEX_RULE,
                 left=accent, sz='18')
    _cell_margins(cell, top=90, bottom=90, left=140, right=110)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=3)
    _run_font(p.add_run(label.upper()), size=8.5, bold=True, color=label_color,
              font=FONT)

    for i, parts in enumerate(paragraphs):
        pp = cell.add_paragraph()
        _p_spacing(pp, before=0, after=0 if i == len(paragraphs) - 1 else 4)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 10)
    return tbl


def step(doc, num, title, paragraphs, accent=HEX_RED):
    """Passo numerato: numero in pastiglia colorata + contenuto."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_width(tbl, 0, 1.0)
    _set_col_width(tbl, 1, 15.5)

    nc = tbl.rows[0].cells[0]
    _cell_shd(nc, accent)
    _cell_margins(nc, top=60, bottom=60, left=30, right=30)
    _cell_valign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run('%02d' % num), size=12, bold=True, color=WHITE, font=FONT)

    cc = tbl.rows[0].cells[1]
    _cell_margins(cc, top=50, bottom=50, left=150, right=60)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=2)
    _run_font(pt.add_run(title), size=11.5, bold=True, color=DARK, font=FONT)

    for parts in paragraphs:
        pp = cc.add_paragraph()
        _p_spacing(pp, before=0, after=3)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 6)
    return tbl


def table_grid(doc, headers, rows, widths, head_fill=HEX_NAVY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, w in enumerate(widths):
        _set_col_width(tbl, i, w)

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _cell_shd(c, head_fill)
        _cell_margins(c, top=60, bottom=60, left=110, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(h.upper()), size=8.5, bold=True, color=WHITE, font=FONT)

    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            c = tbl.rows[r].cells[i]
            _cell_shd(c, HEX_LIGHT if r % 2 else 'FFFFFF')
            _cell_border(c, bottom=HEX_RULE, sz='4')
            _cell_margins(c, top=55, bottom=55, left=110, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            text, color, bold = val if isinstance(val, tuple) else (val, DGRAY, False)
            _run_font(p.add_run(text), size=10, bold=bold, color=color)
    spacer(doc, 10)
    return tbl


# ── Riferimenti temporali del servizio ───────────────────────────────────────
ORA_PRIMO_SLOT = '11:45'       # primo slot di ritiro
ORA_FINE_SERVIZIO = '13:30'    # ultimo slot di ritiro
ORA_CHIUSURA_BANCO = '17:00'   # il banco chiude al pubblico
ORA_FINE_LAVORO = '17:30'      # si finisce di lavorare: dopo, niente


# ── Blocco con ora limite ────────────────────────────────────────────────────

def fascia(doc, ora, titolo, paragrafi, motivo=None, accent=HEX_RED,
           critico=False):
    """Riga di procedura: ora limite in evidenza + attivita' + motivo dell'ora."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_width(tbl, 0, 2.0)
    _set_col_width(tbl, 1, 14.5)

    nc = tbl.rows[0].cells[0]
    _cell_shd(nc, accent)
    _cell_margins(nc, top=70, bottom=70, left=40, right=40)
    _cell_valign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run(ora), size=15, bold=True, color=WHITE, font=FONT)
    if critico:
        pc = nc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(pc, before=0, after=0)
        _run_font(pc.add_run('NON OLTRE'), size=6.5, bold=True, color=WHITE,
                  font=FONT)

    cc = tbl.rows[0].cells[1]
    _cell_margins(cc, top=55, bottom=55, left=160, right=60)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=2)
    _run_font(pt.add_run(titolo), size=11.5, bold=True, color=DARK, font=FONT)

    for parts in paragrafi:
        pp = cc.add_paragraph()
        _p_spacing(pp, before=0, after=3)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)

    if motivo:
        pm = cc.add_paragraph()
        _p_spacing(pm, before=2, after=0)
        _run_font(pm.add_run('Perche quest\'ora:  '), size=8.5, bold=True,
                  color=RED, font=FONT)
        _run_font(pm.add_run(motivo), size=9, italic=True, color=GRAY)

    spacer(doc, 7)
    return tbl


# ── Documento ────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    imposta_stili(doc)

    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)

    # ══ Copertina ═════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, HEX_NAVY)
    _cell_margins(cell, top=300, bottom=300, left=320, right=320)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=4)
    _run_font(p.add_run('QUICKLUNCH  ·  MANUALE DEL GESTORE'), size=9.5,
              bold=True, color=RGBColor(0xb2, 0xc2, 0xd9), font=FONT)

    p2 = cell.add_paragraph()
    _p_spacing(p2, before=0, after=6)
    _run_font(p2.add_run('La giornata, ora per ora'), size=30, bold=True,
              color=WHITE, font=FONT)

    p3 = cell.add_paragraph()
    _p_spacing(p3, before=0, after=10)
    _run_font(p3.add_run(
        'Dal pomeriggio precedente al primo ordine servito. Ogni attivita ha '
        'un\'ora limite e la ragione per cui e fissata a quell\'ora: e la '
        'sequenza che rende il servizio del giorno dopo prevedibile. La '
        'giornata di lavoro finisce alle %s.' % ORA_FINE_LAVORO), size=11,
        color=RGBColor(0xd6, 0xdf, 0xea))

    p4 = cell.add_paragraph()
    _p_spacing(p4, before=0, after=0)
    _run_font(p4.add_run('Assistenza:  '), size=10.5, bold=True, color=WHITE,
              font=FONT)
    _run_font(p4.add_run(CONTATTI), size=10.5, bold=True,
              color=RGBColor(0xff, 0xd7, 0xdf), font=FONT)
    spacer(doc, 16)

    # ══ Come si legge ═════════════════════════════════════════════════════
    heading(doc, 'Come si legge questo manuale', 1)
    rich(doc, [
        ('Gli orari indicati sono ', False), ('limiti, non appuntamenti', True),
        (': "15:45" significa "entro le 15:45, prima e meglio". Sono tarati su '
         'un ritiro dei pasti dalle ', False), (ORA_PRIMO_SLOT, True),
        (' alle ', False), (ORA_FINE_SERVIZIO, True),
        (', su un banco che chiude al pubblico alle ', False),
        (ORA_CHIUSURA_BANCO, True),
        (' e su una giornata di lavoro che finisce alle ', False),
        (ORA_FINE_LAVORO, True),
        (': dopo quell\'ora la scaletta non prevede nulla. Se cambiate gli '
         'slot o gli orari di chiusura, spostate di conseguenza tutta la '
         'scaletta.', False),
    ])
    rich(doc, [
        ('Le voci contrassegnate ', False), ('NON OLTRE', True),
        (' sono quelle che, se salti, si vedono: il servizio parte in ritardo o '
         'parte incompleto. Le altre si possono recuperare.', False),
    ])
    spacer(doc, 4)

    box(doc, 'Portafoglio prepagato: attivo o spento', [
        [('Il portafoglio e una funzione attivabile (Impostazioni › '
          'Funzionalita). Se lavorate ', False),
         ('senza portafoglio', True),
         (', i clienti pagano alla cassa: saldi, ricariche, fidi, punti e '
          'bonus non esistono, ma ordini, banco e cesto registrano comunque '
          'le vendite nei report.', False)],
        [('Le attivita di questo manuale contrassegnate con ', False),
         ('"solo con portafoglio attivo"', True),
         (' in quel caso si saltano: la scaletta resta valida per tutto il '
          'resto.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)
    spacer(doc, 2)

    box(doc, 'Le tre parti della giornata', [
        [('Parte 1 — Pomeriggio precedente', True),
         (': si decide cosa si vendera domani. E la parte che determina tutto il '
          'resto, ed e anche quella che viene salta più spesso.', False)],
        [('Parte 2 — Mattina fino all\'apertura', True),
         (': si mette in condizione la macchina di funzionare.', False)],
        [('Parte 3 — Il primo ordine', True),
         (': la verifica che tutto sia davvero a posto.', False)],
        [('Parte 4 — Oltre la giornata', True),
         (': quello che non si fa ogni giorno ma va fatto — fine mese, backup, '
          'dati di prova.', False)],
    ], accent=HEX_NAVY, fill=HEX_LIGHT, label_color=NAVY)

    _page_break(doc)

    # ══ Quadro d'insieme ══════════════════════════════════════════════════
    heading(doc, 'Quadro d\'insieme', 1, before=0)
    body(doc, 'La pagina da tenere davanti. Il dettaglio di ogni riga e nelle '
              'pagine seguenti.', color=GRAY)
    spacer(doc, 4)

    heading(doc, 'Pomeriggio precedente', 3, color=PURPLE)
    table_grid(doc,
               ['Entro', 'Attivita', 'Dove'],
               [
                   [('14:45', DARK, True), 'Fabbisogno dei giorni successivi',
                    'Cucina › Prenotazioni'],
                   [('15:15', DARK, True), 'Giacenze e ordini ai fornitori',
                    'Magazzino'],
                   [('15:45', RED, True), 'Pasto aziendale di domani pubblicato',
                    'Convenzioni › Pasto del giorno'],
                   [('16:00', DARK, True), 'Sondaggio del menu, se usato',
                    'Sondaggi'],
                   [('16:15', RED, True), 'Clienti in attesa attivati',
                    'Clienti'],
                   [('16:30', DARK, True), 'Ricariche e fidi (solo con '
                    'portafoglio attivo)', 'Clienti › Ricarica'],
                   [('16:45', DARK, True), 'Listino e quantita di domani',
                    'Prodotti'],
                   [(ORA_CHIUSURA_BANCO, RED, True), 'Chiusura del banco e '
                    'ritiro dell\'invenduto', 'Cesto Cucina'],
                   [('17:15', DARK, True), 'Lettura dei numeri del giorno',
                    'Report'],
                   [(ORA_FINE_LAVORO, RED, True), 'Postazioni spente: fine '
                    'della giornata', '—'],
               ],
               widths=[2.2, 8.6, 5.2], head_fill=HEX_PURPLE)

    heading(doc, 'Mattina', 3, color=BLUE)
    table_grid(doc,
               ['Entro', 'Attivita', 'Dove'],
               [
                   [('07:30', DARK, True), 'Postazioni accese e collegate', '—'],
                   [('07:45', DARK, True), 'Stampanti verificate con prova',
                    '—'],
                   [('08:00', DARK, True), 'Pasto del giorno visibile ai clienti',
                    'Convenzioni › Pasto del giorno'],
                   [('08:15', DARK, True), 'Primo sguardo agli ordini',
                    'Cucina'],
                   [('09:00', DARK, True), 'Disponibilita del giorno corrette',
                    'Stock'],
                   [('09:30', DARK, True), 'Pezzi del cesto preparati', '—'],
                   [('10:15', RED, True), 'Etichette del cesto generate',
                    'Cucina › Cesto Cucina'],
                   [('10:45', DARK, True), 'Cesto esposto ed etichettato', '—'],
                   [('11:00', RED, True), 'Lista pasti aziendali stampata',
                    'Convenzioni › Report'],
                   [('11:15', DARK, True), 'Display presidiato, audio attivo',
                    'Cucina'],
                   [('11:30', DARK, True), 'Ultimo controllo prima del servizio',
                    'Cucina'],
                   [(ORA_PRIMO_SLOT, RED, True), 'Primo slot di ritiro',
                    '—'],
               ],
               widths=[2.2, 8.6, 5.2], head_fill=HEX_BLUE)

    _page_break(doc)

    # ══ PARTE 1 ═══════════════════════════════════════════════════════════
    heading(doc, 'Parte 1 — Il pomeriggio precedente', 1, color=PURPLE, before=0)
    body(doc, 'Il banco resta aperto al pubblico fino alle %s: le attivita di '
              'questa parte si svolgono quasi tutte a servizio ancora in corso, '
              'negli spazi fra un cliente e l\'altro. L\'ultima mezz\'ora — '
              'dalle %s alle %s — e a locale chiuso: ritiro dell\'invenduto, '
              'lettura dei numeri, postazioni. Alle %s si finisce di lavorare: '
              'la scaletta e costruita perche a quell\'ora non resti nulla da '
              'fare. Qui si decide cosa il cliente potra ordinare domani, e si '
              'sbloccano i clienti che altrimenti domani non potrebbero '
              'ordinare affatto.'
              % (ORA_CHIUSURA_BANCO, ORA_CHIUSURA_BANCO, ORA_FINE_LAVORO,
                 ORA_FINE_LAVORO), color=GRAY)
    rule(doc)

    fascia(doc, '14:45', 'Calcola il fabbisogno dei giorni successivi', [
        [('Cucina › Prenotazioni', True),
         (': la pagina elenca le prenotazioni dei clienti per i giorni a venire, '
          'raggruppate per data, e per ogni prodotto indica ', False),
         ('quanti pezzi servono e quanti sono già nel cesto', True), ('.', False)],
        [('E la lista della spesa di domani: quello che manca va preparato o '
          'acquistato.', False)],
    ], motivo='serve prima di contattare i fornitori, non dopo.',
        accent=HEX_PURPLE)

    fascia(doc, '15:15', 'Controlla le giacenze e ordina', [
        [('Magazzino', True),
         (': i materiali sotto soglia sono evidenziati e il sistema puo inviare '
          'la richiesta di riordino via email al fornitore.', False)],
        [('Verifica anche le scorte degli ingredienti del builder: un '
          'ingrediente esaurito significa un panino che il cliente componeva e '
          'domani non potra piu comporre.', False)],
    ], motivo='i fornitori chiudono nel pomeriggio; oltre le 16 la consegna '
              'slitta di un giorno.',
        accent=HEX_PURPLE)

    fascia(doc, '15:45', 'Pubblica il pasto aziendale di domani', [
        [('Convenzioni › scegli l\'azienda › Pasto del giorno', True),
         (', con data di domani. Compila primo, secondo, contorno, bevanda, '
          'caffe, gli allergeni, il prezzo e il numero massimo di prenotazioni.',
          False)],
        [('Se il menu si ripete, richiamalo da ', False),
         ('Configurazioni', True),
         (': sono modelli salvati che compilano il modulo in un colpo.', False)],
    ], motivo='i dipendenti prenotano nel tardo pomeriggio e la sera. Ogni ora '
              'di ritardo nella pubblicazione sono prenotazioni che non '
              'arrivano, e coperti che non vendi.',
        accent=HEX_PURPLE, critico=True)

    fascia(doc, '16:00', 'Lancia il sondaggio del menu, se lo usi', [
        [('Sondaggi', True),
         (': crea il sondaggio con la data di domani e le opzioni, poi invialo '
          'su Telegram o per email. I clienti votano e tu sai cosa preparare.',
          False)],
    ], motivo='va inviato insieme al menu aziendale, in un\'unica '
              'comunicazione serale.',
        accent=HEX_PURPLE)

    fascia(doc, '16:15', 'Attiva i clienti in attesa', [
        [('Clienti', True),
         (': chi si e registrato oggi ha un account ', False),
         ('non attivo', True),
         (' e non puo accedere. Il contatore dei clienti da attivare e sulla '
          'dashboard.', False)],
        [('Premendo Attiva il cliente riceve un\'email che gli dice che puo '
          'entrare. Se appartiene a un\'azienda convenzionata, associalo alla '
          'convenzione in questo stesso momento.', False)],
    ], motivo='un cliente attivato stasera puo prenotare il pasto e ordinare '
              'domani mattina. Attivato domani a mezzogiorno, no.',
        accent=HEX_PURPLE, critico=True)

    fascia(doc, '16:30', 'Sistema ricariche e fidi '
           '(solo con portafoglio attivo)', [
        [('Il portafoglio e prepagato: senza credito l\'ordine viene rifiutato. '
          'Ricarica chi ha lasciato contante in cassa e verifica i saldi vicini '
          'allo zero.', False)],
        [('Per il personale interno o per chi salda a fine mese, imposta un ',
          False), ('fido', True),
         (' dalla scheda del cliente: consente di ordinare andando in rosso fino '
          'alla soglia scelta.', False)],
        [('Se il portafoglio prepagato e disattivato nelle Impostazioni, '
          'questo passo non esiste: i clienti pagano alla cassa e si passa '
          'direttamente all\'attivita successiva.', False)],
    ], motivo='il cliente a saldo zero se ne accorge davanti al carrello, '
              'nell\'ora di punta, e in quel momento non hai tempo.',
        accent=HEX_PURPLE)

    fascia(doc, '16:45', 'Aggiorna il listino e le quantita di domani', [
        [('Prodotti', True),
         (': correggi la quantita giornaliera di quello che produrrai domani, e '
          'disattiva i prodotti che non farai. Un prodotto disattivato sparisce '
          'dal menu; uno con quantita a zero resta visibile ma non ordinabile.',
          False)],
        [('Ritocca i prezzi ora, negli ultimi minuti di servizio: il prezzo '
          'viene fotografato sull\'ordine al momento della conferma, quindi '
          'cambiarlo a carrelli aperti crea discussioni al banco.', False)],
    ], motivo='il menu di domani deve essere definitivo prima che i primi '
              'clienti lo guardino, la sera.',
        accent=HEX_PURPLE)

    fascia(doc, ORA_CHIUSURA_BANCO, 'Chiudi il banco e ritira l\'invenduto', [
        [('Fine del servizio al pubblico. Ritira dal cesto i pezzi rimasti e '
          'annulla le loro etichette una per una da ', False),
         ('Cucina › Cesto Cucina', True),
         (': un\'etichetta lasciata attiva resta acquistabile fino alle 24 ore '
          'dalla generazione.', False)],
        [('Non usare "Annulla tutto": cancella dal registro anche le vendite '
          'della giornata.', False)],
    ], motivo='chiudere alle %s lascia la mezz\'ora che serve per invenduto, '
              'numeri e postazioni entro le %s, senza sforare.'
              % (ORA_CHIUSURA_BANCO, ORA_FINE_LAVORO),
        accent=HEX_PURPLE, critico=True)

    fascia(doc, '17:15', 'Leggi i numeri della giornata', [
        [('Report', True),
         (': incasso, ordini per prodotto, andamento. Ora che il servizio e '
          'chiuso i dati sono definitivi.', False)],
        [('Confronta l\'invenduto ritirato dal cesto con i pezzi generati la '
          'mattina: e la correzione piu utile per le quantita di domani.',
          False)],
    ], motivo='subito dopo la chiusura, con i numeri fermi e la giornata '
              'ancora in mente.',
        accent=HEX_PURPLE)

    fascia(doc, ORA_FINE_LAVORO, 'Chiudi le postazioni: fine della giornata', [
        [('Tablet e display in carica, stampanti spente, carta rifornita per '
          'domani. Alle ', False), (ORA_FINE_LAVORO, True),
         (' si esce: quello che non e stato fatto si recupera domani mattina, '
          'non stasera.', False)],
    ], motivo='e l\'ora in cui si finisce di lavorare; una batteria lasciata '
              'scarica costa mezz\'ora alle 7:30 di domani.',
        accent=HEX_PURPLE, critico=True)

    _page_break(doc)

    # ══ PARTE 2 ═══════════════════════════════════════════════════════════
    heading(doc, 'Parte 2 — La mattina, fino all\'apertura', 1, color=BLUE,
            before=0)
    body(doc, 'Quattro ore per passare da locale chiuso a servizio pronto. '
              'L\'ordine delle attivita non e casuale: ogni passo dipende dal '
              'precedente.', color=GRAY)
    rule(doc)

    fascia(doc, '07:30', 'Accendi e verifica le postazioni', [
        [('Le tre postazioni sono il display in cucina, il tablet al banco e il '
          'PC del backoffice. Su ciascuna: accesso effettuato, pagina giusta '
          'aperta, rete funzionante.', False)],
        [('Sul display cucina apri ', False), ('Cucina', True),
         (' a schermo pieno e ', False), ('lascia l\'audio attivo', True),
         (': l\'avviso dei nuovi ordini e sonoro.', False)],
    ], motivo='se qualcosa non si accende, a quest\'ora hai ancora quattro ore '
              'per rimediare.',
        accent=HEX_BLUE)

    fascia(doc, '07:45', 'Prova le stampanti', [
        [('Carta termica nella stampante dei tagliandi, fogli A4 in quella delle '
          'etichette. Fai una ', False), ('prova di stampa reale', True),
         (', non solo un\'occhiata alla spia.', False)],
        [('Verifica anche il registratore di cassa: e lui che emette lo scontrino '
          'da allegare al prodotto.', False)],
    ], motivo='la carta finisce sempre al primo ordine, non al ventesimo.',
        accent=HEX_BLUE)

    fascia(doc, '08:00', 'Verifica che il pasto del giorno sia visibile', [
        [('Apri l\'applicazione come la vede un dipendente convenzionato e '
          'controlla che il pasto di oggi ci sia, con il prezzo giusto e i posti '
          'disponibili.', False)],
        [('Se ieri non l\'hai pubblicato, pubblicalo adesso: e recuperabile, ma '
          'avrai meno prenotazioni.', False)],
    ], motivo='i dipendenti prenotano appena arrivati in ufficio, fra le 8 e le '
              '9.',
        accent=HEX_BLUE)

    fascia(doc, '08:15', 'Guarda il display cucina per la prima volta', [
        [('Gli ordini ', False), ('possono essere già arrivati', True),
         (': i clienti ordinano anche la sera prima o a colazione, scegliendo '
          'uno slot di ritiro di mezzogiorno.', False)],
        [('Non prepararli ora: servono freschi. Ma guardali, perche ti dicono '
          'cosa produrre per primo.', False)],
    ], motivo='conoscere il carico entro le 8:30 permette di ridistribuire il '
              'lavoro della mattina.',
        accent=HEX_BLUE)

    fascia(doc, '09:00', 'Correggi le disponibilita del giorno', [
        [('Stock', True),
         (': la pagina mostra la disponibilita di oggi per ogni prodotto, '
          'inizializzata con la quantita giornaliera del listino. Correggila con '
          'quello che hai davvero prodotto.', False)],
        [('E l\'unico numero che impedisce di vendere piu pezzi di quelli che '
          'esistono.', False)],
    ], motivo='va fatto dopo la produzione del mattino e prima che il grosso '
              'degli ordini arrivi.',
        accent=HEX_BLUE)

    fascia(doc, '09:30', 'Prepara i pezzi del cesto', [
        [('Panini e tramezzini destinati alla vendita a libero servizio, in lotti '
          'omogenei: un lotto per tipo di prodotto.', False)],
    ], motivo='vanno fatti prima di generare le etichette, che vanno applicate '
              'a prodotto finito.',
        accent=HEX_BLUE)

    fascia(doc, '10:15', 'Genera e stampa le etichette del cesto', [
        [('Cucina › Cesto Cucina', True),
         (': scegli il prodotto, indica la quantita, premi Genera. La pagina di '
          'stampa si apre da sola. Stampa, taglia, applica un\'etichetta per '
          'pezzo con il QR ben visibile.', False)],
        [('Le etichette valgono ', False), ('24 ore dalla generazione', True),
         (': generate stamattina coprono tutto il servizio di oggi. Generate '
          'ieri mattina, scadrebbero durante il servizio.', False)],
    ], motivo='e il limite tecnico delle etichette: entro le 10:15 hai anche il '
              'margine per ristampare se il QR non si legge.',
        accent=HEX_BLUE, critico=True)

    fascia(doc, '10:45', 'Esponi il cesto', [
        [('Pezzi etichettati nel cesto, QR verso l\'alto e non deformato dalla '
          'piega. Da questo momento la vendita e automatica: il cliente inquadra '
          'e paga dal telefono.', False)],
    ], motivo='il cesto vende anche prima di mezzogiorno, a chi passa per il '
              'caffe.',
        accent=HEX_BLUE)

    fascia(doc, '11:00', 'Stampa la lista dei pasti aziendali', [
        [('Convenzioni › Report', True),
         (', sulla data di oggi: elenco nominativo per azienda, ordinato per '
          'cognome, con le porzioni per persona. E la lista di produzione e la '
          'base per etichettare le porzioni.', False)],
        [('I nomi sono stampati come ', False), ('"Mario R."', True),
         (' — nome per esteso e cognome puntato: la lista resta appendibile '
          'in cucina senza esporre l\'anagrafica dei clienti. Vale per tutte '
          'le stampe e per il display.', False)],
        [('Il conteggio e ', False), ('definitivo solo 30 minuti prima', True),
         (' di ogni orario di ritiro: fino a quel momento un dipendente puo '
          'ancora disdire. Stampa alle 11:00 e tieni conto che qualche numero '
          'puo scendere.', False)],
    ], motivo='alle 11:00 il grosso delle prenotazioni e chiuso e la cucina ha '
              'ancora 45 minuti per impiattare.',
        accent=HEX_BLUE, critico=True)

    fascia(doc, '11:15', 'Presidia il display e prepara il banco', [
        [('Dalle 11:15 il display cucina non va lasciato solo: si ricarica da se '
          'ogni 20 secondi e segnala i nuovi ordini con un suono, ma qualcuno '
          'deve sentirlo.', False)],
        [('Al banco: tablet acceso sulla pagina del banco, schermo orientabile '
          'verso il cliente, registratore di cassa pronto.', False)],
    ], motivo='mezz\'ora prima del primo ritiro arrivano gli ordini '
              '"per subito".',
        accent=HEX_BLUE)

    fascia(doc, '11:30', 'Ultimo controllo, 15 minuti prima', [
        [('Quattro domande: il menu di oggi e giusto? Le disponibilita sono '
          'aggiornate? Il cesto e esposto? La lista dei pasti e in cucina?',
          False)],
        [('Se una risposta e no, hai ancora un quarto d\'ora.', False)],
    ], motivo='e l\'ultimo momento in cui un errore si corregge senza che il '
              'cliente lo veda.',
        accent=HEX_BLUE)

    fascia(doc, ORA_PRIMO_SLOT, 'Si apre il primo slot di ritiro', [
        [('Da adesso i clienti si presentano al banco per ritirare. Il servizio '
          'prosegue a intervalli di un quarto d\'ora fino alle ', False),
         (ORA_FINE_SERVIZIO, True), ('.', False)],
    ], motivo='e l\'orario del primo slot configurato nell\'applicazione.',
        accent=HEX_RED, critico=True)

    _page_break(doc)

    # ══ PARTE 3 ═══════════════════════════════════════════════════════════
    heading(doc, 'Parte 3 — Il primo ordine', 1, color=GREEN, before=0)
    body(doc, 'Il primo ordine della giornata e un collaudo: se attraversa tutti '
              'i passaggi senza intoppi, la macchina funziona. Seguilo di '
              'persona.', color=GRAY)
    rule(doc)

    step(doc, 1, 'L\'ordine compare sul display', [
        [('Nella colonna ', False), ('Da preparare', True),
         (' con il codice ordine, l\'orario di ritiro in evidenza, il nome del '
          'cliente e il dettaglio: articoli di menu e, per i prodotti composti '
          'dal cliente, l\'elenco completo degli ingredienti.', False)],
        [('L\'ordine arriva ', False), ('già pagato', True),
         (': l\'importo e stato scalato dal portafoglio alla conferma. Al banco '
          'non si incassa nulla.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'La cucina lo prende in carico', [
        [('Premendo ', False), ('In preparazione', True),
         (' il cliente riceve la notifica che il suo ordine e in lavorazione. Se '
          'c\'e un panino da scaldare, sul canale del personale compare anche '
          'l\'avviso della piastra.', False)],
        [('Lavora in ordine di orario di ritiro, non di arrivo.', False)],
    ], accent=HEX_GREEN)

    step(doc, 3, 'Si prepara seguendo la scheda', [
        [('Ingrediente per ingrediente, senza sostituzioni di iniziativa: '
          'l\'ordine e già pagato e composto dal cliente. Se manca qualcosa, si '
          'avvisa il banco.', False)],
    ], accent=HEX_GREEN)

    step(doc, 4, 'Scontrino ed etichetta sul prodotto', [
        [('Sul prodotto vanno lo ', False), ('scontrino di cassa', True),
         (' e l\'etichetta con il codice ordine, leggibili senza aprire '
          'l\'incarto.', False)],
        [('Il tagliando stampabile dal display cucina e un documento interno di '
          'preparazione: utile alla cucina, non sostituisce lo scontrino.',
          False)],
    ], accent=HEX_RED)

    step(doc, 5, 'Si dichiara pronto', [
        [('Premendo ', False), ('Pronto', True),
         (' il cliente riceve la notifica sul telefono. Da qui l\'ordine attende '
          'al banco.', False)],
    ], accent=HEX_GREEN)

    step(doc, 6, 'Consegna e chiusura', [
        [('Alla consegna in mano si premi ', False), ('Consegnato', True),
         (': l\'ordine esce dal display. Se il cliente non si presenta, ', False),
         ('Sollecita ritiro', True), (' gli manda un promemoria.', False)],
        [('Primo ordine chiuso. Da qui il servizio cammina da solo: il tuo posto '
          'e al banco, non davanti allo schermo.', False)],
    ], accent=HEX_GREEN)

    _page_break(doc)

    # ══ PARTE 4 · OLTRE LA GIORNATA ═══════════════════════════════════════
    heading(doc, 'Parte 4 — Oltre la giornata', 1, color=NAVY, before=0)
    body(doc, 'Tre attivita che non appartengono al ciclo quotidiano ma che, se '
              'trascurate, si fanno sentire: la chiusura del mese con le aziende '
              'convenzionate, la copia di sicurezza dei dati e gli strumenti per '
              'provare il sistema.', color=GRAY)
    rule(doc)

    heading(doc, 'A fine mese: la fattura alle aziende', 3, color=DARK)
    fascia(doc, 'Fine\nmese', 'Scarica il riepilogo mensile in PDF', [
        [('Convenzioni Aziendali', True),
         (': su ogni scheda azienda, in fondo, scegli il mese e premi ', False),
         ('Scarica PDF', True),
         ('. L\'icona rossa in alto nella scheda fa lo stesso per il mese in '
          'corso.', False)],
        [('Il documento riporta i pasti per ogni dipendente con i giorni di '
          'presenza, il dettaglio giorno per giorno e il ', False),
         ('totale fatturabile', True),
         (': e pensato per essere allegato alla fattura. I pasti annullati sono '
          'esclusi, gli importi sono al netto di IVA.', False)],
        [('Scaricalo a mese chiuso: fino all\'ultimo giorno i numeri possono '
          'ancora cambiare.', False)],
    ], motivo='il conteggio e definitivo solo dopo l\'ultimo servizio del mese.',
        accent=HEX_NAVY)

    heading(doc, 'Ogni settimana: la copia di sicurezza', 3, color=DARK)
    fascia(doc, 'Ogni\nvenerdì', 'Scarica il backup del database', [
        [('Impostazioni › Dati › ', False), ('Scarica il backup', True),
         (': ottieni un unico file con tutto il contenuto del database — '
          'clienti, ordini, movimenti, catalogo, convenzioni, impostazioni.',
          False)],
        [('Il file contiene ', False),
         ('dati personali dei clienti e credenziali di servizio', True),
         (' (token Telegram, password applicativa Gmail): conservalo come '
          'conserveresti un registro contabile, non su un disco condiviso.',
          False)],
        [('Tienine almeno le ultime quattro copie: un errore ci si accorge '
          'spesso con qualche giorno di ritardo.', False)],
    ], motivo='a fine settimana il lavoro dei cinque giorni e completo e il '
              'locale e chiuso.',
        accent=HEX_NAVY)

    fascia(doc, 'Quando\nserve', 'Ripristina da un backup', [
        [('Impostazioni › Dati › ', False), ('Ripristina dal file', True),
         (': carica un backup e il contenuto attuale del database viene '
          'sostituito. Serve digitare RIPRISTINA per confermare, e al termine '
          'si viene disconnessi perche anche gli utenti sono stati sostituiti.',
          False)],
        [('E un\'operazione da fare a locale chiuso: tutto cio che e stato '
          'registrato dopo il backup va perso.', False)],
    ], motivo='mai durante il servizio: gli ordini in corso spariscono.',
        accent=HEX_NAVY, critico=True)

    heading(doc, 'All\'occasione: provare il sistema', 3, color=DARK)
    fascia(doc, 'Prima\ndell\'avvio', 'Carica un mese di dati di prova', [
        [('Impostazioni › Dati › ', False),
         ('Carico mensile di dati di prova', True),
         (': genera un mese intero di attivita verosimile — pasti aziendali, '
          'ordini, caffe al banco, prodotti del builder — con le quantita '
          'giornaliere che decidi tu.', False)],
        [('Serve per vedere come si comportano report e andamenti prima di '
          'avere dati veri. Ogni carico resta elencato e si ', False),
         ('elimina per intero', True),
         (' con un pulsante: non lascia residui. I saldi dei portafogli non '
          'vengono toccati.', False)],
    ], motivo='da fare in fase di avvio o di formazione, non a locale aperto.',
        accent=HEX_NAVY)

    box(doc, 'Azzeramento completo: da usare una volta sola', [
        [('Impostazioni › Dati › ', False), ('Azzera tutto il database', True),
         (' svuota ogni tabella e ricrea i soli dati di base. Serve una volta, '
          'per passare dalla fase di prova a quella reale.', False)],
        [('Chiede di scrivere AZZERA, e ', False),
         ('le credenziali del super admin tornano quelle predefinite', True),
         (': cambiale subito dopo. Scarica un backup prima di procedere.',
          False)],
    ], accent=HEX_RED, fill=HEX_STOP, label_color=RED)

    _page_break(doc)

    # ══ Vincoli da conoscere ══════════════════════════════════════════════
    heading(doc, 'I quattro vincoli da conoscere', 1, before=0)
    body(doc, 'Comportamenti verificati dell\'applicazione che spiegano perche '
              'gli orari sono quelli. Conoscerli evita di prendere per guasto '
              'quello che e una regola.', color=GRAY)
    spacer(doc, 6)

    box(doc, '1 · Le etichette del cesto valgono 24 ore', [
        [('Un\'etichetta piu vecchia di 24 ore scade da sola alla prima '
          'scansione e il pezzo non e piu vendibile. Per un servizio che chiude '
          'alle ', False), (ORA_FINE_SERVIZIO, True),
         (', qualsiasi etichetta generata prima di quell\'ora del giorno '
          'precedente scade durante il servizio. Genera la mattina stessa.',
          False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    box(doc, '2 · I pasti aziendali si contano davvero solo 30 minuti prima', [
        [('Un dipendente puo modificare o disdire la prenotazione fino a ', False),
         ('30 minuti prima', True),
         (' dell\'orario di ritiro che ha scelto. La lista stampata alle 11:00 e '
          'un\'ottima previsione, non un numero chiuso: tienine conto sulle '
          'ultime porzioni.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    box(doc, '3 · I promemoria automatici non sono un orologio', [
        [('Gli avvisi di ritiro partono con un controllo che gira ', False),
         ('solo quando qualcuno sta usando l\'applicazione', True),
         (', al massimo una volta al minuto. Nelle ore morte, con nessuno '
          'collegato, possono ritardare.', False)],
        [('Il promemoria raggiunge il cliente ', False),
         ('su Telegram se lo ha collegato, altrimenti per email', True),
         ('. Perche la seconda strada funzioni serve Gmail configurata in '
          'Impostazioni: senza ne l\'uno ne l\'altra, nessun avviso parte.',
          False)],
        [('Non contarci per le comunicazioni che devono arrivare a un\'ora '
          'precisa: quelle mandale tu.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    box(doc, '4 · Si puo ordinare anche per uno slot già passato', [
        [('L\'applicazione controlla che lo slot sia attivo e non pieno, ma ',
          False), ('non che l\'orario sia ancora futuro', True),
         (': un cliente distratto puo confermare a mezzogiorno e mezzo un ordine '
          'per lo slot delle 11:45.', False)],
        [('Se vedi comparire nel display un ordine per un orario passato, e '
          'questo: trattalo come "per subito" e avvisa il cliente.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    _page_break(doc)

    # ══ Checklist ═════════════════════════════════════════════════════════
    heading(doc, 'Checklist da stampare', 1, before=0)
    body(doc, 'Una copia per giornata. Le caselle si spuntano a penna.',
         color=GRAY)
    spacer(doc, 6)

    for titolo, colore, righe in [
        ('Pomeriggio precedente', HEX_PURPLE, [
            ('14:45', 'Fabbisogno dei giorni successivi calcolato'),
            ('15:15', 'Giacenze controllate e ordini inviati'),
            ('15:45', 'Pasto aziendale di domani pubblicato'),
            ('16:00', 'Sondaggio inviato (se usato)'),
            ('16:15', 'Clienti in attesa attivati'),
            ('16:30', 'Ricariche e fidi sistemati (se il portafoglio '
                      'e attivo)'),
            ('16:45', 'Listino e quantita di domani aggiornati'),
            (ORA_CHIUSURA_BANCO, 'Banco chiuso, invenduto ritirato dal cesto'),
            ('17:15', 'Numeri della giornata letti'),
            (ORA_FINE_LAVORO, 'Postazioni in carica: si chiude'),
        ]),
        ('Mattina', HEX_BLUE, [
            ('07:30', 'Tre postazioni accese e collegate'),
            ('07:45', 'Stampanti provate, registratore di cassa pronto'),
            ('08:00', 'Pasto del giorno visibile ai clienti'),
            ('08:15', 'Display cucina controllato'),
            ('09:00', 'Disponibilita del giorno corrette'),
            ('09:30', 'Pezzi del cesto preparati'),
            ('10:15', 'Etichette generate, stampate e applicate'),
            ('10:45', 'Cesto esposto'),
            ('11:00', 'Lista pasti aziendali stampata e in cucina'),
            ('11:15', 'Display presidiato, audio attivo, banco pronto'),
            ('11:30', 'Ultimo controllo eseguito'),
            (ORA_PRIMO_SLOT, 'Servizio aperto'),
        ]),
    ]:
        heading(doc, titolo, 3, color=DARK)
        tblc = doc.add_table(rows=len(righe), cols=3)
        _no_borders(tblc)
        _set_col_width(tblc, 0, 1.2)
        _set_col_width(tblc, 1, 2.0)
        _set_col_width(tblc, 2, 12.8)
        for i, (ora, testo) in enumerate(righe):
            cbox = tblc.rows[i].cells[0]
            _cell_border(cbox, top=HEX_RULE, bottom=HEX_RULE, left=HEX_RULE,
                         right=HEX_RULE, sz='6')
            _cell_margins(cbox, top=70, bottom=70, left=60, right=60)
            _p_spacing(cbox.paragraphs[0], before=0, after=0)
            _run_font(cbox.paragraphs[0].add_run('  '), size=11)

            cora = tblc.rows[i].cells[1]
            _cell_margins(cora, top=70, bottom=70, left=100, right=40)
            _p_spacing(cora.paragraphs[0], before=0, after=0)
            _run_font(cora.paragraphs[0].add_run(ora), size=10, bold=True,
                      color=DARK, font=FONT)

            ctxt = tblc.rows[i].cells[2]
            _cell_margins(ctxt, top=70, bottom=70, left=60, right=60)
            _p_spacing(ctxt.paragraphs[0], before=0, after=0)
            _run_font(ctxt.paragraphs[0].add_run(testo), size=10, color=DGRAY)
        spacer(doc, 12)

    heading(doc, 'Periodiche', 3, color=DARK)
    tblp = doc.add_table(rows=4, cols=3)
    _no_borders(tblp)
    _set_col_width(tblp, 0, 1.2)
    _set_col_width(tblp, 1, 3.4)
    _set_col_width(tblp, 2, 11.4)
    for i, (quando, testo) in enumerate([
        ('Fine mese', 'PDF mensile scaricato per ogni convenzione'),
        ('Ogni venerdì', 'Backup del database scaricato e archiviato'),
        ('Ogni mese', 'Vecchi backup ruotati, ultime quattro copie tenute'),
        ('All\'avvio', 'Dati di prova eliminati prima di partire con i clienti'),
    ]):
        cbox = tblp.rows[i].cells[0]
        _cell_border(cbox, top=HEX_RULE, bottom=HEX_RULE, left=HEX_RULE,
                     right=HEX_RULE, sz='6')
        _cell_margins(cbox, top=70, bottom=70, left=60, right=60)
        _p_spacing(cbox.paragraphs[0], before=0, after=0)
        _run_font(cbox.paragraphs[0].add_run('  '), size=11)
        cq = tblp.rows[i].cells[1]
        _cell_margins(cq, top=70, bottom=70, left=100, right=40)
        _p_spacing(cq.paragraphs[0], before=0, after=0)
        _run_font(cq.paragraphs[0].add_run(quando), size=10, bold=True,
                  color=DARK, font=FONT)
        ct = tblp.rows[i].cells[2]
        _cell_margins(ct, top=70, bottom=70, left=60, right=60)
        _p_spacing(ct.paragraphs[0], before=0, after=0)
        _run_font(ct.paragraphs[0].add_run(testo), size=10, color=DGRAY)
    spacer(doc, 12)

    rule(doc, color=HEX_NAVY)
    box(doc, 'Assistenza e contatti', [
        [('Per qualsiasi dubbio su questo manuale o sull\'applicazione:',
          False)],
        [('Daniele Speziale — DS Consulting', True)],
        [('Email:  ', False), (EMAIL, True),
         ('        Cellulare:  ', False), (CELL, True)],
    ], accent=HEX_RED, fill=HEX_LIGHT, label_color=RED)

    body(doc, 'QuickLunch · La giornata del gestore · Orari tarati su ritiro '
              'pasti %s-%s, banco chiuso al pubblico alle %s, fine del lavoro '
              'alle %s: se cambiano, va ritarata la scaletta.'
              % (ORA_PRIMO_SLOT, ORA_FINE_SERVIZIO, ORA_CHIUSURA_BANCO,
                 ORA_FINE_LAVORO),
         size=8.5, color=GRAY, after=2)
    body(doc, '© 2024–26 DS Consulting  ·  %s' % CONTATTI, size=8.5, color=GRAY)

    return doc


def main():
    doc = build()
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] Documento salvato in: {OUT}')
    print('     Ritiro %s-%s, banco chiuso %s, fine lavoro %s'
          % (ORA_PRIMO_SLOT, ORA_FINE_SERVIZIO, ORA_CHIUSURA_BANCO,
             ORA_FINE_LAVORO))
    print('     10 attivita il pomeriggio, 12 la mattina')


if __name__ == '__main__':
    main()
