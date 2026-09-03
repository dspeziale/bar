#!/usr/bin/env python3
"""Genera docs/manuale_proprietario.docx con PT Sans Narrow."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
RED    = RGBColor(0xe9, 0x45, 0x60)
NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x16, 0x21, 0x3e)
DGRAY  = RGBColor(0x34, 0x3a, 0x40)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)

HEX_RED   = 'E94560'
HEX_NAVY  = '0F3460'
HEX_DARK  = '16213E'
HEX_LIGHT = 'F8F9FA'
HEX_WHITE = 'FFFFFF'

FONT = 'PT Sans Narrow'
OUT  = os.path.join(os.path.dirname(__file__), 'manuali',
                    'manuale_proprietario.docx')


# ── XML helpers ───────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old shd if present
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None,
                 sz='8', space='0'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, color in sides.items():
        el = OxmlElement(f'w:{side}')
        if color:
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    sz)
            el.set(qn('w:color'), color.lstrip('#'))
            el.set(qn('w:space'), space)
        else:
            el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_vAlign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(Cm(width_cm).emu / 914.4)))
        tcW.set(qn('w:type'), 'dxa')
        for old in tcPr.findall(qn('w:tcW')):
            tcPr.remove(old)
        tcPr.append(tcW)


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK'])
                  .WD_BREAK.PAGE)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        bdr.append(el)
    tblPr.append(bdr)


def _table_width(table, width_cm):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(Cm(width_cm).emu / 914.4)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _para_border(p, *, bottom_color=None, bottom_sz='6'):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)
    pBdr = OxmlElement('w:pBdr')
    if bottom_color:
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    bottom_sz)
        bot.set(qn('w:color'), bottom_color.lstrip('#'))
        pBdr.append(bot)
    pPr.append(pBdr)


def _run_font(run, size=13, bold=False, color=None, italic=False, font=FONT):
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font)


def _p_spacing(p, before=0, after=6, line=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if line:
        p.paragraph_format.line_spacing = Pt(line)


# ── Document helpers ──────────────────────────────────────────────────────────

def body_para(doc, text='', color=None, size=14, bold=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, after=4, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    _p_spacing(p, before=before, after=after)
    if text:
        run = p.add_run(text)
        _run_font(run, size=size, bold=bold, color=color or DGRAY)
    return p


def h1(doc, number, title, icon=''):
    """Heading sezione — grande con bordo inferiore rosso."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _p_spacing(p, before=8, after=3)

    # Label sezione piccolo in rosso
    label = p.add_run(f'SEZIONE {number}   ')
    _run_font(label, size=9, bold=True, color=RED)

    # Icona + titolo
    if icon:
        ic = p.add_run(icon + '  ')
        _run_font(ic, size=20)
    run = p.add_run(title)
    _run_font(run, size=21, bold=True, color=NAVY)

    _para_border(p, bottom_color=HEX_RED, bottom_sz='8')
    return p


def h2(doc, text):
    """Sottotitolo sezione."""
    p = doc.add_paragraph()
    _p_spacing(p, before=7, after=2)
    run = p.add_run(text)
    _run_font(run, size=15, bold=True, color=DARK)
    return p


def info_box(doc, text, style='tip', label=None):
    """Riquadro brand-consistent: bordo sinistro rosso, sfondo grigio chiaro."""
    IC = {'tip': '💡', 'warning': '⚠️', 'success': '✅'}
    ic = IC.get(style, '💡')

    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 18.0)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, 'F5F5F5')
    _cell_border(cell, top='E8E8E8', bottom='E8E8E8', right='E8E8E8', left=HEX_RED, sz='14')
    _cell_margins(cell, top=60, bottom=60, left=120, right=80)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=0)

    prefix_text = f'{ic}  '
    if label:
        prefix_text += f'{label}  '
    prefix = p.add_run(prefix_text)
    _run_font(prefix, size=13, bold=bool(label), color=RED)

    run = p.add_run(text)
    _run_font(run, size=13, color=DGRAY)

    return tbl


def step_row(doc, num, title, text):
    """Passo numerato: pallino colorato | contenuto."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    _table_width(tbl, 16.5)

    # ── Numero ──
    nc = tbl.rows[0].cells[0]
    _set_col_width(tbl, 0, 1.1)
    _cell_shd(nc, HEX_RED)
    _cell_margins(nc, top=60, bottom=60, left=40, right=40)
    _cell_vAlign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    rn = pn.add_run(str(num))
    _run_font(rn, size=14, bold=True, color=WHITE)

    # ── Contenuto ──
    cc = tbl.rows[0].cells[1]
    _cell_shd(cc, 'F0F4F8')
    _cell_margins(cc, top=60, bottom=60, left=100, right=80)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=3)
    rt = pt.add_run(title + '\n')
    _run_font(rt, size=11, bold=True, color=DARK)
    rb = pt.add_run(text)
    _run_font(rb, size=10, color=GRAY)

    return tbl


def data_table(doc, headers, rows, col_widths=None):
    """Tabella dati con header navy."""
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    _table_width(tbl, 18.0)

    # header
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        _cell_shd(c, HEX_RED)
        _cell_border(c)
        _cell_margins(c, top=60, bottom=60, left=80, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        r = p.add_run(h)
        _run_font(r, size=10, bold=True, color=WHITE)
    if col_widths:
        for i, w in enumerate(col_widths):
            _set_col_width(tbl, i, w)

    # data
    for ri, row_data in enumerate(rows):
        bg = HEX_WHITE if ri % 2 == 0 else 'F8F9FA'
        dr = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c = dr.cells[ci]
            _cell_shd(c, bg)
            _cell_border(c, bottom='DEE2E6')
            _cell_margins(c, top=50, bottom=50, left=80, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            bold = (ci == 0)
            r = p.add_run(val)
            _run_font(r, size=11, bold=bold, color=DARK if bold else GRAY)

    return tbl


def cred_box(doc, title, rows):
    """Box credenziali su sfondo scuro."""
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    _no_borders(tbl)
    _table_width(tbl, 18.0)
    _set_col_width(tbl, 0, 4.5)
    _set_col_width(tbl, 1, 12.0)

    # titolo su riga merger
    title_cell = tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
    _cell_shd(title_cell, HEX_DARK)
    _cell_margins(title_cell, top=80, bottom=60, left=100, right=100)
    pt = title_cell.paragraphs[0]
    _p_spacing(pt, before=0, after=0)
    rt = pt.add_run('🔑  ' + title)
    _run_font(rt, size=10, bold=True, color=RED)

    for ri, (key, val, highlight) in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci in range(2):
            _cell_shd(row.cells[ci], '1A1A2E' if ri % 2 == 0 else HEX_DARK)
            _cell_margins(row.cells[ci], top=50, bottom=50, left=100, right=100)

        pk = row.cells[0].paragraphs[0]
        _p_spacing(pk, before=0, after=0)
        _run_font(pk.add_run(key), size=10,
                  color=RGBColor(0x90, 0xa8, 0xbe))

        pv = row.cells[1].paragraphs[0]
        _p_spacing(pv, before=0, after=0)
        vc = RGBColor(0xff, 0xd7, 0x00) if highlight else WHITE
        _run_font(pv.add_run(val), size=11, bold=highlight, color=vc)

    return tbl


def role_table(doc, roles):
    """Tabella ruoli: badge colorato | permessi."""
    tbl = doc.add_table(rows=len(roles), cols=2)
    _no_borders(tbl)
    _table_width(tbl, 18.0)
    _set_col_width(tbl, 0, 4.0)
    _set_col_width(tbl, 1, 12.5)

    for ri, (badge_bg, badge_txt, name, perms) in enumerate(roles):
        row = tbl.rows[ri]

        # badge cella
        bc = row.cells[0]
        _cell_shd(bc, badge_bg)
        _cell_margins(bc, top=80, bottom=80, left=100, right=60)
        _cell_vAlign(bc, 'center')
        pb = bc.paragraphs[0]
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(pb, before=0, after=0)
        _run_font(pb.add_run(name), size=11, bold=True,
                  color=RGBColor(*bytes.fromhex(badge_txt)))

        # permessi cella
        pc = row.cells[1]
        _cell_shd(pc, HEX_WHITE if ri % 2 == 0 else 'F8F9FA')
        _cell_margins(pc, top=80, bottom=80, left=100, right=80)
        pp = pc.paragraphs[0]
        _p_spacing(pp, before=0, after=0)
        _run_font(pp.add_run(perms), size=11, color=GRAY)

    return tbl


def workflow_table(doc, steps):
    """Flusso orizzontale di step."""
    ncols = len(steps)
    tbl = doc.add_table(rows=1, cols=ncols)
    _table_width(tbl, 18.0)

    for i, (icon, title, desc) in enumerate(steps):
        c = tbl.rows[0].cells[i]
        bg = HEX_NAVY if i % 2 == 0 else '1A2E50'
        _cell_shd(c, bg)
        _cell_border(c, right='16213E')
        _cell_margins(c, top=100, bottom=100, left=80, right=80)

        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=0, after=4)
        _run_font(p.add_run(icon + '\n'), size=18)
        _run_font(p.add_run(title + '\n'), size=10, bold=True, color=WHITE)
        _run_font(p.add_run(desc), size=9, color=RGBColor(0xb0, 0xc4, 0xd8))

    return tbl


# ── Cover ─────────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Impostazioni margini minimi per la cover
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Tabella che copre tutta la pagina: 3 righe (top spacer, content, bottom spacer)
    tbl = doc.add_table(rows=3, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 21.0)

    def _row_height(row, h_cm):
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), str(int(Cm(h_cm).emu / 914.4)))
        trH.set(qn('w:hRule'), 'exact')
        trPr.append(trH)

    _row_height(tbl.rows[0], 5.0)
    _row_height(tbl.rows[1], 17.5)
    _row_height(tbl.rows[2], 5.0)

    for row in tbl.rows:
        _cell_shd(row.cells[0], HEX_DARK)
        _cell_margins(row.cells[0], top=0, bottom=0, left=0, right=0)

    # ── Riga contenuto ──
    cc = tbl.rows[1].cells[0]
    _cell_shd(cc, HEX_DARK)
    _cell_margins(cc, top=100, bottom=100, left=300, right=300)
    _cell_vAlign(cc, 'center')

    def cp(text='', size=11, bold=False, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, after=10):
        p = cc.add_paragraph()
        p.alignment = align
        _p_spacing(p, before=0, after=after)
        if text:
            r = p.add_run(text)
            _run_font(r, size=size, bold=bold, color=color)
        return p

    # prima del contenuto c'è già un paragrafo vuoto nella cella
    p0 = cc.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p0, before=0, after=20)
    r0 = p0.add_run('🍽️')
    _run_font(r0, size=54, color=WHITE)

    cp('QuickLunch',    size=52, bold=True, color=RED,   after=0)
    cp('PRANZO', size=52, bold=True, color=WHITE, after=16)
    cp('Sistema di gestione per bar, mense e caffetterie',
       size=14, color=RGBColor(0xa0, 0xb8, 0xd0), after=32)

    # Feature pills come elenco centrato
    features = ['📦 Ordini & Cucina', '💳 Wallet Digitale', '⭐ Fidelizzazione',
                '🪑 Prenotazione Tavoli', '📊 Report & Statistiche', '📱 Notifiche Telegram',
                '🏧 Banco POS — QR']
    pf = cc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pf, before=0, after=32)
    for i, feat in enumerate(features):
        r = pf.add_run(feat)
        _run_font(r, size=11, color=RGBColor(0xc8, 0xd8, 0xe8))
        if i < len(features) - 1:
            sep = pf.add_run('   ·   ')
            _run_font(sep, size=11, color=RGBColor(0x40, 0x55, 0x70))

    cp('Manuale del Proprietario  ·  Versione 2.0  ·  Giugno 2026  ·  © 2024–26 DS Consulting',
       size=9, color=RGBColor(0x40, 0x55, 0x70), after=0)


# ── Sezioni ───────────────────────────────────────────────────────────────────

def build_toc(doc):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=12)
    r = p.add_run('Indice dei contenuti')
    _run_font(r, size=22, bold=True, color=NAVY)
    _para_border(p, bottom_color=HEX_RED, bottom_sz='8')

    toc_items = [
        ('1',  '🍽️',  "Cos'è QuickLunch"),
        ('2',  '⚡',  'Funzionalità principali'),
        ('3',  '🚀',  'Setup iniziale'),
        ('4',  '📋',  'Gestione menu, prodotti e allergeni'),
        ('5',  '🥪',  'Builder: Panino, Insalata & Poke Bowl'),
        ('6',  '📦',  'Gestione ordini e slot orari'),
        ('7',  '🏧',  'Banco POS — Cassa Rapida QR'),
        ('8',  '💳',  'Wallet digitale e fidelizzazione'),
        ('9',  '🪑',  'Tavoli e prenotazioni'),
        ('10', '🏢',  'Pasto aziendale convenzionato'),
        ('11', '📦',  'Magazzino e stock giornaliero'),
        ('12', '🔐',  'Personale: ruoli e permessi'),
        ('13', '📣',  'Notifiche'),
        ('14', '📊',  'Sondaggi'),
        ('15', '📈',  'Report e statistiche'),
        ('16', '🏬',  'Multi-sede (Multi-tenant)'),
        ('17', '🔑',  'Accesso con Google (OAuth)'),
        ('18', '❓',  'Domande frequenti'),
        ('19', '🔑',  'Credenziali di test'),
        ('A',  '🖥️',  'Pagine principali del sito (mockup)'),
        ('B',  '⚖️',  'Gestione Wallet — Aspetti Fiscali'),
        ('C',  '💰',  'Modello SaaS — Prezzi e Metriche'),
        ('D',  '🏗️',  'Layout Fisici — Cucina, Sala, Cassa'),
    ]

    tbl = doc.add_table(rows=len(toc_items), cols=2)
    _no_borders(tbl)
    _table_width(tbl, 16.5)
    _set_col_width(tbl, 0, 1.0)
    _set_col_width(tbl, 1, 15.5)

    for ri, (num, icon, title) in enumerate(toc_items):
        bg = HEX_WHITE if ri % 2 == 0 else 'F0F4F8'
        row = tbl.rows[ri]

        nc = row.cells[0]
        _cell_shd(nc, HEX_RED)
        _cell_margins(nc, top=50, bottom=50, left=60, right=60)
        pn = nc.paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(pn, before=0, after=0)
        _run_font(pn.add_run(num), size=10, bold=True, color=WHITE)

        tc = row.cells[1]
        _cell_shd(tc, bg)
        _cell_margins(tc, top=50, bottom=50, left=100, right=60)
        pt = tc.paragraphs[0]
        _p_spacing(pt, before=0, after=0)
        _run_font(pt.add_run(icon + '  '), size=10)
        _run_font(pt.add_run(title), size=11, bold=False, color=NAVY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def s01(doc):
    h1(doc, 1, "Cos'è QuickLunch", '🍽️')
    body_para(doc,
        "QuickLunch è una piattaforma web multi-tenant per la gestione completa di bar aziendali, "
        "mense scolastiche e ristoranti con prenotazione. Funziona direttamente dal browser — "
        "nessuna installazione, nessuna app — ed è accessibile da qualsiasi dispositivo: "
        "computer, tablet o smartphone.")
    body_para(doc,
        "I clienti ordinano online, compongono pasti personalizzati e prenotano tavoli con fasce orarie. "
        "Il personale riceve ordini in tempo reale su un pannello KDS dedicato. "
        "Il proprietario monitora incassi, stock, statistiche e comunicazioni da un unico pannello admin.")
    info_box(doc,
        "Zero carta, zero code, zero incomprensioni: ogni ordine è digitale, tracciato e confermato "
        "automaticamente. Il wallet prepagato elimina la gestione del contante.",
        style='success', label='Vantaggio chiave:')

    h2(doc, 'Casi d\'uso')
    data_table(doc,
        ['Tipo di attività', 'Caso d\'uso principale'],
        [
            ['🏢 Bar aziendale / Mensa',
             'Preordini mattutini, slot di ritiro, wallet prepagato, pasto aziendale convenzionato'],
            ['🏫 Mensa scolastica',
             'Prenotazione pasti, gestione diete, allergeni, aziende convenzionate'],
            ['☕ Ristorante / Caffetteria',
             'Prenotazione tavoli con fasce orarie, builder panini/insalate, banco POS QR'],
            ['🏬 Catena multi-sede',
             'Ogni punto vendita indipendente: menu, clienti, staff e report separati'],
        ],
        col_widths=[5.5, 11.0])

    h2(doc, 'Tecnologie')
    data_table(doc,
        ['Componente', 'Tecnologia'],
        [
            ['Backend',    'Python / Flask (blueprints modulari)'],
            ['Frontend',   'Bootstrap 4 / AdminLTE — responsive mobile-first'],
            ['Database',   'SQLite (sviluppo) / PostgreSQL — Neon (produzione)'],
            ['Bot',        'Telegram Bot API per notifiche push individuali e broadcast'],
            ['Auth',       'Session Flask + Google OAuth 2.0 (opzionale)'],
            ['Hosting',    'Vercel (serverless) — HTTPS obbligatorio'],
        ],
        col_widths=[4.0, 13.5])


def s02(doc):
    h1(doc, 2, 'Funzionalità Principali', '⚡')
    body_para(doc,
        "QuickLunch offre un set completo di strumenti per la gestione quotidiana del locale. "
        "Tutte le funzionalità sono accessibili da browser, senza installazioni aggiuntive.")
    data_table(doc,
        ['Funzionalità', 'Descrizione'],
        [
            ['📋 Ordini digitali con slot di ritiro',
             'I clienti ordinano da smartphone, scelgono lo slot orario e ricevono conferma immediata'],
            ['🥪 Builder personalizzato (Panino / Insalata / Poke Bowl)',
             'Composizione passo-passo con ingredienti a scelta; prezzi calcolati in automatico'],
            ['🏧 Banco POS con pagamento QR',
             'Vendita rapida al bancone: griglia articoli touch, QR in un clic, wallet scalato < 10 sec'],
            ['💳 Wallet digitale prepagato',
             'Ogni cliente ha un portafoglio digitale; la ricarica avviene dal pannello admin'],
            ['⭐ Programma fedeltà a punti',
             '10 punti per ogni euro speso; ogni 100 punti = 1,00 € accreditato sul wallet'],
            ['🪑 Prenotazione tavoli con fasce orarie',
             'Slot tavoli configurabili, panoramica real-time, check-in con timer sessione'],
            ['🏢 Pasto aziendale convenzionato',
             'Menu del giorno, prenotazione per dipendenti convenzionati, fatturazione per azienda'],
            ['📦 Magazzino e stock giornaliero',
             'Quantità disponibili impostabili ogni mattina, alert automatico sotto soglia minima'],
            ['🌾 Allergeni sui prodotti',
             '14 allergeni europei (Reg. UE 1169/2011) gestiti per ogni prodotto e ingrediente builder'],
            ['🔐 Ruoli e permessi granulari (RBAC)',
             'Super Admin, Manager, Cassiere, Cuoco, Utente — 13 permessi configurabili singolarmente'],
            ['📣 Notifiche Telegram e Email',
             'Messaggi automatici per ogni evento: ordine confermato, in preparazione, pronto, annullato'],
            ['📊 Sondaggi menu',
             'Chiedi ai clienti cosa vogliono domani; risultati in tempo reale con percentuali'],
            ['📈 Report e statistiche',
             'Incasso 30 giorni, trend ordini, top 10 prodotti, KPI dashboard giornaliera'],
            ['🏬 Multi-tenant (multi-sede)',
             'Ogni sede ha dati completamente isolati; admin globale (tenant_id=None) accede a tutto'],
            ['🔑 Accesso con Google (OAuth)',
             'Clienti e staff accedono con account Google / Google Workspace aziendale'],
        ],
        col_widths=[6.5, 11.0])


def s03(doc):
    h1(doc, 3, 'Setup Iniziale', '🚀')
    body_para(doc,
        "Segui questi passaggi per configurare QuickLunch dal primo accesso. "
        "Tutti i passaggi sono eseguibili dal pannello admin senza assistenza tecnica.")

    h2(doc, 'Credenziali di accesso iniziali')
    data_table(doc,
        ['Campo', 'Valore'],
        [
            ['URL admin',  '/admin/dashboard'],
            ['Email',      'admin@bar.local'],
            ['Password',   'admin123'],
        ],
        col_widths=[4.0, 13.5])
    info_box(doc, "Cambia la password al primo accesso da Profilo → Modifica password.",
             style='warning', label='Sicurezza:')

    h2(doc, 'Passi di configurazione')
    steps = [
        ('Crea il primo tenant (sede)',
         'Vai su Admin → Tenant → Nuovo tenant. Inserisci nome e slug URL (es. "bar-centrale"). '
         'Il tenant isola i dati della tua sede da eventuali altri locali sullo stesso sistema.'),
        ('Configura SMTP per email',
         'Vai su Admin → Impostazioni → Email. Inserisci SMTP server, porta, email mittente '
         'e App Password Gmail (non la password normale: abilita 2FA, poi "Password per app" '
         'da Google Account → Sicurezza).'),
        ('Configura il bot Telegram',
         'Crea un bot su @BotFather e ottieni il token. '
         'Vai su Admin → Impostazioni → Telegram: inserisci Token Bot e Chat ID del gruppo/canale broadcast. '
         'Usa "Test connessione" per verificare.'),
        ('Configura Google OAuth (opzionale)',
         'Per consentire il login con Google: crea un OAuth 2.0 Client ID su Google Cloud Console. '
         'Aggiungi in config.py: GOOGLE_CLIENT_ID e GOOGLE_CLIENT_SECRET. '
         'Callback URL da registrare: https://tuodominio.com/auth/google/callback'),
        ('Aggiungi categorie e prodotti',
         'Vai su Admin → Prodotti → Categorie. Crea le macro-categorie (Panini, Primi, Bevande, ecc.) '
         'con icona Font Awesome e colore Bootstrap. '
         'Poi aggiungi i prodotti da Admin → Prodotti con nome, prezzo e allergeni.'),
        ('Configura slot orari e fasce tavoli',
         'Vai su Admin → Tavoli → Tab "Slot Ordini": aggiungi fasce di ritiro (es. 12:00, 12:15) '
         'con capienza massima. Tab "Fasce orarie": crea blocchi di prenotazione tavolo '
         '(es. 11:30–13:30, 30 min/sessione).'),
        ('Crea utenti staff con ruoli',
         'Vai su Admin → Persone → Personale → Crea nuovo. '
         'Assegna i ruoli: Cassiere per il banco POS, Cuoco per il KDS cucina, '
         'Manager per la supervisione completa.'),
    ]
    for i, (title, text) in enumerate(steps, 1):
        step_row(doc, i, title, text)

    info_box(doc,
        "Genera il QR code del link /t/slug/register e stampalo sul bancone o sui tavoli: "
        "i clienti si registrano in 30 secondi dal cellulare.",
        style='tip', label='Registrazione rapida clienti:')


def s04(doc):
    h1(doc, 4, 'Gestione Menu, Prodotti e Allergeni', '📋')

    h2(doc, 'Categorie prodotti')
    body_para(doc,
        "Ogni prodotto appartiene a una categoria. Le categorie si creano da Admin → Prodotti → Categorie.")
    data_table(doc,
        ['Campo', 'Descrizione', 'Esempio'],
        [
            ['Nome',   'Etichetta visibile nel menu clienti',          'Panini'],
            ['Icona',  'Classe Font Awesome (senza prefisso fa-)',      'utensils'],
            ['Colore', 'Classe Bootstrap per il badge della categoria', 'warning'],
        ],
        col_widths=[3.5, 9.5, 4.5])

    h2(doc, 'Campi di un prodotto')
    data_table(doc,
        ['Campo', 'Obbligatorio', 'Descrizione'],
        [
            ['Nome',         '✅', 'Nome visualizzato nel menu clienti'],
            ['Prezzo',       '✅', 'In euro con decimali (es. 4,50)'],
            ['Categoria',    '✅', 'A quale sezione del menu appartiene'],
            ['Descrizione',  'No', 'Ingredienti, note, testo libero allergeni'],
            ['is_active',    '—',  'Attiva/disattiva il prodotto senza eliminare lo storico'],
        ],
        col_widths=[4.0, 2.5, 11.0])

    h2(doc, 'Allergeni europei supportati (14 — Reg. UE 1169/2011)')
    body_para(doc,
        "I 14 allergeni regolamentati sono gestibili su ogni prodotto e su ogni ingrediente del builder:")
    data_table(doc,
        ['#', 'Allergene', '#', 'Allergene'],
        [
            ['1', 'Glutine (cereali contenenti glutine)', '8',  'Frutta a guscio'],
            ['2', 'Crostacei',                            '9',  'Sedano'],
            ['3', 'Uova',                                 '10', 'Senape'],
            ['4', 'Pesce',                                '11', 'Lupini'],
            ['5', 'Arachidi',                             '12', 'Molluschi'],
            ['6', 'Soia',                                 '13', 'Sesamo'],
            ['7', 'Latte',                                '14', 'Anidride solforosa e solfiti'],
        ],
        col_widths=[0.7, 5.8, 0.7, 5.8])

    h2(doc, 'Stock giornaliero')
    body_para(doc,
        "Ogni mattina vai su Admin → Stock giornaliero e imposta le quantità disponibili per il giorno. "
        "Quando i pezzi si esauriscono, il prodotto scompare automaticamente dal menu. "
        "La dashboard mostra in rosso i prodotti con meno di 3 pezzi rimasti (alert automatico sotto soglia).")
    step_row(doc, 1, 'Imposta le quantità al mattino',
             "Apri Admin → Stock giornaliero. Inserisci la disponibilità di ogni prodotto. "
             "Per prodotti senza limite (caffè, acqua) usa 999.")
    step_row(doc, 2, 'Alert automatico sotto soglia',
             "Quando la quantità scende sotto la soglia minima configurata per il prodotto, "
             "il sistema mostra un alert in dashboard e può inviare una notifica al fornitore.")
    step_row(doc, 3, 'Attivare/disattivare senza eliminare',
             "Usa l'interruttore is_active in Admin → Prodotti. "
             "Il prodotto sparisce dal menu clienti ma rimane nello storico ordini passati.")
    info_box(doc,
        "Non eliminare mai un prodotto che ha ordini storici: disattivalo con is_active = False. "
        "L'eliminazione renderebbe incompleto lo storico ordini.",
        style='warning', label='Importante:')


def s05(doc):
    h1(doc, 5, 'Builder: Panino, Insalata & Poke Bowl', '🥪')
    body_para(doc,
        "Il builder guida il cliente passo-passo nella composizione di un pasto personalizzato. "
        "Il prezzo finale si calcola sommando il prezzo base agli extra scelti. "
        "L'interfaccia ha stile kiosk touch, pensata anche per smartphone.")

    h2(doc, 'Tre tipologie con prezzi base configurabili')
    data_table(doc,
        ['Tipo', 'Prezzo base (config.py)', 'Step del flusso'],
        [
            ['🥪 Panino',    '3,50 € → BUILDER_PRICES["panino"]',   'Pane → Proteina → Verdure → Salse → Extra → Riepilogo'],
            ['🥗 Insalata',  '3,00 € → BUILDER_PRICES["insalata"]', 'Base → Proteina → Verdure → Condimento → Topping → Riepilogo'],
            ['🍱 Poke Bowl', '4,00 € → BUILDER_PRICES["poke"]',     'Base riso → Proteina → Verdure → Salsa fusion → Extra → Riepilogo'],
        ],
        col_widths=[2.5, 5.5, 9.5])

    h2(doc, 'Step obbligatori e opzionali')
    body_para(doc,
        "Ogni step può essere marcato come obbligatorio (★) o opzionale. "
        "Se il cliente preme «Avanti» su uno step obbligatorio senza scegliere, "
        "l'interfaccia mostra un'animazione di errore (shake) e blocca il proseguimento.")

    h2(doc, 'Opzione griglia/piastra 🔥 — solo Panino (+0.30 €)')
    body_para(doc,
        "Nella schermata di riepilogo del Panino, il cliente può attivare «Alla piastra» (+0.30 €). "
        "L'indicazione appare nel KDS cucina con il badge 🔥 visibile su tutte e tre le colonne "
        "(Da preparare, In preparazione, Pronti) per evitare consegne fredde.")

    h2(doc, 'Configurazione ingredienti — Admin → Builder → Ingredienti')
    data_table(doc,
        ['Campo', 'Descrizione'],
        [
            ['Nome',           'Es. «Riso bianco», «Salmone», «Salsa ponzu»'],
            ['Tipo builder',   '«panino», «insalata», «poke» o «tutti»'],
            ['Categoria step', 'A quale step del flusso appartiene (es. Proteina, Salsa...)'],
            ['Prezzo extra',   'Sovrapprezzo rispetto al base (0 = incluso nel prezzo base)'],
            ['Vegetariano',    'Mostra il simbolo 🌿 nell\'interfaccia cliente'],
            ['Allergeni',      'Testo libero: «pesce», «sesamo», «crostacei» ecc.'],
            ['Attivo',         'Disattiva temporaneamente se esaurito senza eliminarlo'],
        ],
        col_widths=[4.0, 12.5])

    h2(doc, 'Come aggiungere nuovi ingredienti e categorie')
    step_row(doc, 1, 'Vai su Admin → Builder → Ingredienti',
             "Premi «Nuovo ingrediente». Scegli il tipo builder e la categoria step.")
    step_row(doc, 2, 'Compila il form',
             "Nome, tipo, categoria step, prezzo extra (0 se incluso), "
             "flag vegetariano, allergeni, attivo.")
    step_row(doc, 3, 'Salva e verifica',
             "L'ingrediente appare immediatamente nel builder per i nuovi ordini. "
             "Testa aprendo /t/slug/builder/visual come cliente.")


def s06(doc):
    h1(doc, 6, 'Gestione Ordini e Slot Orari', '📦')

    h2(doc, 'Flusso di un ordine')
    workflow_table(doc, [
        ('📱', 'Cliente ordina', 'Sceglie prodotti e slot dal browser'),
        ('💳', 'Pagamento wallet', 'Il costo scala dal wallet automaticamente'),
        ('👨‍🍳', 'KDS cucina', 'L\'ordine appare nel pannello KDS in real-time'),
        ('🔔', 'Stato pronto', 'Staff avanza → cliente notificato via Telegram'),
        ('✅', 'Completato', 'Cassiere segna consegnato, punti fedeltà accreditati'),
    ])

    h2(doc, 'Slot di ritiro — Admin → Tavoli → Tab «Slot Ordini»')
    body_para(doc,
        "Gli slot definiscono le finestre orarie in cui i clienti possono ritirare il cibo.")
    data_table(doc,
        ['Operazione', 'Come fare'],
        [
            ['Aggiungere uno slot',
             "Seleziona l'orario con il time picker, imposta la capienza massima ordini, premi «Aggiungi»"],
            ['Rimuovere uno slot',
             "Premi «Elimina» sulla riga: gli ordini già prenotati NON vengono cancellati retroattivamente"],
            ['Disattivare uno slot',
             "Usa l'interruttore attivo/inattivo: i clienti non possono più selezionarlo per nuovi ordini"],
        ],
        col_widths=[4.5, 13.0])

    h2(doc, '«Adesso al banco» — ritiro immediato senza slot')
    body_para(doc,
        "Il cliente può scegliere «Adesso al banco» anziché uno slot programmato. "
        "L'ordine entra nel KDS con codice BANCO e viene processato immediatamente, "
        "senza attendere nessuna fascia oraria.")

    h2(doc, 'Pannello Cucina (KDS) — /admin/kds')
    body_para(doc,
        "Il pannello KDS mostra gli ordini del giorno in tre colonne: "
        "Da preparare, In preparazione, Pronti. "
        "Il cuoco avanza ogni ordine con un clic. Non serve tastiera: ideale per tablet in cucina. "
        "Gli ordini appaiono in tempo reale non appena il cliente completa il checkout.")

    h2(doc, 'Annullamento ordine e rimborso automatico')
    step_row(doc, 1, 'Rimborso wallet automatico',
             "L'importo pagato viene riaccreditato istantaneamente sul wallet del cliente.")
    step_row(doc, 2, 'Storno punti fedeltà',
             "I punti accumulati con quell'ordine vengono rimossi dal totale del cliente.")
    step_row(doc, 3, 'Notifica Telegram',
             "Il cliente riceve un messaggio automatico con l'importo rimborsato.")
    info_box(doc,
        "Se l'ordine è già completato e il cliente chiede rimborso, usa la ricarica manuale "
        "wallet (Admin → Utenti → Ricarica wallet) con causale «Rimborso ordine #XXX».",
        style='tip', label='Ordine già completato:')


def s07(doc):
    h1(doc, 7, 'Banco POS — Cassa Rapida QR', '🏧')
    body_para(doc,
        "Il Banco POS è il sistema di vendita rapida al bancone: permette al cassiere "
        "di incassare consumazioni veloci (caffè, brioche, bevande) senza che il cliente "
        "debba navigare nel menu o effettuare un preordine. "
        "Tutto avviene tramite QR code: zero contanti, zero attese, tutto tracciato digitalmente.")
    info_box(doc,
        "Dal QR generato alla conferma di pagamento: meno di 10 secondi. "
        "Il wallet del cliente viene scalato automaticamente e la transazione appare nella "
        "cronologia con data, ora e articoli acquistati.",
        style='success', label='Velocità:')

    h2(doc, 'Configurazione articoli banco — /admin/banco/items')
    body_para(doc, "Gli articoli si configurano da /admin/banco/items senza toccare il codice.")
    data_table(doc,
        ['Campo', 'Descrizione', 'Esempio'],
        [
            ['Nome',   'Etichetta mostrata sulla griglia banco',              'Caffè Espresso'],
            ['Prezzo', 'Importo scalato dal wallet del cliente',              '0.80 €'],
            ['Icona',  'Emoji visualizzata sulla griglia per riconoscimento', '☕'],
            ['Colore', 'Colore della card sulla griglia touch',               'danger / primary'],
            ['Attivo', "Attiva/disattiva senza eliminare dallo storico",      'Sì / No'],
        ],
        col_widths=[2.5, 9.5, 5.5])

    h2(doc, 'Flusso staff — /admin/banco')
    for i, (t, tx) in enumerate([
        ('Seleziona gli articoli',
         "Tocca ogni articolo sulla griglia touch. Il badge sull'articolo mostra la quantità selezionata. "
         "Ogni tocco aggiunge 1 unità."),
        ('Rivedi il carrello',
         "Il riepilogo laterale mostra articoli, quantità e subtotale. "
         "Usa «−» per rimuovere una quantità."),
        ('Genera il QR',
         "Premi «Genera QR»: il sistema crea una sessione di pagamento valida 10 minuti "
         "con countdown visibile nel modal."),
        ('Attendi il pagamento',
         "Il modal rimane aperto. Quando il cliente completa il pagamento, "
         "compare «✓ Pagato da [nome]\". Nessun contante, nessun cambio."),
    ], 1):
        step_row(doc, i, t, tx)

    h2(doc, 'Flusso cliente — smartphone')
    for i, (t, tx) in enumerate([
        ('Apri QuickLunch e premi «Paga al Banco»',
         "Visibile nella dashboard cliente. Si avvia lo scanner QR con la fotocamera posteriore."),
        ('Inquadra il QR mostrato dal cassiere',
         "Lo scanner legge il codice e apre automaticamente la pagina di conferma pagamento."),
        ('Verifica e conferma',
         "La pagina mostra articoli, totale e saldo wallet disponibile. Premi «Paga ora»."),
        ('Ricevuta automatica',
         "Wallet scalato istantaneamente. La transazione appare nella cronologia del cliente."),
    ], 1):
        step_row(doc, i, t, tx)

    h2(doc, 'Gestione sessioni scadute o annullate')
    body_para(doc,
        "Se il cliente non scansiona entro 10 minuti, la sessione QR scade automaticamente "
        "e nessun addebito avviene. Il cassiere può generare un nuovo QR. "
        "Se il cassiere chiude il modal prima del pagamento, la sessione viene annullata.")

    h2(doc, 'Vantaggi per il proprietario')
    data_table(doc,
        ['Beneficio', 'Dettaglio'],
        [
            ['Zero cash al bancone',
             'Ogni transazione è digitale, tracciata e abbinata a un cliente reale. '
             'Nessun rischio errori di resto o ammanchi di cassa.'],
            ['Report transazioni banco',
             'Le vendite banco appaiono nelle statistiche come categoria separata: '
             'volume giornaliero, settimanale e mensile in Admin → Report.'],
            ['Articoli configurabili senza codice',
             'Aggiungi articoli, modifica prezzi o disattiva voci stagionali '
             'direttamente dalla UI in qualsiasi momento.'],
            ['Banco POS vs Ordini tradizionali',
             'I due sistemi coesistono. Usa il banco per consumazioni veloci al bancone; '
             'usa gli ordini tradizionali per pasti con slot e cucina.'],
        ],
        col_widths=[5.0, 12.5])


def s08(doc):
    h1(doc, 8, 'Wallet Digitale e Fidelizzazione', '💳')

    body_para(doc,
        "Il wallet è una funzione attivabile: dal tab Funzionalità delle "
        "Impostazioni (interruttore \"Portafoglio prepagato e punti "
        "fedeltà\") si può disattivare per lavorare con il solo pagamento "
        "in cassa. A portafoglio spento l'applicazione non muove denaro — "
        "niente saldi, ricariche, fidi, punti o bonus, e ogni riferimento "
        "sparisce dalle pagine — ma ordini, banco e cesto continuano a "
        "registrare le vendite nei report e nei guadagni. Questo capitolo "
        "descrive il funzionamento a portafoglio attivo.")

    h2(doc, 'Wallet digitale prepagato (wallet_balance)')
    body_para(doc,
        "Ogni utente ha un campo wallet_balance che rappresenta il suo credito in euro. "
        "Prima di ordinare il cliente deve ricaricare il wallet. "
        "Il pagamento avviene automaticamente al checkout scalando il saldo. "
        "Se il saldo è insufficiente, l'ordine non viene accettato.")

    h2(doc, 'Ricarica wallet — solo da pannello admin')
    body_para(doc, "Non esiste autopagamento online: la ricarica avviene esclusivamente dal pannello admin.")
    for i, (t, tx) in enumerate([
        ('Vai su Admin → Persone → Utenti (Clienti → Ricarica wallet)',
         "Cerca l'utente nella lista."),
        ('Clicca su «+ Ricarica wallet»',
         "Si apre un popup. Inserisci l'importo e la causale (es. «Ricarica mensile marzo»)."),
        ('Conferma',
         "Il saldo viene aggiornato immediatamente e la transazione è tracciata nello storico."),
    ], 1):
        step_row(doc, i, t, tx)

    h2(doc, 'Pagamenti e rimborsi automatici')
    data_table(doc,
        ['Evento', 'Effetto sul wallet'],
        [
            ['Checkout ordine',          'Importo scalato automaticamente al momento dell\'ordine'],
            ['Pagamento banco POS QR',   'Importo scalato alla conferma del QR da parte del cliente'],
            ['Annullamento ordine',      'Rimborso automatico + storno punti fedeltà accumulati'],
            ['Riscatto punti fedeltà',   '+1,00 € accreditato sul wallet (configurabile)'],
            ['Ricarica admin',           'Saldo aumentato con tracciatura causale'],
        ],
        col_widths=[5.5, 12.0])

    h2(doc, 'Storico transazioni')
    body_para(doc,
        "Ogni movimento del wallet è tracciato con: tipo (acquisto, ricarica, rimborso, riscatto, banco), "
        "importo, data, ora e riferimento ordine. "
        "Il cliente vede il proprio storico nel profilo; l'admin lo vede in Admin → Utenti → Dettaglio.")

    h2(doc, 'Programma fedeltà a punti')
    data_table(doc,
        ['Parametro', 'Valore predefinito', 'Come cambiarlo'],
        [
            ['Punti accumulati per ogni € speso', '10 punti/€',  'config.py → LOYALTY_POINTS_PER_EURO'],
            ['Punti necessari per il riscatto',    '100 punti',  'config.py → LOYALTY_REWARD_POINTS'],
            ['Valore del premio riscattato',        '1,00 €',    'config.py → LOYALTY_REWARD_AMOUNT'],
        ],
        col_widths=[6.0, 3.5, 8.0])
    info_box(doc,
        "Con i valori predefiniti: spendendo 10 € si accumulano 100 punti e si ottiene 1 € di buono "
        "(sconto del 10%). Il riscatto è automatico quando si raggiunge la soglia.",
        style='tip', label='Esempio pratico:')


def s09(doc):
    h1(doc, 9, 'Tavoli e Prenotazioni', '🪑')
    body_para(doc,
        "Il modulo tavoli si gestisce da Admin → Tavoli (pagina con quattro tab). "
        "Slot ordini e fasce orarie tavoli sono sistemi INDIPENDENTI: "
        "gli slot regolano il ritiro del cibo; le fasce orarie regolano la prenotazione del posto.")

    h2(doc, 'Tavoli — anagrafica')
    body_para(doc,
        "Ogni tavolo ha: numero identificativo, numero di posti, zona/descrizione "
        "(es. Finestra, Centro, Terrazzo). "
        "Un tavolo disattivato non appare nella panoramica prenotazioni e non è prenotabile.")

    h2(doc, 'Fasce orarie tavoli (TableTimeBand) — risorse GLOBALI')
    body_para(doc,
        "Le fasce orarie definiscono i blocchi di tempo disponibili per la prenotazione. "
        "Sono risorse globali: si applicano a tutti i tavoli del locale.")
    data_table(doc,
        ['Campo', 'Descrizione', 'Esempio'],
        [
            ['Nome fascia',           'Etichetta identificativa',                   'Pranzo'],
            ['Orario inizio',         'Apertura del blocco prenotazioni',            '11:30'],
            ['Orario fine',           'Chiusura del blocco prenotazioni',            '14:00'],
            ['Durata sessione (min)', 'Tempo che ogni gruppo occupa il tavolo',      '30'],
        ],
        col_widths=[4.5, 8.0, 4.0])
    body_para(doc,
        "Il sistema calcola automaticamente le sessioni: "
        "una fascia 11:30–13:30 con 30 minuti genera le sessioni 11:30, 12:00, 12:30 e 13:00.")
    info_box(doc,
        "La prenotazione tavolo è INDIPENDENTE dall'ordine pasto. "
        "Un cliente può prenotare un posto senza ordinare online, e viceversa.",
        style='tip')

    h2(doc, 'Panoramica disponibilità in tempo reale')
    body_para(doc,
        "La panoramica mostra la disponibilità giornaliera per data (navigazione prev/next). "
        "Chip colorati: verde = libero, rosso = occupato (con nome cliente), "
        "blu = prenotazione del cliente corrente.")

    h2(doc, 'Check-in e timer sessione')
    step_row(doc, 1, 'Staff registra il check-in',
             "Da Admin → Tavoli → Panoramica, clicca il chip del tavolo prenotato → «Check-in». "
             "Il timer di sessione parte automaticamente.")
    step_row(doc, 2, 'Alert Telegram 10 minuti prima della scadenza',
             "Il sistema invia automaticamente una notifica Telegram al cliente "
             "per avvisarlo che la sessione sta per terminare.")
    step_row(doc, 3, 'Fine sessione',
             "Allo scadere del tempo il tavolo torna verde nella panoramica. "
             "Lo staff può registrare manualmente la liberazione anticipata.")


def s10(doc):
    h1(doc, 10, 'Pasto Aziendale Convenzionato', '🏢')
    body_para(doc,
        "Il modulo pasto aziendale permette alle aziende convenzionate di offrire il pranzo "
        "ai propri dipendenti con un menu dedicato e fatturazione mensile centralizzata.")

    h2(doc, 'Aziende convenzionate — Admin → Aziende')
    data_table(doc,
        ['Campo', 'Descrizione'],
        [
            ['Nome azienda',    'Ragione sociale del cliente aziendale'],
            ['Codice',          'Codice breve per identificazione interna (es. TECH01)'],
            ['Email contatto',  'Per comunicazioni e riepilogo fatturazione mensile'],
        ],
        col_widths=[4.0, 13.5])

    h2(doc, 'Associazione dipendenti all\'azienda')
    body_para(doc,
        "Vai su Admin → Persone → Clienti → Modifica utente: "
        "associa l'utente all'azienda convenzionata dal campo «Azienda». "
        "Solo gli utenti associati vedono il menu del pasto aziendale.")

    h2(doc, 'Pasto del giorno — Admin → Pasto Aziendale → Nuovo pasto')
    data_table(doc,
        ['Campo', 'Descrizione'],
        [
            ['Data',          'Il giorno di erogazione del pasto'],
            ['Primo',         'Nome del primo piatto'],
            ['Secondo',       'Nome del secondo piatto'],
            ['Contorno',      'Nome del contorno'],
            ['Bevanda',       'Tipo di bevanda inclusa'],
            ['Caffè incluso', 'Sì / No'],
            ['Allergeni',     'Elenco testuale degli allergeni presenti nel pasto'],
            ['Max porzioni',  'Tetto massimo di prenotazioni accettate (0 = illimitato)'],
            ['Prezzo',        'Costo del pasto completo per la fatturazione aziendale'],
        ],
        col_widths=[3.5, 14.0])

    h2(doc, 'Prenotazione da parte del dipendente')
    body_para(doc,
        "Il dipendente vede il pasto del giorno nella propria dashboard e preme «Prenota». "
        "La prenotazione scala immediatamente dal contatore delle porzioni disponibili. "
        "Una volta raggiunto il numero massimo, il pulsante viene disabilitato.")

    h2(doc, 'Regola annullamento — 30 minuti (tassativa)')
    body_para(doc,
        "Il dipendente può annullare solo se mancano PIÙ di 30 minuti all'orario di ritiro. "
        "Superata questa soglia, il pulsante «Annulla» scompare automaticamente "
        "sia nell'interfaccia cliente che nel controllo lato server.")
    info_box(doc,
        "La regola dei 30 minuti non può essere bypassata dal cliente. "
        "Solo l'admin può annullare d'ufficio oltre questo limite.",
        style='warning', label='Regola tassativa:')

    h2(doc, 'Dashboard admin — card pasti aziendali')
    body_para(doc,
        "La dashboard mostra una card con le prenotazioni per opzione del menu aziendale: "
        "nome opzione, numero prenotazioni, barra di avanzamento e posti residui. "
        "Se l'admin è anche dipendente convenzionato, la card evidenzia il proprio pasto prenotato.")

    h2(doc, 'Calcolo totale fatturabile per azienda')
    body_para(doc,
        "Vai su Admin → Pasto Aziendale → Report mensile: il sistema calcola il totale "
        "fatturabile per ogni azienda (numero prenotazioni × prezzo pasto). "
        "Il riepilogo è esportabile per la fatturazione.")


def s11(doc):
    h1(doc, 11, 'Magazzino e Stock Giornaliero', '📦')

    h2(doc, 'Stock giornaliero prodotti — Admin → Stock giornaliero')
    body_para(doc,
        "Ogni mattina imposta la disponibilità di ogni prodotto. "
        "Quando le unità si esauriscono, il prodotto scompare automaticamente dal menu clienti.")
    step_row(doc, 1, 'Imposta le quantità al mattino',
             "Vai su Admin → Stock giornaliero. Inserisci la disponibilità di ogni prodotto. "
             "Per prodotti senza limite (caffè, acqua) usa 999.")
    step_row(doc, 2, 'Soglia minima e alert automatico',
             "Configura la soglia minima per ogni prodotto. "
             "Quando la quantità scende sotto soglia, il sistema mostra un alert in rosso in dashboard "
             "e invia un'email al fornitore associato.")
    step_row(doc, 3, 'Monitoraggio real-time',
             "La dashboard mostra in rosso tutti i prodotti con meno di 3 pezzi rimasti. "
             "Lo staff può aggiornare le quantità in tempo reale durante il servizio.")

    h2(doc, 'Magazzino consumabili — Admin → Magazzino')
    data_table(doc,
        ['Campo', 'Descrizione', 'Esempio'],
        [
            ['Nome articolo',    'Descrizione del consumabile',       'Farina 00'],
            ['Quantità',         'Scorta attuale in magazzino',        '25'],
            ['Unità di misura',  'kg, litri, pezzi, confezioni',       'kg'],
            ['Fornitore',        'Fornitore associato per il riordino', 'Molini Rossi SRL'],
        ],
        col_widths=[4.0, 8.5, 5.0])

    h2(doc, 'Fornitori — Admin → Magazzino → Fornitori')
    data_table(doc,
        ['Campo', 'Descrizione'],
        [
            ['Nome',     'Ragione sociale del fornitore'],
            ['Email',    'Per invio automatico ordini di riassortimento'],
            ['Telefono', 'Contatto diretto per ordini urgenti'],
            ['Contatto', 'Riferimento personale per gli ordini'],
        ],
        col_widths=[4.0, 13.5])
    info_box(doc,
        "Quando la soglia minima di un articolo viene superata, il sistema può inviare "
        "automaticamente un'email al fornitore con l'elenco degli articoli da riordinare.",
        style='tip', label='Ordini automatici:')


def _old_s09_tail_UNUSED(doc):
    h2(doc, 'Ipotesi di organico per turno')
    body_para(doc,
        "Le tabelle seguenti mostrano come distribuire il personale tra le tre aree operative "
        "— Cassa, Cucina e Sala — in funzione del numero di persone disponibili. "
        "Le ipotesi sono calibrate su un servizio mensa/bar con 50–150 coperti a turno "
        "e picco concentrato nelle fasce 12:00–13:30.")

    # ── Tabella riassuntiva ──────────────────────────────────────────────────
    data_table(doc,
        ['Organico', '🏪 Cassa', '👨‍🍳 Cucina', '🍽️ Sala', 'Scenario tipico'],
        [
            ['4 persone', '1', '2', '1', 'Apertura / turno leggero'],
            ['5 persone', '1', '2', '2', 'Servizio standard ridotto'],
            ['6 persone', '1', '3', '2', 'Servizio standard completo'],
            ['8 persone', '2', '4', '2', 'Picco / evento / lancio menu'],
        ],
        col_widths=[3.0, 2.0, 2.5, 2.0, 7.0])

    info_box(doc,
        "La colonna «Cucina» include sia la preparazione dei piatti caldi sia "
        "l'assemblaggio dei builder (panino, insalata, poke). "
        "Con QuickLunch il cuoco vede gli ordini sul tablet KDS senza che la cassa "
        "debba comunicare nulla verbalmente.",
        style='tip', label='Come QuickLunch riduce la coordinazione verbale:')

    # ── Dettaglio per scenario ────────────────────────────────────────────────
    body_para(doc, '', after=4)

    _scenarios = [
        ('4 persone — Turno minimo', 'E67E22',
         [('🏪 Cassa  ×1',
           'Gestisce le ricariche wallet, risponde alle domande al banco e supervisiona '
           'gli ordini completati. Con QuickLunch la cassa non prende ordini verbali: '
           'li riceve già confermati dal sistema.'),
          ('👨‍🍳 Cucina  ×2',
           '1 addetto ai piatti caldi (primo/secondo), 1 addetto ai builder '
           '(panini, insalate, poke). Entrambi lavorano sul pannello KDS. '
           'Il volume ridotto permette a una sola persona di gestire il builder '
           'in autonomia.'),
          ('🍽️ Sala  ×1',
           'Distribuisce i vassoi ai tavoli, gestisce le prenotazioni, '
           'risponde ai clienti. In caso di picco improvviso supporta la cassa.')]),

        ('5 persone — Turno standard ridotto', '2980B9',
         [('🏪 Cassa  ×1',
           'Stessa funzione del turno a 4. Con un addetto in più in sala '
           'la cassa può dedicarsi alle ricariche e alla gestione wallet '
           'senza doversi occupare della distribuzione.'),
          ('👨‍🍳 Cucina  ×2',
           'Configurazione identica al turno a 4. Adatta se il menu del giorno '
           'non prevede builder complessi o poke (tipicamente metà settimana).'),
          ('🍽️ Sala  ×2',
           '1 addetto alla distribuzione vassoi, 1 addetto ai tavoli e all\'accoglienza. '
           'La doppia presenza in sala riduce sensibilmente i tempi di attesa '
           'percepiti dal cliente.')]),

        ('6 persone — Turno standard completo', '27AE60',
         [('🏪 Cassa  ×1',
           'Con tre persone in cucina e due in sala, la cassa può '
           'concentrarsi sulle operazioni amministrative: wallet, '
           'report del giorno, gestione stock.'),
          ('👨‍🍳 Cucina  ×3',
           '1 addetto piatti caldi, 1 addetto panini/insalate, '
           '1 addetto poke bowl e piatti freddi. '
           'Ogni builder type ha un responsabile dedicato: la preparazione '
           'alla piastra (🔥) viene gestita senza rallentare gli altri ordini.'),
          ('🍽️ Sala  ×2',
           'Turno completo: un addetto alla distribuzione e uno dedicato '
           'ai tavoli con prenotazione. Configurazione ottimale per un '
           'servizio mensa aziendale da 80–120 coperti.')]),

        ('8 persone — Turno pieno / picco', 'E94560',
         [('🏪 Cassa  ×2',
           '1 cassiere principale (wallet, ordini speciali, gestione code), '
           '1 cassiere di supporto (ricariche veloci, accoglienza, '
           'smistamento clienti verso il self-order se disponibile). '
           'Utile in presenza di clienti non digitali che ordinano al banco.'),
          ('👨‍🍳 Cucina  ×4',
           '1 coordinatore KDS (monitora gli slot, avanza gli stati), '
           '1 addetto piatti caldi, 1 addetto builder panino/insalata '
           '(incluse piastre 🔥), 1 addetto poke bowl e dessert. '
           'Con 4 persone in cucina si può gestire un picco di 150+ coperti '
           'mantenendo i tempi di preparazione sotto i 10 minuti per slot.'),
          ('🍽️ Sala  ×2',
           'Invariato rispetto al turno a 6: la sala non scala linearmente '
           'con i coperti grazie ai preordini. I clienti ritirano '
           'autonomamente leggendo il codice ordine sul display.')]),
    ]

    for title, hdr_color, roles in _scenarios:
        # Titolo scenario
        p = doc.add_paragraph()
        _p_spacing(p, before=10, after=4)
        _run_font(p.add_run(title), size=11, bold=True,
                  color=RGBColor(int(hdr_color[0:2], 16),
                                 int(hdr_color[2:4], 16),
                                 int(hdr_color[4:6], 16)))

        # Tabella ruoli dello scenario
        tbl = doc.add_table(rows=len(roles), cols=2)
        _no_borders(tbl)
        _table_width(tbl, 16.5)
        _set_col_width(tbl, 0, 3.5)
        _set_col_width(tbl, 1, 13.0)
        for ri, (role_lbl, role_desc) in enumerate(roles):
            bg = 'FFFFFF' if ri % 2 == 0 else 'F8F9FA'
            lc = tbl.rows[ri].cells[0]
            rc = tbl.rows[ri].cells[1]
            _cell_shd(lc, hdr_color)
            _cell_shd(rc, bg)
            _cell_margins(lc, top=60, bottom=60, left=80, right=60)
            _cell_margins(rc, top=60, bottom=60, left=80, right=60)
            _cell_border(lc, bottom='FFFFFF', right='FFFFFF')
            pl = lc.paragraphs[0]
            _p_spacing(pl, before=0, after=0)
            _run_font(pl.add_run(role_lbl), size=9, bold=True, color=WHITE)
            pr_ = rc.paragraphs[0]
            _p_spacing(pr_, before=0, after=0)
            _run_font(pr_.add_run(role_desc), size=9, color=DGRAY)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    info_box(doc,
        "Il Cassiere gestisce anche il Banco POS (/admin/banco): vendita rapida al bancone "
        "tramite QR code, senza contanti. Ogni consumazione viene scalata direttamente dal "
        "wallet del cliente e tracciata nelle statistiche. "
        "Per i dettagli vedi la sezione «Banco POS — Cassa Rapida QR».",
        style='tip', label='Ruolo Cassiere — Banco POS:')


def s12(doc):
    h1(doc, 12, 'Personale: Ruoli e Permessi', '🔐')
    body_para(doc,
        "Il sistema di controllo accessi (RBAC) garantisce che ogni membro dello staff "
        "veda solo le sezioni di sua competenza. I ruoli sono preconfigurati ma personalizzabili.")

    h2(doc, 'Ruoli di sistema')
    role_table(doc, [
        ('E94560', HEX_WHITE, '👑  Super Admin',
         'Accesso totale. Gestione tenant, utenti globali, configurazione sistema. '
         'L\'attributo is_admin=True bypassa tutti i controlli permessi.'),
        (HEX_NAVY, HEX_WHITE, '🏢  Manager',
         'Ordini, cucina, prodotti, stock, tavoli, slot, report, pasto aziendale.'),
        ('E67E22', HEX_WHITE, '💰  Cassiere',
         'Gestione ordini, wallet clienti, Banco POS (/admin/banco) per vendite QR al bancone.'),
        ('27AE60', HEX_WHITE, '👨‍🍳  Cuoco',
         'Solo pannello cucina KDS (/admin/kds): visualizza e avanza stati ordini.'),
        ('6C757D', HEX_WHITE, '👤  Utente',
         'Nessun accesso admin: solo area clienti per ordinare, prenotare e vedere lo storico.'),
    ])

    h2(doc, 'Permessi granulari (13 permessi)')
    data_table(doc,
        ['Permesso', 'Cosa controlla'],
        [
            ['can_manage_orders',       'Visualizza e modifica tutti gli ordini'],
            ['can_manage_products',     'CRUD prodotti e categorie'],
            ['can_manage_stock',        'Imposta lo stock giornaliero'],
            ['can_manage_users',        'Crea e modifica utenti clienti'],
            ['can_manage_staff',        'Crea e modifica utenti staff'],
            ['can_manage_tables',       'Gestisce tavoli, fasce orarie e prenotazioni'],
            ['can_manage_wallet',       'Ricarica e visualizza wallet utenti'],
            ['can_view_reports',        'Accede alle statistiche e ai report'],
            ['can_manage_tenant',       'Configura impostazioni del tenant'],
            ['can_manage_builder',      'Gestisce ingredienti del builder'],
            ['can_manage_surveys',      'Crea e gestisce sondaggi'],
            ['can_send_notifications',  'Invia notifiche broadcast Telegram/Email'],
            ['can_manage_banco',        'Accede al Banco POS (/admin/banco)'],
        ],
        col_widths=[5.5, 12.0])

    h2(doc, 'Creazione utenti staff — Admin → Persone → Personale')
    step_row(doc, 1, 'Crea il nuovo utente',
             "Premi «Crea nuovo utente staff». Compila email, nome e password temporanea.")
    step_row(doc, 2, 'Assegna uno o più ruoli',
             "Un utente può avere più ruoli: es. un manager può essere anche cassiere.")
    step_row(doc, 3, 'Ruolo personalizzato (opzionale)',
             "Vai su Admin → Persone → Ruoli & Permessi → Nuovo ruolo e seleziona i permessi "
             "dalla lista granulare. Esempio: «Responsabile Comunicazioni» con only "
             "can_manage_surveys + can_send_notifications.")
    info_box(doc,
        "L'attributo is_admin=True (Super Admin) bypassa TUTTI i permessi granulari. "
        "Usalo solo per il proprietario o per figure tecniche di fiducia.",
        style='warning', label='Super Admin:')


def s13(doc):
    h1(doc, 13, 'Notifiche', '📣')

    h2(doc, 'Configurazione bot Telegram')
    step_row(doc, 1, 'Crea il bot su @BotFather',
             "Su Telegram, apri @BotFather, invia /newbot, scegli nome e username. "
             "Ricevi il token API (formato: 123456:ABC-DEF...).")
    step_row(doc, 2, 'Inserisci token e Chat ID in Admin → Impostazioni → Telegram',
             "Chat ID broadcast: ID del gruppo/canale per le comunicazioni a tutti. "
             "Usa @userinfobot per ottenere il Chat ID di qualsiasi utente o gruppo.")
    step_row(doc, 3, 'Testa la connessione',
             "Usa il pulsante «Test connessione» per verificare che il bot risponda.")

    h2(doc, 'Come il cliente comunica il proprio Chat ID')
    step_row(doc, 1, 'Il cliente avvia una chat con il bot',
             "Trova il bot su Telegram (per username) e preme «Start».")
    step_row(doc, 2, 'Ottiene il Chat ID personale',
             "Il cliente invia /start a @userinfobot: riceve il proprio Chat ID numerico.")
    step_row(doc, 3, 'Lo staff inserisce il Chat ID nel profilo',
             "Admin → Persone → Clienti → Modifica → campo «Telegram Chat ID».")

    h2(doc, 'Eventi notificati automaticamente')
    data_table(doc,
        ['Evento', 'Canale', 'Messaggio inviato al cliente'],
        [
            ['Ordine confermato',          'Telegram individuale',
             '✅ «Ordine #XXX confermato! Ritiro alle HH:MM. Totale: X,XX €»'],
            ['Ordine in preparazione',     'Telegram individuale',
             '👨‍🍳 «Il tuo ordine è in preparazione. Pronto a breve!»'],
            ['Ordine pronto',              'Telegram individuale',
             '🔔 «Il tuo ordine è PRONTO per il ritiro!»'],
            ['Ordine annullato',           'Telegram individuale',
             '❌ «Ordine annullato. Rimborso di X,XX € sul tuo wallet.»'],
            ['Pasto aziendale prenotato',  'Telegram individuale',
             '🍽️ «Pasto del giorno prenotato per [data].»'],
            ['Scadenza sessione tavolo',   'Telegram individuale',
             '⏰ «La tua sessione al tavolo scade tra 10 minuti.»'],
        ],
        col_widths=[4.5, 3.5, 9.5])

    h2(doc, 'Configurazione email SMTP — Admin → Impostazioni → Email')
    data_table(doc,
        ['Parametro', 'Valore consigliato (Gmail)'],
        [
            ['SMTP Server', 'smtp.gmail.com'],
            ['Porta',       '587 (TLS)'],
            ['Email',       'tuamail@gmail.com'],
            ['Password',    'App Password — NON la password normale. '
                            'Abilita 2FA → Google Account → Sicurezza → Password per app'],
        ],
        col_widths=[4.5, 13.0])


def s14(doc):
    h1(doc, 14, 'Sondaggi', '📊')
    body_para(doc,
        "I sondaggi permettono di raccogliere le preferenze dei clienti "
        "(es. «Cosa vuoi mangiare domani?») e pianificare la produzione in anticipo.")

    h2(doc, 'Creazione e gestione sondaggio')
    step_row(doc, 1, 'Vai su Admin → Comunicazioni → Sondaggi → Nuovo sondaggio',
             "Inserisci la domanda e aggiungi le opzioni di risposta (es. nomi dei piatti del giorno).")
    step_row(doc, 2, 'Apri il sondaggio',
             "Il sondaggio inizia in stato «chiuso». Premi «Apri» per renderlo votabile dai clienti.")
    step_row(doc, 3, 'Condividi il link',
             "Invia il link del sondaggio via notifica Telegram broadcast o email.")
    step_row(doc, 4, 'Consulta i risultati in tempo reale',
             "Admin → Sondaggi → Dettaglio: voti e percentuali si aggiornano ad ogni nuovo voto.")
    step_row(doc, 5, 'Chiudi il sondaggio',
             "Premi «Chiudi»: il sondaggio diventa sola lettura, nessun nuovo voto possibile.")

    h2(doc, 'Regole')
    data_table(doc,
        ['Regola', 'Comportamento'],
        [
            ['Voto unico per utente',      'Ogni cliente può votare una sola volta per sondaggio'],
            ['Apertura/chiusura manuale',  'Solo l\'admin può aprire e chiudere il sondaggio'],
            ['Risultati in tempo reale',   'Le percentuali si aggiornano a ogni nuovo voto'],
            ['Sondaggi multipli',          'Puoi avere più sondaggi aperti contemporaneamente'],
        ],
        col_widths=[5.5, 12.0])
    info_box(doc,
        "Usa i sondaggi per pianificare la produzione: con i voti del giorno prima "
        "sai quante porzioni preparare per ogni piatto, riducendo sprechi.",
        style='success', label='Suggerimento produzione:')


def s15(doc):
    h1(doc, 15, 'Report e Statistiche', '📈')

    h2(doc, 'Dashboard KPI — Admin → Dashboard')
    body_para(doc, "La dashboard è il cruscotto quotidiano con i dati del giorno in corso:")
    data_table(doc,
        ['KPI', 'Descrizione'],
        [
            ['💰 Incasso oggi',            'Totale ordini completati del giorno (non annullati)'],
            ['📋 Ordini aperti',           'Ordini in attesa, confermati o in preparazione'],
            ['👥 Clienti registrati',      'Totale utenti attivi nel tenant'],
            ['⭐ Prodotti in esaurimento', 'Prodotti con quantità < 3 (alert in rosso)'],
            ['💳 Wallet totale',           'Somma di tutti i saldi wallet dei clienti'],
            ['🍽️ Pasti aziendali oggi',   'Prenotazioni per opzione con barra avanzamento'],
        ],
        col_widths=[5.0, 12.5])

    h2(doc, 'Report ultimi 30 giorni — Admin → Report')
    data_table(doc,
        ['Report', 'Cosa mostra'],
        [
            ['Incasso giornaliero',    'Grafico a barre degli ultimi 30 giorni di incasso'],
            ['Trend ordini',           'Volume ordini giorno per giorno'],
            ['Top 10 prodotti',        'I prodotti più venduti per quantità totale'],
            ['Transazioni banco',      'Volume vendite Banco POS giornaliero, settimanale e mensile'],
        ],
        col_widths=[5.0, 12.5])

    h2(doc, 'Storico ordini completo — Admin → Ordini')
    body_para(doc,
        "Filtra per data, stato (in attesa, confermato, in preparazione, pronto, completato, annullato) "
        "e per cliente. Ogni ordine ha una scheda di prelievo stampabile. "
        "Gli ordini annullati mostrano la causale e il rimborso wallet effettuato.")


def s16(doc):
    h1(doc, 16, 'Multi-sede (Multi-tenant)', '🏬')
    body_para(doc,
        "QuickLunch supporta nativamente la gestione di più sedi (tenant) sullo stesso sistema. "
        "Ogni sede ha dati completamente isolati: menu, clienti, staff, ordini e configurazioni "
        "non si mescolano mai tra tenant diversi.")

    h2(doc, 'Isolamento dati per tenant')
    data_table(doc,
        ['Dato', 'Isolamento'],
        [
            ['Menu e prodotti',    'Ogni tenant ha il proprio catalogo indipendente'],
            ['Clienti e staff',    'Ogni utente appartiene a un solo tenant (tenant_id)'],
            ['Ordini e wallet',    'Non visibili dagli altri tenant'],
            ['Configurazioni',     'SMTP, Telegram, slot orari: separati per tenant'],
        ],
        col_widths=[5.0, 12.5])

    h2(doc, 'Admin globale vs Admin tenant')
    step_row(doc, '👑', 'Super Admin (tenant_id = None)',
             "Account globale che accede a tutti i tenant. "
             "Crea e gestisce le sedi da Admin → Tenant. "
             "Non appartiene a nessuna sede specifica.")
    step_row(doc, '🏢', 'Admin Tenant',
             "Gestisce menu, personale e clienti della propria sede. "
             "Non può accedere agli altri tenant né alla gestione globale.")

    h2(doc, 'URL separati per sede')
    body_para(doc,
        "Ogni sede ha il proprio slug nel URL: /t/<slug>/. "
        "Esempio: /t/bar-centrale/menu per il menu del Bar Centrale, "
        "/t/mensa-tech/menu per la mensa aziendale.")

    h2(doc, 'Creare una nuova sede')
    step_row(doc, 1, 'Accedi come Super Admin e vai su Admin → Tenant → Nuovo tenant',
             "Inserisci nome, slug URL (minuscolo, senza spazi) e colore primario.")
    step_row(doc, 2, 'Crea l\'admin della sede',
             "Clicca «Crea admin» nella riga del tenant: il sistema genera la password "
             "e la mostra UNA SOLA VOLTA. Annotala e consegnala al responsabile.")
    info_box(doc,
        "La password dell'admin tenant viene mostrata una sola volta. Non è recuperabile.",
        style='warning')


def s17(doc):
    h1(doc, 17, 'Accesso con Google (OAuth)', '🔑')
    body_para(doc,
        "QuickLunch supporta il login con Google OAuth 2.0. "
        "I clienti (e lo staff) possono registrarsi e accedere con il proprio account Google "
        "senza dover creare una password separata. "
        "Compatibile con Google Workspace (ex GSuite) aziendale.")

    h2(doc, 'Configurazione')
    step_row(doc, 1, 'Crea un OAuth 2.0 Client ID su Google Cloud Console',
             "Vai su console.cloud.google.com → APIs & Services → Credentials → "
             "Create credentials → OAuth 2.0 Client ID. "
             "Tipo applicazione: Web application.")
    step_row(doc, 2, 'Aggiungi le variabili in config.py',
             "GOOGLE_CLIENT_ID = 'xxx.apps.googleusercontent.com'\n"
             "GOOGLE_CLIENT_SECRET = 'xxx'")
    step_row(doc, 3, 'Registra la Callback URL in Google Cloud Console',
             "Nella sezione «Authorized redirect URIs» aggiungi: "
             "https://tuodominio.com/auth/google/callback")
    step_row(doc, 4, 'Verifica il funzionamento',
             "Vai su /login: deve comparire il pulsante «Accedi con Google». "
             "Il cliente viene reindirizzato a Google e ritorna autenticato.")

    h2(doc, 'Parametri di configurazione')
    data_table(doc,
        ['Parametro', 'Descrizione'],
        [
            ['GOOGLE_CLIENT_ID',     'ID client OAuth ottenuto da Google Cloud Console'],
            ['GOOGLE_CLIENT_SECRET', 'Secret client OAuth (non condividerlo pubblicamente)'],
            ['Callback URL',         'https://tuodominio.com/auth/google/callback — da registrare in GCC'],
        ],
        col_widths=[4.5, 13.0])
    info_box(doc,
        "Se non configuri Google OAuth, il login avviene solo con email e password. "
        "Google OAuth è opzionale e non influisce sulle altre funzionalità.",
        style='tip')


def s18(doc):
    h1(doc, 18, 'Domande Frequenti', '❓')
    faqs = [
        ('Un cliente ha dimenticato la password. Come faccio?',
         "Non esiste il reset automatico via email. Vai su Admin → Persone → Clienti, "
         "clicca «Modifica» e inserisci una nuova password temporanea. "
         "Il cliente la cambierà dal suo profilo al prossimo accesso."),
        ('Un cliente ha il saldo insufficiente. Come ricarico il wallet?',
         "Vai su Admin → Persone → Utenti, trova il cliente, clicca «+ Ricarica wallet». "
         "Inserisci l'importo e la causale. Il saldo è aggiornato immediatamente."),
        ('Come si annulla un ordine e si rimborsa il cliente?',
         "Se l'ordine è ancora aperto: Admin → Ordini → clicca l'ordine → «Annulla ordine». "
         "Rimborso wallet e storno punti avvengono automaticamente. "
         "Se l'ordine è già completato, usa la ricarica manuale wallet con causale «Rimborso #XXX»."),
        ('Uno slot è pieno e un cliente vuole comunque ordinare. Cosa faccio?',
         "Aumenta la capienza dello slot da Admin → Tavoli → Tab Slot Ordini → Modifica, "
         "oppure apri uno slot aggiuntivo. "
         "In alternativa il cliente può usare «Adesso al banco» senza slot."),
        ('Un prodotto è esaurito a metà servizio. Come lo gestisco?',
         "Vai su Admin → Stock giornaliero e imposta la quantità a 0 per quel prodotto: "
         "scompare immediatamente dal menu. "
         "In alternativa disattivalo da Admin → Prodotti con l'interruttore is_active."),
        ('Un tavolo risulta occupato ma il cliente è andato via. Come lo libero?',
         "Vai su Admin → Tavoli → Panoramica, clicca il chip del tavolo occupato "
         "e usa «Termina sessione» per liberarlo manualmente."),
        ('Il pasto aziendale non è visibile al dipendente. Cosa verifico?',
         "Controlla: 1) il dipendente è associato all'azienda convenzionata nel profilo; "
         "2) il pasto del giorno è stato creato per la data corretta; "
         "3) non è stato raggiunto il limite massimo di porzioni."),
        ('Come funziona il pagamento al banco senza ordine tradizionale?',
         "Il cassiere usa il Banco POS (/admin/banco): seleziona articoli dalla griglia, "
         "preme «Genera QR» e mostra il codice al cliente. "
         "Il cliente apre QuickLunch, preme «Paga al Banco» e scansiona il QR: "
         "wallet scalato in meno di 10 secondi, senza contanti."),
        ('Come vedo gli allergeni di un prodotto?',
         "Gli allergeni si configurano da Admin → Prodotti → Modifica → campo Allergeni. "
         "Sono visibili ai clienti nella scheda prodotto e nel riepilogo ordine. "
         "Per il builder, ogni ingrediente ha il proprio campo allergeni."),
        ('Come configuro Google OAuth per il login con Google?',
         "Crea un OAuth 2.0 Client ID su Google Cloud Console, aggiungi GOOGLE_CLIENT_ID e "
         "GOOGLE_CLIENT_SECRET in config.py, e registra la callback URL "
         "https://tuodominio.com/auth/google/callback nel pannello Google."),
        ('Posso avere più admin per la stessa sede?',
         "Sì. Crea un utente staff e assegnagli il ruolo Super Admin (o imposta is_admin=True). "
         "Avrà gli stessi permessi dell'account principale."),
        ('Come disattivo un prodotto senza eliminarlo?',
         "Vai su Admin → Prodotti: usa l'interruttore is_active. "
         "Sparisce dal menu clienti ma rimane nel sistema con tutto lo storico. "
         "Riattivalo quando vuoi senza perdere nessun dato."),
        ('Cosa succede se Internet va giù durante il servizio?',
         "Il sistema è cloud-based. Se la connessione locale cade non sarà temporaneamente accessibile. "
         "Consigliato: tenere una SIM dati di backup per ambienti critici."),
        ('Come chiudo il servizio per un giorno festivo?',
         "Disattiva tutti gli slot orari da Admin → Tavoli → Tab Slot Ordini. "
         "Senza slot attivi i clienti non possono ordinare. "
         "Ricordati di riattivarli il giorno prima della riapertura."),
    ]

    tbl = doc.add_table(rows=len(faqs), cols=2)
    _no_borders(tbl)
    _table_width(tbl, 16.5)
    _set_col_width(tbl, 0, 0.6)
    _set_col_width(tbl, 1, 15.9)

    for ri, (q, a) in enumerate(faqs):
        bg = HEX_WHITE if ri % 2 == 0 else 'F4F6F9'
        row = tbl.rows[ri]

        dc = row.cells[0]
        _cell_shd(dc, HEX_RED if ri % 2 == 0 else 'C73452')
        _cell_margins(dc, top=60, bottom=0, left=60, right=60)
        pd = dc.paragraphs[0]
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(pd, before=0, after=0)
        _run_font(pd.add_run('D'), size=9, bold=True, color=WHITE)

        qc = row.cells[1]
        _cell_shd(qc, bg)
        _cell_margins(qc, top=60, bottom=60, left=100, right=80)
        pq = qc.paragraphs[0]
        _p_spacing(pq, before=0, after=4)
        _run_font(pq.add_run(q), size=10, bold=True, color=NAVY)
        pa = qc.add_paragraph()
        _p_spacing(pa, before=0, after=0)
        _run_font(pa.add_run(a), size=10, color=GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _win_chrome(tbl, url, row_idx=0):
    """Barra browser-like nella riga indicata di una tabella."""
    cell = tbl.rows[row_idx].cells[0]
    _cell_shd(cell, '1E2A38')
    _cell_margins(cell, top=45, bottom=45, left=80, right=80)
    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run('● ● ●   '), size=8, color=RGBColor(0x6c, 0x82, 0x96))
    _run_font(p.add_run(url), size=8, color=RGBColor(0x90, 0xb8, 0xd8))


def _win_row(tbl, row_idx, text, bg, fg=None, size=9, bold=False, center=False):
    cell = tbl.rows[row_idx].cells[0]
    _cell_shd(cell, bg)
    _cell_margins(cell, top=40, bottom=40, left=80, right=80)
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(text), size=size, bold=bold, color=fg or DARK)


def _ingredient_cards(doc, cards, selected_idx=None, bg='F8F9FA'):
    """Riga di carte ingrediente come tabella multi-colonna."""
    n = len(cards)
    tbl = doc.add_table(rows=1, cols=n)
    _no_borders(tbl)
    _table_width(tbl, 16.5)

    for ci, (emoji, name) in enumerate(cards):
        is_sel = (ci == selected_idx)
        c = tbl.rows[0].cells[ci]
        cbg = 'FCE4EC' if is_sel else 'FFFFFF'
        _cell_shd(c, cbg)
        _cell_margins(c, top=70, bottom=70, left=30, right=30)
        _cell_border(c,
                     top='E94560' if is_sel else 'DEE2E6',
                     bottom='E94560' if is_sel else 'DEE2E6',
                     left='E94560' if is_sel else 'DEE2E6',
                     right='E94560' if is_sel else 'DEE2E6')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=0, after=3)
        _run_font(p.add_run(emoji + '\n'), size=14)
        nt = p.add_run(name)
        _run_font(nt, size=7, bold=True, color=RED if is_sel else DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def _progress_steps(doc, steps, active):
    """Barra step numerata colorata."""
    n = len(steps)
    tbl = doc.add_table(rows=1, cols=n)
    _no_borders(tbl)
    _table_width(tbl, 16.5)

    for i, label in enumerate(steps):
        c = tbl.rows[0].cells[i]
        if i < active:
            bg, tc = '28A745', HEX_WHITE
        elif i == active:
            bg, tc = HEX_RED, HEX_WHITE
        else:
            bg, tc = 'DEE2E6', '6C757D'
        _cell_shd(c, bg)
        _cell_margins(c, top=45, bottom=45, left=10, right=10)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=0, after=2)
        num = '✓' if i < active else (str(i + 1) if i < n - 1 else '⊙')
        _run_font(p.add_run(num + '\n'), size=9, bold=True,
                  color=RGBColor(*bytes.fromhex(tc)))
        _run_font(p.add_run(label), size=6.5,
                  color=RGBColor(*bytes.fromhex(tc)))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def _step_header(doc, step_n, total, emoji, title, hint, required=False):
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 16.5)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, HEX_WHITE)
    _cell_margins(cell, top=100, bottom=80, left=80, right=80)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p, before=0, after=4)
    _run_font(p.add_run(f'STEP {step_n} DI {total}  '), size=7,
              bold=True, color=RGBColor(0x6c, 0x75, 0x7d))
    if required:
        _run_font(p.add_run('★ OBBLIGATORIO'), size=7,
                  bold=True, color=RED)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p2, before=4, after=4)
    _run_font(p2.add_run(emoji + '  '), size=20)
    _run_font(p2.add_run(title), size=16, bold=True, color=NAVY)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p3, before=0, after=0)
    _run_font(p3.add_run(hint), size=9, color=GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def _bottom_nav(doc, price_str, show_prev=True, show_next=True, show_add=False):
    tbl = doc.add_table(rows=1, cols=3)
    _no_borders(tbl)
    _table_width(tbl, 16.5)
    _set_col_width(tbl, 0, 3.5)
    _set_col_width(tbl, 1, 9.0)
    _set_col_width(tbl, 2, 4.0)

    for cell in tbl.rows[0].cells:
        _cell_shd(cell, HEX_WHITE)
        _cell_border(cell, top='DEE2E6')
        _cell_margins(cell, top=70, bottom=70, left=80, right=80)

    # price
    pl = tbl.rows[0].cells[0].paragraphs[0]
    _p_spacing(pl, before=0, after=0)
    _run_font(pl.add_run('Totale\n'), size=7, color=GRAY)
    _run_font(pl.add_run(price_str), size=13, bold=True, color=RED)

    # spacer
    ps = tbl.rows[0].cells[1].paragraphs[0]
    _p_spacing(ps, before=0, after=0)

    # buttons
    pr = tbl.rows[0].cells[2].paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _p_spacing(pr, before=0, after=0)
    if show_prev:
        _run_font(pr.add_run('← Indietro   '), size=9, color=GRAY)
    if show_next:
        _run_font(pr.add_run('AVANTI →'), size=9, bold=True, color=WHITE)
        # simulate button bg inline (can't set bg on run; use a small table trick)
    if show_add:
        _run_font(pr.add_run('🛒 Aggiungi'), size=9, bold=True, color=WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return tbl


def _kpi_card_row(doc, cards):
    """Riga di KPI card (max 4 colonne)."""
    n = len(cards)
    tbl = doc.add_table(rows=1, cols=n)
    _no_borders(tbl)
    _table_width(tbl, 16.5)
    for i, (icon, val, label, bg) in enumerate(cards):
        c = tbl.rows[0].cells[i]
        _cell_shd(c, bg)
        _cell_margins(c, top=80, bottom=80, left=100, right=100)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=0, after=3)
        _run_font(p.add_run(icon + '\n'), size=18)
        _run_font(p.add_run(val + '\n'), size=14, bold=True, color=WHITE)
        _run_font(p.add_run(label), size=8, color=RGBColor(0xc0, 0xd5, 0xe8))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def s_appendix(doc):
    h1(doc, 'A', 'Pagine principali del sito', '🖥️')
    body_para(doc,
        "Le rappresentazioni seguenti mostrano le schermate chiave che clienti e personale "
        "utilizzano ogni giorno. L'interfaccia è accessibile da qualsiasi browser, senza "
        "installare alcuna app. La versione mobile adatta automaticamente tutti i layout.")

    # ─────────────────────────────────────────────
    # 1. MENU DI OGGI
    # ─────────────────────────────────────────────
    h2(doc, '1 · Menu di oggi  —  /t/{slug}/menu')
    body_para(doc,
        "I clienti sfogliano i prodotti disponibili oggi, filtrano per categoria "
        "con le pillole di scelta rapida e aggiungono al carrello con la quantità desiderata.")

    n_rows = 7
    m = doc.add_table(rows=n_rows, cols=1)
    _no_borders(m)
    _table_width(m, 16.5)
    _win_chrome(m, 'app.bsrpranzo.it/t/bar-centrale/menu', 0)
    _win_row(m, 1,
             '🔵 Tutto   🔴 Panini   🟢 Primi   🟠 Dolci   🔵 Bevande   🟣 Insalate',
             'F0F4F8', fg=DARK, size=9)
    _win_row(m, 2,
             '─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─ ─',
             'FFFFFF', fg=RGBColor(0xde, 0xe2, 0xe6), size=6)
    prod_rows = [
        ('🥪', 'Panino Classico',       '4.50 €'),
        ('🍝', 'Spaghetti al Pomodoro', '5.00 €'),
        ('🧁', 'Tiramisù',              '2.50 €'),
    ]
    for ri, (ic, name, price) in enumerate(prod_rows, 3):
        bg = HEX_WHITE if ri % 2 == 0 else 'F8F9FA'
        c = m.rows[ri].cells[0]
        _cell_shd(c, bg)
        _cell_margins(c, top=50, bottom=50, left=80, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(ic + '  '), size=11)
        _run_font(p.add_run(f'{name:<32}'), size=9, color=DARK)
        _run_font(p.add_run(price), size=9, bold=True, color=RED)
        _run_font(p.add_run('   [+ Aggiungi]'), size=8, color=GRAY)
    _win_row(m, 6,
             '🛒 Carrello (2 prodotti)   |   💳 Wallet: 12.50 €   |   ⭐ 145 punti',
             '1E2A38', fg=RGBColor(0x90, 0xb8, 0xd8), size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    info_box(doc,
        "Il menu mostra solo i prodotti attivi con quantità > 0 aggiornate dal pannello admin. "
        "I prodotti esauriti scompaiono automaticamente.",
        style='tip')

    # ─────────────────────────────────────────────
    # 2. BUILDER — SCELTA TIPO
    # ─────────────────────────────────────────────
    h2(doc, '2 · Builder — scelta tipo piatto  —  /t/{slug}/builder/visual')
    body_para(doc,
        "Il cliente sceglie tra Panino, Insalata o Poke Bowl. Ogni opzione mostra il prezzo "
        "di partenza e una breve descrizione degli ingredienti che potrà personalizzare "
        "nei passi successivi.")

    tc = doc.add_table(rows=2, cols=1)
    _no_borders(tc)
    _table_width(tc, 16.5)
    _win_chrome(tc, 'app.quicklunch.it/t/bar-centrale/builder/visual', 0)

    cards_cell = tc.rows[1].cells[0]
    _cell_shd(cards_cell, 'F8F9FA')
    _cell_margins(cards_cell, top=80, bottom=80, left=80, right=80)
    p_title = cards_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p_title, before=0, after=10)
    _run_font(p_title.add_run('🍴  Componi il tuo piatto'),
              size=14, bold=True, color=NAVY)
    _run_font(p_title.add_run('\nScegli e personalizza passo per passo'),
              size=9, color=GRAY)

    _card_colors = [
        # emoji, label, desc, price, bg, border_hex, price_rgb
        ('🥪', 'PANINO',     'Pane · Proteina · Verdure · Salse',
         'da 3.50 €', 'FFF5F7', 'E94560', RED),
        ('🥗', 'INSALATA',   'Base · Proteina · Verdure · Condimento',
         'da 3.00 €', 'F0FFF4', '28A745', RGBColor(0x28, 0xa7, 0x45)),
        ('🍱', 'POKE BOWL',  'Base riso · Proteina · Verdure · Salsa fusion',
         'da 4.00 €', 'EBF8FC', '00B4D8', RGBColor(0x00, 0xb4, 0xd8)),
    ]
    inner = cards_cell.add_table(rows=1, cols=3)
    _no_borders(inner)
    for ci, (emoji, label, desc, price, bg, border, price_rgb) in enumerate(_card_colors):
        c = inner.rows[0].cells[ci]
        _cell_shd(c, bg)
        _cell_border(c, top=border, bottom=border, left=border, right=border, sz='12')
        _cell_margins(c, top=100, bottom=100, left=60, right=60)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=0, after=6)
        _run_font(p.add_run(emoji + '\n'), size=22)
        _run_font(p.add_run(label + '\n'), size=11, bold=True, color=DARK)
        _run_font(p.add_run(desc + '\n'), size=7, color=GRAY)
        _run_font(p.add_run(price), size=9, bold=True, color=price_rgb)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ─────────────────────────────────────────────
    # 3. BUILDER — STEP INGREDIENTI (KIOSK STYLE)
    # ─────────────────────────────────────────────
    h2(doc, '3 · Builder — step ingredienti (stile kiosk McDonald\'s)')
    body_para(doc,
        "Il cliente è guidato passo per passo. Ogni categoria di ingredienti occupa "
        "uno step distinto. Le carte sono grandi e touch-friendly. "
        "La carta selezionata si evidenzia in rosso. "
        "L'utente non può passare allo step successivo senza aver scelto in uno step obbligatorio.")

    # -- Chrome
    chrome_tbl = doc.add_table(rows=1, cols=1)
    _no_borders(chrome_tbl)
    _table_width(chrome_tbl, 16.5)
    _win_chrome(chrome_tbl,
                'app.bsrpranzo.it/t/bar-centrale/builder/visual?type=panino', 0)

    # -- Progress bar
    _progress_steps(doc, ['Pane', 'Proteina', 'Verdure', 'Salse', 'Fine'], active=0)

    # -- Step header
    _step_header(doc, 1, 4, '🍞', 'Scegli il Pane',
                 'Scegli 1 opzione', required=True)

    # -- Ingredient cards
    _ingredient_cards(doc, [
        ('🍞', 'Bianco'),
        ('🌾', 'Integrale'),
        ('🥖', 'Ciabatta'),
        ('🫓', 'Rosetta'),
        ('✨', 'S. Glutine'),
    ], selected_idx=None)

    # -- Bottom nav
    _bottom_nav(doc, '3.50 €', show_prev=False, show_next=True)

    info_box(doc,
        "Se lo step è obbligatorio e il cliente preme «Avanti» senza scegliere, "
        "il pannello si agita (animazione shake) e il titolo diventa rosso.",
        style='warning', label='Validazione:')

    # 3b — Step con selezione attiva
    body_para(doc, 'Dopo aver selezionato un ingrediente, la carta si evidenzia e compare '
                   'il segno di spunta. Il totale nella barra inferiore si aggiorna in tempo reale.')

    chrome2 = doc.add_table(rows=1, cols=1)
    _no_borders(chrome2)
    _table_width(chrome2, 16.5)
    _win_chrome(chrome2,
                'app.bsrpranzo.it/t/bar-centrale/builder/visual?type=panino', 0)

    _progress_steps(doc, ['Pane', 'Proteina', 'Verdure', 'Salse', 'Fine'], active=1)
    _step_header(doc, 2, 4, '🥩', 'Scegli la Proteina', 'Scegli 1 opzione', required=True)
    _ingredient_cards(doc, [
        ('🥩', 'Prosciutto Cotto'),
        ('🍖', 'Prosciutto Crudo'),
        ('🐟', 'Tonno'),
        ('🧀', 'Mozzarella'),
        ('🥩', 'Bresaola ✓'),
    ], selected_idx=4)
    _bottom_nav(doc, '3.50 €', show_prev=True, show_next=True)

    # ─────────────────────────────────────────────
    # 4. BUILDER — RIEPILOGO FINALE
    # ─────────────────────────────────────────────
    h2(doc, '4 · Builder — riepilogo e aggiunta al carrello')
    body_para(doc,
        "All'ultimo step il cliente vede tutte le sue scelte in un riepilogo visivo "
        "con emoji e prezzi. Il totale è evidenziato in grande con sfondo in gradiente. "
        "Il pulsante verde «Aggiungi al carrello» invia l'ordine.")

    chrome3 = doc.add_table(rows=1, cols=1)
    _no_borders(chrome3)
    _table_width(chrome3, 16.5)
    _win_chrome(chrome3,
                'app.bsrpranzo.it/t/bar-centrale/builder/visual?type=panino', 0)

    _progress_steps(doc, ['✓ Pane', '✓ Proteina', '✓ Verdure', '✓ Salse', 'Fine'], active=4)

    sum_tbl = doc.add_table(rows=1, cols=1)
    _no_borders(sum_tbl)
    _table_width(sum_tbl, 16.5)
    sc = sum_tbl.rows[0].cells[0]
    _cell_shd(sc, HEX_WHITE)
    _cell_margins(sc, top=80, bottom=60, left=80, right=80)
    ph = sc.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(ph, before=0, after=8)
    _run_font(ph.add_run('🎉\n'), size=22)
    _run_font(ph.add_run('Il tuo panino\n'), size=14, bold=True, color=NAVY)
    _run_font(ph.add_run('Controlla gli ingredienti e aggiungi al carrello'),
              size=9, color=GRAY)

    si = sc.add_table(rows=2, cols=4)
    _no_borders(si)
    sum_items = [
        ('Base', '🥖', 'Ciabatta', '3.50 €'),
        ('Proteina', '🥩', 'Bresaola', '+0.00 €'),
        ('Verdure', '🥗', 'Rucola · Pomodoro', '+0.00 €'),
        ('Salse', '🌿', 'Pesto', '+0.30 €'),
        ('Extra', '🧀', 'Parmigiano', '+0.50 €'),
        ('Extra', '🥓', 'Bacon', '+0.70 €'),
        ('',      '',   '',          ''),
        ('',      '',   '',          ''),
    ]
    for ri in range(2):
        for ci in range(4):
            idx = ri * 4 + ci
            if idx >= len(sum_items):
                break
            cat, em_s, name_s, price_s = sum_items[idx]
            c = si.rows[ri].cells[ci]
            _cell_shd(c, 'F8F9FA' if ri % 2 == 0 else HEX_WHITE)
            _cell_border(c, top='DEE2E6', bottom='DEE2E6',
                         left='DEE2E6', right='DEE2E6')
            _cell_margins(c, top=60, bottom=60, left=30, right=30)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _p_spacing(p, before=0, after=2)
            _run_font(p.add_run((cat + '\n') if cat else ''), size=6.5, color=GRAY)
            _run_font(p.add_run((em_s + '\n') if em_s else ''), size=12)
            _run_font(p.add_run((name_s + '\n') if name_s else ''), size=7.5, bold=True, color=DARK)
            _run_font(p.add_run(price_s if price_s else ''), size=7.5, color=RED)

    total_tbl = doc.add_table(rows=1, cols=1)
    _no_borders(total_tbl)
    _table_width(total_tbl, 16.5)
    tc2 = total_tbl.rows[0].cells[0]
    _cell_shd(tc2, HEX_RED)
    _cell_margins(tc2, top=80, bottom=80, left=100, right=100)
    pt2 = tc2.paragraphs[0]
    pt2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _p_spacing(pt2, before=0, after=0)
    _run_font(pt2.add_run('Totale    '), size=9, color=RGBColor(0xff, 0xd0, 0xd8))
    _run_font(pt2.add_run('5.00 €'), size=18, bold=True, color=WHITE)
    _run_font(pt2.add_run('         🛒 Aggiungi al carrello'),
              size=11, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ─────────────────────────────────────────────
    # 5. ADMIN — DASHBOARD
    # ─────────────────────────────────────────────
    h2(doc, '5 · Pannello Admin — Dashboard  —  /admin/dashboard')
    body_para(doc,
        "Il pannello admin è accessibile solo al personale autorizzato. "
        "La dashboard mostra in tempo reale i dati del giorno: "
        "incasso, ordini, clienti e alert prodotti in esaurimento.")

    chrome_adm = doc.add_table(rows=1, cols=1)
    _no_borders(chrome_adm)
    _table_width(chrome_adm, 16.5)
    _win_chrome(chrome_adm, 'app.bsrpranzo.it/admin/dashboard', 0)

    # nav bar
    nav_tbl = doc.add_table(rows=1, cols=1)
    _no_borders(nav_tbl)
    _table_width(nav_tbl, 16.5)
    nav_c = nav_tbl.rows[0].cells[0]
    _cell_shd(nav_c, HEX_DARK)
    _cell_margins(nav_c, top=50, bottom=50, left=80, right=80)
    pn = nav_c.paragraphs[0]
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run('🍽️ QuickLunch   '), size=10, bold=True, color=RED)
    _run_font(pn.add_run('│ Dashboard │ Ordini │ Prodotti │ Personale │ Report │ ...'),
              size=8, color=RGBColor(0x90, 0xa8, 0xc0))

    # KPI cards
    _kpi_card_row(doc, [
        ('💰', '€ 248.50',   'Incasso oggi',      HEX_RED),
        ('📋', '23',         'Ordini aperti',     HEX_NAVY),
        ('👥', '148',        'Clienti registrati','1A4A70'),
        ('⭐', '4 prodotti', 'In esaurimento',    '5D4037'),
    ])

    # orders table
    ord_tbl = doc.add_table(rows=4, cols=4)
    _table_width(ord_tbl, 16.5)
    _set_col_width(ord_tbl, 0, 1.0)
    _set_col_width(ord_tbl, 1, 5.5)
    _set_col_width(ord_tbl, 2, 5.0)
    _set_col_width(ord_tbl, 3, 5.0)
    header_row = ord_tbl.rows[0]
    for ci, hdr in enumerate(['#', 'Cliente', 'Prodotti', 'Stato']):
        c = header_row.cells[ci]
        _cell_shd(c, HEX_NAVY)
        _cell_margins(c, top=45, bottom=45, left=60, right=60)
        _cell_border(c)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(hdr), size=9, bold=True, color=WHITE)
    order_data = [
        ('#101', 'Marco Rossi',    'Panino Classico × 1', '🟡 In preparazione'),
        ('#102', 'Giulia Ferrari', 'Insalata Greca × 2',  '🟢 Pronto'),
        ('#103', 'Luca Bianchi',   'Pasta al Pesto × 1',  '🔵 In attesa'),
    ]
    for ri, (n, cl, pr, st) in enumerate(order_data, 1):
        bg = HEX_WHITE if ri % 2 == 0 else 'F8F9FA'
        for ci, val in enumerate([n, cl, pr, st]):
            c = ord_tbl.rows[ri].cells[ci]
            _cell_shd(c, bg)
            _cell_margins(c, top=40, bottom=40, left=60, right=60)
            _cell_border(c, bottom='DEE2E6')
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            _run_font(p.add_run(val), size=9, bold=(ci == 0), color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ─────────────────────────────────────────────
    # 6. CUCINA / KDS
    # ─────────────────────────────────────────────
    h2(doc, '6 · Pannello Cucina (KDS)  —  /admin/kds')
    body_para(doc,
        "Il pannello cucina mostra gli ordini del giorno in tre colonne. "
        "Un clic fa avanzare ogni ordine allo stato successivo, senza tastiera e senza mouse: "
        "ideale per un tablet montato in cucina.")

    chrome_kds = doc.add_table(rows=1, cols=1)
    _no_borders(chrome_kds)
    _table_width(chrome_kds, 16.5)
    _win_chrome(chrome_kds, 'app.bsrpranzo.it/admin/kds', 0)

    kds = doc.add_table(rows=1, cols=3)
    _table_width(kds, 16.5)
    kds_cols = [
        ('🔵 Da preparare (4)', [('#101 · Mario B.', '🥪 Panino Classico'),
                                  ('#104 · Anna V.',  '🍝 Pasta al Pesto')]),
        ('🟡 In preparazione (2)', [('#102 · Giulia F.', '🥗 Insalata Greca × 2')]),
        ('🟢 Pronti (3)', [('#099 · Luca M.',  '☕ Caffè + 🧁 Brioche'),
                            ('#100 · Sara C.',  '🥪 Panino Vegano')]),
    ]
    col_bgs = ['EBF5FB', 'FFF8E1', 'E8F5E9']
    col_hdrs = ['1565C0', 'F57C00', '2E7D32']
    for ci, (col_title, items) in enumerate(kds_cols):
        c = kds.rows[0].cells[ci]
        _cell_shd(c, col_bgs[ci])
        _cell_margins(c, top=60, bottom=60, left=60, right=60)
        ph = c.paragraphs[0]
        _p_spacing(ph, before=0, after=8)
        _run_font(ph.add_run(col_title), size=9, bold=True,
                  color=RGBColor(*bytes.fromhex(col_hdrs[ci])))
        for order_id, dishes in items:
            po = c.add_paragraph()
            _p_spacing(po, before=4, after=4)
            _cell_border(c, bottom='CCCCCC')
            _run_font(po.add_run(order_id + '\n'), size=8, bold=True, color=DARK)
            _run_font(po.add_run(dishes), size=8, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    info_box(doc,
        "Il pannello KDS non richiede aggiornamento manuale: gli ordini appaiono "
        "in tempo reale appena il cliente completa il checkout.",
        style='success', label='Aggiornamento real-time:')


def s_appendix_wallet(doc):
    h1(doc, 'B', 'Gestione Wallet — Aspetti Fiscali', '⚖️')
    body_para(doc,
        "Documento riepilogativo delle possibili modalità di gestione del wallet ricaricabile "
        "di QuickLunch, utilizzato per l'acquisto di pranzi e consumazioni in bar e mense. "
        "Esistono due approcci distinti con implicazioni fiscali differenti. "
        "L'appendice riguarda la sola modalità prepagata: con il wallet "
        "disattivato non c'è raccolta anticipata di denaro e il tema del "
        "credito residuo non si pone.")

    info_box(doc,
        "Prima dell'adozione operativa del sistema è opportuno richiedere una verifica da parte "
        "del commercialista o consulente fiscale incaricato, poiché il corretto inquadramento "
        "dipende dalle specifiche modalità di utilizzo del wallet.",
        style='warning', label='Nota legale:')

    # ── Tabella di confronto rapido ──
    h2(doc, 'Confronto tra le due soluzioni')
    data_table(doc,
        ['Aspetto', 'Sol. 1 — Voucher multiuso', 'Sol. 2 — Voucher monouso'],
        [
            ['Momento fiscale',
             'Al consumo: scontrino a ogni acquisto',
             'Alla ricarica: scontrino unico anticipato'],
            ['Quando si applica',
             'Prodotti diversi, IVA non determinabile alla ricarica',
             'Beni/servizi noti e aliquota IVA identificabile subito'],
            ['Flessibilità',
             'Alta — credito spendibile su qualsiasi prodotto',
             'Bassa — legata a pasti o pacchetti predefiniti'],
            ['Complessità gestionale',
             'Maggiore (scontrino a ogni transazione)',
             'Minore (scontrino solo alla ricarica)'],
            ['Caso d\'uso tipico',
             'Bar aziendale / mensa con menu variabile',
             'Abbonamento pasto a pacchetto fisso (es. 10 pasti)'],
        ],
        col_widths=[4.2, 6.15, 6.15])

    # ══════════════════════════════════════════════
    # SOLUZIONE 1
    # ══════════════════════════════════════════════
    h2(doc, 'Soluzione 1 — Wallet come voucher multiuso  (generalmente consigliata)')
    body_para(doc,
        "Il cliente ricarica un credito (es. 50 €) utilizzabile successivamente per acquistare "
        "prodotti diversi del bar. La ricarica non rappresenta ancora una cessione di beni o "
        "servizi, poiché la natura degli acquisti futuri non è determinata al momento del versamento.")

    h2(doc, 'Flusso operativo — Soluzione 1')
    step_row(doc, 1, 'Ricarica wallet',
             "Il cliente versa l'importo scelto. Il sistema registra il credito e produce "
             "un documento/ricevuta di ricarica. Nessuno scontrino fiscale viene emesso in questa fase.")
    step_row(doc, 2, 'Nessuna fiscalizzazione alla ricarica',
             "La ricarica è un acconto su prestazioni future non ancora determinate. "
             "L'IVA non è esigibile e non si emette documento commerciale.")
    step_row(doc, 3, 'Al consumo — documento commerciale obbligatorio',
             "Ogni volta che il cliente acquista un prodotto, il sistema scala il credito dal wallet "
             "ed è obbligatoria l'emissione del documento commerciale (scontrino elettronico) "
             "relativo alla vendita effettiva.")

    h2(doc, 'Riferimenti normativi — Soluzione 1')
    data_table(doc,
        ['Norma', 'Contenuto rilevante per il wallet multiuso'],
        [
            ['DPR 633/1972  art. 6-bis',
             'Voucher monouso — esigibilità IVA al momento dell\'emissione del voucher'],
            ['DPR 633/1972  art. 6-ter',
             'Voucher multiuso — IVA esigibile al momento del riscatto (consumo effettivo)'],
            ['DPR 633/1972  art. 6-quater',
             'Distribuzione di voucher tramite intermediari — trattamento IVA'],
            ['D.Lgs. 141/2018',
             'Recepimento Direttiva UE 2016/1065 relativa al trattamento IVA dei voucher'],
        ],
        col_widths=[4.5, 12.0])

    info_box(doc,
        "Con la Soluzione 1 il registratore di cassa (o il sistema di cassa) deve emettere "
        "scontrino elettronico a ogni singola transazione di consumo. "
        "QuickLunch registra ogni ordine con importo e data, facilitando la riconciliazione contabile.",
        style='success', label='Come QuickLunch supporta questa soluzione:')

    # ══════════════════════════════════════════════
    # SOLUZIONE 2
    # ══════════════════════════════════════════════
    h2(doc, 'Soluzione 2 — Voucher monouso fiscalizzato alla ricarica')
    body_para(doc,
        "Il credito acquistato corrisponde a beni o servizi già determinati e fiscalmente "
        "identificabili al momento del pagamento. Questa soluzione è applicabile solo quando "
        "la natura della prestazione e il trattamento IVA sono già noti alla ricarica.")

    body_para(doc, 'Esempio pratico:')
    step_row(doc, 1, 'Acquisto pacchetto pasti predefinito',
             "Il cliente acquista 10 pasti completi (es. primo + secondo + acqua a prezzo fisso). "
             "Natura della prestazione e aliquota IVA sono già note e determinate.")
    step_row(doc, 2, 'Emissione scontrino alla ricarica',
             "Lo scontrino fiscale viene emesso al momento dell'acquisto del voucher/pacchetto, "
             "per l'intero importo versato, con indicazione dell'IVA applicabile.")
    step_row(doc, 3, 'Utilizzo senza nuova fiscalizzazione',
             "I pasti successivi vengono erogati senza emettere ulteriori documenti commerciali, "
             "poiché la vendita era già stata registrata e fiscalizzata integralmente.")

    info_box(doc,
        "Questa soluzione è applicabile solo quando prodotti, quantità e aliquota IVA sono "
        "predeterminati. Non è adatta a bar con menu variabile o credito a utilizzo libero.",
        style='tip', label='Quando NON applicare la Soluzione 2:')

    # ── Flusso a confronto ──
    h2(doc, 'Flusso operativo a confronto')
    workflow_table(doc, [
        ('💳', 'Ricarica wallet',   'Cliente versa il credito'),
        ('🧾', '→ Scontrino?',      'Sol. 1: NO\nSol. 2: SÌ (immediato)'),
        ('🍽️', 'Consumo / acquisto', 'Credito scalato da wallet'),
        ('🧾', '→ Scontrino?',      'Sol. 1: SÌ (obbligatorio)\nSol. 2: NO'),
        ('✅', 'Chiusura',          'Credito azzerato o esaurito'),
    ])

    info_box(doc,
        "Indipendentemente dalla soluzione adottata, QuickLunch conserva lo storico completo "
        "di tutte le ricariche e transazioni. Questo registro è utile per le verifiche fiscali "
        "e per la riconciliazione con il registratore di cassa.",
        style='tip', label='Registro transazioni QuickLunch:')


def s19(doc):
    h1(doc, 19, 'Credenziali di Test', '🔑')
    info_box(doc,
        "Cambia tutte le password predefinite prima di andare in produzione. "
        "Le credenziali seguenti sono per l'ambiente di test e demo.",
        style='warning', label='Attenzione:')

    h2(doc, 'Account di sistema')
    data_table(doc,
        ['Email', 'Password', 'Ruolo', 'Utilizzo'],
        [
            ['admin@bar.local',    'admin123',    'Super Admin',
             'Pannello completo, gestione tenant, configurazione sistema'],
            ['banco@bar.local',    'Banco2024!',  'Cassiere',
             'Tablet Banco POS (/admin/banco), gestione ordini e wallet'],
            ['cucina@bar.local',   'Cucina2024!', 'Cuoco',
             'Display Cucina KDS (/admin/kds), avanzamento stati ordini'],
            ['sala@bar.local',     'Sala2024!',   'Manager',
             'Tablet Staff Sala: ordini, tavoli, prodotti, stock, report'],
            ['cliente1@bar.local', 'Cliente1!',   'Utente (luca_verdi)',
             'Area clienti: ordini, wallet 30 €, punti fedeltà'],
            ['cliente2@bar.local', 'Cliente2!',   'Utente (anna_rossi)',
             'Area clienti: ordini, wallet 15 €, prenotazioni tavoli'],
        ],
        col_widths=[4.5, 2.8, 2.5, 7.7])

    h2(doc, 'Link di registrazione clienti')
    data_table(doc,
        ['Sede', 'Link di registrazione'],
        [
            ['Bar Centrale',      '/t/bar-centrale/register'],
            ['Mensa AziendaTech', '/t/mensa-tech/register'],
            ['Caffetteria Duomo', '/t/caffetteria-duomo/register'],
        ],
        col_widths=[5.5, 11.0])

    info_box(doc,
        "Per creare credenziali per un nuovo admin tenant (produzione), usa "
        "Admin → Tenant → Crea admin. La password viene mostrata una sola volta.",
        style='success')


# ── Appendice C — Modello SaaS ───────────────────────────────────────────────

def s_appendix_saas(doc):
    h1(doc, 'C', 'Modello SaaS — Prezzi, Commissioni e Metriche', '💰')
    body_para(doc,
        "QuickLunch è venduto come Software as a Service (SaaS) con un modello ibrido: "
        "un canone fisso mensile a copertura dei costi di infrastruttura e una commissione "
        "variabile dell'1,5% sul valore degli scontrini elaborati (IVA esclusa). "
        "Più il cliente guadagna, più il fornitore guadagna — allineamento di interessi totale.")

    h2(doc, 'Struttura tariffaria')
    data_table(doc,
        ['Voce', 'Importo', 'Frequenza', 'Dettaglio'],
        [
            ['Fee di attivazione', '500 €', 'Una tantum',
             'Setup sede, config menu, formazione staff (2–4 h), migrazione dati'],
            ['Canone fisso', '50 €/mese', 'Mensile',
             'Hosting Vercel, DB PostgreSQL, backup, aggiornamenti, supporto email'],
            ['Commissione transazionale', '1,5 % del transato', 'Mensile',
             'Sul valore totale degli scontrini elaborati dal sistema, IVA esclusa'],
        ],
        col_widths=[3.8, 3.0, 2.5, 8.7])

    h2(doc, 'Ricavo mensile per tipologia di cliente')
    body_para(doc,
        "Il ricavo mensile varia con il volume del cliente. La commissione dell'1,5% "
        "sul transato (GMV) si somma al canone fisso di 50 €. "
        "Formula: Ricavo mensile = 50 € + (coperti/giorno × scontrino medio ex-IVA × 22 gg × 1,5%)")
    data_table(doc,
        ['Tipologia', 'Coperti/giorno', 'Scontrino medio (IVA esclusa)',
         'GMV mensile', 'Commissione 1,5%', 'Ricavo mensile totale'],
        [
            ['Bar piccolo',      '30',  '4,50 €', '2.970 €',  '45 €',  '95 €'],
            ['Bar medio',        '60',  '6,50 €', '8.580 €',  '129 €', '179 €'],
            ['Mensa aziendale',  '120', '8,00 €', '21.120 €', '317 €', '367 €'],
            ['Grande struttura', '200', '10,00 €','44.000 €', '660 €', '710 €'],
        ],
        col_widths=[3.2, 2.8, 3.8, 2.8, 2.8, 2.6])
    info_box(doc,
        "22 giorni lavorativi/mese come riferimento. "
        "Il GMV (Gross Merchandise Value) è il valore lordo degli ordini elaborati, "
        "IVA esclusa, su cui si applica la commissione.",
        label='Ipotesi di calcolo:')

    h2(doc, 'Metriche SaaS chiave')
    body_para(doc,
        "Le metriche di riferimento sono calcolate sul profilo «Bar medio» "
        "(60 coperti/giorno, 6,50 € scontrino medio ex-IVA) come cliente rappresentativo.")
    data_table(doc,
        ['Metrica', 'Valore (bar medio)', 'Calcolo / Note'],
        [
            ['MRR (Monthly Recurring Revenue)',
             '179 €/cliente',
             '50 € canone + 8.580 € GMV × 1,5% = 50 + 129 = 179 €'],
            ['ARR (Annual Recurring Revenue)',
             '2.148 €/cliente/anno',
             'MRR 179 € × 12 mesi'],
            ['GMV annuo gestito (bar medio)',
             '102.960 €/anno',
             '8.580 € × 12 mesi di transato elaborato'],
            ['LTV — Lifetime Value (3 anni)',
             '6.944 €/cliente',
             '500 (setup) + 36 mesi × 179 € = 6.444 + 500'],
            ['LTV — Lifetime Value (5 anni)',
             '11.240 €/cliente',
             '500 (setup) + 60 mesi × 179 € = 10.740 + 500'],
            ['CAC (Customer Acquisition Cost)',
             'obiettivo < 200 €',
             'Demo + contratto + onboarding: ~3 h lavoro commerciale'],
            ['LTV / CAC ratio (3 anni)',
             '35× (eccellente)',
             '6.944 / 200 = 34,7 — soglia minima accettabile: 3×'],
            ['Payback period',
             '~1 mese',
             'CAC 200 € / MRR 179 € = 1,1 mesi per rientrare'],
            ['Churn rate target',
             '< 3 % mensile',
             'Vita media cliente: 1 / 3% = 33 mesi. Sopra 5%: non sostenibile'],
            ['Take rate effettiva',
             '~2,1 % del GMV',
             '(50 + 129) / 8.580 = 2,1% — include canone fisso'],
        ],
        col_widths=[5.5, 3.5, 9.0])

    h2(doc, 'Break-even — quanti clienti per coprire i costi operativi')
    body_para(doc,
        "Con MRR medio di 179 €/cliente (bar medio), il break-even si raggiunge "
        "con pochi clienti attivi grazie alla componente variabile sul transato.")
    data_table(doc,
        ['Scenario costi mensili', 'Costi stimati', 'Clienti break-even', 'MRR al break-even'],
        [
            ['Minimal (solo infrastruttura)',    '~100 €/mese', '1 cliente',  '179 €'],
            ['Standard (infrastruttura + 5h)',   '~350 €/mese', '2 clienti',  '358 €'],
            ['Growth (infrastruttura + 15h)',    '~750 €/mese', '5 clienti',  '895 €'],
            ['Scale (team part-time dedicato)',  '~1.500 €/mese','9 clienti', '1.611 €'],
        ],
        col_widths=[5.5, 3.5, 3.0, 6.0])
    info_box(doc,
        "Obiettivo realistico anno 1: 8–12 clienti. "
        "Con 10 bar medi attivi → MRR = 1.790 €, ARR = 21.480 €, GMV gestito ≈ 1 M€/anno. "
        "Il break-even operativo si raggiunge già al 5° cliente (scenario Growth).",
        label='Target anno 1:')

    h2(doc, 'ROI per il bar cliente')
    body_para(doc,
        "Il costo mensile per il bar è 50 € + 1,5% del transato. "
        "Per un bar medio con 8.580 € GMV mensile, il costo totale è 179 €/mese. "
        "Ecco i benefici concreti che giustificano abbondantemente questa spesa:")
    data_table(doc,
        ['Beneficio', 'Stima mensile', 'Come si calcola'],
        [
            ['Risparmio tempo gestione ordini',
             '+200 – 400 €/mese',
             '1–2 h/giorno × 22 gg × 10 €/h → cassa/operatore liberato'],
            ['Riduzione errori ordine',
             '+30 – 80 €/mese',
             '1–2 ordini errati/giorno × 2–5 € costo rifacimento'],
            ['Aumento scontrino medio (fidelizzazione)',
             '+50 – 200 €/mese',
             '+5–10% grazie a punti fedeltà, preordini e suggerimenti'],
            ['Riduzione sprechi (preordini anticipati)',
             '+30 – 100 €/mese',
             'Menu pianificato su dati reali → -10% cibo non venduto'],
            ['Totale beneficio stimato',
             '310 – 780 €/mese',
             'ROI per il bar: da 1,7× a 4,4× il costo del servizio (179 €)'],
        ],
        col_widths=[5.0, 3.5, 9.5])
    info_box(doc,
        "Script commerciale: «Con QuickLunch risparmia almeno 1 ora al giorno "
        "di gestione. Sono 200 € al mese di valore del tuo tempo. "
        "Il sistema costa 179 € al mese ma ne genera oltre 500 di beneficio misurabile.»",
        label='Argomento di vendita:')

    h2(doc, 'Proiezione ricavi — scenari di crescita')
    body_para(doc,
        "Ipotesi: mix di clienti (60% bar medi a 179 €/mese, 30% bar piccoli a 95 €, "
        "10% grandi strutture a 367 €/mese). MRR medio ponderato: ~179 €/cliente.")
    data_table(doc,
        ['Periodo', 'Clienti', 'MRR totale', 'ARR', 'GMV gestito/anno', 'Entrate cumul. (setup incluso)'],
        [
            ['Mese 3',   '3',  '537 €',   '6.444 €',    '309.000 €',  '3.111 €'],
            ['Mese 6',   '6',  '1.074 €', '12.888 €',   '618.000 €',  '9.444 €'],
            ['Mese 12',  '10', '1.790 €', '21.480 €',   '1.030.000 €','16.800 €'],
            ['Anno 2',   '16', '2.864 €', '34.368 €',   '1.648.000 €','54.000 €'],
            ['Anno 3',   '22', '3.938 €', '47.256 €',   '2.265.000 €','104.000 €'],
        ],
        col_widths=[2.2, 1.8, 2.5, 2.5, 4.5, 4.5])
    info_box(doc,
        "Il GMV gestito cresce proporzionalmente: con 22 clienti al 3° anno, "
        "QuickLunch elabora oltre 2,2 M€ di transazioni annue. "
        "La commissione dell'1,5% genera da sola ~34.000 € ARR, "
        "a cui si aggiungono i canoni fissi (~13.200 €) per un totale di ~47.000 € ARR.",
        label='Valore del transato:')


# ── Appendice D — Layout Fisici ───────────────────────────────────────────────

def _layout_zone(cell, label, bg, fg=None, size=10, bold=True):
    """Riempie una cella con colore di zona e testo."""
    _cell_shd(cell, bg)
    _cell_margins(cell, top=60, bottom=60, left=50, right=50)
    _cell_vAlign(cell, 'center')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(label), size=size, bold=bold,
              color=fg or RGBColor(0xff, 0xff, 0xff))


def _layout_title(doc, title, subtitle):
    p = doc.add_paragraph()
    _p_spacing(p, before=6, after=1)
    _run_font(p.add_run(title), size=12, bold=True, color=NAVY)
    _run_font(p.add_run('  —  ' + subtitle), size=10, color=GRAY)


def s_appendix_layouts(doc):
    h1(doc, 'D', 'Layout Fisici — Cucina, Sala, Cassa', '🏗️')
    body_para(doc,
        "Le planimetrie seguenti rappresentano tre configurazioni tipo per l'allestimento "
        "fisico di un locale che adotta QuickLunch. Ogni layout indica la posizione "
        "ottimale dei tablet KDS, della cassa, del bancone di distribuzione e delle zone di "
        "preparazione. I colori identificano le zone funzionali:")
    data_table(doc,
        ['Colore', 'Zona', 'Funzione'],
        [
            ['🟥 Rosso scuro',   'Cassa',      'Registratore di cassa, tablet gestione wallet, ricariche'],
            ['🟧 Arancio',       'Cucina',      'Preparazione piatti: caldo, freddo, builder, poke'],
            ['🟩 Verde',         'Sala',        'Tavoli clienti, zona ritiro vassoi, self-service bevande'],
            ['🟦 Blu scuro',     'Bancone',     'Distribuzione vassoi, ritiro ordini prenotati, display ordini'],
            ['⬜ Grigio chiaro', 'Ingresso/Corridoio', 'Flusso clienti, attesa, accesso'],
        ],
        col_widths=[3.0, 3.0, 12.0])

    # ── LAYOUT A ──────────────────────────────────────────────────────────────
    _layout_title(doc, 'Layout A — Bar compatto', '4 persone · 40–60 coperti · superficie ~80 m²')
    body_para(doc,
        "Ideale per un bar aziendale piccolo o una caffetteria con servizio pranzo. "
        "La cucina è unica e aperta, il builder è gestito dallo stesso addetto. "
        "Un solo tablet KDS è sufficiente. La cassa si trova vicino all'ingresso.")

    # Griglia planimetrica Layout A: 5 righe × 4 colonne
    tA = doc.add_table(rows=5, cols=4)
    _no_borders(tA)
    _table_width(tA, 18.0)
    for r in tA.rows:
        for c in r.cells:
            _cell_border(c, top='FFFFFF', bottom='FFFFFF', left='FFFFFF', right='FFFFFF', sz='4')

    # Riga 0: INGRESSO (merge 4 celle)
    ing = tA.rows[0].cells[0].merge(tA.rows[0].cells[3])
    _layout_zone(ing, '🚪  INGRESSO / FLUSSO CLIENTI', 'CCCCCC',
                 fg=RGBColor(0x33, 0x33, 0x33), size=10)

    # Riga 1: CASSA | BANCONE RITIRO (1+3)
    _layout_zone(tA.rows[1].cells[0], '🏪  CASSA\n(Cassiere)', 'C0392B', size=10)
    banc_a = tA.rows[1].cells[1].merge(tA.rows[1].cells[3])
    _layout_zone(banc_a, '📋  BANCONE DISTRIBUZIONE / RITIRO ORDINI\n(display: ordine pronto → cliente ritira)', '1F618D', size=10)

    # Riga 2: CUCINA piatti caldi (2) + BUILDER (2)
    cuc1 = tA.rows[2].cells[0].merge(tA.rows[2].cells[1])
    _layout_zone(cuc1, '🍳  CUCINA PIATTI CALDI\n(Cuoco 1 · KDS tablet)', 'E67E22', size=10)
    cuc2 = tA.rows[2].cells[2].merge(tA.rows[2].cells[3])
    _layout_zone(cuc2, '🥪  BUILDER + POKE\n(Cuoco 2 · piastra 🔥)', 'D35400', size=10)

    # Riga 3: magazzino/dispensa (1) + spazio (3)
    _layout_zone(tA.rows[3].cells[0], '📦\nDispensa', '7F8C8D', size=9)
    sp_a = tA.rows[3].cells[1].merge(tA.rows[3].cells[3])
    _layout_zone(sp_a, '', 'F0F0F0', fg=RGBColor(0x99, 0x99, 0x99))

    # Riga 4: SALA (merge 4)
    sala_a = tA.rows[4].cells[0].merge(tA.rows[4].cells[3])
    _layout_zone(sala_a,
        '🍽️  SALA — 8–10 TAVOLI (40–60 coperti)\n'
        'Addetto sala 1: distribuzione vassoi e risposta clienti',
        '1E8449', size=10)

    body_para(doc, '')
    info_box(doc,
        "Con 4 persone il KDS unico in cucina è visibile a entrambi i cuochi. "
        "Il tablet deve essere posizionato al centro del bancone di lavoro, "
        "a circa 80 cm di altezza, orientato in paesaggio.",
        label='Posizionamento tablet KDS:')

    # ── LAYOUT B ──────────────────────────────────────────────────────────────
    _layout_title(doc, 'Layout B — Mensa standard', '6 persone · 80–120 coperti · superficie ~150 m²')
    body_para(doc,
        "Configurazione tipica per mensa aziendale o bar di medie dimensioni. "
        "La cucina è divisa in due zone distinte (caldo / freddo-builder). "
        "Due tablet KDS: uno per i piatti caldi, uno per il builder e poke. "
        "La cassa gestisce anche le ricariche wallet al mattino.")

    tB = doc.add_table(rows=6, cols=6)
    _no_borders(tB)
    _table_width(tB, 18.0)
    for r in tB.rows:
        for c in r.cells:
            _cell_border(c, top='FFFFFF', bottom='FFFFFF', left='FFFFFF', right='FFFFFF', sz='4')

    # Riga 0: INGRESSO
    ing_b = tB.rows[0].cells[0].merge(tB.rows[0].cells[5])
    _layout_zone(ing_b, '🚪  INGRESSO — flusso cliente: cassa → bancone → sala', 'BBBBBB',
                 fg=RGBColor(0x22, 0x22, 0x22), size=10)

    # Riga 1: CASSA (2) | BANCONE DISTRIBUZIONE (4)
    cassa_b = tB.rows[1].cells[0].merge(tB.rows[1].cells[1])
    _layout_zone(cassa_b, '🏪  CASSA\n(Cassiere + ricariche)', 'C0392B', size=10)
    banc_b = tB.rows[1].cells[2].merge(tB.rows[1].cells[5])
    _layout_zone(banc_b,
        '📋  BANCONE DISTRIBUZIONE — 4 m\n'
        '(Addetto sala 1: consegna vassoi · Display ordini pronti)',
        '1F618D', size=10)

    # Riga 2: CUCINA CALDO (3) | CUCINA BUILDER (3)
    cc_b = tB.rows[2].cells[0].merge(tB.rows[2].cells[2])
    _layout_zone(cc_b, '🍳  CUCINA PIATTI CALDI\nCuoco 1 + Cuoco 2 · KDS tablet 1', 'E67E22', size=10)
    cb_b = tB.rows[2].cells[3].merge(tB.rows[2].cells[5])
    _layout_zone(cb_b, '🥪🍱  BUILDER + POKE BOWL\nCuoco 3 · piastra 🔥 · KDS tablet 2', 'D35400', size=10)

    # Riga 3: passavivande + lavaggio (2) + dispensa (2) + uscita cucina (2)
    _layout_zone(tB.rows[3].cells[0].merge(tB.rows[3].cells[1]),
                 '🚿  Lavaggio', '7F8C8D', size=9)
    _layout_zone(tB.rows[3].cells[2].merge(tB.rows[3].cells[3]),
                 '📦  Dispensa / Frigo', '7F8C8D', size=9)
    _layout_zone(tB.rows[3].cells[4].merge(tB.rows[3].cells[5]),
                 '↔️  Passavivande', 'A0A0A0', size=9)

    # Riga 4: corridoio
    corr_b = tB.rows[4].cells[0].merge(tB.rows[4].cells[5])
    _layout_zone(corr_b, 'corridoio di servizio', 'E8E8E8', fg=RGBColor(0x99, 0x99, 0x99), size=8, bold=False)

    # Riga 5: SALA
    sala_b = tB.rows[5].cells[0].merge(tB.rows[5].cells[5])
    _layout_zone(sala_b,
        '🍽️  SALA — 12–15 TAVOLI (80–120 coperti)\n'
        'Addetto sala 1: distribuzione · Addetto sala 2: tavoli e accoglienza',
        '1E8449', size=10)

    body_para(doc, '')
    info_box(doc,
        "Il KDS tablet 1 (piatti caldi) deve essere visibile a entrambi i cuochi del caldo. "
        "Il KDS tablet 2 (builder) va montato sul bancone del builder a portata di mano. "
        "Consigliato: supporto a parete con inclinazione 30°.",
        label='Posizionamento KDS — Layout B:')

    # ── LAYOUT C ──────────────────────────────────────────────────────────────
    _layout_title(doc, 'Layout C — Mensa grande / picco', '8 persone · 150+ coperti · superficie ~250 m²')
    body_para(doc,
        "Configurazione per picchi di servizio, eventi aziendali o mense con più linee. "
        "Due casse separate per ridurre le code, quattro zone cucina indipendenti, "
        "due tablet KDS, sala divisa in zone A e B per gestire il flusso.")

    tC = doc.add_table(rows=7, cols=8)
    _no_borders(tC)
    _table_width(tC, 18.0)
    for r in tC.rows:
        for c in r.cells:
            _cell_border(c, top='FFFFFF', bottom='FFFFFF', left='FFFFFF', right='FFFFFF', sz='4')

    # Riga 0: INGRESSO
    ing_c = tC.rows[0].cells[0].merge(tC.rows[0].cells[7])
    _layout_zone(ing_c, '🚪  INGRESSO PRINCIPALE — doppio flusso (cassa 1 / cassa 2)', 'AAAAAA',
                 fg=RGBColor(0x22, 0x22, 0x22), size=10)

    # Riga 1: CASSA 1 (2) | CASSA 2 (2) | BANCONE DISTRIBUZIONE (4)
    _layout_zone(tC.rows[1].cells[0].merge(tC.rows[1].cells[1]),
                 '🏪  CASSA 1\n(Cassiere principale)', 'C0392B', size=9)
    _layout_zone(tC.rows[1].cells[2].merge(tC.rows[1].cells[3]),
                 '🏪  CASSA 2\n(Ricariche wallet / supporto)', '922B21', size=9)
    _layout_zone(tC.rows[1].cells[4].merge(tC.rows[1].cells[7]),
                 '📋  BANCONE DISTRIBUZIONE — 6 m\n(Display ordini · Ritiro vassoi · Self-service bevande)',
                 '1F618D', size=9)

    # Riga 2: 4 zone cucina
    _layout_zone(tC.rows[2].cells[0].merge(tC.rows[2].cells[1]),
                 '🍳  CUCINA CALDO\nCoordinatore + Cuoco 1\nKDS tablet 1', 'E67E22', size=9)
    _layout_zone(tC.rows[2].cells[2].merge(tC.rows[2].cells[3]),
                 '🥘  CUCINA FREDDA\nCuoco 2\n(antipasti, insalate)', 'CA6F1E', size=9)
    _layout_zone(tC.rows[2].cells[4].merge(tC.rows[2].cells[5]),
                 '🥪  BUILDER\nCuoco 3 · piastra 🔥\nKDS tablet 2', 'D35400', size=9)
    _layout_zone(tC.rows[2].cells[6].merge(tC.rows[2].cells[7]),
                 '🍱  POKE & DESSERT\nCuoco 4', 'BA4A00', size=9)

    # Riga 3: retro cucina
    _layout_zone(tC.rows[3].cells[0].merge(tC.rows[3].cells[1]), '🚿  Lavaggio', '7F8C8D', size=9)
    _layout_zone(tC.rows[3].cells[2].merge(tC.rows[3].cells[3]), '📦  Dispensa', '7F8C8D', size=9)
    _layout_zone(tC.rows[3].cells[4].merge(tC.rows[3].cells[5]), '🧊  Cella frigo', '7F8C8D', size=9)
    _layout_zone(tC.rows[3].cells[6].merge(tC.rows[3].cells[7]), '↔️  Uscita servizio', 'A0A0A0', size=9)

    # Riga 4: corridoio
    _layout_zone(tC.rows[4].cells[0].merge(tC.rows[4].cells[7]),
                 'corridoio di servizio', 'E8E8E8', fg=RGBColor(0x99, 0x99, 0x99), size=8, bold=False)

    # Riga 5: SALA A (4) | SALA B (4)
    _layout_zone(tC.rows[5].cells[0].merge(tC.rows[5].cells[3]),
                 '🍽️  SALA A — 12 tavoli\nAddetto sala 1: distribuzione', '1E8449', size=9)
    _layout_zone(tC.rows[5].cells[4].merge(tC.rows[5].cells[7]),
                 '🍽️  SALA B — 8 tavoli\nAddetto sala 2: accoglienza e tavoli', '196F3D', size=9)

    # Riga 6: self-service
    _layout_zone(tC.rows[6].cells[0].merge(tC.rows[6].cells[7]),
                 '☕  ZONA SELF-SERVICE BEVANDE — acqua · caffè · bibite · distributore',
                 '2E86C1', size=9)

    body_para(doc, '')
    info_box(doc,
        "Con 8 persone il coordinatore KDS non cucina: monitora i tablet, "
        "avanza gli stati degli ordini e coordina i tempi tra le 4 zone. "
        "È la figura chiave per mantenere i tempi sotto i 10 minuti per slot "
        "anche con 150+ coperti.",
        label='Ruolo coordinatore KDS (Layout C):')

    h2(doc, 'Checklist allestimento tecnologico')
    data_table(doc,
        ['Elemento', 'Layout A', 'Layout B', 'Layout C', 'Note'],
        [
            ['Tablet KDS cucina',       '1',  '2',  '2',  'Consigliato: 10" min., schermo opaco anti-riflesso'],
            ['Tablet cassa/admin',      '1',  '1',  '2',  'Accesso admin, gestione wallet, report'],
            ['Tablet Banco POS',        '1',  '1',  '1',  'Tablet dedicato al bancone per il Cassiere: /admin/banco'],
            ['Schermo display sala',    '0',  '1',  '1',  'Monitor o TV 32" con ordini pronti visibili ai clienti'],
            ['Router WiFi cucina',      '1',  '1',  '2',  'Rete separata da quella clienti, segnale stabile'],
            ['Stampante scontrini',     '1',  '1',  '2',  'Termica, collegata al tablet cassa'],
            ['Supporti tablet a parete','0',  '2',  '4',  'Altezza 80 cm, inclinazione 30°, protezione IP54'],
        ],
        col_widths=[5.0, 2.0, 2.0, 2.0, 7.0])


# ── Impostazioni documento ────────────────────────────────────────────────────

def set_document_defaults(doc):
    # Margini ridotti per più testo per pagina
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)

    # Stile di default: PT Sans Narrow 14pt
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(14)

    # Header con nome documento
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _p_spacing(p, before=0, after=4)
        _para_border(p, bottom_color='DEE2E6', bottom_sz='4')
        r1 = p.add_run('QuickLunch  ')
        _run_font(r1, size=8, bold=True, color=RED)
        r2 = p.add_run('· Manuale del Proprietario')
        _run_font(r2, size=8, color=GRAY)

    # Footer con numero di pagina
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=4, after=0)
        _para_border(p, bottom_color=None)
        r = p.add_run('— ')
        _run_font(r, size=8, color=GRAY)
        # Campo numero pagina
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_el = p.add_run()._r
        run_el.append(fldChar1)
        run_el.append(instrText)
        run_el.append(fldChar3)
        _run_font(p.runs[-1], size=8, color=GRAY)
        r2 = p.add_run(' —')
        _run_font(r2, size=8, color=GRAY)


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_document_defaults(doc)

    build_cover(doc)
    build_toc(doc)

    sections = [s01, s02, s03, s04, s05, s06, s07,
                s08, s09, s10, s11, s12, s13, s14, s15,
                s16, s17, s18, s19,
                s_appendix, s_appendix_wallet,
                s_appendix_saas, s_appendix_layouts]

    for fn in sections:
        fn(doc)

    # Footer finale
    p = doc.add_paragraph()
    _p_spacing(p, before=20, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border(p, bottom_color=None)
    r = p.add_run('QuickLunch  ·  Sistema di gestione bar e mense  ·  v2.0  ·  Giugno 2026  ·  © 2024–26 DS Consulting')
    _run_font(r, size=9, color=GRAY)

    p = doc.add_paragraph()
    _p_spacing(p, before=4, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Assistenza:  Daniele Speziale — DS Consulting  ·  '
                  'dspeziale@gmail.com  ·  +39 352 0150489')
    _run_font(r, size=9.5, bold=True, color=RED)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] Documento salvato in: {OUT}')


if __name__ == '__main__':
    main()
