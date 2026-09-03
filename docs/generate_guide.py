#!/usr/bin/env python3
"""
Genera docs/guida_utente.docx — Guida Utente Completa per QuickLunch.

Ruoli coperti:
  1. Super Admin
  2. Admin Tenant (gestore locale)
  3. Cassiere / Cassa
  4. Cucina / KDS
  5. Sala (gestione tavoli)
  6. Cliente
  7. Dipendente Aziendale (convenzione pasto fisso)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
RED   = RGBColor(0xe9, 0x45, 0x60)
NAVY  = RGBColor(0x0f, 0x34, 0x60)
DARK  = RGBColor(0x16, 0x21, 0x3e)
DGRAY = RGBColor(0x34, 0x3a, 0x40)
GRAY  = RGBColor(0x6c, 0x75, 0x7d)
WHITE = RGBColor(0xff, 0xff, 0xff)
GREEN = RGBColor(0x27, 0xae, 0x60)
PURPL = RGBColor(0x8e, 0x44, 0xad)
TEAL  = RGBColor(0x16, 0xa0, 0x85)
ORNG  = RGBColor(0xe6, 0x7e, 0x22)

HEX_RED   = 'E94560'
HEX_NAVY  = '0F3460'
HEX_DARK  = '16213E'
HEX_LIGHT = 'F8F9FA'
HEX_WHITE = 'FFFFFF'
HEX_GREEN = '27AE60'
HEX_PURPL = '8E44AD'
HEX_TEAL  = '16A085'
HEX_ORNG  = 'E67E22'

FONT = 'PT Sans Narrow'
OUT  = os.path.join(os.path.dirname(__file__), 'manuali', 'guida_utente.docx')


# ── XML helpers ───────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None,
                 sz='8', space='0'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    sides = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, color in sides.items():
        el = OxmlElement(f'w:{side}')
        if color:
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    sz)
            el.set(qn('w:color'), color.lstrip('#'))
            el.set(qn('w:space'), space)
        else:
            el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_vAlign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(Cm(width_cm).emu / 914.4)))
        tcW.set(qn('w:type'), 'dxa')
        for old in tcPr.findall(qn('w:tcW')):
            tcPr.remove(old)
        tcPr.append(tcW)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        bdr.append(el)
    tblPr.append(bdr)


def _table_width(table, width_cm):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(Cm(width_cm).emu / 914.4)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _para_border(p, *, bottom_color=None, bottom_sz='6'):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)
    pBdr = OxmlElement('w:pBdr')
    if bottom_color:
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    bottom_sz)
        bot.set(qn('w:color'), bottom_color.lstrip('#'))
        pBdr.append(bot)
    pPr.append(pBdr)


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK'])
                  .WD_BREAK.PAGE)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _row_height(row, h_cm):
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(Cm(h_cm).emu / 914.4)))
    trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)


# ── Run helpers ───────────────────────────────────────────────────────────────

def _run_font(run, size=14, bold=False, color=None, italic=False, font=FONT):
    run.font.name   = font
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font)


def _p_spacing(p, before=0, after=6, line=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if line:
        p.paragraph_format.line_spacing = Pt(line)


# ── Document-level helpers ─────────────────────────────────────────────────────

def set_document_defaults(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(14)
    normal.font.color.rgb = DGRAY

    sect = doc.sections[0]
    sect.page_width   = Cm(21.0)
    sect.page_height  = Cm(29.7)
    sect.top_margin    = Cm(2.0)
    sect.bottom_margin = Cm(2.0)
    sect.left_margin   = Cm(2.2)
    sect.right_margin  = Cm(2.2)


def body_para(doc, text='', color=None, size=14, bold=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, after=4, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    _p_spacing(p, before=before, after=after)
    if text:
        run = p.add_run(text)
        _run_font(run, size=size, bold=bold, color=color or DGRAY)
    return p


def h1(doc, number, title, icon='', accent=HEX_RED):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _p_spacing(p, before=10, after=4)
    label = p.add_run(f'SEZIONE {number}   ')
    _run_font(label, size=9, bold=True, color=RED)
    if icon:
        ic = p.add_run(icon + '  ')
        _run_font(ic, size=20)
    run = p.add_run(title)
    _run_font(run, size=22, bold=True, color=NAVY)
    _para_border(p, bottom_color=accent, bottom_sz='8')
    return p


def h2(doc, text, color=None):
    p = doc.add_paragraph()
    _p_spacing(p, before=8, after=3)
    run = p.add_run(text)
    _run_font(run, size=16, bold=True, color=color or DARK)
    return p


def h3(doc, text, color=None):
    p = doc.add_paragraph()
    _p_spacing(p, before=5, after=2)
    run = p.add_run(text)
    _run_font(run, size=14, bold=True, color=color or NAVY)
    return p


def info_box(doc, text, style='tip', label=None):
    IC = {'tip': '💡', 'warning': '⚠️', 'success': '✅', 'info': 'ℹ️'}
    ic = IC.get(style, '💡')
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 17.6)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, 'F5F5F5')
    _cell_border(cell, top='E8E8E8', bottom='E8E8E8', right='E8E8E8', left=HEX_RED, sz='14')
    _cell_margins(cell, top=70, bottom=70, left=130, right=90)
    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=0)
    prefix_text = f'{ic}  '
    if label:
        prefix_text += f'{label}  '
    prefix = p.add_run(prefix_text)
    _run_font(prefix, size=13, bold=bool(label), color=RED)
    run = p.add_run(text)
    _run_font(run, size=13, color=DGRAY)
    return tbl


def info_box_color(doc, text, bg='EBF5FB', border=HEX_TEAL, icon='ℹ️'):
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 17.6)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, bg)
    _cell_border(cell, top=border, bottom=border, right=border, left=border, sz='6')
    _cell_margins(cell, top=70, bottom=70, left=130, right=90)
    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=0)
    r = p.add_run(f'{icon}  {text}')
    _run_font(r, size=13, color=DGRAY)
    return tbl


def step_row(doc, num, title, text, accent=HEX_RED):
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    _table_width(tbl, 17.6)

    nc = tbl.rows[0].cells[0]
    _set_col_width(tbl, 0, 1.2)
    _cell_shd(nc, accent)
    _cell_margins(nc, top=70, bottom=70, left=40, right=40)
    _cell_vAlign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    rn = pn.add_run(str(num))
    _run_font(rn, size=15, bold=True, color=WHITE)

    cc = tbl.rows[0].cells[1]
    _cell_shd(cc, 'F0F4F8')
    _cell_margins(cc, top=70, bottom=70, left=110, right=90)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=3)
    rt = pt.add_run(title + '\n')
    _run_font(rt, size=13, bold=True, color=DARK)
    rb = pt.add_run(text)
    _run_font(rb, size=12, color=GRAY)
    return tbl


def data_table(doc, headers, rows, col_widths=None, accent=HEX_RED):
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    _table_width(tbl, 17.6)

    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        _cell_shd(c, accent)
        _cell_border(c)
        _cell_margins(c, top=60, bottom=60, left=80, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        r = p.add_run(h)
        _run_font(r, size=11, bold=True, color=WHITE)
    if col_widths:
        for i, w in enumerate(col_widths):
            _set_col_width(tbl, i, w)

    for ri, row_data in enumerate(rows):
        bg = HEX_WHITE if ri % 2 == 0 else 'F8F9FA'
        dr = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c = dr.cells[ci]
            _cell_shd(c, bg)
            _cell_border(c, bottom='DEE2E6')
            _cell_margins(c, top=55, bottom=55, left=80, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            bold = (ci == 0)
            r = p.add_run(val)
            _run_font(r, size=11, bold=bold, color=DARK if bold else GRAY)
    return tbl


def role_badge(doc, icon, role_name, desc, bg_hex, text_hex='FFFFFF'):
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    _table_width(tbl, 17.6)
    _set_col_width(tbl, 0, 3.5)
    _set_col_width(tbl, 1, 14.1)

    bc = tbl.rows[0].cells[0]
    _cell_shd(bc, bg_hex)
    _cell_margins(bc, top=80, bottom=80, left=60, right=60)
    _cell_vAlign(bc, 'center')
    pb = bc.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pb, before=0, after=0)
    _run_font(pb.add_run(icon + '\n'), size=22)
    _run_font(pb.add_run(role_name), size=10, bold=True,
              color=RGBColor(*bytes.fromhex(text_hex)))

    dc = tbl.rows[0].cells[1]
    _cell_shd(dc, HEX_LIGHT)
    _cell_margins(dc, top=80, bottom=80, left=120, right=90)
    _cell_vAlign(dc, 'center')
    pd = dc.paragraphs[0]
    _p_spacing(pd, before=0, after=0)
    _run_font(pd.add_run(desc), size=13, color=DGRAY)
    return tbl


def workflow_table(doc, steps, accent=HEX_NAVY):
    ncols = len(steps)
    tbl = doc.add_table(rows=1, cols=ncols)
    _table_width(tbl, 17.6)
    for i, (icon, title, desc) in enumerate(steps):
        c = tbl.rows[0].cells[i]
        bg = accent if i % 2 == 0 else '1A2E50'
        _cell_shd(c, bg)
        _cell_border(c, right='16213E')
        _cell_margins(c, top=100, bottom=100, left=80, right=80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(p, before=0, after=4)
        _run_font(p.add_run(icon + '\n'), size=18)
        _run_font(p.add_run(title + '\n'), size=10, bold=True, color=WHITE)
        _run_font(p.add_run(desc), size=9, color=RGBColor(0xb0, 0xc4, 0xd8))
    return tbl


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=0)
    p.paragraph_format.space_after = Pt(pts)
    return p


def divider(doc):
    p = doc.add_paragraph()
    _p_spacing(p, before=4, after=4)
    _para_border(p, bottom_color='DEE2E6', bottom_sz='4')
    return p


# ── Cover page ────────────────────────────────────────────────────────────────

def build_cover(doc):
    tbl = doc.add_table(rows=3, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 21.0)
    _row_height(tbl.rows[0], 4.5)
    _row_height(tbl.rows[1], 19.0)
    _row_height(tbl.rows[2], 6.2)

    for row in tbl.rows:
        _cell_shd(row.cells[0], HEX_DARK)
        _cell_margins(row.cells[0], top=0, bottom=0, left=0, right=0)

    # ── Riga centrale ──
    cc = tbl.rows[1].cells[0]
    _cell_shd(cc, HEX_DARK)
    _cell_margins(cc, top=0, bottom=0, left=300, right=300)
    _cell_vAlign(cc, 'center')

    def cp(text='', size=14, bold=False, color=WHITE,
           align=WD_ALIGN_PARAGRAPH.CENTER, after=10):
        p = cc.add_paragraph()
        p.alignment = align
        _p_spacing(p, before=0, after=after)
        if text:
            r = p.add_run(text)
            _run_font(r, size=size, bold=bold, color=color)
        return p

    p0 = cc.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p0, before=0, after=16)
    _run_font(p0.add_run('🍽️'), size=56, color=WHITE)

    cp('QuickLunch', size=54, bold=True, color=RED, after=0)
    cp('PRANZO', size=54, bold=True, color=WHITE, after=8)
    cp('GUIDA UTENTE COMPLETA', size=16, bold=True,
       color=RGBColor(0xc0, 0xd0, 0xe0), after=24)
    cp('Sistema di gestione per bar, mense e caffetterie',
       size=13, color=RGBColor(0x90, 0xa8, 0xc0), after=28)

    # Ruoli
    roles = [
        ('👑 Super Admin', HEX_RED),
        ('🏢 Admin Tenant', HEX_NAVY),
        ('💳 Cassiere', HEX_GREEN),
        ('🍳 Cucina / KDS', HEX_ORNG),
        ('🪑 Sala', HEX_TEAL),
        ('👤 Cliente', HEX_PURPL),
        ('🏭 Dipendente Aziendale', '5D6D7E'),
    ]
    pr = cc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pr, before=0, after=20)
    for i, (name, _) in enumerate(roles):
        _run_font(pr.add_run(name), size=11, color=RGBColor(0xb0, 0xc8, 0xe0))
        if i < len(roles) - 1:
            _run_font(pr.add_run('  ·  '), size=10, color=RGBColor(0x50, 0x60, 0x70))

    cp('Versione 2.0  ·  2025', size=10,
       color=RGBColor(0x60, 0x70, 0x80), after=0)

    # ── Riga inferiore ──
    bc = tbl.rows[2].cells[0]
    _cell_shd(bc, HEX_RED)
    _cell_margins(bc, top=60, bottom=60, left=300, right=300)
    _cell_vAlign(bc, 'center')
    pb = bc.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pb, before=0, after=6)
    _run_font(pb.add_run('⚠️  DOCUMENTO RISERVATO — USO INTERNO'), size=12,
              bold=True, color=WHITE)
    pb2 = bc.add_paragraph()
    pb2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pb2, before=0, after=0)
    _run_font(pb2.add_run(
        'Questo documento contiene le procedure operative di QuickLunch. '
        'Non distribuire all\'esterno senza autorizzazione.'),
        size=10, color=RGBColor(0xff, 0xcc, 0xcc))


# ── Table of contents ─────────────────────────────────────────────────────────

def build_toc(doc):
    h2(doc, '📋  Indice dei contenuti')
    spacer(doc, 4)

    toc_items = [
        ('1', '👑  Super Admin', 'Configurazione piattaforma, creazione tenant, utenti staff'),
        ('2', '🏢  Admin Tenant', 'Gestione prodotti, slot, magazzino, convenzioni, report'),
        ('3', '💳  Cassiere / Cassa', 'Punto vendita, pagamenti, wallet — come pagare un caffè'),
        ('4', '🍳  Cucina / KDS', 'Schermo cucina, gestione ordini, stati preparazione'),
        ('5', '🪑  Sala', 'Prenotazioni tavoli, check-in, monitoraggio tempo sosta'),
        ('6', '👤  Cliente', 'Auto-ordine, wallet, fidelizzazione, prenotazione tavoli'),
        ('7', '🏭  Dipendente Aziendale', 'Prenotazione pasto fisso convenzione aziendale'),
    ]

    tbl = doc.add_table(rows=len(toc_items), cols=3)
    _no_borders(tbl)
    _table_width(tbl, 17.6)
    _set_col_width(tbl, 0, 1.2)
    _set_col_width(tbl, 1, 5.8)
    _set_col_width(tbl, 2, 10.6)

    for ri, (num, title, desc) in enumerate(toc_items):
        bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        row = tbl.rows[ri]
        _cell_shd(row.cells[0], HEX_RED)
        _cell_margins(row.cells[0], top=70, bottom=70, left=60, right=60)
        _cell_vAlign(row.cells[0], 'center')
        pn = row.cells[0].paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(pn, before=0, after=0)
        _run_font(pn.add_run(num), size=14, bold=True, color=WHITE)

        _cell_shd(row.cells[1], bg)
        _cell_margins(row.cells[1], top=70, bottom=70, left=100, right=60)
        _cell_border(row.cells[1], bottom='DEE2E6')
        pt = row.cells[1].paragraphs[0]
        _p_spacing(pt, before=0, after=0)
        _run_font(pt.add_run(title), size=12, bold=True, color=DARK)

        _cell_shd(row.cells[2], bg)
        _cell_margins(row.cells[2], top=70, bottom=70, left=100, right=60)
        _cell_border(row.cells[2], bottom='DEE2E6')
        pd = row.cells[2].paragraphs[0]
        _p_spacing(pd, before=0, after=0)
        _run_font(pd.add_run(desc), size=11, color=GRAY)


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 1 — SUPER ADMIN
# ══════════════════════════════════════════════════════════════════════════════

def s_super_admin(doc):
    h1(doc, '1', 'Super Admin', '👑')

    role_badge(doc, '👑', 'Super Admin',
               'Il Super Admin ha accesso illimitato alla piattaforma. '
               'Gestisce tenant, utenti globali, OAuth, bot Telegram, SMTP e configurazioni di sistema. '
               'Accede alla stessa interfaccia Admin ma con visibilità su tutti i locali.',
               HEX_RED)
    spacer(doc, 8)

    # ── 1.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.1  Primo accesso')
    body_para(doc, 'Le credenziali predefinite per il primo avvio della piattaforma sono:')
    spacer(doc, 4)

    data_table(doc,
        ['Campo', 'Valore predefinito'],
        [
            ['Username / Email', 'admin@bar.local'],
            ['Password',         'admin123'],
        ],
        col_widths=[5.0, 12.6])
    spacer(doc, 6)

    info_box(doc, 'Cambia subito la password predefinita dopo il primo accesso: '
             'Admin → Impostazioni → Modifica password. '
             'Non lasciare mai admin123 in produzione.', style='warning')
    spacer(doc, 8)

    step_row(doc, 1, 'Apri il browser', 'Vai all\'URL della piattaforma (es. https://pranzo.miodominio.it)')
    spacer(doc, 4)
    step_row(doc, 2, 'Inserisci le credenziali', 'Email admin@bar.local e password admin123 (o quella aggiornata)')
    spacer(doc, 4)
    step_row(doc, 3, 'Dashboard principale', 'Vedrai la Dashboard con KPI globali: ordini, wallet, prodotti attivi, avvisi magazzino')
    spacer(doc, 8)

    # ── 1.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.2  Creazione tenant (nuovo locale)')
    body_para(doc, 'Ogni locale (bar, mensa, caffetteria) è un "tenant" indipendente con propri '
              'prodotti, utenti e impostazioni. Percorso: Admin → Tenant → + Nuovo Tenant.')
    spacer(doc, 4)

    data_table(doc,
        ['Campo', 'Descrizione', 'Esempio'],
        [
            ['Nome locale',  'Nome commerciale visualizzato nell\'app', 'Bar Centrale'],
            ['Slug / Codice', 'Identificativo URL, solo lettere minuscole e trattini', 'bar-centrale'],
            ['Colore tema',  'Colore esadecimale del brand (usato nei badge e nella UI)', '#E94560'],
            ['Logo',         'File immagine PNG/JPG del logo (opzionale)', 'logo.png'],
            ['Email contatto', 'Email del gestore principale per notifiche di sistema', 'info@barcentrale.it'],
        ],
        col_widths=[3.5, 8.0, 6.1])
    spacer(doc, 8)

    info_box(doc, 'Lo slug deve essere univoco: viene usato come prefisso URL e nei cookie di sessione. '
             'Una volta impostato non modificarlo senza aggiornare anche le configurazioni del server.',
             style='warning')
    spacer(doc, 8)

    # ── 1.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.3  Configurazione Google OAuth')
    body_para(doc, 'Per abilitare il login con Google, inserisci le credenziali OAuth 2.0 ottenute '
              'dalla Google Cloud Console (APIs & Services → Credentials).')
    spacer(doc, 6)

    data_table(doc,
        ['Variabile / Campo', 'Dove inserirla', 'Valore da inserire'],
        [
            ['GOOGLE_CLIENT_ID',     'Admin → Impostazioni → OAuth Google', 'ID client dalla Google Console'],
            ['GOOGLE_CLIENT_SECRET', 'Admin → Impostazioni → OAuth Google', 'Secret client dalla Google Console'],
            ['Callback URL',         'Google Console → URI autorizzati',    '/auth/google/callback'],
        ],
        col_widths=[4.5, 6.5, 6.6])
    spacer(doc, 6)

    info_box_color(doc,
                   'URI di reindirizzamento autorizzato da configurare in Google Console:\n'
                   'https://tuodominio.it/auth/google/callback',
                   bg='EBF5FB', border=HEX_TEAL, icon='🔗')
    spacer(doc, 8)

    # ── 1.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.4  Configurazione Telegram Bot')
    body_para(doc, 'Le notifiche Telegram (nuovi ordini, avvisi magazzino, scadenze tavoli) '
              'richiedono un bot Telegram per tenant. Usa @BotFather su Telegram per creare il bot '
              'e ottenere il token.')
    spacer(doc, 4)

    step_row(doc, 1, 'Crea il bot con @BotFather', 'Su Telegram: /newbot → scegli nome → ottieni token (es. 123456:ABCdef…)')
    spacer(doc, 4)
    step_row(doc, 2, 'Inserisci il token nel pannello', 'Admin → Impostazioni → Telegram Bot Token → Salva')
    spacer(doc, 4)
    step_row(doc, 3, 'Aggiungi il bot al gruppo notifiche', 'Crea un gruppo Telegram, aggiungici il bot e ottieni il Chat ID')
    spacer(doc, 4)
    step_row(doc, 4, 'Inserisci il Chat ID', 'Admin → Impostazioni → Telegram Chat ID → Salva')
    spacer(doc, 8)

    # ── 1.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.5  Configurazione SMTP Email')
    body_para(doc, 'Le email automatiche (avvisi magazzino al fornitore, conferme ordine) '
              'richiedono la configurazione SMTP. Percorso: Admin → Impostazioni → Email SMTP.')
    spacer(doc, 4)

    data_table(doc,
        ['Campo', 'Esempio'],
        [
            ['SMTP Host',     'smtp.gmail.com'],
            ['SMTP Port',     '587 (TLS) / 465 (SSL)'],
            ['Username',      'notifiche@tuodominio.it'],
            ['Password',      'app-password generata dal provider'],
            ['Email mittente', 'QuickLunch <notifiche@tuodominio.it>'],
        ],
        col_widths=[4.5, 13.1])
    spacer(doc, 8)

    # ── 1.6 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.6  Gestione admin globali')
    body_para(doc, 'Per aggiungere o modificare un amministratore globale: Admin → Utenti Staff.')
    spacer(doc, 4)

    step_row(doc, 1, 'Modifica password admin esistente', 'Admin → Utenti Staff → trova l\'admin → "Modifica" → nuovo campo password → Salva')
    spacer(doc, 4)
    step_row(doc, 2, 'Aggiungi nuovo Super Admin', 'Admin → Utenti Staff → + Nuovo Staff → compila email, username, password → spunta "is_admin" e "superadmin" → Salva')
    spacer(doc, 8)

    info_box(doc, 'Il flag "superadmin" garantisce visibilità su tutti i tenant. '
             'Il solo flag "is_admin" senza "superadmin" crea un Admin Tenant limitato al proprio locale.',
             style='info')
    spacer(doc, 8)

    # ── 1.7 ──────────────────────────────────────────────────────────────────
    h2(doc, '1.7  CLI: flask seed-demo')
    body_para(doc, 'Il comando seed-demo resetta il database e ricarica dati di dimostrazione '
              'preconfigurati: prodotti, utenti, slot, convenzioni e ordini di esempio. '
              'Utile per ambienti di test o demo commerciali.')
    spacer(doc, 4)

    info_box_color(doc,
                   'Esegui dalla directory radice del progetto:\n'
                   'flask seed-demo\n\n'
                   'ATTENZIONE: cancella TUTTI i dati esistenti prima di reinserire i dati demo.',
                   bg='FEF9E7', border=HEX_ORNG, icon='⚠️')
    spacer(doc, 6)

    data_table(doc,
        ['Cosa crea', 'Dettaglio'],
        [
            ['Tenant demo',        'Un locale configurato con prodotti, slot e fasce orarie'],
            ['Utenti staff',       'Admin, cassiere, cuoco, sala (vedi Appendice A.3 per le credenziali)'],
            ['Clienti demo',       'cliente1 e cliente2 con wallet precaricato'],
            ['Convenzione demo',   'Azienda ACME con dipendente associato e pasto del giorno'],
            ['Ordini e transazioni', 'Storico degli ultimi 7 giorni per report e test'],
        ],
        col_widths=[4.5, 13.1])
    spacer(doc, 8)

    info_box(doc, 'Prima di ogni aggiornamento di produzione, esegui sempre il backup del database. '
             'I run di migrazione automatici (_ensure) aggiungono colonne in sicurezza senza perdere dati.',
             style='warning')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 2 — ADMIN TENANT
# ══════════════════════════════════════════════════════════════════════════════

def s_admin_tenant(doc):
    h1(doc, '2', 'Admin / Manager del Locale', '🏢')

    role_badge(doc, '🏢', 'Admin Tenant',
               'Gestisce tutto ciò che riguarda il proprio locale: menu, prodotti, slot ordini, '
               'fasce orarie tavoli, stock giornaliero, magazzino, staff, clienti, '
               'convenzioni aziendali, sondaggi e report economici.',
               HEX_NAVY)
    spacer(doc, 8)

    # ── 2.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.1  Dashboard KPI')
    body_para(doc, 'La Dashboard è la prima schermata dopo il login Admin. '
              'Mostra in tempo reale tutti i KPI del locale.')
    spacer(doc, 6)

    data_table(doc,
        ['KPI', 'Icona', 'Cosa indica'],
        [
            ['Ordini oggi',           '🛒', 'Numero di ordini ricevuti oggi'],
            ['Incasso del giorno',    '💰', 'Totale incassato oggi (contanti + wallet)'],
            ['Wallet totale clienti', '💜', 'Somma di tutti i saldi wallet dei clienti'],
            ['Prenotazioni oggi',     '🪑', 'Tavoli prenotati per la giornata corrente'],
            ['Alert magazzino',       '🔴', 'Materiali consumabili sotto la soglia minima'],
            ['Prodotti attivi',       '📦', 'Prodotti disponibili e visibili nel menu'],
            ['Pasti Aziendali Oggi',  '🏭', 'Card per ogni convenzione: barra avanzamento prenotazioni/posti, '
                                            'posti residui e (per il dipendente loggato) il proprio pasto prenotato'],
        ],
        col_widths=[4.8, 1.5, 11.3])
    spacer(doc, 8)

    # ── 2.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.2  Gestione Prodotti e Categorie')
    workflow_table(doc, [
        ('📝', 'Crea prodotto', 'Nome, prezzo, categoria'),
        ('🏷️', 'Allergeni', '14 tipologie selezionabili'),
        ('✅', 'Attiva/Disattiva', 'Senza eliminare'),
        ('📦', 'Verifica', 'Appare nel menu cliente'),
    ])
    spacer(doc, 8)

    h3(doc, 'Aggiungere o modificare un prodotto')
    step_row(doc, 1, 'Vai a Menu → Prodotti', 'Nel menu laterale Admin, sezione "Menu"')
    spacer(doc, 4)
    step_row(doc, 2, 'Clicca "+ Nuovo Prodotto"', 'Si apre il form di creazione')
    spacer(doc, 4)
    step_row(doc, 3, 'Compila i campi obbligatori', 'Nome, prezzo, categoria. Il campo "is_active" controlla la visibilità nel menu')
    spacer(doc, 4)
    step_row(doc, 4, 'Gestisci gli allergeni', 'Spunta le 14 tipologie applicabili (glutine, lattosio, arachidi…)')
    spacer(doc, 4)
    step_row(doc, 5, 'Salva e verifica', 'Il prodotto appare immediatamente nel menu cliente se is_active = true')
    spacer(doc, 8)

    info_box(doc, 'I prodotti possono essere disattivati (is_active = false) senza eliminarli: '
             'lo storico ordini rimane integro. Utile per prodotti stagionali o temporaneamente esauriti.',
             style='tip')
    spacer(doc, 8)

    # ── 2.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.3  Stock Giornaliero')
    body_para(doc, 'Ogni mattina l\'Admin imposta la disponibilità giornaliera dei prodotti '
              '(quante porzioni sono preparate). Quando esaurito, il prodotto si disattiva automaticamente.')
    spacer(doc, 6)

    step_row(doc, 1, 'Vai a Menu → Stock Giornaliero', 'Sezione dedicata nel menu Admin')
    spacer(doc, 4)
    step_row(doc, 2, 'Imposta le porzioni per ogni prodotto', 'Inserisci il numero di porzioni disponibili per oggi')
    spacer(doc, 4)
    step_row(doc, 3, 'Configura soglia minima magazzino', 'Magazzino → Consumabili: imposta la soglia sotto cui scatta l\'alert fornitore')
    spacer(doc, 4)
    step_row(doc, 4, 'Aggiungi articoli consumabili', 'Magazzino → + Nuovo: nome, unità (pz/kg/lt), soglia minima, fornitore associato')
    spacer(doc, 4)
    step_row(doc, 5, 'Aggiungi fornitori', 'Magazzino → Fornitori → + Nuovo: nome e email. Riceveranno avvisi automatici sotto soglia')
    spacer(doc, 8)

    info_box(doc, 'Il sistema invia UNA SOLA email di avviso per ogni carenza. '
             'Una seconda email parte solo dopo che le scorte sono tornate sopra soglia '
             'e poi sono nuovamente scese. Questo evita spam al fornitore.', style='info')
    spacer(doc, 8)

    # ── 2.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.4  Slot Ordini e Fasce Orarie Tavoli')

    info_box_color(doc,
                   'ATTENZIONE — Questi sono DUE SISTEMI SEPARATI e indipendenti:\n'
                   '• SLOT ORDINI: orari di ritiro del cibo, capienza cucina\n'
                   '• FASCE ORARIE TAVOLI: prenotazione posti a sedere, sessioni tavolo\n'
                   'Non si sovrappongono né si influenzano a vicenda.',
                   bg='FEF9E7', border=HEX_ORNG, icon='⚠️')
    spacer(doc, 8)

    data_table(doc,
        ['Sistema', 'Scopo', 'Percorso', 'Contenuto'],
        [
            ['Slot Ordini',
             'Orario di ritiro del cibo. Capacità max ordini per non sovraccaricare la cucina.',
             '/admin/tavoli → tab "Slot Ordini"',
             'Orario (time picker) + capacità massima ordini'],
            ['Fasce Orarie Tavoli\n(TableTimeBand)',
             'Blocchi di tempo per prenotare un posto a sedere. Risorse GLOBALI del locale.',
             '/admin/tavoli → tab "Fasce Orarie"',
             'Nome, orario inizio/fine, sort_order per ordinamento'],
        ],
        col_widths=[3.5, 5.5, 4.5, 4.1])
    spacer(doc, 8)

    h3(doc, 'Aggiungere e rimuovere Slot Ordini')
    step_row(doc, 1, 'Apri il tab "Slot Ordini"', 'Admin → Tavoli → tab Slot Ordini')
    spacer(doc, 4)
    step_row(doc, 2, 'Compila il form in cima alla pagina', 'Inserisci l\'orario con il time picker e la capacità massima ordini per quello slot')
    spacer(doc, 4)
    step_row(doc, 3, 'Clicca "Aggiungi slot"', 'Lo slot viene creato immediatamente e appare nella lista sottostante')
    spacer(doc, 4)
    step_row(doc, 4, 'Eliminare uno slot esistente', 'Clicca il pulsante "Elimina" a fianco dello slot — viene rimosso senza conferma aggiuntiva')
    spacer(doc, 8)

    h3(doc, 'Configurare Fasce Orarie Tavoli')
    body_para(doc, 'Le fasce orarie generano sessioni prenotabili: una fascia 11:25–12:30 con 30 min '
              'a seduta crea le sessioni 11:25, 11:55 e 12:25. Percorso: Admin → Tavoli → tab "Fasce Orarie".')
    spacer(doc, 4)

    info_box(doc, 'La durata di permanenza (campo nella fascia oraria) determina quando viene inviata '
             'la notifica Telegram al responsabile sala: l\'avviso parte 10 minuti prima della scadenza.',
             style='warning')
    spacer(doc, 8)

    # ── 2.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.5  Articoli Banco POS')
    body_para(doc, 'Il Banco POS è la griglia di articoli rapidi per pagamenti al bancone via QR. '
              'L\'Admin configura gli articoli disponibili dalla pagina /admin/banco/items '
              '(oppure Admin → Banco → pulsante "Gestisci").')
    spacer(doc, 6)

    data_table(doc,
        ['Azione', 'Come fare', 'Nota'],
        [
            ['Aggiungere articolo',
             'Form inline: nome, prezzo, icona Font Awesome (es. fa-coffee), colore hex → "Salva"',
             'Appare subito nella griglia del Banco'],
            ['Modificare articolo',
             'Clicca sull\'articolo nella lista → si apre modal di modifica → "Salva"',
             'Modifica immediata senza ricaricare la pagina'],
            ['Disattivare articolo',
             'Pulsante "Elimina" a fianco dell\'articolo',
             'Soft-delete: is_active=false, dati storici conservati'],
        ],
        col_widths=[3.8, 8.2, 5.6])
    spacer(doc, 8)

    # ── 2.6 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.6  Personale Staff')
    body_para(doc, 'Percorso: Admin → Utenti Staff. Da qui si gestisce tutto il personale del locale.')
    spacer(doc, 4)

    step_row(doc, 1, 'Aggiungi nuovo membro staff', 'Admin → Utenti Staff → + Nuovo Staff: inserisci email, username, password e ruolo')
    spacer(doc, 4)
    step_row(doc, 2, 'Assegna ruoli e permessi', 'Spunta i flag appropriati: manage_orders (cassiere), manage_kitchen (cucina), manage_reservations (sala)')
    spacer(doc, 4)
    step_row(doc, 3, 'Modifica profilo cliente (anagrafica estesa)', 'Admin → Clienti → trova utente → Modifica: nome, cognome, telefono, data di nascita')
    spacer(doc, 8)

    data_table(doc,
        ['Ruolo', 'Flag da attivare', 'Accesso'],
        [
            ['Admin Tenant',    'is_admin',              'Tutta l\'area Admin del locale'],
            ['Cassiere',        'manage_orders',         'Pannello ordini, banco POS, wallet clienti'],
            ['Cucina / KDS',    'manage_kitchen',        'Schermata KDS, stati ordine'],
            ['Sala',            'manage_reservations',   'Prenotazioni tavoli, check-in, ping-alerts'],
        ],
        col_widths=[3.5, 4.5, 9.6])
    spacer(doc, 8)

    # ── 2.7 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.7  Gestione Clienti')
    step_row(doc, 1, 'Ricarica wallet cliente', 'Admin → Clienti → trova cliente → "Ricarica Wallet" → inserisci importo → Conferma')
    spacer(doc, 4)
    step_row(doc, 2, 'Visualizza storico transazioni', 'Admin → Clienti → trova cliente → tab "Transazioni"')
    spacer(doc, 4)
    step_row(doc, 3, 'Visualizza anagrafica', 'Admin → Clienti → trova cliente: nome, cognome, email, telefono, data nascita, saldo wallet, punti fedeltà')
    spacer(doc, 8)

    # ── 2.8 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.8  Convenzioni Aziendali (Corporate Meals)')
    body_para(doc, 'Le convenzioni permettono a dipendenti di aziende convenzionate di prenotare '
              'un pasto fisso giornaliero a prezzo concordato, separato dal menu standard.')
    spacer(doc, 4)

    step_row(doc, 1, 'Aggiungi azienda', 'Convenzioni → Aziende → + Nuova: nome azienda, codice identificativo, email di contatto')
    spacer(doc, 4)
    step_row(doc, 2, 'Associa dipendenti', 'Nella scheda azienda, spunta i clienti già registrati da aggiungere come dipendenti')
    spacer(doc, 4)
    step_row(doc, 3, 'Crea il pasto del giorno', 'Convenzioni → nome azienda → "Pasto del giorno": '
             'primo, secondo, contorno, bevanda, caffè, allergeni, max porzioni, prezzo, data')
    spacer(doc, 4)
    step_row(doc, 4, 'Monitora dashboard pasti', 'La scheda mostra per ogni opzione: '
             'prenotazioni effettuate, barra avanzamento, posti residui')
    spacer(doc, 4)
    step_row(doc, 5, 'Segna come consumato', 'Lista prenotazioni → "✔ Consumato" per ogni dipendente che ha ritirato il pasto')
    spacer(doc, 4)
    step_row(doc, 6, 'Verifica totale fatturabile', 'In fondo alla lista prenotazioni: N coperti × prezzo = importo da fatturare all\'azienda')
    spacer(doc, 8)

    # ── 2.9 ──────────────────────────────────────────────────────────────────
    h2(doc, '2.9  Sondaggi')
    body_para(doc, 'I sondaggi permettono di raccogliere preferenze dai clienti sul menu. '
              'Percorso: Admin → Sondaggi.')
    spacer(doc, 4)

    step_row(doc, 1, 'Crea sondaggio', 'Admin → Sondaggi → + Nuovo: inserisci domanda e almeno 2 opzioni di risposta')
    spacer(doc, 4)
    step_row(doc, 2, 'Apertura manuale', 'Nella lista sondaggi, clicca "Apri" per renderlo visibile ai clienti')
    spacer(doc, 4)
    step_row(doc, 3, 'Chiusura manuale', 'Clicca "Chiudi" per smettere di raccogliere voti')
    spacer(doc, 4)
    step_row(doc, 4, 'Risultati in tempo reale', 'Nella scheda sondaggio: barre percentuale per ogni opzione, aggiornate ad ogni voto')
    spacer(doc, 8)

    # ── 2.10 ─────────────────────────────────────────────────────────────────
    h2(doc, '2.10  Report')
    body_para(doc, 'I report coprono gli ultimi 30 giorni e includono vendite, prodotti più richiesti '
              'e andamento dell\'incasso giornaliero. Percorso: Admin → Report.')
    spacer(doc, 4)

    data_table(doc,
        ['Report', 'Percorso', 'Contenuto'],
        [
            ['Ultimi 30 giorni',      'Admin → Report → Mensile',      'Totale ordini e incasso per ciascuno degli ultimi 30 giorni'],
            ['Top prodotti',          'Admin → Report → Top prodotti',  'Ranking prodotti per quantità venduta e fatturato'],
            ['Incasso giornaliero',   'Admin → Report → Giornaliero',   'Dettaglio ordini e metodi di pagamento del giorno'],
            ['Wallet e fedeltà',      'Admin → Clienti',                'Saldo wallet e punti per ogni cliente'],
            ['Magazzino',             'Admin → Magazzino',              'Stock attuale e storico movimenti consumabili'],
        ],
        col_widths=[4.0, 5.5, 8.1])
    spacer(doc, 8)


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 3 — CASSIERE / CASSA
# ══════════════════════════════════════════════════════════════════════════════

def s_cassiere(doc):
    h1(doc, '3', 'Cassiere / Cassa', '💳')

    role_badge(doc, '💳', 'Cassiere',
               'Il Cassiere gestisce il punto vendita fisico: riceve ordini walk-in, '
               'gestisce i pagamenti tramite Banco POS con QR, ricarica i wallet dei clienti '
               'e annulla ordini se necessario. È la figura più operativa del sistema.',
               HEX_GREEN)
    spacer(doc, 8)

    info_box(doc,
             'FOCUS OPERATIVO: padroneggia tre azioni fondamentali: '
             '1) ordini walk-in al pannello ordini, '
             '2) pagamenti rapidi con il Banco POS / QR, '
             '3) ricarica wallet cliente.',
             style='warning', label='FOCUS OPERATIVO')
    spacer(doc, 10)

    # ── 3.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '3.1  Accesso pannello ordini')
    step_row(doc, 1, 'Apri il pannello ordini', 'Vai su /admin/ordini oppure Admin → Ordini nel menu laterale')
    spacer(doc, 4)
    step_row(doc, 2, 'Login con credenziali cassiere', 'Username e password forniti dall\'Admin (vedi account di test in Appendice A.3)')
    spacer(doc, 4)
    step_row(doc, 3, 'Panoramica ordini', 'Vedi tutti gli ordini del giorno: in attesa, in preparazione, pronti, consegnati')
    spacer(doc, 8)

    # ── 3.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '3.2  Gestione ordini walk-in al banco')
    body_para(doc, 'Per i clienti che ordinano direttamente al bancone senza usare l\'app cliente:')
    spacer(doc, 4)

    step_row(doc, 1, 'Apri nuovo ordine manuale', 'Admin → Ordini → "+ Nuovo ordine" — non richiede che il cliente sia registrato')
    spacer(doc, 4)
    step_row(doc, 2, 'Seleziona prodotti dal catalogo', 'Clicca sui prodotti. Usa la barra di ricerca per trovare velocemente')
    spacer(doc, 4)
    step_row(doc, 3, 'Modifica quantità', 'Clicca sul prodotto nel riepilogo per aumentare o diminuire la quantità')
    spacer(doc, 4)
    step_row(doc, 4, 'Aggiungi note speciali (opzionale)', 'Es. "senza zucchero", "al latte di soia" — appariranno visibili in cucina')
    spacer(doc, 4)
    step_row(doc, 5, 'Invia in cucina', 'Clicca "Invia ordine" — l\'ordine appare sullo schermo KDS della cucina')
    spacer(doc, 4)
    step_row(doc, 6, 'Incassa quando il cliente ritira', 'Seleziona metodo di pagamento: CONTANTI, WALLET o MISTO')
    spacer(doc, 8)

    # ── 3.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '3.3  Banco POS — Pagamento QR (8 step)')
    body_para(doc, 'Il Banco POS è la modalità di cassa rapida: lo staff compone un carrello '
              'veloce su una griglia di articoli e genera un QR che il cliente scansiona '
              'con la sua app per pagare istantaneamente dal wallet. '
              'Percorso: Admin → Banco (icona tazza nel menu laterale).')
    spacer(doc, 6)

    info_box_color(doc,
                   'SCENARIO: Un cliente chiede un caffè e una brioche. '
                   'Lo staff tocca gli articoli nella griglia, genera il QR '
                   'e il cliente paga con il telefono in pochi secondi.',
                   bg='EAF7EA', border=HEX_GREEN, icon='☕')
    spacer(doc, 8)

    step_row(doc, 1, 'Aprire /admin/banco', 'Vai in Admin → Banco. La schermata è divisa: griglia articoli a sinistra, riepilogo + QR a destra.', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca gli articoli nella griglia', 'Ogni tocco aggiunge 1 unità. Il badge numerico sul pulsante mostra la quantità corrente.', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 3, 'Rimuovi articoli se necessario', 'Nel riepilogo a destra, usa il pulsante "−" accanto all\'articolo per sottrarre unità.', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 4, 'Verifica il totale', 'Il totale aggiornato in tempo reale appare sotto il riepilogo.', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 5, 'Tocca "Genera QR"', 'Si apre un modal con: QR code, totale in grande, countdown 10 minuti, status "In attesa…"', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 6, 'Mostra il QR al cliente', 'Il modal resta aperto in attesa — il cliente apre l\'app e scansiona il QR.', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 7, 'Pagamento confermato', 'Il modal mostra "✓ Pagato da [Nome Cliente]" — il carrello si svuota automaticamente.', accent=HEX_GREEN)
    spacer(doc, 4)
    step_row(doc, 8, 'QR scaduto: rigenerare', 'Se il countdown arriva a zero prima della scansione, chiudi il modal e tocca di nuovo "Genera QR".', accent=HEX_GREEN)
    spacer(doc, 8)

    info_box(doc, 'Il QR scade dopo 10 minuti (countdown visibile nel modal). '
             'Per annullare la sessione in corso senza attendere la scadenza, '
             'usa il pulsante "Annulla sessione" nel modal.', style='warning')
    spacer(doc, 8)

    # ── 3.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '3.4  Situazioni speciali al banco')
    data_table(doc,
        ['Situazione', 'Cosa fare', 'Nota'],
        [
            ['Cliente senza app',
             'Usa il pannello ordini tradizionale (sezione 3.2)',
             'Il cliente non ha bisogno di essere registrato'],
            ['QR scaduto',
             'Chiudi il modal → tocca "Genera QR" di nuovo',
             'Il carrello rimane invariato'],
            ['Saldo insufficiente',
             'Il cliente non riesce a pagare il QR: proponi ricarica wallet (sezione 3.5) o pagamento contanti',
             'Il banco POS scala dal wallet del cliente'],
            ['Errore di rete / pagamento non arriva',
             'Attendi 30 secondi, poi tocca "Annulla sessione" e rigenerai il QR',
             'Verifica connessione internet del dispositivo banco'],
        ],
        col_widths=[3.8, 7.2, 6.6])
    spacer(doc, 8)

    # ── 3.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '3.5  Ricarica wallet cliente')
    step_row(doc, 1, 'Vai in Admin → Clienti', 'Oppure usa la ricerca rapida nella barra in alto')
    spacer(doc, 4)
    step_row(doc, 2, 'Trova il cliente', 'Cerca per nome o email')
    spacer(doc, 4)
    step_row(doc, 3, 'Clicca "Ricarica Wallet"', 'Si apre il pannello di ricarica')
    spacer(doc, 4)
    step_row(doc, 4, 'Inserisci l\'importo', 'Es. 20,00 € — il cliente paga in contanti o carta fisica')
    spacer(doc, 4)
    step_row(doc, 5, 'Conferma', 'Il saldo viene aggiornato immediatamente. Il cliente riceve notifica Telegram (se configurato)')
    spacer(doc, 8)

    # ── 3.6 ──────────────────────────────────────────────────────────────────
    h2(doc, '3.6  Annullamento ordine dal pannello')
    step_row(doc, 1, 'Apri Admin → Ordini', 'Vai al pannello ordini')
    spacer(doc, 4)
    step_row(doc, 2, 'Trova l\'ordine da annullare', 'Cerca per numero ordine o nome cliente')
    spacer(doc, 4)
    step_row(doc, 3, 'Clicca "Annulla"', 'Pulsante disponibile solo se l\'ordine è ancora in stato "In attesa" o "In preparazione"')
    spacer(doc, 4)
    step_row(doc, 4, 'Conferma', 'L\'ordine viene marcato come "Annullato" — eventuale addebito wallet viene stornato automaticamente')
    spacer(doc, 8)

    info_box(doc, 'Non è possibile annullare un ordine già in stato "Pronto" o "Consegnato". '
             'In quel caso contatta l\'Admin per una correzione manuale.', style='warning')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 4 — CUCINA / KDS
# ══════════════════════════════════════════════════════════════════════════════

def s_cucina(doc):
    h1(doc, '4', 'Cucina / KDS — Kitchen Display System', '🍳')

    role_badge(doc, '🍳', 'Cucina / KDS',
               'La cucina riceve gli ordini in tempo reale sullo schermo KDS. '
               'L\'obiettivo è preparare gli ordini nell\'ordine corretto e aggiornare '
               'lo stato perché sala e cassiere sappiano sempre cosa è pronto.',
               HEX_ORNG)
    spacer(doc, 8)

    # ── 4.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '4.1  Accesso schermata KDS')
    step_row(doc, 1, 'Apri il browser sullo schermo cucina', 'Vai su /admin/ordini (oppure Admin → Cucina/KDS nel menu)', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 2, 'Login con credenziali cucina', 'Username e password del profilo Cucina forniti dall\'Admin', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 3, 'Attiva la modalità full screen', 'Premi F11 (o il pulsante "Full Screen" se disponibile) per massimizzare la leggibilità', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 4, 'Schermo sempre attivo', 'La pagina KDS si aggiorna automaticamente ogni 30 secondi. Non chiudere il browser durante il servizio.', accent=HEX_ORNG)
    spacer(doc, 8)

    # ── 4.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '4.2  Lettura card ordini — colori per stato')
    body_para(doc, 'Gli ordini appaiono come card colorate. Il colore indica immediatamente lo stato:')
    spacer(doc, 4)

    data_table(doc,
        ['Colore card', 'Stato', 'Significato'],
        [
            ['⬜ Grigio',   'NUOVO',           'Ordine appena arrivato — nessuno lo ha preso in carico'],
            ['🟡 Giallo',   'NUOVO (alert)',    'Ordine in attesa da oltre 15 minuti — prioritizza'],
            ['🔵 Blu',      'IN PREPARAZIONE', 'Cucina ha preso in carico — piatto in lavorazione'],
            ['🟢 Verde',    'PRONTO',           'Piatto pronto per essere servito o ritirato'],
        ],
        col_widths=[3.0, 4.0, 10.6])
    spacer(doc, 6)

    data_table(doc,
        ['Elemento card', 'Significato'],
        [
            ['🕐 Orario ordine',  'Quando è stato inviato l\'ordine dal cliente o dalla cassa'],
            ['👤 Cliente/Tavolo', 'Chi ha ordinato e (se prenotato) a quale tavolo'],
            ['📋 Lista prodotti', 'Ogni prodotto con quantità e note speciali in corsivo'],
            ['BANCO nel codice',  'Ordine "Adesso al banco" — richiede preparazione immediata'],
        ],
        col_widths=[4.5, 13.1])
    spacer(doc, 8)

    # ── 4.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '4.3  Flusso operativo — 5 step')
    workflow_table(doc, [
        ('🟡', 'VEDI ORDINE', 'Leggi card grigia/gialla'),
        ('👆', 'IN PREP.', 'Clicca "In preparazione"'),
        ('🍳', 'PREPARA', 'Cucina il piatto'),
        ('✅', 'PRONTO', 'Clicca "Pronto"'),
        ('🔔', 'NOTIFICA', 'Cliente/sala avvisati'),
    ], accent=HEX_ORNG)
    spacer(doc, 8)

    step_row(doc, 1, 'Vedi ordine — card grigia/gialla', 'Leggi prodotti e note. Se hai dubbi, chiedi alla cassa prima di iniziare.', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 2, 'Clicca "In preparazione"', 'La card diventa blu: segnale che la cucina ha preso in carico l\'ordine. Non saltare questo step.', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 3, 'Prepara i piatti', 'Segui l\'ordine di arrivo. Per lo stesso cliente, prepara tutto prima di segnare "pronto".', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 4, 'Clicca "Pronto"', 'La card diventa verde. Se configurato, il cliente riceve notifica Telegram automatica.', accent=HEX_ORNG)
    spacer(doc, 4)
    step_row(doc, 5, 'L\'ordine viene archiviato', 'Dopo il ritiro confermato dalla sala/cassiere, l\'ordine scompare dal KDS.', accent=HEX_ORNG)
    spacer(doc, 8)

    info_box(doc, 'Non saltare lo step "In preparazione": serve alla cassa per sapere '
             'che la cucina ha visto l\'ordine e non rischiare di inviarlo di nuovo.',
             style='warning')
    spacer(doc, 8)

    # ── 4.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '4.4  Gestione note speciali e ordini builder')
    body_para(doc, 'Gli ordini possono contenere personalizzazioni (ordini "builder") '
              'o note scritte dal cliente. Leggile sempre prima di iniziare a preparare.')
    spacer(doc, 4)

    data_table(doc,
        ['Tipo', 'Dove appare', 'Come gestirlo'],
        [
            ['Note testo (es. "senza glutine")',
             'Sotto il prodotto, in corsivo',
             'Segui l\'indicazione alla lettera — potrebbe essere un\'allergia'],
            ['Ordine builder (personalizzato)',
             'Card con sezione "Personalizzazioni"',
             'Prepara ogni componente come indicato nell\'elenco opzioni'],
            ['Prodotto non disponibile',
             '—',
             'Avvisa immediatamente la cassa via interfono — la cassa gestirà il cliente'],
        ],
        col_widths=[4.0, 4.5, 9.1])
    spacer(doc, 8)

    # ── 4.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '4.5  Ordini "Adesso al banco" (BANCO)')
    body_para(doc, 'Quando il codice ordine contiene il tag BANCO (es. QL-250706-BANCO-0042), '
              'significa che il cliente è fisicamente al bancone e l\'ordine deve essere '
              'preparato immediatamente, senza attendere uno slot orario.')
    spacer(doc, 4)

    info_box_color(doc,
                   'Riconosci un ordine BANCO dal codice in alto nella card: include la parola BANCO.\n'
                   'Priorità alta: il cliente è già davanti al bancone e aspetta.',
                   bg='FEF9E7', border=HEX_ORNG, icon='⚡')
    spacer(doc, 8)

    # ── 4.6 ──────────────────────────────────────────────────────────────────
    h2(doc, '4.6  Notifiche Telegram cucina')
    body_para(doc, 'Se l\'Admin ha configurato il bot Telegram per la cucina, '
              'questi eventi generano notifiche automatiche sul gruppo cucina:')
    spacer(doc, 4)

    data_table(doc,
        ['Evento', 'Messaggio Telegram'],
        [
            ['Nuovo ordine ricevuto',    '🛒 Nuovo ordine #123 — Mario Rossi: 1x Risotto, 2x Acqua'],
            ['Ordine in attesa >15 min', '⚠️ Ordine #120 in attesa da 20 minuti — prioritizza!'],
            ['Ordine annullato',         '❌ Ordine #121 ANNULLATO — Mario Rossi'],
        ],
        col_widths=[5.5, 12.1])
    spacer(doc, 8)

    info_box(doc, 'Le notifiche Telegram sono in aggiunta al KDS, non in sostituzione. '
             'Tieni sempre aperto lo schermo KDS durante il servizio.', style='tip')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 5 — SALA
# ══════════════════════════════════════════════════════════════════════════════

def s_sala(doc):
    h1(doc, '5', 'Sala — Gestione Tavoli e Prenotazioni', '🪑')

    role_badge(doc, '🪑', 'Sala',
               'Il personale sala gestisce le prenotazioni dei tavoli, effettua il check-in '
               'dei clienti all\'arrivo, monitora le scadenze di sessione e annulla prenotazioni '
               'se necessario. Riceve avvisi Telegram automatici per i tavoli in scadenza.',
               HEX_TEAL)
    spacer(doc, 8)

    # ── 5.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '5.1  Panoramica prenotazioni del giorno')
    body_para(doc, 'La panoramica mostra tutte le prenotazioni del giorno suddivise per tavolo '
              'e fascia oraria tramite chip colorati. Percorso: Admin → Tavoli → tab Panoramica.')
    spacer(doc, 4)

    step_row(doc, 1, 'Vai in Admin → Tavoli → tab Panoramica', 'Vedi tutte le fasce orarie e i tavoli del giorno corrente')
    spacer(doc, 4)
    step_row(doc, 2, 'Naviga tra i giorni', 'Usa le frecce ‹ › o il selettore data per vedere un giorno diverso')
    spacer(doc, 4)
    step_row(doc, 3, 'Leggi i chip colorati', 'Chip verde = tavolo libero in quella sessione; chip rosso = occupato con nome cliente')
    spacer(doc, 8)

    info_box(doc, 'Le FASCE ORARIE TAVOLI sono separate dagli SLOT ORDINI (ritiro cibo). '
             'Un cliente può prenotare un tavolo e ordinare separatamente dal menu — '
             'sono due azioni indipendenti.', style='info')
    spacer(doc, 8)

    # ── 5.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '5.2  Check-in cliente al tavolo')
    body_para(doc, 'Quando un cliente con prenotazione arriva fisicamente, '
              'registra il check-in per avviare il timer di permanenza.')
    spacer(doc, 4)

    step_row(doc, 1, 'Trova la prenotazione', 'Cerca il cliente per nome o orario sessione nella Panoramica o nella lista')
    spacer(doc, 4)
    step_row(doc, 2, 'Clicca il pulsante check-in (icona porta)', 'Il sistema registra l\'orario esatto di arrivo del cliente')
    spacer(doc, 4)
    step_row(doc, 3, 'Il timer parte', 'Accanto alla prenotazione appare l\'orario di check-in e il conto alla rovescia')
    spacer(doc, 4)
    step_row(doc, 4, 'Accompagna il cliente al tavolo', 'Il numero tavolo è visibile sulla scheda prenotazione')
    spacer(doc, 8)

    info_box(doc, 'Il check-in è obbligatorio per far funzionare gli avvisi Telegram. '
             'Senza check-in il timer non parte e nessuna notifica di scadenza verrà inviata.',
             style='warning')
    spacer(doc, 8)

    # ── 5.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '5.3  Monitoraggio scadenza sessioni')
    body_para(doc, 'Quando un cliente è in sala da (durata fascia - 10 minuti), '
              'il sistema invia automaticamente un avviso Telegram al canale sala.')
    spacer(doc, 4)

    info_box_color(doc,
                   'Esempio: Fascia 11:25–12:30 con 30 min di durata.\n'
                   'Cliente fa check-in alle 11:25.\n'
                   'Alle 11:45 (20 min dopo, cioè 10 min prima della scadenza) → avviso Telegram:\n'
                   '⏰ Tavolo 3 — Mario Rossi. Tempo rimasto: ~10 min.',
                   bg='FEF9E7', border=HEX_ORNG, icon='⏰')
    spacer(doc, 8)

    data_table(doc,
        ['Stato tavolo', 'Tempo rimasto', 'Azione sistema'],
        [
            ['🟢 Tranquillo', '> 15 minuti',  'Nessuna azione'],
            ['🟡 Attenzione', '10–15 minuti', 'Telegram avviso preparazione liberazione'],
            ['🔴 Urgente',    '< 10 minuti',  'Telegram avviso liberazione immediata'],
            ['⬛ Scaduto',    '0 o negativo', 'Telegram avviso immediato + verifica manuale'],
        ],
        col_widths=[3.5, 3.5, 10.6])
    spacer(doc, 8)

    # ── 5.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '5.4  Annullamento prenotazione')
    step_row(doc, 1, 'Trova la prenotazione da annullare', 'Admin → Tavoli → lista prenotazioni o Panoramica → cerca per nome o sessione')
    spacer(doc, 4)
    step_row(doc, 2, 'Clicca "Annulla prenotazione"', 'Il pulsante è visibile nella scheda della prenotazione')
    spacer(doc, 4)
    step_row(doc, 3, 'Conferma l\'annullamento', 'Il popup chiede conferma — clicca OK')
    spacer(doc, 4)
    step_row(doc, 4, 'Il tavolo si libera', 'Il chip torna verde nella Panoramica e il posto è di nuovo prenotabile')
    spacer(doc, 8)

    # ── 5.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '5.5  Configurazione ping-alerts (polling automatico)')
    body_para(doc, 'Il sistema controlla automaticamente ogni 60 secondi se ci sono tavoli '
              'con poco tempo rimasto. Il polling avviene tramite l\'endpoint '
              '/admin/tavoli/ping-alerts, integrato nella pagina KDS.')
    spacer(doc, 4)

    info_box(doc, 'Per ricevere gli avvisi automatici, mantieni aperta la pagina KDS '
             '(o Admin → Tavoli) su almeno uno schermo durante il servizio. '
             'Il polling ogni 60 secondi invia i Telegram necessari senza intervento manuale.',
             style='tip')
    spacer(doc, 8)


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 6 — CLIENTE
# ══════════════════════════════════════════════════════════════════════════════

def s_cliente(doc):
    h1(doc, '6', 'Cliente — App e Auto-ordine', '👤')

    role_badge(doc, '👤', 'Cliente',
               'Il cliente usa l\'app web per ordinare in autonomia, gestire il wallet, '
               'accumulare punti fedeltà, prenotare un tavolo, votare il menu e '
               'pagare rapidamente al bancone tramite QR. '
               'Non serve scaricare nulla: funziona dal browser dello smartphone.',
               HEX_PURPL)
    spacer(doc, 8)

    info_box_color(doc,
                   'DISTINZIONE IMPORTANTE:\n'
                   '• "Adesso al banco" = il CLIENTE ordina dal menu e NON sceglie uno slot '
                   '(ritiro immediato al bancone, cucina prepara subito)\n'
                   '• "Paga al Banco" = lo STAFF genera un QR per una vendita rapida '
                   '(caffè, brioche…) e il cliente lo scansiona per pagare dal wallet',
                   bg='EBF5FB', border=HEX_NAVY, icon='ℹ️')
    spacer(doc, 10)

    # ── 6.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.1  Registrazione e accesso')
    step_row(doc, 1, 'Apri il browser sul telefono', 'Vai all\'URL del locale (es. https://pranzo.barcentrale.it)')
    spacer(doc, 4)
    step_row(doc, 2, 'Registrazione email/password', 'Clicca "Registrati": inserisci nome, email e una password sicura')
    spacer(doc, 4)
    step_row(doc, 3, 'In alternativa: Google OAuth', 'Clicca "Accedi con Google" per registrarti o accedere con il tuo account Google (se l\'Admin ha configurato OAuth)')
    spacer(doc, 4)
    step_row(doc, 4, 'Accesso ricordato 30 giorni', 'Il sistema mantiene il login attivo per 30 giorni — non devi reinserire la password ogni volta')
    spacer(doc, 4)
    step_row(doc, 5, 'Aggiungi al telefono (opzionale)', 'Nel browser tocca "Aggiungi alla schermata home" per usarla come app')
    spacer(doc, 8)

    # ── 6.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.2  Home dashboard')
    body_para(doc, 'Dopo il login la dashboard mostra KPI personali e le azioni rapide principali.')
    spacer(doc, 4)

    data_table(doc,
        ['Elemento dashboard', 'Descrizione'],
        [
            ['KPI personali',         'Saldo wallet, punti fedeltà, ultimi ordini'],
            ['Menu',                  'Azione rapida → vai al menu del giorno'],
            ['Builder',               'Azione rapida → crea un ordine personalizzato (testo)'],
            ['Builder Visuale',       'Azione rapida → crea un ordine con interfaccia grafica'],
            ['Prenota Tavolo',        'Azione rapida → scegli data, fascia oraria e tavolo'],
            ['Paga al Banco',         'Azione rapida → apri scanner QR per pagare la sessione generata dallo staff'],
        ],
        col_widths=[4.5, 13.1])
    spacer(doc, 8)

    # ── 6.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.3  Flusso ordine dal menu')
    workflow_table(doc, [
        ('🍽️', 'Apri Menu', 'Sfoglia categorie'),
        ('➕', 'Aggiungi', 'Tocca prodotto → +'),
        ('🛒', 'Carrello', 'Rivedi e scegli slot'),
        ('💳', 'Paga', 'Wallet o cassa'),
        ('⏳', 'Attendi', 'Notifica quando pronto'),
    ], accent=HEX_PURPL)
    spacer(doc, 8)

    step_row(doc, 1, 'Vai in Menu', 'Dal menu laterale o dal pulsante rapido in dashboard')
    spacer(doc, 4)
    step_row(doc, 2, 'Sfoglia per categoria', 'Primo, Secondo, Contorno, Bevande, Dolci — tocca la categoria desiderata')
    spacer(doc, 4)
    step_row(doc, 3, 'Aggiungi prodotti al carrello', 'Tocca il prodotto → "+" per aggiungere. Ripeti per più prodotti')
    spacer(doc, 4)
    step_row(doc, 4, 'Vai al carrello', 'Tocca l\'icona carrello o vai in Ordina → Carrello')
    spacer(doc, 4)
    step_row(doc, 5, 'Scegli lo slot di ritiro', 'Seleziona l\'orario di ritiro disponibile (slot configurati dall\'Admin)')
    spacer(doc, 4)
    step_row(doc, 6, 'Scegli come pagare', 'Se hai saldo wallet: seleziona "Wallet". Altrimenti "Paga alla cassa" al ritiro')
    spacer(doc, 4)
    step_row(doc, 7, 'Conferma e ricevi numero ordine', 'Tocca "Invia ordine". Ricevi conferma con numero ordine')
    spacer(doc, 8)

    # ── 6.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.4  "Adesso al banco" nel carrello')
    body_para(doc, 'Se sei fisicamente al bancone e vuoi che il tuo ordine venga preparato SUBITO, '
              'scegli "Adesso al banco" nel carrello invece di uno slot orario. '
              'La cucina prepara l\'ordine immediatamente.')
    spacer(doc, 6)

    step_row(doc, 1, 'Componi l\'ordine dal menu', 'Aggiungi i prodotti al carrello normalmente')
    spacer(doc, 4)
    step_row(doc, 2, 'Vai al carrello', 'Tocca l\'icona carrello o Ordina → Carrello')
    spacer(doc, 4)
    step_row(doc, 3, 'Seleziona "Adesso al banco"', 'Nella sezione slot, trovi come prima opzione "Adesso al banco" — selezionala (nessuno slot richiesto)')
    spacer(doc, 4)
    step_row(doc, 4, 'Conferma l\'ordine', 'L\'ordine viene inviato in cucina immediatamente con tag BANCO nel codice')
    spacer(doc, 6)

    info_box_color(doc,
                   'Il codice ordine generato con "Adesso al banco" include il tag BANCO:\n'
                   'QL-AAMMGG-BANCO-NNNN\n'
                   'Mostralo allo staff al bancone per il ritiro immediato.',
                   bg='EBF5FB', border=HEX_TEAL, icon='ℹ️')
    spacer(doc, 10)

    # ── 6.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.5  "Paga al Banco" — scanner QR')
    body_para(doc, 'Quando lo STAFF al bancone genera un QR per una vendita rapida '
              '(caffè, brioche, ecc.), il cliente lo scansiona dalla sua dashboard '
              'per pagare istantaneamente dal wallet. '
              'Il cliente NON deve ordinare nulla: è lo staff che ha composto il carrello.')
    spacer(doc, 6)

    step_row(doc, 1, 'Vai nella tua Dashboard', 'Dalla home o dal menu laterale dell\'app')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca "Paga al Banco"', 'Il pulsante si trova nelle azioni rapide della dashboard')
    spacer(doc, 4)
    step_row(doc, 3, 'Autorizza la fotocamera', 'Alla prima apertura il browser chiede il permesso — consenti')
    spacer(doc, 4)
    step_row(doc, 4, 'Inquadra il QR dello staff', 'Punta la fotocamera sul QR mostrato dallo staff al bancone')
    spacer(doc, 4)
    step_row(doc, 5, 'Rivedi l\'importo e conferma', 'Vedi il totale da pagare e tocca "Conferma pagamento"')
    spacer(doc, 4)
    step_row(doc, 6, 'Pagamento completato', 'Lo schermo del banco si aggiorna mostrando "✓ Pagato da [Il tuo nome]"')
    spacer(doc, 8)

    info_box(doc, 'Il QR generato dallo staff ha una validità di 10 minuti. '
             'Se scade prima della scansione, chiedi allo staff di generarne uno nuovo.', style='warning')
    spacer(doc, 8)

    # ── 6.6 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.6  Prenotazione tavolo')
    step_row(doc, 1, 'Vai in Tavoli → Prenota Tavolo', 'Dal menu laterale o dal pulsante rapido in dashboard')
    spacer(doc, 4)
    step_row(doc, 2, 'Scegli la data e la fascia oraria', 'Sono mostrate solo le fasce con posti disponibili')
    spacer(doc, 4)
    step_row(doc, 3, 'Seleziona il tavolo', 'Scegli tra i tavoli liberi nella fascia prescelta')
    spacer(doc, 4)
    step_row(doc, 4, 'Conferma prenotazione', 'Ricevi codice prenotazione e notifica Telegram di conferma')
    spacer(doc, 4)
    step_row(doc, 5, 'All\'arrivo al locale', 'Dichiara il tuo nome al personale sala per il check-in')
    spacer(doc, 8)

    # ── 6.7 ──────────────────────────────────────────────────────────────────
    h2(doc, '6.7  Wallet e punti fedeltà')
    body_para(doc, 'Il wallet è il portafoglio digitale nel locale. '
              'Ogni acquisto con wallet accumula punti fedeltà. '
              'Sezione: Wallet & Fedeltà nel menu laterale. '
              'Il wallet è una funzione disattivabile dalle Impostazioni: se '
              'il locale lavora con il pagamento in cassa, le voci e le '
              'pagine di questa sezione non compaiono.')
    spacer(doc, 4)

    data_table(doc,
        ['Funzione', 'Come si usa', 'Nota'],
        [
            ['Saldo wallet',        'Wallet & Fedeltà → Saldo attuale',  'Pagamenti rapidi senza contanti'],
            ['Storico transazioni', 'Wallet & Fedeltà → Movimenti',       'Ogni spesa tracciata con data e importo'],
            ['Punti fedeltà',       'Wallet & Fedeltà → I tuoi punti',    'Accumulati ad ogni acquisto wallet'],
            ['Ricarica wallet',     'Chiedi al cassiere al bancone',       'Il cassiere aggiorna il saldo in tempo reale'],
        ],
        col_widths=[4.0, 6.5, 7.1])
    spacer(doc, 8)


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 7 — DIPENDENTE AZIENDALE
# ══════════════════════════════════════════════════════════════════════════════

def s_dipendente(doc):
    h1(doc, '7', 'Dipendente Aziendale — Pasto Fisso Convenzionato', '🏭')

    role_badge(doc, '🏭', 'Dipendente Aziendale',
               'Sei un dipendente di un\'azienda convenzionata con questo locale. '
               'L\'azienda ha concordato un pasto completo a prezzo fisso ogni giorno. '
               'Il tuo compito è semplicissimo: prenotare il pasto del giorno e ritirarlo.',
               '5D6D7E')
    spacer(doc, 8)

    info_box_color(doc,
                   'Come funziona in breve:\n'
                   '1. L\'Admin crea ogni mattina il "pasto del giorno" per la tua azienda\n'
                   '2. Tu lo prenoti scegliendo uno slot orario\n'
                   '3. Arrivi al locale nello slot scelto e ritiri il pasto\n'
                   '4. Non paghi nulla: paga direttamente la tua azienda a fine mese',
                   bg='EBF5FB', border=HEX_TEAL, icon='🏭')
    spacer(doc, 10)

    # ── 7.1 ──────────────────────────────────────────────────────────────────
    h2(doc, '7.1  Accesso (stesse credenziali cliente)')
    body_para(doc, 'Il dipendente aziendale usa le stesse credenziali di un normale cliente. '
              'Non esiste un\'app separata. La sezione "Pasto Aziendale" appare nel menu laterale '
              'solo dopo che l\'Admin ti ha associato alla convenzione aziendale.')
    spacer(doc, 4)

    step_row(doc, 1, 'Registrati come cliente normale', 'Vai all\'URL del locale, clicca "Registrati" e crea il tuo account email/password')
    spacer(doc, 4)
    step_row(doc, 2, 'Comunica la tua email all\'Admin', 'L\'Admin deve associarti alla convenzione aziendale — invia email o username')
    spacer(doc, 4)
    step_row(doc, 3, 'Attendi l\'associazione', 'L\'Admin ti aggiunge alla convenzione. Di solito avviene entro poche ore')
    spacer(doc, 4)
    step_row(doc, 4, 'Comparirà "Pasto Aziendale"', 'La voce appare nel menu laterale solo per i dipendenti convenzionati')
    spacer(doc, 8)

    # ── 7.2 ──────────────────────────────────────────────────────────────────
    h2(doc, '7.2  Prenotazione pasto fisso convenzionato — 6 step')
    workflow_table(doc, [
        ('📱', 'Apri app', 'Accedi con le tue credenziali'),
        ('🍽️', 'Pasto Aziendale', 'Menu laterale → Pasto Aziendale'),
        ('👁️', 'Vedi menu', 'Leggi pasto, allergeni, posti'),
        ('🕐', 'Scegli slot', 'Seleziona l\'orario di ritiro'),
        ('✅', 'Prenota', 'Conferma la prenotazione'),
        ('🏃', 'Ritira', 'Presentati allo slot scelto'),
    ], accent=HEX_TEAL)
    spacer(doc, 8)

    step_row(doc, 1, 'Vai in "Pasto Aziendale"', 'Menu laterale → icona edificio "Pasto Aziendale"')
    spacer(doc, 4)
    step_row(doc, 2, 'Leggi il pasto del giorno', 'Nome, descrizione, composizione (primo/secondo/contorno) e allergeni')
    spacer(doc, 4)
    step_row(doc, 3, 'Controlla i posti rimasti', 'Badge in alto: verde >50%, giallo 20–50%, rosso <20%. Se è 0, posti esauriti.')
    spacer(doc, 4)
    step_row(doc, 4, 'Scegli lo slot orario', 'Dal menu a tendina seleziona l\'orario di ritiro desiderato')
    spacer(doc, 4)
    step_row(doc, 5, 'Clicca "Prenota il mio pasto"', 'Prenotazione confermata. Ricevi notifica Telegram con riepilogo.')
    spacer(doc, 4)
    step_row(doc, 6, 'Arriva al locale nello slot scelto', 'Dichiara il tuo nome — il cassiere trova la prenotazione e segna "Consumato"')
    spacer(doc, 8)

    # ── 7.3 ──────────────────────────────────────────────────────────────────
    h2(doc, '7.3  Annullamento prenotazione')

    info_box(doc,
             'REGOLA DI ANNULLAMENTO: puoi annullare solo se mancano PIU\' DI 30 MINUTI '
             'all\'orario dello slot prenotato. '
             'Il pulsante "Annulla prenotazione" sparisce automaticamente '
             'quando rimangono 30 minuti o meno allo slot.',
             style='warning', label='LIMITE 30 MINUTI')
    spacer(doc, 8)

    step_row(doc, 1, 'Vai in "Pasto Aziendale"', 'La tua prenotazione attiva è mostrata con badge verde "Prenotato"')
    spacer(doc, 4)
    step_row(doc, 2, 'Verifica che manchino >30 minuti allo slot', 'Se mancano 30 minuti o meno il pulsante non è visibile — annullamento non consentito')
    spacer(doc, 4)
    step_row(doc, 3, 'Clicca "Annulla prenotazione"', 'Pulsante rosso sotto la scheda prenotazione — visibile solo se il termine non è scaduto')
    spacer(doc, 4)
    step_row(doc, 4, 'Conferma l\'annullamento', 'Il popup chiede conferma — clicca OK')
    spacer(doc, 4)
    step_row(doc, 5, 'Il posto si libera', 'Il posto torna disponibile per un collega. Puoi riprenotare scegliendo un altro slot.')
    spacer(doc, 8)

    # ── 7.4 ──────────────────────────────────────────────────────────────────
    h2(doc, '7.4  Lettura interfaccia Pasto Aziendale')
    data_table(doc,
        ['Elemento', 'Significato'],
        [
            ['Nome azienda (in alto)',     'Conferma che sei associato alla convenzione corretta'],
            ['Banner del pasto',            'Nome, descrizione e composizione del pasto di oggi'],
            ['Badge posti rimasti (verde)', '>50% posti disponibili — prenotazione sicura'],
            ['Badge posti rimasti (giallo)', '20–50% posti rimasti — prenota presto'],
            ['Badge posti rimasti (rosso)', '<20% posti rimasti — ultimi posti disponibili'],
            ['Badge stato prenotazione',    'Verde = Prenotato, Arancio = Consumato, Grigio = Annullato'],
            ['Slot orario prenotato',       'L\'orario scelto — rispettalo o annulla con >30 min di anticipo'],
            ['"Nessun pasto disponibile"',  'L\'Admin non ha ancora pubblicato il menu oggi — ricontrolla più tardi'],
        ],
        col_widths=[5.0, 12.6])
    spacer(doc, 8)

    info_box(doc, 'Se non vedi "Pasto Aziendale" nel menu, non sei ancora associato '
             'a nessuna convenzione. Contatta l\'Admin del locale per farti aggiungere.',
             style='warning')
    spacer(doc, 8)

    # ── 7.5 ──────────────────────────────────────────────────────────────────
    h2(doc, '7.5  FAQ Dipendente Aziendale')
    data_table(doc,
        ['Domanda', 'Risposta'],
        [
            ['Posso usare anche il menu normale?',
             'Sì, il pasto aziendale è aggiuntivo. Puoi ordinare anche dal menu standard e pagare normalmente.'],
            ['Cosa succede se non mi presento?',
             'La prenotazione resta come "prenotata". Nessuna penale automatica, ma il posto è andato perso.'],
            ['Il prezzo varia ogni giorno?',
             'No, il prezzo è fisso per contratto. Cambia solo se l\'azienda rinnova la convenzione.'],
            ['Posso annullare anche all\'ultimo momento?',
             'No: puoi annullare SOLO se mancano PIU\' DI 30 MINUTI allo slot. '
             'Il pulsante sparisce automaticamente entro i 30 minuti.'],
            ['Ho dimenticato la password',
             'Clicca "Password dimenticata?" nella pagina di login oppure contatta l\'Admin del locale.'],
        ],
        col_widths=[5.8, 11.8])
    spacer(doc, 8)


# ══════════════════════════════════════════════════════════════════════════════
# APPENDICE
# ══════════════════════════════════════════════════════════════════════════════

def s_appendice(doc):
    h1(doc, 'A', 'Appendice — Permessi, Flussi e Account di Test', '📎')

    # ── A.1 ──────────────────────────────────────────────────────────────────
    h2(doc, 'A.1  Matrice permessi per ruolo')
    data_table(doc,
        ['Funzione', '👑\nSuper', '🏢\nAdmin', '💳\nCassa', '🍳\nKDS', '🪑\nSala', '👤\nCliente', '🏭\nDip.'],
        [
            ['Login area Admin',         '✅', '✅', '✅', '✅', '✅', '❌', '❌'],
            ['Gestione prodotti/menu',   '✅', '✅', '❌', '❌', '❌', '❌', '❌'],
            ['Creazione ordini',         '✅', '✅', '✅', '❌', '❌', '✅', '❌'],
            ['Incasso pagamenti',        '✅', '✅', '✅', '❌', '❌', '❌', '❌'],
            ['Banco POS (genera QR)',    '✅', '✅', '✅', '❌', '❌', '❌', '❌'],
            ['Paga al Banco (scan QR)',  '❌', '❌', '❌', '❌', '❌', '✅', '✅'],
            ['Stato ordine (KDS)',       '✅', '✅', '❌', '✅', '❌', '❌', '❌'],
            ['Prenotazioni tavoli',      '✅', '✅', '❌', '❌', '✅', '✅', '❌'],
            ['Check-in tavolo',          '✅', '✅', '❌', '❌', '✅', '❌', '❌'],
            ['Wallet clienti',           '✅', '✅', '✅', '❌', '❌', '✅', '❌'],
            ['Magazzino consumabili',    '✅', '✅', '❌', '❌', '❌', '❌', '❌'],
            ['Convenzioni aziendali',    '✅', '✅', '❌', '❌', '❌', '❌', '❌'],
            ['Pasto aziendale fisso',    '✅', '✅', '❌', '❌', '❌', '❌', '✅'],
            ['Report e statistiche',     '✅', '✅', '❌', '❌', '❌', '❌', '❌'],
            ['Sondaggi (crea/chiudi)',   '✅', '✅', '❌', '❌', '❌', '❌', '❌'],
            ['Voto menu (sondaggio)',    '❌', '✅', '❌', '❌', '❌', '✅', '❌'],
        ],
        col_widths=[5.0, 1.8, 1.8, 1.8, 1.8, 1.8, 2.0, 2.6])
    spacer(doc, 10)

    # ── A.2 ──────────────────────────────────────────────────────────────────
    h2(doc, 'A.2  Flusso giornaliero tipo — bar/mensa')
    workflow_table(doc, [
        ('🌅', 'Apertura', 'Admin: stock, pasto aziendale, slot ordini'),
        ('💳', 'Pre-pranzo', 'Cassa: ordini walk-in, ricariche wallet'),
        ('🍳', 'Servizio', 'KDS: gestione ordini; sala: check-in tavoli'),
        ('⏰', 'Avvisi', 'Telegram: tavoli in scadenza, stock basso'),
        ('📊', 'Chiusura', 'Admin: report 30 gg, aggiorna stock'),
    ])
    spacer(doc, 10)

    # ── A.3 ──────────────────────────────────────────────────────────────────
    h2(doc, 'A.3  Account di test (seed-demo)')
    body_para(doc, 'Dopo l\'esecuzione di "flask seed-demo" sono disponibili questi account preconfigurati:')
    spacer(doc, 4)

    data_table(doc,
        ['Email', 'Password', 'Ruolo', 'Note'],
        [
            ['admin@bar.local',    'admin123',    'Super Admin',  'Accesso completo a tutti i tenant'],
            ['banco@bar.local',    'Banco2024!',  'Cassiere',     'Pannello ordini + Banco POS'],
            ['cucina@bar.local',   'Cucina2024!', 'Cuoco',        'Schermata KDS, stati ordine'],
            ['sala@bar.local',     'Sala2024!',   'Manager Sala', 'Prenotazioni tavoli, check-in'],
            ['cliente1@bar.local', 'Cliente1!',   'Cliente',      'Wallet: 30 €, 50 punti fedeltà'],
            ['cliente2@bar.local', 'Cliente2!',   'Cliente',      'Wallet: 15 €, dipendente ACME'],
        ],
        col_widths=[4.5, 3.2, 3.5, 6.4])
    spacer(doc, 8)

    info_box(doc, 'Non usare gli account di test in produzione. '
             'Cambia subito le password predefinite dopo il primo avvio in ambiente reale.',
             style='warning')


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════


def _footer_copyright(doc):
    """Nota di copyright in chiusura del documento."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    p = doc.add_paragraph()
    p.alignment = _AL.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('QuickLunch  \u00b7  © 2024–26 DS Consulting')
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    r.font.name = FONT

    p2 = doc.add_paragraph()
    p2.alignment = _AL.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run('Assistenza:  Daniele Speziale — DS Consulting  \u00b7  '
                    'dspeziale@gmail.com  \u00b7  +39 352 0150489')
    r2.font.size = Pt(9.5)
    r2.font.bold = True
    r2.font.color.rgb = RED
    r2.font.name = FONT

def main():
    doc = Document()
    set_document_defaults(doc)

    # ── Cover ──
    build_cover(doc)
    _page_break(doc)

    # ── Indice ──
    build_toc(doc)
    _page_break(doc)

    # ── Sezioni ──
    s_super_admin(doc)
    _page_break(doc)

    s_admin_tenant(doc)
    _page_break(doc)

    s_cassiere(doc)
    _page_break(doc)

    s_cucina(doc)
    _page_break(doc)

    s_sala(doc)
    _page_break(doc)

    s_cliente(doc)
    _page_break(doc)

    s_dipendente(doc)
    _page_break(doc)

    s_appendice(doc)

    _footer_copyright(doc)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'OK  Guida utente salvata in: {OUT}')
    print(f'    Sezioni: Super Admin · Admin Tenant · Cassiere · Cucina/KDS · Sala · Cliente · Dipendente · Appendice')


if __name__ == '__main__':
    main()
