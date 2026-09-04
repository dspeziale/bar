#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera docs/manuali/presentazione_quicklunch.docx.

La presentazione professionale dell'applicazione da consegnare al potenziale
cliente (il gestore di bar o mensa aziendale): che cos'e', come funziona, i
moduli, la dotazione necessaria, la tecnologia, il percorso di avvio.
Le condizioni economiche non sono qui: stanno nell'offerta commerciale
(generate_offerta_saas.py), che questo documento richiama.

    python docs/generate_presentazione_quicklunch.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Palette (coerente con gli altri documenti in docs/) ──────────────────────
RED    = RGBColor(0xe9, 0x45, 0x60)
NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x16, 0x21, 0x3e)
DGRAY  = RGBColor(0x34, 0x3a, 0x40)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREEN  = RGBColor(0x27, 0xae, 0x60)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
BLUE   = RGBColor(0x34, 0x98, 0xdb)
PURPLE = RGBColor(0x8e, 0x44, 0xad)

HEX_RED    = 'E94560'
HEX_NAVY   = '0F3460'
HEX_GREEN  = '27AE60'
HEX_ORANGE = 'E67E22'
HEX_BLUE   = '3498DB'
HEX_PURPLE = '8E44AD'
HEX_LIGHT  = 'F8F9FA'
HEX_RULE   = 'E8EBEE'
HEX_WARN   = 'FDF6E7'

FONT = 'PT Sans Narrow'

EMAIL = 'dspeziale@gmail.com'
CELL = '+39 352 0150489'
CONTATTI = 'DS Consulting  ·  Daniele Speziale  ·  %s  ·  %s' % (EMAIL, CELL)

OUT = os.path.join(os.path.dirname(__file__), 'manuali',
                   'presentazione_quicklunch.docx')


# ── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None, sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        if color:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color.lstrip('#'))
        else:
            el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), val)
    tcPr.append(el)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _p_spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def _run_font(run, size=10.5, bold=False, italic=False, color=None, font=FONT):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    _p_spacing(p, before=0, after=0)


# ── Fogli di stile ───────────────────────────────────────────────────────────

def imposta_stili(doc):
    """Definisce gli stili del documento, cosi' che in Word si possa ritoccare
    l'aspetto da un punto solo invece che paragrafo per paragrafo."""
    from docx.enum.style import WD_STYLE_TYPE

    normale = doc.styles['Normal']
    normale.font.name = FONT
    normale.font.size = Pt(10.5)
    normale.font.color.rgb = DGRAY
    normale.paragraph_format.space_after = Pt(6)
    normale.paragraph_format.line_spacing = 1.08
    rpr = normale.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), FONT)

    for nome, dim, colore, prima, dopo in [
        ('Heading 1', 19, NAVY, 16, 5),
        ('Heading 2', 14.5, NAVY, 14, 4),
        ('Heading 3', 12, DARK, 12, 3),
    ]:
        try:
            st = doc.styles[nome]
        except KeyError:
            st = doc.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st.font.size = Pt(dim)
        st.font.bold = True
        st.font.color.rgb = colore
        st.paragraph_format.space_before = Pt(prima)
        st.paragraph_format.space_after = Pt(dopo)
        st.paragraph_format.keep_with_next = True

    try:
        nota = doc.styles['Nota QuickLunch']
    except KeyError:
        nota = doc.styles.add_style('Nota QuickLunch', WD_STYLE_TYPE.PARAGRAPH)
    nota.base_style = doc.styles['Normal']
    nota.font.name = FONT
    nota.font.size = Pt(9.5)
    nota.font.color.rgb = GRAY


# ── Blocchi di contenuto ─────────────────────────────────────────────────────

def heading(doc, text, level=1, color=None, before=None):
    p = doc.add_paragraph(style='Heading %d' % min(max(level, 1), 3))
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    if color is not None:
        r.font.color.rgb = color
    return p


def body(doc, text, size=10.5, color=DGRAY, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    _run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def rich(doc, parts, size=10.5, after=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def bullet(doc, parts, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    _p_spacing(p, before=0, after=3)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def rule(doc, color=HEX_RULE, after=8):
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_border(cell, bottom=color, sz='8')
    _cell_margins(cell, top=0, bottom=0, left=0, right=0)
    _p_spacing(cell.paragraphs[0], before=0, after=0)
    _run_font(cell.paragraphs[0].add_run(''), size=1)
    spacer(doc, after)


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(''), size=pts / 2)


def box(doc, label, paragraphs, accent=HEX_RED, fill=HEX_LIGHT, label_color=RED):
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, fill)
    _cell_border(cell, top=HEX_RULE, bottom=HEX_RULE, right=HEX_RULE,
                 left=accent, sz='18')
    _cell_margins(cell, top=90, bottom=90, left=140, right=110)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=3)
    _run_font(p.add_run(label.upper()), size=8.5, bold=True, color=label_color)

    for i, parts in enumerate(paragraphs):
        pp = cell.add_paragraph()
        _p_spacing(pp, before=0, after=0 if i == len(paragraphs) - 1 else 4)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 10)
    return tbl


def step(doc, num, title, paragraphs, accent=HEX_RED):
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_width(tbl, 0, 1.0)
    _set_col_width(tbl, 1, 15.5)

    nc = tbl.rows[0].cells[0]
    _cell_shd(nc, accent)
    _cell_margins(nc, top=60, bottom=60, left=30, right=30)
    _cell_valign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run('%02d' % num), size=12, bold=True, color=WHITE)

    cc = tbl.rows[0].cells[1]
    _cell_margins(cc, top=50, bottom=50, left=150, right=60)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=2)
    _run_font(pt.add_run(title), size=11.5, bold=True, color=DARK)

    for parts in paragraphs:
        pp = cc.add_paragraph()
        _p_spacing(pp, before=0, after=3)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=10, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 6)
    return tbl


def table_grid(doc, headers, rows, widths, head_fill=HEX_NAVY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, w in enumerate(widths):
        _set_col_width(tbl, i, w)

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _cell_shd(c, head_fill)
        _cell_margins(c, top=60, bottom=60, left=110, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(h.upper()), size=8.5, bold=True, color=WHITE)

    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            c = tbl.rows[r].cells[i]
            _cell_shd(c, HEX_LIGHT if r % 2 else 'FFFFFF')
            _cell_border(c, bottom=HEX_RULE, sz='4')
            _cell_margins(c, top=55, bottom=55, left=110, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            text, color, bold = val if isinstance(val, tuple) else (val, DGRAY, False)
            _run_font(p.add_run(text), size=10, bold=bold, color=color)
    spacer(doc, 10)
    return tbl


def modulo(doc, titolo, colore, descrizione, punti):
    """Scheda di un modulo: titolo colorato, descrizione, punti chiave."""
    heading(doc, titolo, 3, color=colore)
    body(doc, descrizione, after=4)
    for parts in punti:
        bullet(doc, parts)
    spacer(doc, 6)


# ── Documento ────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    imposta_stili(doc)

    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)

    # ══ Copertina ═════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, HEX_NAVY)
    _cell_margins(cell, top=300, bottom=300, left=320, right=320)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=4)
    _run_font(p.add_run('QUICKLUNCH  ·  PRESENTAZIONE DEL PRODOTTO'), size=9.5,
              bold=True, color=RGBColor(0xb2, 0xc2, 0xd9))

    p2 = cell.add_paragraph()
    _p_spacing(p2, before=0, after=6)
    _run_font(p2.add_run('Il bar aziendale, digitale'), size=30, bold=True,
              color=WHITE)

    p3 = cell.add_paragraph()
    _p_spacing(p3, before=0, after=10)
    _run_font(p3.add_run(
        'La piattaforma che fa ordinare, pagare e ritirare senza code e senza '
        'contante: ordini dal telefono con ritiro su appuntamento, portafoglio '
        'prepagato con punti fedeltà, cassa veloce con QR, vendita a libero '
        'servizio e pasti aziendali convenzionati. Tutto in un unico sistema, '
        'fornito come servizio: niente da installare, niente server da '
        'mantenere.'), size=11.5, color=RGBColor(0xd6, 0xdf, 0xea))

    p4 = cell.add_paragraph()
    _p_spacing(p4, before=10, after=0)
    _run_font(p4.add_run('Il vostro referente:  '), size=11, bold=True,
              color=WHITE)
    _run_font(p4.add_run('Daniele Speziale  ·  %s  ·  %s' % (EMAIL, CELL)),
              size=11, bold=True, color=RGBColor(0xff, 0xd7, 0xdf))
    spacer(doc, 16)

    # ══ In una pagina ═════════════════════════════════════════════════════
    heading(doc, 'QuickLunch in una pagina', 1)
    rich(doc, [
        ('QuickLunch è la piattaforma di gestione per ', False),
        ('bar e mense aziendali', True),
        (' che porta su un unico sistema tutto quello che oggi passa per '
         'telefonate, bigliettini e code alla cassa: il cliente ordina dal '
         'telefono, paga da un borsellino prepagato e ritira all\'orario che '
         'ha scelto; il gestore vede tutto in tempo reale su un display in '
         'cucina e chiude la giornata con i numeri già pronti.', False),
    ])
    rich(doc, [
        ('È pensata per il ', False), ('bar interno o convenzionato', True),
        (' che serve una comunità stabile — un\'azienda, un polo di uffici, '
         'un condominio di imprese — dove le stesse persone tornano ogni '
         'giorno e la puntualità del servizio conta più di tutto.', False),
    ])
    spacer(doc, 4)

    table_grid(doc,
               ['Cosa', 'Come'],
               [
                   [('Tre canali di vendita', DARK, True),
                    'Ordini su appuntamento, cassa al banco con QR, '
                    'cesto a libero servizio'],
                   [('Zero contante al ritiro', DARK, True),
                    'Tutto è già pagato dal portafoglio prepagato: '
                    'al banco si consegna e basta'],
                   [('Pasti aziendali', DARK, True),
                    'Menu del giorno, prenotazioni dei dipendenti e '
                    'report mensile pronto per la fattura'],
                   [('Fedeltà integrata', DARK, True),
                    'Punti su ogni acquisto e premi automatici, senza '
                    'tessere di cartone'],
                   [('Notifiche ai clienti', DARK, True),
                    'Telegram, email e notifiche sul telefono: '
                    '"il tuo ordine è pronto"'],
                   [('Nessuna installazione', DARK, True),
                    'Software come servizio: si usa dal browser, '
                    'aggiornamenti e backup inclusi'],
               ],
               widths=[4.6, 11.4])

    box(doc, 'Il principio', [
        [('Il cliente fa da solo tutto quello che non richiede le vostre '
          'mani: ordina, paga, prenota, riceve gli avvisi. A voi restano le '
          'due cose che contano — ', False),
         ('preparare bene e consegnare in orario', True),
         ('. L\'applicazione organizza il resto.', False)],
    ], accent=HEX_NAVY, fill=HEX_LIGHT, label_color=NAVY)

    box(doc, 'Due modalità di incasso, a vostra scelta', [
        [('Modalità prepagata (consigliata)', True),
         (': i clienti caricano un borsellino e ogni acquisto scala dal '
          'saldo — zero contante al ritiro, punti fedeltà e bonus di '
          'benvenuto attivi.', False)],
        [('Modalità pagamento in cassa', True),
         (': il portafoglio si disattiva dalle Impostazioni e l\'app non '
          'muove denaro — i clienti ordinano, prenotano e ricevono gli '
          'avvisi come sempre, ma pagano alla cassa al ritiro. Le vendite '
          'restano registrate nei report.', False)],
        [('Si passa dall\'una all\'altra con un interruttore, anche dopo '
          'l\'avvio.', False)],
    ], accent=HEX_GREEN, fill=HEX_LIGHT, label_color=GREEN)

    _page_break(doc)

    # ══ Come funziona: il giro completo ═══════════════════════════════════
    heading(doc, 'Come funziona: il giro completo del cliente', 1, before=0)
    body(doc, 'Dal primo contatto al ritiro del primo ordine, il percorso del '
              'cliente è questo. Nessun passaggio richiede assistenza da parte '
              'vostra, tranne l\'approvazione iniziale. Il percorso descrive la '
              'modalità prepagata: con il pagamento in cassa i passi 2 e 5 non '
              'esistono e il cliente salda al ritiro.', color=GRAY)
    rule(doc)

    step(doc, 1, 'Si registra inquadrando un QR', [
        [('Una locandina con il QR di registrazione, appesa vicino alla '
          'cassa, porta alla pagina di iscrizione: email e password, oppure '
          'direttamente con l\'account ', False), ('Google', True),
         ('. Voi approvate il nuovo cliente con un clic e lui riceve '
          'l\'email di benvenuto con un ', False),
         ('bonus di primo utilizzo', True), (' già caricato.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'Ricarica il portafoglio', [
        [('Il cliente lascia il contante o paga con carta alla cassa una '
          'volta ogni tanto, e voi caricate l\'importo sul suo borsellino. '
          'Da quel momento ogni acquisto scala dal saldo, con lo storico '
          'sempre consultabile. Per i clienti di fiducia si può impostare '
          'un ', False), ('fido', True),
         (': continuano a ordinare anche a saldo zero, entro la soglia che '
          'decidete voi.', False)],
    ], accent=HEX_GREEN)

    step(doc, 3, 'Ordina dal telefono', [
        [('Menu con foto, prezzi e allergeni. Il cliente sceglie i prodotti '
          'oppure ', False), ('compone il suo panino, insalata o poke', True),
         (' ingrediente per ingrediente, con il prezzo che si aggiorna in '
          'diretta. Alla conferma sceglie l\'orario di ritiro fra gli slot '
          'liberi: l\'importo viene scalato subito e l\'ordine arriva in '
          'cucina già pagato.', False)],
    ], accent=HEX_GREEN)

    step(doc, 4, 'Riceve l\'avviso e ritira', [
        [('Quando la cucina segna l\'ordine pronto, il telefono del cliente '
          'suona: Telegram se collegato, altrimenti email. Al banco ritira '
          'il pacchetto con il tagliando e lo scontrino: ', False),
         ('nessun pagamento, nessuna coda', True), ('.', False)],
    ], accent=HEX_GREEN)

    step(doc, 5, 'Accumula punti e premi', [
        [('Ogni euro speso vale punti; alla soglia che impostate, il premio '
          'si trasforma in credito. Il cliente vede punti e soglia dalla sua '
          'pagina: la fedeltà lavora da sola.', False)],
    ], accent=HEX_GREEN)

    _page_break(doc)

    # ══ I tre canali di vendita ═══════════════════════════════════════════
    heading(doc, 'I tre canali di vendita', 1, before=0)
    body(doc, 'Tre modi diversi di vendere, un solo magazzino e un solo '
              'incasso. Ogni canale copre un momento diverso della giornata.',
         color=GRAY)
    rule(doc)

    modulo(doc, '1 · Ordini su appuntamento (il canale principale)', PURPLE,
           'Il cliente ordina dal telefono — anche la sera prima o dalla '
           'scrivania — e sceglie un orario di ritiro fra gli slot che avete '
           'configurato (ad esempio ogni quarto d\'ora, dalle 11:45 alle '
           '13:30, con un tetto di ordini per slot). La cucina riceve tutto '
           'su un display, in ordine di ritiro.',
           [
               [('Slot con capienza: ', True),
                ('mai più ordini di quelli che la cucina può evadere.', False)],
               [('Builder visuale: ', True),
                ('panino, insalata e poke componibili ingrediente per '
                 'ingrediente, con scorte per singolo ingrediente.', False)],
               [('Disponibilità giornaliere: ', True),
                ('ogni prodotto ha i pezzi del giorno; finiti quelli, '
                 'l\'ordine si blocca da solo.', False)],
           ])

    modulo(doc, '2 · Il banco, per chi passa (caffè e cassa veloce)', BLUE,
           'Per il caffè e l\'acquisto d\'impulso c\'è il POS da banco: '
           'componete il conto sul tablet, il cliente inquadra il QR che '
           'compare sullo schermo e il conto si chiude dal suo portafoglio. '
           'Niente resto, niente battitura doppia.',
           [
               [('Articoli rapidi configurabili', True),
                (' (caffè, cappuccino, brioche...): il conto tipico si '
                 'compone in due tocchi.', False)],
               [('Il pagamento QR ', True),
                ('scala dal borsellino e accumula punti come ogni altro '
                 'acquisto.', False)],
           ])

    modulo(doc, '3 · Il cesto, a libero servizio', ORANGE,
           'I pezzi pre-preparati (tramezzini, panini, dolci) vanno in un '
           'cesto vicino alla cassa, ognuno con la sua etichetta QR generata '
           'dall\'applicazione. Il cliente prende il pezzo, inquadra il QR e '
           'paga da solo: la vendita si registra senza che nessuno tocchi la '
           'cassa.',
           [
               [('Etichette con prezzo, allergeni e scadenza', True),
                (' (24 ore), stampate in lotti dalla pagina della cucina.',
                 False)],
               [('L\'invenduto si ritira a fine giornata', True),
                (' annullando le etichette: il conteggio guida le quantità '
                 'del giorno dopo.', False)],
               [('Funzione attivabile', True),
                (': se il cesto non vi serve, si spegne dalle Impostazioni '
                 'e sparisce da tutta l\'applicazione.', False)],
           ])

    _page_break(doc)

    # ══ Pasti aziendali ═══════════════════════════════════════════════════
    heading(doc, 'I pasti aziendali convenzionati', 1, before=0)
    body(doc, 'Il modulo per servire le aziende del territorio con un menu '
              'del giorno a prezzo concordato. È il canale che trasforma il '
              'bar in una mensa leggera, senza la burocrazia di una mensa.',
         color=GRAY)
    rule(doc)

    step(doc, 1, 'Pubblicate il menu di domani', [
        [('Primo, secondo, contorno, bevanda e caffè, con allergeni, prezzo '
          'e numero massimo di coperti. I menu ricorrenti si salvano come '
          'modelli e si richiamano in un colpo.', False)],
    ], accent=HEX_NAVY)
    step(doc, 2, 'I dipendenti prenotano da soli', [
        [('Ognuno dal proprio telefono, scegliendo l\'orario di ritiro. '
          'Possono modificare o disdire fino a 30 minuti prima: la cucina '
          'lavora su numeri veri, non su stime.', False)],
    ], accent=HEX_NAVY)
    step(doc, 3, 'Servite con la lista alla mano', [
        [('La lista nominativa del giorno, per azienda e ordinata per '
          'cognome, si stampa in un clic. Al ritiro si spunta la presenza: '
          'quello che risulta consumato è quello che fatturerete.', False)],
    ], accent=HEX_NAVY)
    step(doc, 4, 'A fine mese, il conto è già fatto', [
        [('Per ogni azienda: il ', False),
         ('riepilogo mensile in PDF', True),
         (' con i pasti per dipendente, il dettaglio giorno per giorno e il '
          'totale fatturabile — pronto da allegare alla fattura. In più il '
          'registro presenze stampabile e il report del singolo giorno.',
          False)],
    ], accent=HEX_NAVY)

    box(doc, 'Per l\'azienda convenzionata', [
        [('Nessun buono cartaceo, nessun conteggio a mano: i dipendenti '
          'prenotano dal telefono, l\'azienda riceve a fine mese un '
          'documento nominativo e verificabile. La convenzione diventa un '
          'servizio di welfare che non costa tempo a nessuno.', False)],
    ], accent=HEX_GREEN, fill=HEX_LIGHT, label_color=GREEN)

    _page_break(doc)

    # ══ Gli strumenti del gestore ═════════════════════════════════════════
    heading(doc, 'Gli strumenti del gestore', 1, before=0)
    body(doc, 'Il backoffice raccoglie tutto quello che serve per governare '
              'il servizio. Ogni collaboratore vede solo le pagine del suo '
              'ruolo.', color=GRAY)
    rule(doc)

    modulo(doc, 'Il display di cucina', RED,
           'Lo schermo che scandisce il lavoro: gli ordini arrivano in '
           'colonna "Da preparare" con orario di ritiro in evidenza e '
           'dettaglio completo (per i prodotti composti, ingrediente per '
           'ingrediente). Si aggiorna da solo e avvisa con un suono.',
           [
               [('Stati chiari: ', True),
                ('in preparazione, pronto, consegnato — e a ogni passaggio '
                 'il cliente riceve la notifica.', False)],
               [('Prenotazioni future: ', True),
                ('la pagina del fabbisogno dice cosa preparare per i giorni '
                 'a venire, prodotto per prodotto.', False)],
           ])

    modulo(doc, 'Report e andamenti', NAVY,
           'Incassi per giorno e per prodotto, andamento del mese, storico '
           'completo delle transazioni. I numeri della giornata sono pronti '
           'alla chiusura, senza fogli di calcolo.',
           [
               [('Ogni movimento è tracciato', True),
                (': ricariche, acquisti, punti, premi — con data, cliente e '
                 'causale.', False)],
           ])

    modulo(doc, 'Magazzino e riordino', GREEN,
           'I consumabili sotto soglia si evidenziano da soli e la richiesta '
           'di riordino parte via email al fornitore direttamente '
           'dall\'applicazione.',
           [
               [('Scorte anche per gli ingredienti del builder', True),
                (': un ingrediente esaurito sparisce dalle scelte del '
                 'cliente.', False)],
           ])

    modulo(doc, 'Clienti, ruoli e permessi', PURPLE,
           'Anagrafica clienti con saldo, punti, storico e stato; ruoli '
           'separati per cassa, cucina e amministrazione, ognuno con i suoi '
           'permessi.',
           [
               [('Sondaggi sul menu: ', True),
                ('proponete le alternative di domani e i clienti votano dal '
                 'telefono.', False)],
               [('Prenotazione tavoli', True),
                (' a fasce orarie, attivabile solo se vi serve.', False)],
               [('Funzioni attivabili: ', True),
                ('tavoli, cesto e perfino il portafoglio prepagato si '
                 'accendono e spengono dalle Impostazioni — chi preferisce '
                 'incassare solo alla cassa usa tutto il resto senza '
                 'muovere denaro nell\'app, con le vendite comunque '
                 'registrate nei report.', False)],
           ])

    body(doc, 'La giornata operativa — dal pomeriggio precedente al primo '
              'ordine servito, con gli orari limite di ogni attività — è '
              'descritta passo per passo nel Manuale del gestore consegnato '
              'con l\'attivazione.', italic=True, color=GRAY)

    _page_break(doc)

    # ══ Dotazione ═════════════════════════════════════════════════════════
    heading(doc, 'La dotazione necessaria', 1, before=0)
    body(doc, 'Attrezzatura ordinaria, nessun hardware proprietario: se '
              'avete già tablet e stampanti, quasi certamente vanno bene.',
         color=GRAY)
    rule(doc)

    table_grid(doc,
               ['Postazione', 'Dispositivo', 'A cosa serve'],
               [
                   [('Cucina', DARK, True), 'Tablet o monitor con browser',
                    'Display degli ordini, sempre acceso'],
                   [('Banco', DARK, True), 'Tablet',
                    'Cassa veloce con pagamento QR'],
                   [('Backoffice', DARK, True), 'PC o portatile',
                    'Gestione, report, convenzioni'],
                   [('Stampe', DARK, True), 'Stampante termica 80 mm',
                    'Tagliandi ordine'],
                   [('Stampe', DARK, True), 'Stampante A4',
                    'Etichette del cesto, liste, report'],
                   [('Cassa', DARK, True), 'Registratore di cassa',
                    'Scontrino fiscale (apparecchio separato, già vostro)'],
               ],
               widths=[2.8, 5.4, 7.8])

    box(doc, 'Nota', [
        [('QuickLunch non sostituisce il registratore di cassa: lo '
          'scontrino fiscale resta compito dell\'apparecchio che già usate. '
          'La regola operativa è semplice: a ogni prodotto consegnato si '
          'allegano scontrino e tagliando.', False)],
    ], accent=HEX_ORANGE, fill=HEX_WARN, label_color=ORANGE)

    # ══ Tecnologia e sicurezza ════════════════════════════════════════════
    heading(doc, 'Tecnologia, dati e sicurezza', 1)
    body(doc, 'QuickLunch è fornito in modalità SaaS — Software as a '
              'Service: come l\'home banking, si usa dal browser e tutta '
              'l\'infrastruttura è a carico nostro.', color=GRAY)
    spacer(doc, 4)

    for parts in [
        [('Niente da installare: ', True),
         ('funziona su qualsiasi dispositivo con un browser, dal primo '
          'giorno.', False)],
        [('Sempre aggiornato: ', True),
         ('le nuove versioni arrivano da sole, senza interventi e senza '
          'fermi del servizio.', False)],
        [('Copie di sicurezza: ', True),
         ('backup gestiti da noi, più un backup scaricabile in autonomia '
          'dalle Impostazioni, in un formato leggibile.', False)],
        [('Accessi protetti: ', True),
         ('connessione cifrata (HTTPS), accesso con Google, verifica in due '
          'passaggi disponibile per gli account amministrativi.', False)],
        [('Dati personali: ', True),
         ('trattamento conforme al GDPR con nomina a responsabile ex art. '
          '28; alla cessazione i dati vi vengono restituiti in formato '
          'aperto.', False)],
        [('Cognomi non esposti: ', True),
         ('nelle liste, sui tagliandi e nei report i clienti compaiono con '
          'il nome per esteso e il cognome ridotto all\'iniziale — '
          '"Mario R." — così una lista appesa in cucina o consegnata a '
          'un\'azienda resta leggibile per chi deve lavorarci senza '
          'esporre l\'anagrafica.', False)],
        [('I dati sono vostri: ', True),
         ('clienti, storico e listini si esportano e si portano via, in '
          'qualsiasi momento.', False)],
    ]:
        bullet(doc, parts)

    _page_break(doc)

    # ══ Percorso di avvio ═════════════════════════════════════════════════
    heading(doc, 'Il percorso di avvio', 1, before=0)
    body(doc, 'Dall\'accordo al primo ordine reale servono pochi giorni, non '
              'mesi. Il percorso tipo:', color=GRAY)
    rule(doc)

    step(doc, 1, 'Configurazione (1 giorno)', [
        [('Attiviamo l\'ambiente con il vostro nome e i vostri dati: '
          'listino, categorie, slot di ritiro, articoli del banco, '
          'eventuali convenzioni.', False)],
    ], accent=HEX_BLUE)
    step(doc, 2, 'Prova con dati realistici (2-3 giorni)', [
        [('Carichiamo un mese di attività di prova per farvi vedere report '
          'e flussi con numeri veri. Quando avete preso la mano, i dati di '
          'prova si eliminano con un pulsante e si parte puliti.', False)],
    ], accent=HEX_BLUE)
    step(doc, 3, 'Formazione sul posto', [
        [('Mezza giornata con chi sta al banco e in cucina, con i manuali '
          'operativi alla mano: cucina, gestore, onboarding dei clienti.',
          False)],
    ], accent=HEX_BLUE)
    step(doc, 4, 'Primi clienti', [
        [('Locandina QR alla cassa, prime registrazioni, prime ricariche. '
          'I clienti si portano dietro gli altri: il bonus di benvenuto '
          'aiuta.', False)],
    ], accent=HEX_BLUE)
    step(doc, 5, 'Servizio a regime', [
        [('Assistenza continuativa nei giorni lavorativi, dalle 9:00 alle '
          '18:00, con presa in carico entro 4 ore lavorative.', False)],
    ], accent=HEX_BLUE)

    heading(doc, 'La documentazione consegnata', 2)
    for parts in [
        [('Manuale del gestore', True),
         (' — la giornata operativa ora per ora, con gli orari limite.',
          False)],
        [('Manuale operativo di cucina', True),
         (' — i tre flussi di preparazione, passo per passo.', False)],
        [('Onboarding del cliente', True),
         (' — come portare un cliente dalla registrazione al primo ordine.',
          False)],
        [('Guida del cliente', True),
         (' — da distribuire o stampare per i clienti finali.', False)],
        [('Dotazione e postazioni', True),
         (' — il documento tecnico dei dispositivi.', False)],
        [('Catalogo delle stampe', True),
         (' — tutte le stampe del sistema, con esempi reali.', False)],
    ]:
        bullet(doc, parts)

    box(doc, 'Le condizioni economiche', [
        [('Canone, quota sugli incassi e contributo di attivazione sono '
          'dettagliati nell\'', False),
         ('offerta commerciale', True),
         (' che accompagna questa presentazione, insieme alle condizioni '
          'contrattuali complete e alle proiezioni sui vostri volumi.',
          False)],
    ], accent=HEX_NAVY, fill=HEX_LIGHT, label_color=NAVY)

    # ══ Contatti ══════════════════════════════════════════════════════════
    rule(doc, color=HEX_NAVY)
    box(doc, 'Parliamone', [
        [('Una dimostrazione dal vivo richiede mezz\'ora: portiamo noi '
          'l\'ambiente di prova, con dati realistici, sui vostri volumi.',
          False)],
        [('Daniele Speziale — DS Consulting', True)],
        [('Email:  ', False), (EMAIL, True),
         ('        Cellulare:  ', False), (CELL, True)],
    ], accent=HEX_RED, fill=HEX_LIGHT, label_color=RED)

    body(doc, 'QuickLunch · Presentazione del prodotto · I moduli e i flussi '
              'descritti corrispondono all\'applicazione in uso.',
         size=8.5, color=GRAY, after=2)
    body(doc, '© 2024–26 DS Consulting  ·  %s' % CONTATTI, size=8.5,
         color=GRAY)

    return doc


def main():
    doc = build()
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] Documento salvato in: {OUT}')


if __name__ == '__main__':
    main()
