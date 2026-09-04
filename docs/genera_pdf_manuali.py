#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Converte in PDF i manuali .docx del kit, senza riscriverne il contenuto.

Il .docx resta l'unica fonte: qui viene riletto con python-docx e ricomposto
con fpdf2, la stessa libreria degli altri PDF dell'applicazione. Serve perche'
l'email di benvenuto allega la guida del cliente in PDF, e l'invio avviene in
produzione, dove non esistono Word ne' LibreOffice: il PDF va quindi prodotto
qui e versionato.

    python docs/genera_pdf_manuali.py

Uscite (entrambe scritte dallo stesso comando, cosi' non divergono):
    docs/manuali/<nome>.pdf          copia editoriale, accanto al .docx
    app/static/docs/<nome>.pdf       copia che l'applicazione allega
"""

import os
import shutil
import sys

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from fpdf import FPDF

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIR_DOCX = os.path.join(ROOT, 'docs', 'manuali')
DIR_APP = os.path.join(ROOT, 'app', 'static', 'docs')
FONT_DIR = os.path.join(ROOT, 'app', 'static', 'fonts')

# I manuali da avere anche in PDF. Aggiungere qui quando ne serve un altro.
DA_CONVERTIRE = ['guida_cliente', 'manuale_onboarding_cliente']

RED = (233, 69, 96)
NAVY = (15, 52, 96)
DARK = (26, 26, 46)
DGRAY = (80, 80, 95)
LGRAY = (205, 205, 214)
VLIGHT = (247, 248, 250)
FONT = 'PTSansNarrow'

CONTATTI = 'DS Consulting · Daniele Speziale · dspeziale@gmail.com · +39 352 0150489'


class ManualePDF(FPDF):
    """Impaginazione semplice: intestazione con il titolo, numero di pagina."""

    titolo_doc = ''

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font(FONT, '', 8)
        self.set_text_color(*DGRAY)
        self.cell(0, 5, self.titolo_doc, align='L')
        self.ln(6)
        self.set_draw_color(*LGRAY)
        self.set_line_width(0.2)
        self.line(15, self.get_y(), 195, self.get_y())
        self.ln(3)

    def footer(self):
        self.set_y(-14)
        self.set_font(FONT, '', 7.5)
        self.set_text_color(*DGRAY)
        self.cell(0, 4, '© 2024–26 DS Consulting  ·  %s' % CONTATTI,
                  align='C')
        self.ln(3.6)
        self.set_font(FONT, '', 7.5)
        self.cell(0, 4, 'pagina %d' % self.page_no(), align='C')


def _blocchi(doc):
    """Paragrafi e tabelle nell'ordine in cui stanno nel documento.

    Iterare su doc.paragraphs e poi su doc.tables mescolerebbe l'ordine: il
    kit usa tabelle anche per riquadri e passi numerati, che stanno in mezzo
    al testo.
    """
    corpo = doc.element.body
    for figlio in corpo.iterchildren():
        if figlio.tag.endswith('}p'):
            yield Paragraph(figlio, doc)
        elif figlio.tag.endswith('}tbl'):
            yield Table(figlio, doc)


def _pulisci(testo):
    """Via i caratteri che il font non ha, e gli spazi inutili."""
    fuori = {'→': '>', '←': '<', '↵': ' ',
             '•': '-', '': '-',
             '⋮': ':', 'ℹ': 'i', '✓': 'v', '✔': 'v'}
    for vecchio, nuovo in fuori.items():
        testo = testo.replace(vecchio, nuovo)
    # Gli emoji del kit non esistono in PT Sans Narrow: si toglierebbero da
    # soli con un avviso, meglio farlo qui in silenzio.
    return ''.join(c for c in testo if ord(c) < 0x2500).strip()


def _scrivi_titolo(pdf, testo, livello):
    dimensioni = {1: 15, 2: 12.5, 3: 11}
    colori = {1: NAVY, 2: NAVY, 3: DARK}
    if livello == 1 and pdf.get_y() > 200:
        pdf.add_page()
    pdf.ln(3 if livello > 1 else 5)
    pdf.set_font(FONT, 'B', dimensioni.get(livello, 11))
    pdf.set_text_color(*colori.get(livello, DARK))
    pdf.multi_cell(180, dimensioni.get(livello, 11) * 0.48, testo)
    if livello == 1:
        pdf.set_draw_color(*RED)
        pdf.set_line_width(0.5)
        pdf.line(15, pdf.get_y() + 1, 195, pdf.get_y() + 1)
        pdf.ln(3)
    else:
        pdf.ln(1.5)


def _scrivi_paragrafo(pdf, testo, elenco=False):
    pdf.set_font(FONT, '', 10)
    pdf.set_text_color(*DGRAY)
    if elenco:
        pdf.set_x(19)
        pdf.multi_cell(176, 4.8, '·  ' + testo)
    else:
        pdf.set_x(15)
        pdf.multi_cell(180, 4.8, testo)
    pdf.ln(1.2)


def _scrivi_tabella(pdf, tabella):
    """Le tabelle del kit sono riquadri, passi numerati e griglie di dati."""
    righe = []
    for riga in tabella.rows:
        celle = [_pulisci(c.text) for c in riga.cells]
        # Le celle unite si ripetono: si tengono i valori distinti.
        viste, pulite = set(), []
        for c in celle:
            if c and c not in viste:
                viste.add(c)
                pulite.append(c)
        if pulite:
            righe.append(pulite)
    if not righe:
        return

    # Riquadro (una sola colonna): fondo chiaro e barra rossa a sinistra.
    if all(len(r) == 1 for r in righe):
        testo = '\n'.join(r[0] for r in righe)
        pdf.set_x(15)
        y0 = pdf.get_y()
        pdf.set_font(FONT, '', 9.5)
        pdf.set_text_color(*DARK)
        pdf.set_fill_color(*VLIGHT)
        pdf.multi_cell(180, 4.6, testo, fill=True)
        pdf.set_draw_color(*RED)
        pdf.set_line_width(1.2)
        pdf.line(15.4, y0, 15.4, pdf.get_y())
        pdf.ln(2.4)
        return

    for riga in righe:
        pdf.set_x(15)
        pdf.set_font(FONT, 'B', 9.5)
        pdf.set_text_color(*DARK)
        prima = riga[0]
        pdf.multi_cell(180, 4.6, prima)
        for pezzo in riga[1:]:
            pdf.set_x(21)
            pdf.set_font(FONT, '', 9.5)
            pdf.set_text_color(*DGRAY)
            pdf.multi_cell(174, 4.6, pezzo)
        pdf.ln(1)
    pdf.ln(1.4)


def converti(nome):
    percorso_docx = os.path.join(DIR_DOCX, nome + '.docx')
    if not os.path.isfile(percorso_docx):
        raise SystemExit('manca %s: genera prima il .docx' % percorso_docx)
    doc = Document(percorso_docx)

    pdf = ManualePDF(format='A4')
    pdf.add_font(FONT, '', os.path.join(FONT_DIR, 'PTSansNarrow-Regular.ttf'))
    pdf.add_font(FONT, 'B', os.path.join(FONT_DIR, 'PTSansNarrow-Bold.ttf'))
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(True, margin=18)

    # Titolo del documento: il primo testo in grassetto grande della
    # copertina, che nel kit e' una tabella a piena pagina.
    titolo = ''
    for blocco in _blocchi(doc):
        if isinstance(blocco, Table):
            for riga in blocco.rows:
                for cella in riga.cells:
                    for par in cella.paragraphs:
                        for run in par.runs:
                            if (run.font.size and run.font.size.pt >= 20
                                    and _pulisci(run.text)):
                                titolo = _pulisci(run.text)
                                break
                        if titolo:
                            break
                    if titolo:
                        break
            if titolo:
                break
    titolo = titolo or nome.replace('_', ' ').capitalize()
    pdf.titolo_doc = titolo

    pdf.add_page()
    pdf.set_font(FONT, 'B', 9)
    pdf.set_text_color(*RED)
    pdf.cell(0, 5, 'QUICKLUNCH')
    pdf.ln(7)
    pdf.set_font(FONT, 'B', 22)
    pdf.set_text_color(*NAVY)
    pdf.multi_cell(180, 10, titolo)
    pdf.set_draw_color(*RED)
    pdf.set_line_width(0.7)
    pdf.line(15, pdf.get_y() + 2, 195, pdf.get_y() + 2)
    pdf.ln(8)

    saltato_copertina = False
    for blocco in _blocchi(doc):
        if isinstance(blocco, Table):
            if not saltato_copertina:
                # La prima tabella e' la copertina, gia' ricomposta sopra.
                saltato_copertina = True
                continue
            _scrivi_tabella(pdf, blocco)
            continue
        testo = _pulisci(blocco.text)
        if not testo:
            continue
        stile = (blocco.style.name or '') if blocco.style else ''
        if stile.startswith('Heading'):
            try:
                livello = int(stile.split()[-1])
            except ValueError:
                livello = 2
            _scrivi_titolo(pdf, testo, livello)
        elif 'List' in stile:
            _scrivi_paragrafo(pdf, testo, elenco=True)
        else:
            grande = any(r.font.size and r.font.size.pt >= 14
                         for r in blocco.runs)
            if grande:
                _scrivi_titolo(pdf, testo, 2)
            else:
                _scrivi_paragrafo(pdf, testo)

    os.makedirs(DIR_APP, exist_ok=True)
    uscita_app = os.path.join(DIR_APP, nome + '.pdf')
    pdf.output(uscita_app)
    uscita_docs = os.path.join(DIR_DOCX, nome + '.pdf')
    shutil.copyfile(uscita_app, uscita_docs)
    return uscita_app, len(pdf.pages)


def main():
    for nome in DA_CONVERTIRE:
        percorso, pagine = converti(nome)
        print('[OK] %s (%d pagine, %d byte)'
              % (percorso, pagine, os.path.getsize(percorso)))
        print('     copia anche in %s'
              % os.path.join(DIR_DOCX, nome + '.pdf'))


if __name__ == '__main__':
    main()
