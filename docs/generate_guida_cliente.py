#!/usr/bin/env python3
"""
Genera docs/guida_cliente.docx — Guida Cliente Completa per QuickLunch.

Guida dedicata esclusivamente al ruolo Cliente: dopo averla letta, un cliente
deve essere in grado di registrarsi, ordinare, pagare, prenotare un tavolo,
gestire wallet e punti fedeltà, votare i sondaggi e usare il pasto aziendale
(se dipendente convenzionato) senza bisogno di assistenza.
"""

import os
import sys

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generate_guide import (
    set_document_defaults, _page_break, h1, h2, h3, body_para,
    info_box, info_box_color, step_row, data_table, role_badge,
    workflow_table, spacer, divider, _footer_copyright,
    _no_borders, _table_width, _row_height, _cell_shd, _cell_margins,
    _cell_vAlign, _run_font, _p_spacing, _cell_border, _set_col_width,
    HEX_RED, HEX_NAVY, HEX_DARK, HEX_LIGHT, HEX_WHITE, HEX_GREEN,
    HEX_PURPL, HEX_TEAL, HEX_ORNG,
    RED, NAVY, DARK, DGRAY, GRAY, WHITE, GREEN, PURPL, TEAL, ORNG,
)

OUT = os.path.join(os.path.dirname(__file__), 'manuali', 'guida_cliente.docx')


# ── Cover page dedicata al cliente ──────────────────────────────────────────

def build_cover_cliente(doc):
    tbl = doc.add_table(rows=3, cols=1)
    _no_borders(tbl)
    _table_width(tbl, 21.0)
    _row_height(tbl.rows[0], 4.5)
    _row_height(tbl.rows[1], 19.0)
    _row_height(tbl.rows[2], 6.2)

    for row in tbl.rows:
        _cell_shd(row.cells[0], HEX_DARK)
        _cell_margins(row.cells[0], top=0, bottom=0, left=0, right=0)

    cc = tbl.rows[1].cells[0]
    _cell_shd(cc, HEX_DARK)
    _cell_margins(cc, top=0, bottom=0, left=300, right=300)
    _cell_vAlign(cc, 'center')

    def cp(text='', size=14, bold=False, color=WHITE,
           align=WD_ALIGN_PARAGRAPH.CENTER, after=10):
        p = cc.add_paragraph()
        p.alignment = align
        _p_spacing(p, before=0, after=after)
        if text:
            r = p.add_run(text)
            _run_font(r, size=size, bold=bold, color=color)
        return p

    p0 = cc.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(p0, before=0, after=16)
    _run_font(p0.add_run('👤🍽️'), size=56, color=WHITE)

    cp('QuickLunch', size=54, bold=True, color=RGBColor(0x8e, 0x44, 0xad), after=0)
    cp('PRANZO', size=54, bold=True, color=WHITE, after=8)
    cp('GUIDA CLIENTE COMPLETA', size=18, bold=True,
       color=RGBColor(0xc0, 0xd0, 0xe0), after=24)
    cp('Tutto quello che ti serve per ordinare, pagare,\n'
       'prenotare un tavolo e usare i punti fedeltà — da solo, in pochi minuti.',
       size=13, color=RGBColor(0x90, 0xa8, 0xc0), after=28)

    bullets = [
        '🍽️  Ordinare dal menu', '💳  Pagare con il wallet',
        '🥪  Comporre panino/insalata/poke', '🪑  Prenotare un tavolo',
        '⭐  Punti fedeltà e premi', '🗳️  Votare i sondaggi',
    ]
    pr = cc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pr, before=0, after=20)
    for i, name in enumerate(bullets):
        _run_font(pr.add_run(name), size=11, color=RGBColor(0xb0, 0xc8, 0xe0))
        if i < len(bullets) - 1:
            _run_font(pr.add_run('  ·  '), size=10, color=RGBColor(0x50, 0x60, 0x70))

    cp('Versione 1.0  ·  2026', size=10,
       color=RGBColor(0x60, 0x70, 0x80), after=0)

    bc = tbl.rows[2].cells[0]
    _cell_shd(bc, HEX_PURPL)
    _cell_margins(bc, top=60, bottom=60, left=300, right=300)
    _cell_vAlign(bc, 'center')
    pb = bc.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pb, before=0, after=6)
    _run_font(pb.add_run('📱  NON SERVE INSTALLARE NIENTE — FUNZIONA DAL BROWSER'),
              size=12, bold=True, color=WHITE)
    pb2 = bc.add_paragraph()
    pb2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pb2, before=0, after=0)
    _run_font(pb2.add_run(
        'Apri il link del tuo locale dal telefono o dal computer e segui questa guida passo passo.'),
        size=10, color=RGBColor(0xff, 0xff, 0xff))


def build_toc_cliente(doc):
    h2(doc, '📋  Indice dei contenuti')
    spacer(doc, 4)

    toc_items = [
        ('1',  '🚀  Iniziare',                   'Registrazione, Google login, sessione, PWA, Home Dashboard'),
        ('2',  '👤  Il Mio Profilo',              'Dati personali, cambio password, avatar'),
        ('3',  '🍽️  Ordinare dal Menu',           'Sfoglia, allergeni, carrello, slot, Adesso al banco'),
        ('4',  '🥪  Crea il Tuo Piatto',          'Builder panino, insalata, poke bowl su misura'),
        ('5',  '📦  I Miei Ordini',               'Storico, stati, annullamento, codice ordine'),
        ('6',  '💳  Wallet e Punti Fedeltà',      'Saldo, ricarica, storico, riscatto premi'),
        ('7',  '📲  Paga al Banco — QR',          'Pagamento QR istantaneo generato dallo staff'),
        ('8',  '🪑  Prenotare un Tavolo',         'Fasce orarie, selezione tavolo, prenotazioni'),
        ('9',  '🏢  Pasto Aziendale',             'Solo dipendenti con convenzione attiva'),
        ('10', '🗳️  Sondaggi',                   'Vota il menu, risultati in tempo reale'),
        ('11', '🔔  Notifiche Telegram',          'Conferme e avvisi automatici via Telegram'),
        ('12', '❓  Domande Frequenti',           '12 risposte ai problemi più comuni'),
        ('13', '🧭  Scheda Riassuntiva',          'Tutte le azioni principali in una tabella'),
    ]

    tbl = doc.add_table(rows=len(toc_items), cols=3)
    _no_borders(tbl)
    _table_width(tbl, 17.6)
    _set_col_width(tbl, 0, 1.2)
    _set_col_width(tbl, 1, 5.8)
    _set_col_width(tbl, 2, 10.6)

    for ri, (num, title, desc) in enumerate(toc_items):
        bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        row = tbl.rows[ri]
        _cell_shd(row.cells[0], HEX_PURPL)
        _cell_margins(row.cells[0], top=70, bottom=70, left=60, right=60)
        _cell_vAlign(row.cells[0], 'center')
        pn = row.cells[0].paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _p_spacing(pn, before=0, after=0)
        _run_font(pn.add_run(num), size=14, bold=True, color=WHITE)

        _cell_shd(row.cells[1], bg)
        _cell_margins(row.cells[1], top=70, bottom=70, left=100, right=60)
        _cell_border(row.cells[1], bottom='DEE2E6')
        pt = row.cells[1].paragraphs[0]
        _p_spacing(pt, before=0, after=0)
        _run_font(pt.add_run(title), size=12, bold=True, color=DARK)

        _cell_shd(row.cells[2], bg)
        _cell_margins(row.cells[2], top=70, bottom=70, left=100, right=60)
        _cell_border(row.cells[2], bottom='DEE2E6')
        pd = row.cells[2].paragraphs[0]
        _p_spacing(pd, before=0, after=0)
        _run_font(pd.add_run(desc), size=11, color=GRAY)


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 1 — INIZIARE
# ══════════════════════════════════════════════════════════════════════════════

def s1_iniziare(doc):
    h1(doc, '1', 'Iniziare', '🚀', accent=HEX_PURPL)

    role_badge(doc, '👤', 'Cliente',
               'QuickLunch Pranzo è l\'app del tuo locale: ordini il pranzo, lo paghi dal telefono '
               'con il tuo wallet, prenoti un tavolo e accumuli punti fedeltà. Non serve scaricare '
               'nessuna applicazione: funziona direttamente nel browser dello smartphone o del computer.',
               HEX_PURPL)
    spacer(doc, 8)

    h2(doc, '1.1  Registrazione con email e password')
    step_row(doc, 1, 'Apri il link del locale',
             'Il tuo bar/mensa ti fornisce un indirizzo web (es. https://pranzo.barcentrale.it). '
             'Aprilo dal browser del telefono o del PC')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca "Registrati"',
             'Sotto al modulo di accesso trovi il link per creare un nuovo account')
    spacer(doc, 4)
    step_row(doc, 3, 'Inserisci email e password',
             'La password deve avere almeno 6 caratteri. Ripetila nel campo "Conferma password"')
    spacer(doc, 4)
    step_row(doc, 4, 'Tocca "Registrati"',
             'Il sistema crea il tuo account, ti assegna uno username automatico e ti fa entrare '
             'subito, già loggato')
    spacer(doc, 8)

    info_box(doc, 'Se l\'email è già registrata, il sistema te lo segnala: in quel caso usa "Accedi" con la '
             'password che avevi scelto. Se l\'hai dimenticata, chiedi all\'amministratore del locale di '
             'reimpostarla dal pannello clienti.', style='tip')
    spacer(doc, 8)

    h2(doc, '1.2  Accedi con Google (opzione alternativa)')
    body_para(doc,
        'Nella pagina di login trovi anche il pulsante "Accedi con Google". '
        'È compatibile con account Google personali e con account GSuite aziendali. '
        'Tocca il pulsante, scegli il tuo account Google dalla finestra che si apre '
        'e sei subito dentro — senza dover impostare una password separata.')
    spacer(doc, 6)

    info_box_color(doc,
        'Se usi Google per accedere, il tuo avatar Google viene recuperato automaticamente '
        'e mostrato nel tuo profilo. Non devi caricare nessuna foto manualmente.',
        bg='EBF5FB', border=HEX_TEAL, icon='💡')
    spacer(doc, 8)

    h2(doc, '1.3  Login e sessione')
    step_row(doc, 1, 'Apri il link del locale', 'Stesso indirizzo della registrazione')
    spacer(doc, 4)
    step_row(doc, 2, 'Inserisci email e password (oppure tocca "Accedi con Google")',
             'Usa le credenziali scelte in fase di registrazione')
    spacer(doc, 4)
    step_row(doc, 3, 'Tocca "Accedi"',
             'Il sistema ti ricorda per circa 30 giorni: non dovrai rifare il login ogni volta')
    spacer(doc, 8)

    h2(doc, '1.4  Aggiunta a schermata Home (PWA)')
    body_para(doc,
        'QuickLunch funziona come una Progressive Web App (PWA): puoi aggiungerla alla schermata '
        'Home del telefono e aprirla come una vera app, senza dover riscrivere il link ogni volta.')
    spacer(doc, 6)

    data_table(doc,
        ['Dispositivo', 'Come aggiungere alla schermata Home'],
        [
            ['iPhone / iPad (Safari)',
             'Tocca il pulsante Condividi (quadrato con freccia su) → scorri e tocca "Aggiungi a schermata Home" → tocca "Aggiungi"'],
            ['Android (Chrome)',
             'Tocca il menu ⋮ in alto a destra → "Aggiungi a schermata Home" oppure "Installa app" → tocca "Aggiungi"'],
        ],
        col_widths=[4.5, 13.1])
    spacer(doc, 8)

    h2(doc, '1.5  La Home Dashboard')
    body_para(doc, 'La Home (dashboard) è la tua pagina di partenza. Appena entri trovi 4 riquadri KPI:')
    spacer(doc, 4)
    data_table(doc,
        ['Riquadro KPI', 'Cosa mostra'],
        [
            ['Saldo Wallet',       'Il credito disponibile nel tuo portafoglio digitale'],
            ['Punti Fedeltà',      'I punti accumulati e quanto manca alla prossima soglia di riscatto'],
            ['Ordini Oggi',        'Quanti ordini hai fatto oggi e il loro stato corrente'],
            ['Tavoli Prenotati',   'Le prenotazioni tavolo attive per oggi'],
        ],
        col_widths=[4.8, 12.8])
    spacer(doc, 6)

    body_para(doc,
        'Sotto i riquadri KPI trovi le azioni rapide (pulsanti diretti a Menu, Carrello, Paga al Banco, '
        'Tavoli) e le ultime transazioni del wallet. Dal menu laterale (icona ☰ in alto a sinistra su '
        'telefono) raggiungi tutte le sezioni: Menu, Carrello, I Miei Ordini, Wallet & Fedeltà, Tavoli, '
        'Le Mie Prenotazioni, Vota il Menu e — se sei un dipendente con convenzione aziendale attiva — '
        'Pasto Aziendale.')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 2 — IL MIO PROFILO
# ══════════════════════════════════════════════════════════════════════════════

def s2_profilo(doc):
    h1(doc, '2', 'Il Mio Profilo', '👤', accent=HEX_PURPL)

    role_badge(doc, '👤', 'Profilo personale',
               'Dalla sezione Profilo puoi aggiornare i tuoi dati anagrafici, cambiare la password '
               'e controllare le impostazioni del tuo account.',
               HEX_PURPL)
    spacer(doc, 8)

    h2(doc, '2.1  Come accedere alle impostazioni profilo')
    step_row(doc, 1, 'Apri il menu laterale (☰)',
             'Tocca l\'icona in alto a sinistra da smartphone, oppure usa la barra laterale su desktop')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca il tuo nome o l\'icona account',
             'In cima al menu trovi il tuo nome utente e l\'icona profilo: tocca per aprire le impostazioni')
    spacer(doc, 8)

    h2(doc, '2.2  Dati modificabili')
    body_para(doc, 'Dal pannello profilo puoi aggiornare i seguenti campi:')
    spacer(doc, 4)
    data_table(doc,
        ['Campo', 'Note'],
        [
            ['Nome',            'Il nome visualizzato agli altri (es. operatori del locale)'],
            ['Cognome',         'Cognome anagrafico'],
            ['Numero di telefono', 'Utile per eventuali comunicazioni del locale'],
            ['Data di nascita', 'Facoltativa — usata per promozioni e auguri'],
            ['Indirizzo',       'Facoltativo — non richiesto per l\'uso normale dell\'app'],
        ],
        col_widths=[4.8, 12.8])
    spacer(doc, 6)

    body_para(doc,
        'Dopo aver modificato un campo, tocca "Salva" per confermare. '
        'Le modifiche sono immediate e visibili subito nel tuo account.')
    spacer(doc, 8)

    h2(doc, '2.3  Cambio password')
    step_row(doc, 1, 'Vai nelle impostazioni profilo', 'Come descritto al punto 2.1')
    spacer(doc, 4)
    step_row(doc, 2, 'Trova la sezione "Cambia password"',
             'Inserisci la password attuale nel primo campo')
    spacer(doc, 4)
    step_row(doc, 3, 'Inserisci la nuova password',
             'Almeno 6 caratteri. Ripetila nel campo "Conferma nuova password"')
    spacer(doc, 4)
    step_row(doc, 4, 'Tocca "Aggiorna password"',
             'La nuova password è attiva da subito. Al prossimo accesso usa quella nuova')
    spacer(doc, 8)

    info_box(doc,
        'Se hai effettuato l\'accesso tramite Google, il cambio password riguarda solo '
        'l\'accesso email/password. Il login con Google non è influenzato.',
        style='info')
    spacer(doc, 8)

    h2(doc, '2.4  Avatar')
    body_para(doc,
        'Se hai creato l\'account con email e password, nel profilo puoi vedere un avatar generato '
        'automaticamente con le tue iniziali. '
        'Se invece hai eseguito l\'accesso con Google, l\'immagine del profilo Google viene '
        'recuperata automaticamente e usata come avatar in QuickLunch — non devi fare nulla.')
    spacer(doc, 6)

    info_box_color(doc,
        'Il nome visualizzato al locale (cassiere, operatori) corrisponde al nome e cognome '
        'che hai inserito nel profilo. Tienili aggiornati per facilitare il riconoscimento al banco.',
        bg='F4ECF7', border=HEX_PURPL, icon='💡')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 3 — ORDINARE DAL MENU
# ══════════════════════════════════════════════════════════════════════════════

def s3_ordinare(doc):
    h1(doc, '3', 'Ordinare dal Menu', '🍽️', accent=HEX_PURPL)

    workflow_table(doc, [
        ('🍽️', 'Apri Menu', 'Sfoglia categorie'),
        ('➕', 'Aggiungi', 'Tocca il prodotto'),
        ('🛒', 'Carrello', 'Scegli lo slot'),
        ('💳', 'Paga', 'Wallet'),
        ('✅', 'Conferma', 'Ricevi il codice ordine'),
    ], accent=HEX_NAVY)
    spacer(doc, 8)

    h2(doc, '3.1  Sfogliare per categoria e aggiungere al carrello')
    step_row(doc, 1, 'Vai in "Menu"', 'Dal menu laterale, o dal bottone in Home')
    spacer(doc, 4)
    step_row(doc, 2, 'Scegli una categoria',
             'Es. Primi, Secondi, Contorni, Bevande, Dolci, Senza Glutine — '
             'ogni categoria ha un\'icona colorata')
    spacer(doc, 4)
    step_row(doc, 3, 'Imposta la quantità e tocca "Aggiungi"',
             'Il campo quantità è di default 1. Tocca "Aggiungi" per mettere il prodotto nel carrello: '
             'compare un messaggio di conferma in alto')
    spacer(doc, 4)
    step_row(doc, 4, 'Ripeti per altri prodotti',
             'Puoi mischiare prodotti di categorie diverse nello stesso ordine')
    spacer(doc, 8)

    info_box(doc, 'I prodotti con disponibilità "0" per oggi sono esauriti: la disponibilità si rinnova '
             'ogni giorno automaticamente.', style='info')
    spacer(doc, 8)

    h2(doc, '3.2  Allergeni')
    body_para(doc,
        'Ogni scheda prodotto mostra gli allergeni presenti nella preparazione. '
        'Gli allergeni sono elencati sotto il nome e la descrizione del prodotto, '
        'con icone colorate per i più comuni (glutine, lattosio, frutta a guscio, ecc.). '
        'Controlla sempre gli allergeni prima di aggiungere un prodotto al carrello se hai '
        'intolleranze o allergie alimentari.')
    spacer(doc, 6)

    info_box_color(doc,
        'In caso di dubbi sugli ingredienti o su possibili contaminazioni crociate, '
        'chiedi conferma al personale del locale prima di ordinare.',
        bg='FEF9E7', border=HEX_ORNG, icon='⚠️')
    spacer(doc, 8)

    h2(doc, '3.3  Aggiornare quantità o rimuovere prodotti dal carrello')
    step_row(doc, 1, 'Apri "Carrello"', 'Icona carrello nel menu, o link diretto')
    spacer(doc, 4)
    step_row(doc, 2, 'Modifica le quantità',
             'Cambia il numero accanto a ogni riga e conferma: il totale si aggiorna automaticamente')
    spacer(doc, 4)
    step_row(doc, 3, 'Rimuovi un prodotto',
             'Tocca l\'icona cestino sulla riga che vuoi togliere')
    spacer(doc, 4)
    step_row(doc, 4, 'Controlla il totale',
             'In fondo al carrello vedi il totale da pagare e il tuo saldo wallet attuale')
    spacer(doc, 8)

    h2(doc, '3.4  Scegliere lo slot di ritiro')
    body_para(doc,
        'Prima di confermare l\'ordine devi scegliere lo slot orario in cui passerai a ritirarlo. '
        'Ogni slot mostra l\'orario e il numero di posti disponibili. '
        'Se uno slot è pieno, scegli un orario alternativo. '
        'Il codice ordine generato avrà il formato QuickLunch-YYMMDD-HHMM-NNNN '
        '(es. QuickLunch-260630-1145-0042).')
    spacer(doc, 8)

    h2(doc, '3.5  "Adesso al banco" — opzione speciale')
    body_para(doc,
        'Prima della lista degli slot orari trovi una radio "Adesso al banco": selezionandola '
        'non scegli nessuno slot. Il codice ordine generato avrà il formato '
        'QuickLunch-YYMMDD-BANCO-NNNN. Utile se sei già fisicamente al bancone e vuoi pagare '
        'un ordine personalizzato subito, senza attendere uno slot prefissato.')
    spacer(doc, 8)

    info_box_color(doc,
        'DISTINZIONE IMPORTANTE\n\n'
        '"Adesso al banco" (§3.5): hai fatto un ordine dal menu ma lo RITIRI SUBITO al banco '
        'senza prenotare uno slot orario. Sei TU a comporre l\'ordine e a scegliere questa opzione.\n\n'
        '"Paga al Banco" (§7): è LO STAFF a generare un QR code per vendite rapide (es. caffè, brioche). '
        'Tu inquadri il QR e paghi — senza aver scelto nulla dal menu.',
        bg='FEF9E7', border=HEX_ORNG, icon='⚡')
    spacer(doc, 8)

    h2(doc, '3.6  Conferma ordine')
    step_row(doc, 1, 'Tocca "Conferma ordine"',
             'Se il saldo wallet è sufficiente, l\'ordine va a buon fine all\'istante')
    spacer(doc, 4)
    step_row(doc, 2, 'Ricevi il codice ordine',
             'Es. QuickLunch-260630-1145-0042 (oppure BANCO se hai scelto "Adesso al banco"). '
             'Mostralo al cassiere al momento del ritiro')
    spacer(doc, 4)
    step_row(doc, 3, 'Wallet scalato + punti fedeltà assegnati',
             '10 punti per ogni euro speso, accreditati subito. '
             'Se hai Telegram collegato, ricevi una notifica con orario di ritiro e totale pagato')
    spacer(doc, 8)

    info_box(doc, 'Se il messaggio dice "Saldo wallet insufficiente", vieni reindirizzato automaticamente '
             'alla pagina Wallet. Chiedi al cassiere di ricaricare il tuo wallet (vedi Sezione 6), '
             'poi torna nel carrello e conferma di nuovo: i prodotti restano salvati.', style='warning')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 4 — CREA IL TUO PIATTO (BUILDER)
# ══════════════════════════════════════════════════════════════════════════════

def s4_builder(doc):
    h1(doc, '4', 'Crea il Tuo Piatto', '🥪', accent=HEX_PURPL)

    role_badge(doc, '🥪', 'Builder',
               'Oltre ai prodotti già pronti del menu, puoi comporre tu stesso un panino, '
               'un\'insalata o una poke bowl, scegliendo ogni ingrediente. Il prezzo si calcola automaticamente.',
               HEX_ORNG)
    spacer(doc, 8)

    h2(doc, '4.1  I tre tipi disponibili')
    data_table(doc,
        ['Tipo', 'Prezzo base', 'Caratteristica'],
        [
            ['Panino',       '3,50€', 'Scegli pane, proteina, verdure, salsa, extra. Opzione "alla griglia" disponibile'],
            ['Insalata',     '3,00€', 'Base di insalata, verdure, proteina, condimento, extra'],
            ['Poke Bowl',    '4,00€', 'Base di riso o altro, proteina, verdure, topping, salsa'],
        ],
        col_widths=[3.5, 3.0, 11.1])
    spacer(doc, 8)

    h2(doc, '4.2  Selezione step per step')
    step_row(doc, 1, 'Vai in "Componi il tuo piatto"', 'Dal menu laterale o dal pulsante in Menu')
    spacer(doc, 4)
    step_row(doc, 2, 'Scegli il tipo', 'Panino, Insalata o Poke Bowl')
    spacer(doc, 4)
    step_row(doc, 3, 'Seleziona gli ingredienti per categoria',
             'Base/pane, proteine, verdure, condimenti, extra — ogni categoria mostra '
             'quante scelte sono permesse e quali sono obbligatorie')
    spacer(doc, 4)
    step_row(doc, 4, 'Rispetta le categorie obbligatorie',
             'Es. "Scegli il pane" è obbligatoria: senza una scelta non puoi proseguire')
    spacer(doc, 8)

    h2(doc, '4.3  Ingredienti con prezzo extra')
    body_para(doc,
        'Alcuni ingredienti hanno un costo aggiuntivo chiaramente indicato accanto al nome '
        '(es. +0.50€, +1.00€). Il totale si aggiorna automaticamente a ogni selezione, '
        'in modo da avere sempre il prezzo esatto sotto controllo.')
    spacer(doc, 8)

    h2(doc, '4.4  Opzione "alla griglia" (solo panini)')
    body_para(doc,
        'Per i panini è disponibile una casella opzionale 🔥 "Alla griglia" (+0,30€). '
        'Spuntandola il panino verrà scaldato/grigliato prima della consegna. '
        'L\'opzione è visibile solo se attivata dal locale per quel giorno.')
    spacer(doc, 8)

    h2(doc, '4.5  Totale automatico aggiornato a ogni selezione')
    body_para(doc,
        'In fondo alla pagina builder il prezzo totale si aggiorna in tempo reale: '
        'prezzo base + costo extra di ogni ingrediente selezionato + eventuale costo "alla griglia". '
        'Non ci sono sorprese al momento del pagamento.')
    spacer(doc, 8)

    h2(doc, '4.6  Aggiungere al carrello e completare l\'ordine')
    step_row(doc, 1, 'Tocca "Aggiungi al carrello"',
             'Il piatto personalizzato entra nel carrello con un nome descrittivo '
             '(es. "Panino personalizzato 🔥: Pane integrale, Pollo, Lattuga, Maionese")')
    spacer(doc, 4)
    step_row(doc, 2, 'Procedi come un ordine normale',
             'Scegli lo slot di ritiro (o "Adesso al banco"), controlla il totale e tocca '
             '"Conferma ordine". Puoi avere più prodotti — sia builder che dal menu — nello stesso carrello')
    spacer(doc, 8)

    info_box(doc, 'Se una categoria ammette al massimo 2 scelte e ne selezioni 3, il sistema ti blocca '
             'e ti chiede di toglierne una prima di continuare.', style='warning')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 5 — I MIEI ORDINI
# ══════════════════════════════════════════════════════════════════════════════

def s5_ordini(doc):
    h1(doc, '5', 'I Miei Ordini', '📦', accent=HEX_PURPL)

    _nota_privacy_nome(doc)

    h2(doc, '5.1  Accedere allo storico')
    step_row(doc, 1, 'Vai in "I Miei Ordini"', 'Dal menu laterale')
    spacer(doc, 4)
    step_row(doc, 2, 'Scorri la lista',
             'Vedi gli ultimi 50 ordini, dal più recente: data, codice, prodotti, totale, stato')
    spacer(doc, 8)

    h2(doc, '5.2  Stati dell\'ordine')
    data_table(doc,
        ['Stato', 'Significato'],
        [
            ['Confermato',      'L\'ordine è stato pagato ed è in coda per la preparazione'],
            ['In preparazione', 'La cucina ha iniziato a prepararlo'],
            ['Pronto',          'Puoi ritirarlo alla cassa/bancone'],
            ['Consegnato',      'L\'ordine è stato ritirato'],
            ['Annullato',       'L\'ordine è stato cancellato e l\'importo rimborsato'],
        ],
        col_widths=[4.5, 13.1])
    spacer(doc, 8)

    h2(doc, '5.3  Annullamento ordine')
    step_row(doc, 1, 'Apri "I Miei Ordini"', 'Trova l\'ordine da annullare')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca "Annulla"',
             'Disponibile SOLO se l\'ordine è ancora "Confermato". '
             'Non puoi più annullare un ordine già "In preparazione" o "Pronto"')
    spacer(doc, 4)
    step_row(doc, 3, 'Conferma l\'annullamento',
             'Il rimborso è immediato e automatico: l\'importo torna subito sul tuo wallet. '
             'I punti fedeltà guadagnati con quell\'ordine vengono stornati')
    spacer(doc, 8)

    info_box(doc, 'Se hai cambiato idea ma l\'ordine è già "In preparazione", non puoi annullarlo da solo: '
             'rivolgiti subito al cassiere o alla cucina mostrando il codice ordine.', style='warning')
    spacer(doc, 8)

    h2(doc, '5.4  Codice ordine: come usarlo')
    body_para(doc,
        'Il codice ordine ha formato QuickLunch-YYMMDD-HHMM-NNNN (es. QuickLunch-260630-1145-0042) '
        'oppure QuickLunch-YYMMDD-BANCO-NNNN se hai scelto "Adesso al banco". '
        'Mostralo al personale per: ritirare l\'ordine, segnalare un problema, chiedere una correzione. '
        'Con quel codice il cassiere trova e gestisce il tuo ordine in pochi secondi.')

    spacer(doc, 6)
    info_box_color(doc,
        'Se un ordine non è arrivato, è sbagliato o c\'è un addebito che non riconosci, vai in '
        '"I Miei Ordini", individua il numero ordine e mostralo al cassiere.',
        bg='EBF5FB', border=HEX_TEAL, icon='💡')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 6 — WALLET E PUNTI FEDELTÀ
# ══════════════════════════════════════════════════════════════════════════════

def _nota_privacy_nome(doc):
    """Il cliente si vede per intero, gli altri lo vedono abbreviato."""
    body_para(doc,
        'Il tuo nome sulle liste: quando il tuo ordine compare sul display '
        'della cucina, sul tagliando o su una lista stampata, sei indicato '
        'con il nome e la sola iniziale del cognome — per esempio '
        '"Mario R.". Qui nella tua area personale, invece, il nome resta '
        'per intero.')
    spacer(doc, 8)


def s6_wallet(doc):
    h1(doc, '6', 'Wallet e Punti Fedeltà', '💳', accent=HEX_PURPL)

    body_para(doc,
        'Nota: il wallet è una funzione che il locale può disattivare. Se nel '
        'menu non vedi la voce "Wallet & Fedeltà", nel tuo bar si paga alla '
        'cassa al momento del ritiro e questo capitolo non ti riguarda.')
    spacer(doc, 6)

    role_badge(doc, '💳', 'Wallet',
               'Il wallet è il tuo portafoglio digitale nel locale: ci carichi credito e paghi i tuoi ordini '
               'in un tocco, senza contanti o carta ogni volta.',
               HEX_NAVY)
    spacer(doc, 8)

    h2(doc, '6.1  Saldo wallet: dove vederlo')
    body_para(doc,
        'Il saldo wallet è sempre visibile in due punti: '
        'il riquadro KPI "Saldo Wallet" nella Home Dashboard, '
        'e il riquadro in alto nella pagina "Wallet & Fedeltà". '
        'È il credito che puoi spendere ora.')
    spacer(doc, 8)

    h2(doc, '6.2  Come viene scalato')
    body_para(doc,
        'Il wallet viene scalato automaticamente a ogni acquisto confermato: '
        'ordine dal menu, ordine builder, pagamento al banco tramite QR. '
        'L\'addebito avviene nel momento in cui tocchi "Conferma ordine" o "Paga ora" — '
        'mai prima.')
    spacer(doc, 8)

    h2(doc, '6.3  Ricarica: solo tramite personale del locale')
    info_box(doc, 'La ricarica del wallet non si fa dall\'app cliente: chiedi al cassiere di ricaricarti '
             'l\'importo che preferisci (es. 10€, 20€, 50€). Verrà accreditato all\'istante e lo vedrai '
             'subito nello storico transazioni come "Ricarica".', style='info')
    spacer(doc, 8)

    h2(doc, '6.4  Punti fedeltà: 10 punti per ogni euro speso')
    body_para(doc,
        'Per ogni euro speso con il wallet guadagni automaticamente 10 punti fedeltà. '
        'I punti vengono accreditati subito dopo la conferma dell\'ordine. '
        'Se l\'ordine viene annullato, i punti vengono stornati.')
    spacer(doc, 8)

    h2(doc, '6.5  Riscatto punti fedeltà')
    data_table(doc,
        ['Regola', 'Dettaglio'],
        [
            ['Soglia di riscatto', 'Ogni 100 punti puoi riscattare +1,00€ di credito wallet'],
            ['Pulsante "Riscatta"', 'Appare automaticamente quando hai raggiunto almeno 100 punti'],
            ['Riscatti multipli', 'Con 250 punti puoi riscattare 2 blocchi da 100 = +2,00€ (restano 50 punti)'],
            ['Punti residui', 'I punti sotto la soglia non vengono persi: restano per il prossimo riscatto'],
        ],
        col_widths=[4.8, 12.8])
    spacer(doc, 8)

    h2(doc, '6.6  Storico transazioni')
    body_para(doc, 'In "Wallet & Fedeltà" vedi ogni movimento con: data, tipo, descrizione e importo.')
    spacer(doc, 4)
    data_table(doc,
        ['Tipo transazione', 'Quando appare'],
        [
            ['Acquisto',        'Ogni ordine confermato dal menu o builder'],
            ['Ricarica',        'Quando il cassiere accredita credito sul tuo wallet'],
            ['Rimborso',        'Quando un ordine viene annullato'],
            ['Riscatto punti',  'Quando converti i punti in credito wallet'],
            ['Banco',           'Pagamento tramite QR al banco generato dallo staff'],
        ],
        col_widths=[4.8, 12.8])
    spacer(doc, 8)

    h2(doc, '6.7  Saldo insufficiente')
    info_box(doc, 'Se il saldo non basta per confermare un ordine, vieni reindirizzato automaticamente '
             'alla pagina Wallet con un messaggio di avviso. '
             'Chiedi al cassiere una ricarica, poi torna nel carrello: i prodotti restano salvati.',
             style='warning')
    spacer(doc, 6)

    info_box_color(doc,
        'I punti restanti sotto i 100 (es. 50 punti dopo un riscatto da 250) non vengono persi: '
        'restano in conto per il prossimo premio.',
        bg='EBF5FB', border=HEX_TEAL, icon='💡')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 7 — PAGA AL BANCO (QR)
# ══════════════════════════════════════════════════════════════════════════════

def s7_banco(doc):
    h1(doc, '7', 'Paga al Banco — Pagamento QR', '📲', accent=HEX_PURPL)

    role_badge(doc, '📲', 'Paga al Banco',
               'Con "Paga al Banco" puoi pagare dal tuo telefono un conto aperto dallo staff, '
               'senza dover fare un ordine normale dal menu. Lo staff genera un QR code sul tablet: '
               'tu lo inquadri, controlli il riepilogo e confermi il pagamento in un tocco. '
               'Il wallet viene scalato in automatico e il banco riceve conferma istantanea.',
               HEX_GREEN)
    spacer(doc, 6)
    body_para(doc,
        'Se il tuo locale lavora senza wallet, il flusso è identico ma il '
        'pulsante finale si chiama "Conferma": registri la consumazione e '
        'paghi alla cassa.')
    spacer(doc, 8)

    h2(doc, '7.1  Quando usarlo')
    body_para(doc,
        'Usa "Paga al Banco" per acquisti rapidi al bancone (caffè, brioche, acqua, ecc.) '
        'dove è LO STAFF a comporre l\'ordine e generare il QR, '
        'senza che tu debba selezionare nulla dal menu dell\'app.')
    spacer(doc, 8)

    h2(doc, '7.2  Come aprire lo scanner')
    body_para(doc,
        'Tocca il pulsante "Paga al Banco" nelle azioni rapide della Home Dashboard. '
        'Il browser potrebbe chiederti il permesso di accedere alla fotocamera: '
        'tocca "Consenti" per procedere.')
    spacer(doc, 8)

    h2(doc, '7.3  Flusso completo in 5 step')

    workflow_table(doc, [
        ('📱', 'Dashboard', 'Tocca "Paga al Banco"'),
        ('📷', 'Fotocamera', 'Concedi il permesso'),
        ('🔲', 'Inquadra QR', 'QR sul tablet del banco'),
        ('📋', 'Riepilogo', 'Controlla articoli e totale'),
        ('✅', 'Paga ora', 'Wallet scalato istantaneamente'),
    ], accent=HEX_GREEN)
    spacer(doc, 8)

    step_row(doc, 1, 'Tocca "Paga al Banco" nella Home',
             'Il pulsante si trova nelle azioni rapide della dashboard')
    spacer(doc, 4)
    step_row(doc, 2, 'Il browser chiede il permesso fotocamera — concedilo',
             'Senza permesso fotocamera lo scanner non funziona. '
             'Puoi sempre usare il campo manuale come alternativa (vedi §7.5)')
    spacer(doc, 4)
    step_row(doc, 3, 'Punta la fotocamera verso il QR sul tablet del banco',
             'Lo staff ha generato il QR dal loro pannello. '
             'Non appena la fotocamera lo riconosce, si apre automaticamente la pagina di conferma')
    spacer(doc, 4)
    step_row(doc, 4, 'Controlla il riepilogo: articoli + totale',
             'Vedi la lista degli articoli inseriti dallo staff e il totale da pagare. '
             'Verifica che tutto sia corretto prima di procedere')
    spacer(doc, 4)
    step_row(doc, 5, 'Tocca "Paga ora" — wallet scalato automaticamente',
             'Il totale viene addebitato dal wallet. Il banco riceve conferma istantanea sul loro schermo')
    spacer(doc, 8)

    h2(doc, '7.4  Conferma')
    body_para(doc,
        'Dopo il pagamento vedi una schermata di conferma sul tuo telefono. '
        'Sul tablet del banco compare il messaggio "✓ Pagato da [tuo nome]" in tempo reale: '
        'non è necessario mostrare nulla al personale.')
    spacer(doc, 8)

    h2(doc, '7.5  Note e problemi comuni')
    data_table(doc,
        ['Situazione', 'Come comportarsi'],
        [
            ['La fotocamera non funziona',
             'Usa il campo "Inserisci manualmente il codice" nella stessa pagina: '
             'chiedi il codice sessione allo staff a voce e digitalo'],
            ['Sessione QR scaduta (>10 minuti)',
             'Il QR code scade automaticamente dopo 10 minuti se non viene usato. '
             'Chiedi al personale di rigenerarne uno nuovo'],
            ['Saldo wallet insufficiente',
             'Il pagamento non va a buon fine. Chiedi al cassiere una ricarica e riprova'],
            ['Quando viene addebitato il wallet?',
             'Solo quando tocchi "Paga ora". Finché non confermi, nessun importo viene scalato'],
        ],
        col_widths=[5.5, 12.1])
    spacer(doc, 8)

    info_box_color(doc,
        'DISTINZIONE: questa funzione è diversa da "Adesso al banco" in §3.5.\n\n'
        'Qui non hai scelto articoli dal menu: è il PERSONALE che ha composto l\'ordine e '
        'genera il QR. Tu arrivi, inquadri, controlli e paghi.\n\n'
        '"Adesso al banco" (§3.5) invece è un ordine fatto da TE dal menu, senza prenotare uno slot.',
        bg='E9F7EF', border=HEX_GREEN, icon='⚡')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 8 — PRENOTARE UN TAVOLO
# ══════════════════════════════════════════════════════════════════════════════

def s8_tavoli(doc):
    h1(doc, '8', 'Prenotare un Tavolo', '🪑', accent=HEX_PURPL)

    body_para(doc,
        'La prenotazione di un tavolo funziona per FASCE ORARIE: ogni fascia è un blocco di tempo '
        '(es. 11:25–12:30) suddiviso in sessioni della stessa durata (es. 30 min). '
        'Prenoti un tavolo per una sessione specifica.')
    spacer(doc, 6)

    h2(doc, '8.1  Aprire la sezione Tavoli')
    step_row(doc, 1, 'Vai in "Tavoli" dal menu laterale',
             'Si apre la pagina di prenotazione tavoli. Usa le frecce ‹ › per scegliere il giorno, '
             'oppure inserisci la data nel campo in alto. Di default vedi oggi')
    spacer(doc, 8)

    h2(doc, '8.2  Scegliere la fascia oraria')
    body_para(doc,
        'Ogni fascia oraria (es. "11:25 – 12:30 · 30 min a seduta") mostra le sessioni disponibili '
        'con orario di inizio, numero posti liberi e occupati. '
        'Leggi i dati di disponibilità e scegli la sessione che preferisci.')
    spacer(doc, 8)

    h2(doc, '8.3  Selezionare il tavolo')
    step_row(doc, 1, 'Guarda la mappa dei tavoli',
             'Verde = libero (puoi prenotare), rosso = occupato (non disponibile), '
             'blu = già tua prenotazione')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca un tavolo verde',
             'Si apre la scheda di prenotazione con i dettagli della sessione scelta')
    spacer(doc, 8)

    h2(doc, '8.4  Indicare numero persone e note facoltative')
    step_row(doc, 1, 'Inserisci il numero di persone',
             'Non puoi superare i posti disponibili al tavolo (es. tavolo da 4 → massimo 4 persone)')
    spacer(doc, 4)
    step_row(doc, 2, 'Aggiungi eventuali note (facoltativo)',
             'Es. "Seggiolone per bambino", "Vicino alla finestra", "Compleanno"')
    spacer(doc, 8)

    h2(doc, '8.5  Conferma prenotazione')
    step_row(doc, 1, 'Tocca "Prenota"',
             'Conferma immediata: la sessione è riservata per te. '
             'Il tavolo cambia colore in blu nella mappa')
    spacer(doc, 8)

    h2(doc, '8.6  Lista "Le mie prenotazioni"')
    step_row(doc, 1, 'Tocca il pulsante "Le mie prenotazioni" in alto nella pagina Tavoli',
             'Oppure vai in "Le Mie Prenotazioni" dal menu laterale. '
             'Vedi tutte le prenotazioni fatte, dalla più recente, con data, tavolo e sessione')
    spacer(doc, 8)

    h2(doc, '8.7  Annullamento prenotazione')
    step_row(doc, 1, 'Nella lista prenotazioni, trova quella da annullare',
             'Tocca "Annulla": il tavolo torna disponibile per altri clienti. '
             'L\'annullamento è possibile solo prima del check-in')
    spacer(doc, 8)

    h2(doc, '8.8  Vincolo: una prenotazione per fascia oraria')
    info_box(doc, 'Non puoi prenotare due tavoli nella stessa sessione oraria. '
             'Puoi invece prenotare tavoli in sessioni diverse (es. una prenotazione alle 11:25 '
             'e una alle 12:30 sono ammesse).', style='warning')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 9 — PASTO AZIENDALE
# ══════════════════════════════════════════════════════════════════════════════

def s9_pasto_aziendale(doc):
    h1(doc, '9', 'Pasto Aziendale', '🏢', accent=HEX_PURPL)

    role_badge(doc, '🏢', 'Dipendente convenzionato',
               'Se la tua azienda ha una convenzione attiva con il locale, vedrai una voce in più nel menu: '
               '"Pasto Aziendale". Da lì prenoti il pasto fisso del giorno al prezzo speciale aziendale, '
               'senza passare dal menu normale.',
               HEX_PURPL)
    spacer(doc, 8)

    h2(doc, '9.1  Chi può accedere')
    body_para(doc,
        'La sezione "Pasto Aziendale" è visibile solo se il tuo account è stato associato '
        'dall\'admin aziendale a una convenzione attiva. '
        'Se non la vedi nel menu, vedi §9.5 più in basso.')
    spacer(doc, 8)

    h2(doc, '9.2  Prenotazione pasto del giorno (5 step)')
    step_row(doc, 1, 'Vai in "Pasto Aziendale"',
             'Visibile nel menu laterale solo se associato a convenzione attiva')
    spacer(doc, 4)
    step_row(doc, 2, 'Guarda il menu e scegli l\'opzione',
             'Il piatto fisso del giorno con nome, descrizione e posti rimasti. '
             'Se sono presenti più opzioni, seleziona quella che preferisci')
    spacer(doc, 4)
    step_row(doc, 3, 'Imposta la quantità',
             'Di norma 1 pasto per persona. Conferma la quantità')
    spacer(doc, 4)
    step_row(doc, 4, 'Tocca "Prenota il mio pasto"',
             'La prenotazione è immediata: un posto in meno per gli altri colleghi')
    spacer(doc, 4)
    step_row(doc, 5, 'Al momento del ritiro',
             'Vai al locale all\'orario indicato e ritira il pasto mostrando, se richiesto, il tuo nome')
    spacer(doc, 8)

    h2(doc, '9.3  Badge posti rimasti')
    data_table(doc,
        ['Colore badge', 'Posti rimasti', 'Significato'],
        [
            ['Verde',   '>50%',   'Ampia disponibilità'],
            ['Giallo',  '20–50%', 'Disponibilità ridotta — prenota presto'],
            ['Rosso',   '<20%',   'Quasi esaurito'],
        ],
        col_widths=[3.5, 3.0, 11.1])
    spacer(doc, 8)

    h2(doc, '9.4  Annullamento prenotazione pasto aziendale')
    step_row(doc, 1, 'Apri "Pasto Aziendale"', 'Vedi la tua prenotazione attiva per oggi')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca "Annulla prenotazione"',
             'Possibile SOLO se mancano PIÙ di 30 minuti allo slot scelto. '
             'Il pulsante "Annulla" sparisce automaticamente quando mancano meno di 30 minuti: '
             'in quel caso il pasto risulta confermato definitivamente. '
             'Il rimborso, quando consentito, avviene in automatico')
    spacer(doc, 8)

    info_box(doc, 'Deadline annullamento: puoi cancellare la prenotazione fino a 30 minuti prima '
             'dell\'orario dello slot. Oltre quella soglia il pulsante "Annulla" scompare e il pasto '
             'risulta confermato definitivamente.',
             style='warning')
    spacer(doc, 6)

    info_box(doc, 'Se il messaggio dice "Posti esauriti per oggi" significa che il numero massimo di pasti '
             'prenotabili è stato raggiunto: riprova il giorno successivo o contatta l\'amministratore aziendale.',
             style='warning')
    spacer(doc, 6)

    h2(doc, '9.5  Se la sezione "Pasto Aziendale" non appare nel menu')
    info_box_color(doc,
        'Se non vedi la voce "Pasto Aziendale" nel menu, la tua azienda non ha ancora attivato '
        'la convenzione, oppure il tuo account non è stato associato. '
        'Contatta il referente aziendale o l\'amministratore del locale per verificare '
        'e richiedere l\'associazione del tuo account.',
        bg='EBF5FB', border=HEX_TEAL, icon='ℹ️')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 10 — SONDAGGI
# ══════════════════════════════════════════════════════════════════════════════

def s10_sondaggi(doc):
    h1(doc, '10', 'Sondaggi', '🗳️', accent=HEX_PURPL)

    body_para(doc, 'Il locale può aprire sondaggi per capire cosa i clienti vorrebbero trovare nel menu. '
              'Votando, aiuti a scegliere i prossimi piatti.')
    spacer(doc, 6)

    h2(doc, '10.1  Trovare i sondaggi attivi')
    step_row(doc, 1, 'Vai in "Sondaggi" (o "Vota il Menu") dal menu laterale',
             'Se c\'è un sondaggio attivo, si apre direttamente la pagina di voto. '
             'Se non vedi nessun sondaggio, il locale non ne ha aperti al momento: ricontrolla più avanti')
    spacer(doc, 8)

    h2(doc, '10.2  Votare')
    step_row(doc, 1, 'Leggi la domanda e le opzioni',
             'Ogni opzione ha un\'emoji e un testo descrittivo')
    spacer(doc, 4)
    step_row(doc, 2, 'Tocca l\'opzione che preferisci',
             'Il voto è immediato, definitivo e non modificabile: scegli con cura')
    spacer(doc, 8)

    info_box(doc, 'Puoi votare una sola volta per ogni sondaggio: dopo il voto l\'opzione scelta '
             'resta evidenziata e non potrai cambiarla.', style='info')
    spacer(doc, 8)

    h2(doc, '10.3  Risultati in tempo reale')
    step_row(doc, 1, 'Dopo aver votato, guarda i risultati',
             'Vedi le percentuali aggiornate in tempo reale per ogni opzione, '
             'così puoi capire come si sta orientando la preferenza di tutti i clienti')


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 11 — NOTIFICHE TELEGRAM
# ══════════════════════════════════════════════════════════════════════════════

def s11_notifiche(doc):
    h1(doc, '11', 'Notifiche Telegram', '🔔', accent=HEX_PURPL)

    body_para(doc, 'QuickLunch può avvisarti via Telegram in tempo reale per i principali eventi '
              'del tuo account, senza dover tenere l\'app aperta.')
    spacer(doc, 6)

    h2(doc, '11.1  Cosa ricevi')
    data_table(doc,
        ['Evento', 'Contenuto della notifica'],
        [
            ['Conferma ordine',              'Codice ordine, orario di ritiro, totale pagato'],
            ['Ordine in preparazione',       'Codice ordine e messaggio "la cucina ha iniziato"'],
            ['Ordine pronto',                'Codice ordine e messaggio "puoi ritirarlo al banco"'],
            ['Annullamento ordine',          'Numero ordine e importo rimborsato'],
            ['Prenotazione pasto aziendale', 'Nome del piatto, azienda e data di prenotazione'],
        ],
        col_widths=[5.5, 12.1])
    spacer(doc, 8)

    h2(doc, '11.2  Come attivare le notifiche')
    body_para(doc,
        'Il collegamento con Telegram viene impostato dal personale del locale sul tuo profilo cliente. '
        'Per attivarlo, comunica al cassiere o all\'amministratore il tuo Telegram username '
        '(es. @tuonome) oppure il tuo Chat ID numerico.')
    spacer(doc, 6)

    info_box(doc, 'Se non colleghi Telegram, l\'app funziona comunque normalmente: '
             'controlla "I Miei Ordini" per gli aggiornamenti sullo stato.', style='tip')
    spacer(doc, 8)

    h2(doc, '11.3  Come trovare il proprio Chat ID')
    step_row(doc, 1, 'Apri Telegram sul tuo telefono', '')
    spacer(doc, 4)
    step_row(doc, 2, 'Cerca @userinfobot nella barra di ricerca',
             'È un bot ufficiale di Telegram')
    spacer(doc, 4)
    step_row(doc, 3, 'Avvia una conversazione e tocca "Start"',
             '@userinfobot risponde immediatamente con il tuo ID numerico (es. 123456789)')
    spacer(doc, 4)
    step_row(doc, 4, 'Comunica quel numero al cassiere o all\'admin',
             'Il personale lo inserisce nel tuo profilo e le notifiche si attivano subito')
    spacer(doc, 8)

    h2(doc, '11.4  Tabella eventi e notifiche')
    data_table(doc,
        ['Quando', 'Chi genera l\'evento', 'Cosa ricevi su Telegram'],
        [
            ['Ordine confermato',      'Tu (conferma carrello)',       'Codice, slot, totale'],
            ['In preparazione',        'Staff (cucina)',               'Codice ordine + avviso'],
            ['Ordine pronto',          'Staff (banco)',                'Codice ordine + "vieni a ritirare"'],
            ['Ordine annullato',       'Tu o staff',                  'Codice + importo rimborsato'],
            ['Pasto az. prenotato',    'Tu (Pasto Aziendale)',        'Piatto + data + azienda'],
        ],
        col_widths=[4.5, 4.5, 8.6])


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 12 — DOMANDE FREQUENTI
# ══════════════════════════════════════════════════════════════════════════════

def s12_faq(doc):
    h1(doc, '12', 'Domande Frequenti', '❓', accent=HEX_PURPL)

    faqs = [
        ('Ho dimenticato la password — cosa faccio?',
         'L\'app non ha il "recupero password" automatico: chiedi all\'amministratore del locale '
         'di reimpostartela dal pannello clienti. Se accedi con Google, non hai bisogno di una password.'),

        ('Il saldo wallet è insufficiente — come ricarico?',
         'La ricarica non si fa dall\'app cliente: chiedi al cassiere di accreditarti l\'importo che '
         'preferisci (es. 10€, 20€, 50€). Verrà accreditato all\'istante e visibile nello storico.'),

        ('Posso pagare in contanti o con carta dall\'app?',
         'No: i pagamenti nell\'app passano sempre dal wallet. Per pagare in contanti o carta, '
         'rivolgiti direttamente alla cassa con il cassiere.'),

        ('Posso annullare un ordine?',
         'Sì, ma solo se l\'ordine è ancora in stato "Confermato" (vedi §5.3). '
         'Il rimborso è immediato e automatico sul wallet. '
         'Se l\'ordine è già "In preparazione", devi rivolgerti al personale.'),

        ('Il prodotto che voglio non si aggiunge al carrello',
         'Probabilmente è esaurito per oggi (disponibilità a 0). '
         'La disponibilità si aggiorna in tempo reale: riprova il giorno successivo '
         'o scegli un prodotto alternativo.'),

        ('Lo slot di ritiro è pieno — cosa faccio?',
         'Gli slot hanno un numero massimo di ordini. Se lo slot che preferisci è pieno, '
         'scegline uno vicino (prima o dopo) oppure usa "Adesso al banco" se sei già al locale.'),

        ('Il tavolo che volevo è occupato — come faccio?',
         'I tavoli verdi nella mappa sono liberi in quel momento. '
         'Se il tavolo che volevi è rosso, scegli un altro tavolo verde oppure '
         'verifica una fascia oraria diversa.'),

        ('Ho già votato un sondaggio, posso cambiare voto?',
         'No: il voto è definitivo e non modificabile dal cliente. '
         'Contatta l\'amministratore del locale se serve una correzione eccezionale.'),

        ('Non vedo la sezione "Pasto Aziendale" nel menu',
         'La sezione appare solo se il tuo account è associato a una convenzione aziendale attiva. '
         'Contatta il referente aziendale o l\'amministratore del locale per verificare e richiedere '
         'l\'associazione.'),

        ('Come pago al banco se la fotocamera non funziona?',
         'Nella pagina "Paga al Banco" (§7) trovi il campo "Inserisci manualmente il codice": '
         'chiedi il codice sessione al personale a voce e digitalo. '
         'In alternativa chiedi al personale di rigenerare un nuovo QR e riprova.'),

        ('Qual è la differenza tra "Adesso al banco" e "Paga al Banco"?',
         '"Adesso al banco" (§3.5): hai scelto tu i prodotti dal menu, '
         'ma li ritiri subito senza prenotare uno slot. Il codice sarà BANCO-NNNN.\n'
         '"Paga al Banco" (§7): il personale ha composto l\'ordine (es. caffè, brioche) '
         'e genera un QR. Tu inquadri e paghi senza aver fatto nulla dal menu.'),

        ('Il rimborso per un ordine annullato è immediato?',
         'Sì: il rimborso è automatico e immediato. L\'importo torna subito sul wallet e '
         'puoi verificarlo nello storico transazioni in "Wallet & Fedeltà".'),
    ]

    for i, (q, a) in enumerate(faqs):
        h3(doc, f'{i+1}.  {q}', color=PURPL)
        body_para(doc, a, size=13)
        if i < len(faqs) - 1:
            spacer(doc, 4)


# ══════════════════════════════════════════════════════════════════════════════
# SEZIONE 13 — SCHEDA RIASSUNTIVA
# ══════════════════════════════════════════════════════════════════════════════

def s13_riassunto(doc):
    h1(doc, '13', 'Scheda Riassuntiva', '🧭', accent=HEX_PURPL)

    body_para(doc, 'Tieni questa pagina come riferimento veloce: ogni azione, dove trovarla e la sezione di approfondimento.')
    spacer(doc, 6)

    data_table(doc,
        ['Azione', 'Dove', 'Sezione'],
        [
            ['Registrarmi con email e password',     'Pagina login → "Registrati"',              '1.1'],
            ['Accedere con Google',                  'Pagina login → "Accedi con Google"',        '1.2'],
            ['Aggiungere l\'app alla schermata Home','Browser → Aggiungi a schermata Home',       '1.4'],
            ['Aggiornare dati profilo',              'Menu → Profilo → Impostazioni',             '2.2'],
            ['Cambiare password',                    'Menu → Profilo → Cambia password',          '2.3'],
            ['Sfogliare il menu e aggiungere al carrello', 'Menu',                                '3.1'],
            ['Controllare allergeni di un prodotto', 'Menu → scheda prodotto',                   '3.2'],
            ['Modificare quantità / rimuovere dal carrello', 'Carrello',                         '3.3'],
            ['Scegliere lo slot di ritiro',          'Carrello → sezione slot orari',             '3.4'],
            ['Ritirare subito senza slot',           'Carrello → "Adesso al banco"',              '3.5'],
            ['Confermare e pagare l\'ordine',        'Carrello → Conferma ordine',                '3.6'],
            ['Comporre un panino / insalata / poke', 'Menu → Componi il tuo piatto',             '4'],
            ['Vedere gli ordini passati',            'I Miei Ordini',                             '5.1'],
            ['Annullare un ordine',                  'I Miei Ordini → Annulla',                  '5.3'],
            ['Vedere saldo wallet',                  'Home (riquadro KPI) o Wallet & Fedeltà',   '6.1'],
            ['Ricaricare il wallet',                 'Chiedi al cassiere',                        '6.3'],
            ['Riscattare punti fedeltà',             'Wallet & Fedeltà → Riscatta',               '6.5'],
            ['Vedere storico transazioni',           'Wallet & Fedeltà → Storico',                '6.6'],
            ['Pagare al banco tramite QR',           'Home → Paga al Banco',                      '7'],
            ['Prenotare un tavolo',                  'Tavoli',                                    '8'],
            ['Vedere / annullare prenotazioni tavolo', 'Le Mie Prenotazioni',                    '8.6–8.7'],
            ['Prenotare il pasto aziendale',         'Pasto Aziendale',                           '9.2'],
            ['Annullare pasto aziendale',            'Pasto Aziendale → Annulla (entro 30 min)', '9.4'],
            ['Votare un sondaggio',                  'Sondaggi (o Vota il Menu)',                 '10'],
            ['Attivare notifiche Telegram',          'Chiedi al cassiere / admin',                '11.2'],
            ['Trovare il proprio Chat ID Telegram',  'Telegram → @userinfobot',                  '11.3'],
        ],
        col_widths=[6.5, 6.5, 4.6])
    spacer(doc, 10)

    info_box_color(doc,
        'Hai letto tutta la guida? Allora sai già fare tutto quello che serve: ordinare, pagare, '
        'prenotare un tavolo, usare i tuoi punti fedeltà e pagare al banco con il QR — '
        'da solo, senza bisogno di chiedere aiuto. Buon pranzo! 🍽️',
        bg='F4ECF7', border=HEX_PURPL, icon='🎉')


# ── Main ─────────────────────────────────────────────────────────────────────


def main():
    doc = Document()
    set_document_defaults(doc)

    build_cover_cliente(doc)
    _page_break(doc)

    build_toc_cliente(doc)
    _page_break(doc)

    s1_iniziare(doc)
    _page_break(doc)

    s2_profilo(doc)
    _page_break(doc)

    s3_ordinare(doc)
    _page_break(doc)

    s4_builder(doc)
    _page_break(doc)

    s5_ordini(doc)
    _page_break(doc)

    s6_wallet(doc)
    _page_break(doc)

    s7_banco(doc)
    _page_break(doc)

    s8_tavoli(doc)
    _page_break(doc)

    s9_pasto_aziendale(doc)
    _page_break(doc)

    s10_sondaggi(doc)
    _page_break(doc)

    s11_notifiche(doc)
    _page_break(doc)

    s12_faq(doc)
    _page_break(doc)

    s13_riassunto(doc)

    _footer_copyright(doc)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'OK  Guida cliente salvata in: {OUT}')


if __name__ == '__main__':
    main()
