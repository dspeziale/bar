#!/usr/bin/env python3
"""Genera docs/manuali/offerta_saas.docx — offerta commerciale QuickLunch in SaaS.

Modello: canone fisso mensile + percentuale sugli incassi. I due valori sono le
stesse impostazioni che l'applicazione usa per calcolare il dovuto
(tenant_monthly_fee e platform_fee_percentage), quindi cambiarli qui significa
cambiarli anche in Impostazioni.

    python docs/generate_offerta_saas.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Palette (coerente con gli altri documenti in docs/) ──────────────────────
RED    = RGBColor(0xe9, 0x45, 0x60)
NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x16, 0x21, 0x3e)
DGRAY  = RGBColor(0x34, 0x3a, 0x40)
GRAY   = RGBColor(0x6c, 0x75, 0x7d)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREEN  = RGBColor(0x27, 0xae, 0x60)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
BLUE   = RGBColor(0x34, 0x98, 0xdb)
PURPLE = RGBColor(0x8e, 0x44, 0xad)
SLATE  = RGBColor(0x7f, 0x8c, 0x8d)

HEX_RED    = 'E94560'
HEX_NAVY   = '0F3460'
HEX_GREEN  = '27AE60'
HEX_ORANGE = 'E67E22'
HEX_BLUE   = '3498DB'
HEX_PURPLE = '8E44AD'
HEX_SLATE  = '7F8C8D'
HEX_LIGHT  = 'F8F9FA'
HEX_RULE   = 'E8EBEE'
HEX_WARN   = 'FDF6E7'
HEX_STOP   = 'FDEEF0'

FONT = 'PT Sans Narrow'
BODY_FONT = FONT          # tutto il documento nello stesso carattere

OUT = os.path.join(os.path.dirname(__file__), 'manuali', 'offerta_saas.docx')


# ── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill.lstrip('#'))
    tcPr.append(shd)


def _cell_border(cell, *, top=None, bottom=None, left=None, right=None, sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        if color:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color.lstrip('#'))
        else:
            el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:vAlign')
    el.set(qn('w:val'), val)
    tcPr.append(el)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def _set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def _p_spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def _run_font(run, size=11.5, bold=False, italic=False, color=None, font=BODY_FONT):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    _p_spacing(p, before=0, after=0)


# ── Blocchi di contenuto ─────────────────────────────────────────────────────

def heading(doc, text, level=1, color=NAVY, before=None):
    """Titolo che usa lo stile Titolo N: il colore passato lo sovrascrive solo
    quando serve distinguere una sezione."""
    p = doc.add_paragraph(style='Heading %d' % min(max(level, 1), 3))
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    if color is not None:
        r.font.color.rgb = color
    return p


def body(doc, text, size=11.5, color=DGRAY, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    _run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def rich(doc, parts, size=11.5, after=6, indent=None):
    """Paragrafo con porzioni in grassetto: parts = [(testo, bold), ...]."""
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def bullet(doc, parts, size=11.5):
    p = doc.add_paragraph(style='List Bullet')
    _p_spacing(p, before=0, after=3)
    for text, bold in parts:
        _run_font(p.add_run(text), size=size, bold=bold,
                  color=DARK if bold else DGRAY)
    return p


def rule(doc, color=HEX_RULE, after=8):
    """Filetto orizzontale come tabella a una cella."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_border(cell, bottom=color, sz='8')
    _cell_margins(cell, top=0, bottom=0, left=0, right=0)
    _p_spacing(cell.paragraphs[0], before=0, after=0)
    _run_font(cell.paragraphs[0].add_run(''), size=1)
    spacer(doc, after)


def spacer(doc, pts=8):
    p = doc.add_paragraph()
    _p_spacing(p, before=0, after=0)
    _run_font(p.add_run(''), size=pts / 2)


def box(doc, label, paragraphs, accent=HEX_RED, fill=HEX_LIGHT, label_color=RED):
    """Riquadro con bordo sinistro colorato."""
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, fill)
    _cell_border(cell, top=HEX_RULE, bottom=HEX_RULE, right=HEX_RULE,
                 left=accent, sz='18')
    _cell_margins(cell, top=90, bottom=90, left=140, right=110)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=3)
    _run_font(p.add_run(label.upper()), size=9.5, bold=True, color=label_color,
              font=FONT)

    for i, parts in enumerate(paragraphs):
        pp = cell.add_paragraph()
        _p_spacing(pp, before=0, after=0 if i == len(paragraphs) - 1 else 4)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=11, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 10)
    return tbl


def step(doc, num, title, paragraphs, accent=HEX_RED):
    """Passo numerato: numero in pastiglia colorata + contenuto."""
    tbl = doc.add_table(rows=1, cols=2)
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_width(tbl, 0, 1.0)
    _set_col_width(tbl, 1, 15.5)

    nc = tbl.rows[0].cells[0]
    _cell_shd(nc, accent)
    _cell_margins(nc, top=60, bottom=60, left=30, right=30)
    _cell_valign(nc, 'center')
    pn = nc.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_spacing(pn, before=0, after=0)
    _run_font(pn.add_run('%02d' % num), size=13, bold=True, color=WHITE, font=FONT)

    cc = tbl.rows[0].cells[1]
    _cell_margins(cc, top=50, bottom=50, left=150, right=60)
    pt = cc.paragraphs[0]
    _p_spacing(pt, before=0, after=2)
    _run_font(pt.add_run(title), size=12.5, bold=True, color=DARK, font=FONT)

    for parts in paragraphs:
        pp = cc.add_paragraph()
        _p_spacing(pp, before=0, after=3)
        for text, bold in parts:
            _run_font(pp.add_run(text), size=11, bold=bold,
                      color=DARK if bold else DGRAY)
    spacer(doc, 6)
    return tbl


def table_grid(doc, headers, rows, widths, head_fill=HEX_NAVY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _no_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, w in enumerate(widths):
        _set_col_width(tbl, i, w)

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _cell_shd(c, head_fill)
        _cell_margins(c, top=60, bottom=60, left=110, right=80)
        p = c.paragraphs[0]
        _p_spacing(p, before=0, after=0)
        _run_font(p.add_run(h.upper()), size=9.5, bold=True, color=WHITE, font=FONT)

    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            c = tbl.rows[r].cells[i]
            _cell_shd(c, HEX_LIGHT if r % 2 else 'FFFFFF')
            _cell_border(c, bottom=HEX_RULE, sz='4')
            _cell_margins(c, top=55, bottom=55, left=110, right=80)
            p = c.paragraphs[0]
            _p_spacing(p, before=0, after=0)
            text, color, bold = val if isinstance(val, tuple) else (val, DGRAY, False)
            _run_font(p.add_run(text), size=11, bold=bold, color=color)
    spacer(doc, 10)
    return tbl


# ── Parametri commerciali ────────────────────────────────────────────────────
# Canone e percentuale corrispondono alle impostazioni tenant_monthly_fee e
# platform_fee_percentage. Il contributo di ingresso e' una tantum e non ha
# un'impostazione nell'applicazione: si fattura all'attivazione.
FEE_INGRESSO = 500.00      # euro, una tantum all'attivazione
CANONE_MENSILE = 40.00     # euro al mese
FEE_PERCENTUALE = 3.2      # per cento sull'imponibile
IVA = 10.0                 # aliquota scorporata dal calcolo dell'applicazione

# Proiezioni: incasso medio giornaliero e giorni di apertura al mese.
# 22 giorni = apertura dal lunedi al venerdi, tipica di una mensa aziendale.
SCENARI_GIORNO = [500, 750, 1000]
GIORNI_MESE = 22


def _eur(v):
    """Formato italiano: 1.234,56"""
    s = '%0.2f' % v
    interi, dec = s.split('.')
    gruppi = []
    while len(interi) > 3:
        gruppi.insert(0, interi[-3:])
        interi = interi[:-3]
    gruppi.insert(0, interi)
    return '.'.join(gruppi) + ',' + dec


def _pct(v):
    """Percentuale in formato italiano, senza zeri inutili: 3.5 -> '3,5'."""
    return ('%0.2f' % v).rstrip('0').rstrip('.').replace('.', ',')


def _perc2(v):
    """Percentuale con due decimali in formato italiano: 3.41 -> '3,41 %'."""
    return ('%0.2f' % v).replace('.', ',') + ' %'


def _riga_scenario(lordo):
    """Applica esattamente la formula usata dall'applicazione."""
    imponibile = lordo / (1 + IVA / 100.0)
    quota_var = imponibile * FEE_PERCENTUALE / 100.0
    totale = quota_var + CANONE_MENSILE
    incidenza = totale / lordo * 100.0
    return imponibile, quota_var, totale, incidenza


# ── Fogli di stile ───────────────────────────────────────────────────────────

def imposta_stili(doc):
    """Definisce gli stili del documento, cosi' che in Word si possa ritoccare
    l'aspetto da un punto solo invece che paragrafo per paragrafo."""
    from docx.enum.style import WD_STYLE_TYPE

    # Normale: e' la radice da cui tutto eredita
    normale = doc.styles['Normal']
    normale.font.name = FONT
    normale.font.size = Pt(11.5)
    normale.font.color.rgb = DGRAY
    normale.paragraph_format.space_after = Pt(6)
    normale.paragraph_format.line_spacing = 1.08
    # Anche per i caratteri non latini, altrimenti Word ricade su Calibri
    rpr = normale.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), FONT)

    # Titoli: gerarchia esplicita, tutti in PT Sans Narrow
    for nome, dim, colore, prima, dopo in [
        ('Heading 1', 20, NAVY, 16, 5),
        ('Heading 2', 15.5, NAVY, 14, 4),
        ('Heading 3', 13, DARK, 12, 3),
    ]:
        try:
            st = doc.styles[nome]
        except KeyError:
            st = doc.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = FONT
        st.font.size = Pt(dim)
        st.font.bold = True
        st.font.color.rgb = colore
        st.paragraph_format.space_before = Pt(prima)
        st.paragraph_format.space_after = Pt(dopo)
        st.paragraph_format.keep_with_next = True

    # Stile dedicato alle didascalie e alle note
    try:
        nota = doc.styles['Nota QuickLunch']
    except KeyError:
        nota = doc.styles.add_style('Nota QuickLunch', WD_STYLE_TYPE.PARAGRAPH)
    nota.base_style = doc.styles['Normal']
    nota.font.name = FONT
    nota.font.size = Pt(9.5)
    nota.font.color.rgb = GRAY


# ── Documento ────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    imposta_stili(doc)

    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)

    # ══ Copertina ═════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=1)
    _no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _cell_shd(cell, HEX_NAVY)
    _cell_margins(cell, top=300, bottom=300, left=320, right=320)

    p = cell.paragraphs[0]
    _p_spacing(p, before=0, after=4)
    _run_font(p.add_run('QUICKLUNCH  ·  PROPOSTA COMMERCIALE'), size=10.5,
              bold=True, color=RGBColor(0xb2, 0xc2, 0xd9), font=FONT)

    p2 = cell.add_paragraph()
    _p_spacing(p2, before=0, after=6)
    _run_font(p2.add_run('Offerta in abbonamento'), size=31, bold=True,
              color=WHITE, font=FONT)

    p3 = cell.add_paragraph()
    _p_spacing(p3, before=0, after=0)
    _run_font(p3.add_run(
        'La piattaforma di ordinazione, portafoglio prepagato e gestione del '
        'bar aziendale, fornita in modalita SaaS: nessuna installazione, '
        'nessun server da mantenere, aggiornamenti e copie di sicurezza '
        'inclusi nel canone.'), size=12,
        color=RGBColor(0xd6, 0xdf, 0xea))

    p4 = cell.add_paragraph()
    _p_spacing(p4, before=10, after=0)
    _run_font(p4.add_run('Il vostro referente:  '), size=11.5, bold=True,
              color=WHITE, font=FONT)
    _run_font(p4.add_run('Daniele Speziale  ·  dspeziale@gmail.com  ·  +39 352 0150489'),
              size=11.5, bold=True, color=RGBColor(0xff, 0xd7, 0xdf),
              font=FONT)
    spacer(doc, 16)

    # ══ Cappello: la tecnologia ═══════════════════════════════════════════
    heading(doc, 'Software come servizio: che cosa significa', 1)

    body(doc, 'QuickLunch non e un programma da installare sui vostri computer: '
              'e un servizio che funziona su internet, come la posta elettronica '
              'o l\'home banking. Il modello si chiama SaaS, dall\'inglese '
              '"Software as a Service": voi usate l\'applicazione, noi ci '
              'occupiamo di tutto il resto.', color=DGRAY)

    body(doc, 'La differenza pratica e questa: non c\'e un server da comprare in '
              'sala macchine, non ci sono licenze da rinnovare, non c\'e un '
              'tecnico da chiamare per gli aggiornamenti. Si apre il browser, si '
              'entra col proprio accesso e si lavora — dal tablet del banco, dal '
              'display della cucina, dal telefono del cliente.', color=DGRAY)
    spacer(doc, 8)

    heading(doc, 'Come funziona, in concreto', 3, color=DARK)
    for titolo, testo in [
        ('Niente da installare',
         'L\'applicazione vive su internet. Sui vostri dispositivi serve solo un '
         'browser aggiornato: nessun programma, nessuna configurazione, nessuno '
         'spazio occupato.'),
        ('Sempre l\'ultima versione',
         'Gli aggiornamenti li facciamo noi, di notte, senza interruzioni del '
         'servizio. Il giorno dopo trovate le funzioni nuove: non dovete decidere '
         'quando aggiornare ne pagare un passaggio di versione.'),
        ('I dati sono al sicuro e sono vostri',
         'Il database e ospitato su infrastruttura professionale con copie di '
         'sicurezza automatiche e connessione cifrata. In qualsiasi momento '
         'potete scaricare l\'archivio completo dei vostri dati con un clic: se '
         'un giorno decidete di andarvene, li portate con voi.'),
        ('Si lavora da qualsiasi postazione',
         'Cucina, banco e ufficio vedono gli stessi dati nello stesso istante. '
         'Un ordine confermato dal telefono di un cliente compare sul display '
         'della cucina in pochi secondi.'),
        ('Cresce con voi',
         'Un locale o cinque, cinquanta clienti o cinquemila: cambia solo il '
         'volume, non l\'impianto. Nessun intervento tecnico quando il giro '
         'd\'affari aumenta.'),
        ('Costo proporzionato',
         'Non si compra un programma da migliaia di euro sperando di ripagarlo: '
         'si paga un canone contenuto e una quota su quanto si incassa '
         'davvero. Se un mese si lavora poco, si paga poco.'),
    ]:
        rich(doc, [(titolo + '.  ', True), (testo, False)], size=11, after=5)

    spacer(doc, 6)
    box(doc, 'Cosa vi serve davvero', [
        [('Una connessione internet, un browser e i dispositivi che avete già: '
          'un tablet al banco, un display o un tablet in cucina, un computer '
          'per la configurazione. ', False),
         ('Nessun server, nessuna licenza, nessuna manutenzione a carico '
          'vostro.', True)],
    ], accent=HEX_GREEN, fill=HEX_LIGHT, label_color=GREEN)

    _page_break(doc)

    # ══ Il prezzo ═════════════════════════════════════════════════════════
    heading(doc, 'Il prezzo', 1)
    spacer(doc, 4)

    tblp = doc.add_table(rows=1, cols=3)
    _no_borders(tblp)
    _set_col_width(tblp, 0, 5.4)
    _set_col_width(tblp, 1, 5.4)
    _set_col_width(tblp, 2, 5.4)

    for idx, (titolo, valore, sotto) in enumerate([
        ('Contributo di ingresso', '%s €' % _eur(FEE_INGRESSO),
         'una tantum, all\'attivazione'),
        ('Canone mensile', '%s €' % _eur(CANONE_MENSILE),
         'fisso, per singolo locale'),
        ('Quota sugli incassi', '%s %%' % _pct(FEE_PERCENTUALE),
         'sull\'imponibile del venduto'),
    ]):
        c = tblp.rows[0].cells[idx]
        _cell_shd(c, HEX_LIGHT)
        _cell_border(c, top=HEX_RULE, bottom=HEX_RULE, right=HEX_RULE,
                     left=HEX_RED, sz='18')
        _cell_margins(c, top=140, bottom=140, left=160, right=120)
        pa = c.paragraphs[0]
        _p_spacing(pa, before=0, after=2)
        _run_font(pa.add_run(titolo.upper()), size=9.5, bold=True, color=RED,
                  font=FONT)
        pb = c.add_paragraph()
        _p_spacing(pb, before=0, after=2)
        _run_font(pb.add_run(valore), size=23, bold=True, color=NAVY, font=FONT)
        pc = c.add_paragraph()
        _p_spacing(pc, before=0, after=0)
        _run_font(pc.add_run(sotto), size=10, color=GRAY)

    spacer(doc, 14)
    body(doc, 'Il contributo di ingresso si paga una volta sola, all\'attivazione. '
              'Nessun canone di licenza, nessun minimo garantito sulla quota '
              'variabile: se un mese non si incassa, si paga il solo canone.',
         color=DGRAY)

    # ══ Come si calcola la quota ══════════════════════════════════════════
    heading(doc, 'Su cosa si calcola il %s%%' % _pct(FEE_PERCENTUALE), 1)
    body(doc, 'La base di calcolo non e una stima concordata: la misura '
              'l\'applicazione stessa, voce per voce, con un prospetto mensile '
              'consultabile in qualsiasi momento. La quota si calcola allo '
              'stesso modo in entrambe le modalita di incasso: anche con il '
              'portafoglio prepagato disattivato le vendite restano '
              'registrate.', color=GRAY)
    spacer(doc, 4)

    table_grid(doc,
               ['Voce', 'Cosa comprende', 'Nel calcolo'],
               [
                   ['Ordini da menu e builder',
                    'Ordini completati nel mese, al netto degli annullati',
                    ('sì', GREEN, True)],
                   ['Vendite al banco',
                    'Sessioni di pagamento QR effettivamente saldate',
                    ('sì', GREEN, True)],
                   ['Pasti aziendali',
                    'Porzioni prenotate e non annullate, al prezzo del menu',
                    ('sì', GREEN, True)],
                   ['Vendite dal cesto',
                    'Pezzi pre-preparati acquistati inquadrando il QR, '
                    'compresi gli extra aggiunti alla stessa vendita',
                    ('sì', GREEN, True)],
                   ['Ricariche del portafoglio',
                    'Il denaro caricato non e un incasso finche non si spende',
                    ('no', RED, True)],
               ],
               widths=[4.6, 8.2, 3.2])

    box(doc, 'Due precisazioni che vi convengono', [
        [('La quota si applica all\'', False), ('imponibile', True),
         (', non al lordo: dagli incassi viene prima scorporata l\'IVA al %s%%. '
          'Su 1.000 € incassati la base e %s €, non 1.000.'
          % ('%0.0f' % IVA, _eur(1000 / (1 + IVA / 100.0))), False)],
        [('Le ', False), ('ricariche del portafoglio non sono incassi', True),
         (': la quota matura solo quando il credito viene speso, una volta sola '
          'e sul consumo reale.', False)],
    ], accent=HEX_GREEN, fill=HEX_LIGHT, label_color=GREEN)

    _page_break(doc)

    # ══ Simulazione ═══════════════════════════════════════════════════════
    heading(doc, 'Proiezioni sugli incassi', 1, before=0)
    rich(doc, [
        ('Tre ipotesi di incasso medio giornaliero, su ', False),
        ('%d giorni di apertura al mese' % GIORNI_MESE, True),
        (' (lunedi-venerdi). L\'ultima colonna e il dato che conta: quanto pesa '
         'il servizio sul venduto.', False),
    ], after=10)

    righe = []
    for giorno in SCENARI_GIORNO:
        lordo = giorno * GIORNI_MESE
        imponibile, quota_var, totale, incidenza = _riga_scenario(lordo)
        righe.append([
            ('%s €/gg' % _eur(giorno), DARK, True),
            '%s €' % _eur(lordo),
            '%s €' % _eur(imponibile),
            '%s €' % _eur(quota_var),
            ('%s €' % _eur(totale), NAVY, True),
            (_perc2(incidenza), GREEN, True),
        ])

    table_grid(doc,
               ['Incasso medio', 'Al mese (lordo)', 'Imponibile',
                'Quota %s%%' % _pct(FEE_PERCENTUALE),
                'Totale mese', 'Incidenza'],
               righe,
               widths=[2.6, 3.0, 2.7, 2.5, 2.8, 2.4])

    body(doc, 'Il totale mensile comprende la quota variabile e il canone di %s €. '
              'Il contributo di ingresso non e ricorrente: incide solo sul primo '
              'anno.' % _eur(CANONE_MENSILE), color=GRAY, after=14)

    heading(doc, 'Primo anno e anni successivi', 3, color=DARK)
    spacer(doc, 4)

    righe2 = []
    for giorno in SCENARI_GIORNO:
        lordo = giorno * GIORNI_MESE
        _i, _q, totale, _inc = _riga_scenario(lordo)
        anno_pieno = totale * 12
        primo_anno = anno_pieno + FEE_INGRESSO
        inc_primo = primo_anno / (lordo * 12) * 100.0
        righe2.append([
            ('%s €/gg' % _eur(giorno), DARK, True),
            '%s €' % _eur(lordo * 12),
            ('%s €' % _eur(primo_anno), NAVY, True),
            ('%s €' % _eur(anno_pieno), DARK, True),
            (_perc2(inc_primo), GREEN, True),
        ])

    table_grid(doc,
               ['Incasso medio', 'Incassi annui', 'Primo anno',
                'Anni successivi', 'Incidenza 1o anno'],
               righe2,
               widths=[2.9, 3.4, 3.2, 3.2, 3.3])

    _i0, _q0, _t0, _inc0 = _riga_scenario(SCENARI_GIORNO[0] * GIORNI_MESE)
    _i2, _q2, _t2, _inc2 = _riga_scenario(SCENARI_GIORNO[-1] * GIORNI_MESE)
    rich(doc, [
        ('A questi volumi l\'incidenza e sostanzialmente piatta, fra il ', False),
        (_perc2(_inc2), True), (' e il ', False), (_perc2(_inc0), True),
        (': il canone fisso e il contributo di ingresso si diluiscono e il costo '
         'segue quasi solo il venduto. In pratica il servizio costa poco meno '
         'del %s%% degli incassi.' % _pct(FEE_PERCENTUALE), False),
    ])

    # ══ Cosa comprende ════════════════════════════════════════════════════
    heading(doc, 'Cosa comprende il servizio', 1)
    spacer(doc, 4)

    for titolo, voci in [
        ('Per i clienti del locale', [
            'Menu digitale con ordinazione e scelta dell\'orario di ritiro',
            'Composizione libera di panini, insalate e poke, con extra a listino',
            'Portafoglio prepagato con punti fedelta e premi '
            '(attivabile: si puo lavorare anche con il solo pagamento in cassa)',
            'Pagamento al banco inquadrando un QR, senza contanti',
            'Acquisto self-service dal cesto dei prodotti pronti',
            'Prenotazione del pasto aziendale convenzionato',
            'Avvisi su Telegram e notifiche sul telefono quando l\'ordine e pronto',
        ]),
        ('Per il personale', [
            'Display di cucina con avanzamento degli ordini e avviso sonoro',
            'Cassa rapida al banco con generazione del QR di pagamento',
            'Etichette QR per i prodotti preparati in anticipo',
            'Gestione delle convenzioni aziendali e dei menu del giorno',
            'Magazzino dei materiali di consumo con avviso al fornitore',
            'Anagrafica clienti, ricariche, ruoli e permessi per postazione',
            'Report di incasso e prodotti piu venduti',
        ]),
        ('Compreso nel canone', [
            'Aggiornamenti e nuove funzioni, senza costi aggiuntivi',
            'Hosting, backup e certificato HTTPS',
            'Correzione dei malfunzionamenti',
            'Assistenza per la configurazione iniziale del menu e delle postazioni',
        ]),
    ]:
        heading(doc, titolo, 3, color=DARK)
        for v in voci:
            bullet(doc, [(v, False)])
        spacer(doc, 6)

    _page_break(doc)

    # ══ Cosa resta a carico ═══════════════════════════════════════════════
    heading(doc, 'Cosa resta a carico del locale', 1, before=0)
    body(doc, 'Il servizio e software: l\'attrezzatura di sala e gli obblighi '
              'fiscali restano vostri. Il dettaglio delle postazioni e nel '
              'documento "Layout dispositivi".', color=GRAY)
    spacer(doc, 4)

    table_grid(doc,
               ['Voce', 'Nota'],
               [
                   [('Tre postazioni', DARK, True),
                    'Un display o tablet in cucina, un tablet al banco, un PC per '
                    'la configurazione'],
                   [('Registratore di cassa', DARK, True),
                    'Lo scontrino fiscale resta emesso dal vostro registratore: '
                    'QuickLunch non lo sostituisce e non vi si collega'],
                   [('Stampanti', DARK, True),
                    'Una A4 per etichette e liste; una termica 80 mm, consigliata, '
                    'per i tagliandi di cucina'],
                   [('Connettivita', DARK, True),
                    'Rete Wi-Fi che copra cucina, banco e sala'],
                   [('Account per gli avvisi', DARK, True),
                    'Un indirizzo Gmail e un bot Telegram, gratuiti, per le '
                    'notifiche a clienti e personale'],
               ],
               widths=[4.6, 11.4])

    # ══ Trasparenza ═══════════════════════════════════════════════════════
    heading(doc, 'Come verificate quanto dovete', 1)
    body(doc, 'La quota variabile e verificabile da entrambe le parti sugli stessi '
              'dati, senza dichiarazioni da scambiarsi.', color=GRAY)
    spacer(doc, 4)

    step(doc, 1, 'Il prospetto mensile', [
        [('L\'applicazione espone un riepilogo per mese con le tre voci che '
          'compongono la base — ordini, banco, pasti aziendali — l\'imponibile '
          'scorporato e la quota risultante.', False)],
    ], accent=HEX_GREEN)

    step(doc, 2, 'Gli stessi numeri dei vostri report', [
        [('La base di calcolo esce dagli stessi movimenti che vedete nei report di '
          'incasso: nessun conteggio separato, nessuna stima.', False)],
    ], accent=HEX_GREEN)

    step(doc, 3, 'Fatturazione a mese chiuso', [
        [('Canone e quota vengono fatturati a mese concluso, sul dato definitivo. '
          'Gli ordini annullati e rimborsati non fanno base.', False)],
    ], accent=HEX_GREEN)

    # ══ Attivazione ═══════════════════════════════════════════════════════
    heading(doc, 'Attivazione', 1)
    spacer(doc, 4)

    step(doc, 1, 'Predisposizione', [
        [('Creazione del vostro spazio, inserimento del menu con prezzi e '
          'allergeni, orari di ritiro, account per le postazioni.', False)],
        [('E il lavoro coperto dal ', False),
         ('contributo di ingresso di %s €' % _eur(FEE_INGRESSO), True),
         (', che si versa una volta sola.', False)],
    ], accent=HEX_BLUE)

    step(doc, 2, 'Prova con il personale', [
        [('Un turno di prova con cucina e banco sui manuali operativi, per '
          'verificare i flussi prima di aprire ai clienti.', False)],
    ], accent=HEX_BLUE)

    step(doc, 3, 'Apertura ai clienti', [
        [('Affissione del QR di registrazione e primi crediti in cassa. Dal primo '
          'ordine il servizio e a regime.', False)],
    ], accent=HEX_BLUE)

    spacer(doc, 6)
    _page_break(doc)

    # ══ Condizioni contrattuali ═══════════════════════════════════════════
    heading(doc, 'Condizioni contrattuali', 1, before=0)
    body(doc, 'Le condizioni che regolano il servizio. Quanto non previsto qui '
              'sara\' disciplinato dal contratto di abbonamento sottoscritto '
              'all\'attivazione.', color=GRAY)
    spacer(doc, 4)

    CONDIZIONI = [
        ('Corrispettivi',
         'Contributo di ingresso di %s €, una tantum, fatturato all\'attivazione. '
         'Canone di %s € al mese per singolo locale. Quota del %s%% '
         'sull\'imponibile del venduto, calcolata come descritto in questa '
         'offerta. Importi al netto di IVA di legge.'
         % (_eur(FEE_INGRESSO), _eur(CANONE_MENSILE), _pct(FEE_PERCENTUALE))),
        ('Durata e rinnovo',
         'Il servizio ha durata di 12 mesi dall\'attivazione e si rinnova '
         'tacitamente di 12 mesi in 12 mesi, salvo disdetta di una delle parti '
         'con preavviso scritto di 60 giorni.'),
        ('Periodo di prova',
         'Nei primi 60 giorni dall\'attivazione il cliente puo\' recedere in '
         'qualsiasi momento, senza preavviso e senza penali, corrispondendo i '
         'soli canoni e quote maturati fino a quel giorno. Il contributo di '
         'ingresso, che copre il lavoro di predisposizione gia\' svolto, resta '
         'acquisito.'),
        ('Recesso',
         'Superato il periodo di prova, il cliente puo\' recedere in qualsiasi '
         'momento con preavviso scritto di 60 giorni, senza penali. Alla '
         'cessazione viene consegnato l\'archivio completo dei dati.'),
        ('Fatturazione e pagamenti',
         'Canone e quota sono fatturati mensilmente, a mese concluso, sulla base '
         'del prospetto consultabile nell\'applicazione. Pagamento a 30 giorni '
         'data fattura, tramite bonifico bancario.'),
        ('Assistenza',
         'Assistenza inclusa nel canone, nei giorni feriali dalle 9:00 alle '
         '18:00, per email e telefono. I malfunzionamenti che impediscono di '
         'lavorare vengono presi in carico entro 4 ore lavorative; gli altri '
         'entro il giorno lavorativo successivo.'),
        ('Dati personali',
         'Il locale e\' titolare del trattamento dei dati dei propri clienti; '
         'DS Consulting agisce come responsabile del trattamento ai sensi '
         'dell\'art. 28 del GDPR, con nomina contestuale al contratto. Alla '
         'cessazione del servizio i dati vengono consegnati al cliente in '
         'formato aperto e cancellati dai sistemi entro 30 giorni.'),
        ('Adeguamento dei corrispettivi',
         'Canone e contributo possono essere aggiornati una volta l\'anno '
         'sulla base dell\'indice ISTAT FOI, con comunicazione scritta almeno '
         '60 giorni prima. La quota percentuale resta quella pattuita.'),
        ('Referenti',
         'Per DS Consulting il referente e\' Daniele Speziale '
         '(dspeziale@gmail.com, cell. +39 352 0150489). Il cliente indica il proprio referente '
         'operativo e amministrativo all\'attivazione.'),
    ]
    for titolo, testo in CONDIZIONI:
        rich(doc, [(titolo + '.  ', True), (testo, False)], size=11, after=6)

    spacer(doc, 8)

    # ══ Validita' e accettazione ══════════════════════════════════════════
    heading(doc, 'Validita\' e accettazione', 1)
    body(doc, 'La presente offerta e\' valida per 30 giorni dalla data di '
              'emissione. Per accettazione, restituire una copia firmata: '
              'l\'attivazione viene pianificata entro 10 giorni lavorativi '
              'dalla ricezione.', color=DGRAY)
    spacer(doc, 14)

    tblf = doc.add_table(rows=1, cols=2)
    _no_borders(tblf)
    _set_col_width(tblf, 0, 8.2)
    _set_col_width(tblf, 1, 8.2)
    for idx, intest in enumerate(['DS CONSULTING', 'IL CLIENTE,\nPER ACCETTAZIONE']):
        cella = tblf.rows[0].cells[idx]
        _cell_margins(cella, top=60, bottom=60, left=120, right=120)
        pi = cella.paragraphs[0]
        _p_spacing(pi, before=0, after=2)
        _run_font(pi.add_run(intest.split('\n')[0]), size=10, bold=True,
                  color=NAVY, font=FONT)
        if '\n' in intest:
            pj = cella.add_paragraph()
            _p_spacing(pj, before=0, after=2)
            _run_font(pj.add_run(intest.split('\n')[1]), size=10, bold=True,
                      color=NAVY, font=FONT)
        for etich in ('Data', 'Timbro e firma'):
            pr = cella.add_paragraph()
            _p_spacing(pr, before=16, after=0)
            _run_font(pr.add_run(etich + '  '), size=10, color=GRAY)
            _run_font(pr.add_run('_' * 34), size=10, color=GRAY)

    spacer(doc, 12)
    rule(doc, color=HEX_NAVY)
    body(doc, 'Offerta valida 30 giorni dalla data di emissione. Gli importi '
              'sono al netto di IVA di legge.', size=9.5, color=GRAY, after=2)
    body(doc, 'QuickLunch · © 2024–26 DS Consulting', size=9.5, color=GRAY)

    return doc


def main():
    doc = build()
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f'[OK] Documento salvato in: {OUT}')
    print('     Ingresso %s € · canone %s € · quota %s%%'
          % (_eur(FEE_INGRESSO), _eur(CANONE_MENSILE), _pct(FEE_PERCENTUALE)))
    for g in SCENARI_GIORNO:
        lordo = g * GIORNI_MESE
        _i, _q, tot, inc = _riga_scenario(lordo)
        print('     %5d €/gg -> %9s €/mese lordo -> %8s €/mese (%0.2f%%)'
              % (g, _eur(lordo), _eur(tot), inc))


if __name__ == '__main__':
    main()
